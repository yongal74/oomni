# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "C:/Users/장우경/oomni/n8n_cases_full.docx"
doc = Document(PATH)

def sc(doc, text, level=1):
    colors = {1: RGBColor(0x1F,0x49,0x7D), 2: RGBColor(0x2E,0x74,0xB5), 3: RGBColor(0x1F,0x49,0x7D)}
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(sizes[level]); run.font.bold = True; run.font.color.rgb = colors[level]
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after = Pt(6)

def par(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

def bul(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run("  • " + item)
        run.font.size = Pt(10); run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)

def flow(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(9); run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    p._element.pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)

def case(doc, num, title, desc, nodes, flowchart, tips):
    sc(doc, f"[{num}] {title}", 2)
    par(doc, desc)
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run("▶ 주요 노드: "); run.font.bold = True; run.font.size = Pt(10)
    run.font.name = '맑은 고딕'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run2 = p.add_run(nodes); run2.font.size = Pt(10); run2.font.name = '맑은 고딕'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    p2 = doc.add_paragraph(); p2.style = doc.styles['Normal']
    run3 = p2.add_run("▶ 워크플로우 순서도"); run3.font.bold = True; run3.font.size = Pt(10)
    run3.font.name = '맑은 고딕'; run3._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p2.paragraph_format.space_before = Pt(4); p2.paragraph_format.space_after = Pt(2)
    flow(doc, flowchart)
    p3 = doc.add_paragraph(); p3.style = doc.styles['Normal']
    run4 = p3.add_run("▶ 핵심 포인트"); run4.font.bold = True; run4.font.size = Pt(10)
    run4.font.name = '맑은 고딕'; run4._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p3.paragraph_format.space_before = Pt(4); p3.paragraph_format.space_after = Pt(2)
    bul(doc, tips)
    doc.add_paragraph()

# ─── CH3: HR 자동화 (021~030) ─────────────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 3: HR 자동화", 1)
par(doc, "채용, 온보딩, 급여, 연차 관리 등 HR 업무를 자동화하여 인사팀의 행정 부담을 줄이고 직원 경험을 향상시킵니다.")

case(doc, "021", "채용 지원서 자동 접수 및 분류",
"이메일로 수신된 이력서를 AI로 분석하여 적합도를 평가하고 HR 담당자에게 요약 리포트를 발송합니다.",
"Gmail Trigger, OpenAI, Google Sheets, Slack",
"""Gmail Trigger (채용@ 이메일 + 첨부파일)
  │
  ▼
Google Drive (이력서 PDF 저장)
  │
  ▼
OpenAI (이력서 분석)
  │ - 핵심 스킬 추출
  │ - 요구 자격 충족도 평가 (0~100점)
  │ - 강점/약점 요약
  ▼
Google Sheets (지원자 DB 기록)
  │ 컬럼: 이름, 이메일, 점수, 강점, 약점, 상태
  ▼
[IF] 적합도 점수 >= 70?
  │ YES                    │ NO
  ▼                        ▼
Slack (HR팀 알림)         Gmail (불합격 자동 회신)
"📄 우수 지원자: {{name}}"   (2주 후 발송 예약)""",
["OpenAI 분석에 채용 포지션 JD를 컨텍스트로 제공하면 정확도 향상", "불합격 이메일은 즉시 발송보다 2주 후 발송이 더 자연스러움", "지원자 데이터는 GDPR/개인정보보호법에 따라 보관 기간 설정"])

case(doc, "022", "신입 직원 온보딩 자동화",
"신규 입사자 정보 입력 시 IT 계정 신청, 장비 요청, 온보딩 문서 공유 등을 자동화합니다.",
"Google Forms, Slack, Gmail, Jira, Google Drive",
"""Google Forms (신규 입사 정보 입력)
  │ 입력: 이름, 부서, 직책, 입사일, 이메일
  ▼
[병렬 처리]
  ├─ Slack (HR채널: 신규 입사자 소개)
  ├─ Gmail → IT팀 (계정 생성 요청)
  ├─ Jira (온보딩 태스크 생성)
  │   - 노트북 세팅, 출입증 발급
  │   - 사내 시스템 교육 일정
  └─ Google Drive (온보딩 패키지 공유)
       - 사내 규정, 조직도, 업무 매뉴얼
  ▼
Gmail (신규 직원에게 환영 이메일)
  └─ 첫날 일정, 입사 준비물 안내
  ▼
[완료]""",
["온보딩 채크리스트를 Jira/Notion으로 관리하면 진행 상황 추적 용이", "IT 계정 신청은 입사 3일 전에 자동 트리거되도록 설정 권장", "멘토 자동 배정: 같은 부서 1년 이상 직원 중 랜덤 선택"])

case(doc, "023", "연차/휴가 신청 자동 처리",
"직원이 연차 신청 폼을 제출하면 관리자 승인 요청, 캘린더 차단, 팀 공유 캘린더 업데이트를 자동 처리합니다.",
"Google Forms, Gmail, Google Calendar, Slack",
"""Google Forms (연차 신청)
  │ 입력: 신청자, 유형, 시작일, 종료일, 사유
  ▼
Code 노드 (잔여 연차 확인)
  │ - Google Sheets에서 잔여 일수 조회
  │ - 신청 일수 계산
  ▼
[IF] 잔여 연차 충분?
  │ YES                    │ NO
  ▼                        ▼
Gmail (관리자 승인 요청)  Gmail (신청자에게 안내)
  │ - 승인/거절 링크 포함
  ▼
[Wait] 관리자 승인 클릭
  │ 승인             │ 거절
  ▼                  ▼
Google Calendar    Gmail (거절 안내)
(신청자 캘린더 차단)
  ↓
팀 공유 캘린더 업데이트
  ↓
Google Sheets (연차 잔여 차감)
  ↓
Slack (팀채널: 연차 공지)""",
["관리자 승인 링크는 JWT 토큰으로 보안 처리", "반차/경조사 등 유형별 다른 처리 로직 적용", "연차 잔여일 경고: 5일 이하 시 연말 사용 권고 알림"])

case(doc, "024", "급여 명세서 자동 발송",
"매월 급여 지급일 전날 자동으로 각 직원에게 개인화된 급여 명세서 이메일을 발송합니다.",
"Schedule Trigger, Google Sheets, Code, Gmail",
"""Schedule Trigger (매월 24일 오전 9시)
  │
  ▼
Google Sheets (급여 데이터 읽기)
  │ 시트: 당월_급여
  │ 컬럼: 사원번호, 이름, 이메일, 기본급, 각종수당, 공제
  ▼
[Loop] 각 직원 처리
  │
  ▼
Code 노드 (급여 계산)
  │ - 총지급액 계산
  │ - 4대보험/소득세 공제
  │ - 실수령액 계산
  ▼
Gmail (개인화 급여명세서 이메일)
  │ 제목: "[기밀] {{이름}}님 {{월}}월 급여명세서"
  │ 본문: HTML 테이블 형식 명세서
  ▼
Google Sheets (발송 완료 기록)""",
["급여 데이터는 구글 스프레드시트 접근 권한을 최소화하여 보안 강화", "이메일 본문에 HTML 테이블 사용 시 반응형 디자인 고려", "발송 실패 시 즉시 HR 담당자에게 알림 처리"])

case(doc, "025", "직원 생일 자동 축하",
"직원 생일에 Slack으로 팀 전체에 생일 축하 메시지를 발송하고 HR은 소정의 기념품 발송을 안내받습니다.",
"Schedule Trigger, Google Sheets, Slack, Gmail",
"""Schedule Trigger (매일 오전 9시)
  │
  ▼
Google Sheets (오늘 생일 직원 조회)
  │
  ▼
[IF] 생일 직원 존재?
  │ YES
  ▼
[Loop] 각 생일자 처리
  │
  ├─ Slack (#전체채널)
  │   "🎂 {{이름}}님 오늘 생일이에요! 축하해주세요! 🎉"
  │
  └─ Gmail (HR 담당자)
      "{{이름}}님 생일입니다. 기념품 발송 검토 바랍니다."
  ▼
[완료]""",
["주말 생일은 금요일 오전에 미리 발송", "생일 메시지에 해당 직원의 팀 기여 내용 포함하면 더욱 의미 있음", "생일 정보는 Google Sheets 대신 HR 시스템 API 연동 권장"])

case(doc, "026", "성과 리뷰 사이클 자동화",
"반기 성과 리뷰 기간이 되면 자동으로 자기평가 폼을 배포하고 관리자 평가를 수집하며 HR에 완료 현황을 보고합니다.",
"Schedule Trigger, Google Forms, Gmail, Google Sheets, Slack",
"""Schedule Trigger (6월 1일, 12월 1일)
  │
  ▼
Google Sheets (전체 직원 목록 조회)
  │
  ▼
[Loop] 각 직원 처리
  │
  ├─ Gmail (자기평가 폼 링크 발송)
  │   마감: D+14
  └─ Gmail (관리자에게 팀원 평가 링크 발송)
  ▼
Schedule Trigger (D+7: 미완료자 리마인더)
  │
  ▼
Schedule Trigger (D+14: 마감 처리)
  │
  ▼
Code 노드 (완료율 집계)
  ▼
Slack (HR팀: 완료 현황 리포트)
  └─ [성과평가 완료율: {{rate}}%
      미완료: {{pending_list}}]""",
["마감 1일 전 미완료자에게 자동 리마인더 발송", "평가 완료 시 즉시 확인 이메일 발송으로 누락 방지", "데이터 집계 후 대시보드(Google Data Studio)와 연동 권장"])

case(doc, "027", "교육/훈련 이수 추적 자동화",
"직원들의 필수 교육 이수 현황을 추적하고 미이수자에게 자동 리마인더를 발송합니다.",
"Schedule Trigger, Google Sheets, Gmail, Slack",
"""Schedule Trigger (매주 월요일 오전 10시)
  │
  ▼
Google Sheets (교육 이수 현황 조회)
  │ - 필수 교육 목록
  │ - 각 직원별 이수 여부
  │ - 마감일
  ▼
Code 노드 (미이수 + 마감 임박 분류)
  │ - D-7 이하: 긴급
  │ - D-14~8: 주의
  │ - D-30~15: 안내
  ▼
[Loop] 미이수자 처리
  │
  ├─ Gmail (개인 리마인더 이메일)
  │   "{{교육명}} 이수 마감이 {{days}}일 남았습니다"
  └─ Slack (부서장에게 현황 공유)
  ▼
Schedule (D-1: 최종 경고 발송)""",
["교육 이수율을 월별로 추적하여 KPI 대시보드 연동", "마감 초과 시 자동으로 HR에 에스컬레이션", "온라인 교육 플랫폼 API와 연동하면 이수 여부 자동 갱신 가능"])

case(doc, "028", "직원 설문조사 자동화",
"분기별 직원 만족도 설문을 자동 배포하고 결과를 집계하여 경영진 리포트를 생성합니다.",
"Schedule Trigger, Gmail, Google Forms, Code, Slack",
"""Schedule Trigger (분기 첫날)
  │
  ▼
Google Sheets (직원 목록 조회)
  │
  ▼
Gmail (설문 링크 개인화 발송)
  │ - 개인별 고유 링크 (익명 보장)
  │ - 응답 기한: D+14
  ▼
Schedule (D+7: 미응답자 리마인더)
  │
  ▼
Schedule (D+14: 집계 처리)
  │
  ▼
Google Sheets (응답 데이터 분석)
  │
  ▼
Code 노드 (카테고리별 점수 계산)
  │ - eNPS, 직무 만족도, 성장 기회 등
  ▼
Slack (경영진 채널: 결과 요약)
  └─ 전분기 대비 변화 포함""",
["익명성 보장을 위해 개인 식별 정보를 결과에서 분리", "응답률 80% 이상 달성 목표 설정 및 추적", "낮은 점수 항목에 대한 액션 플랜 자동 제안 (AI 활용)"])

case(doc, "029", "퇴사자 오프보딩 체크리스트",
"직원 퇴사 처리 시 IT 계정 비활성화, 장비 반납, 권한 회수 등의 체크리스트를 자동 생성하고 처리합니다.",
"Google Forms, Jira, Gmail, Slack",
"""Google Forms (퇴사 신청 접수)
  │ 입력: 이름, 부서, 퇴사일, 이유
  ▼
Jira (오프보딩 태스크 생성)
  │ 담당자별 태스크:
  ├─ IT팀: 계정 비활성화 (퇴사일)
  ├─ IT팀: 장비 반납 확인
  ├─ HR팀: 퇴직금 정산
  ├─ 재무팀: 경비 정산 확인
  └─ 관리팀: 출입증 반납
  ▼
Gmail (퇴사자에게 오프보딩 안내)
  │ - 반납 절차, 인수인계 일정
  ▼
Slack (각 담당 팀에 알림)
  ▼
Schedule (퇴사일 D-1: 최종 확인)""",
["계정 비활성화는 퇴사일 당일 근무 시간 종료 후 처리", "업무 인수인계 문서화를 Confluence/Notion 페이지 자동 생성", "퇴사 사유 데이터를 누적하여 이직률 트렌드 분석에 활용"])

case(doc, "030", "채용 공고 자동 멀티 플랫폼 게시",
"HR 시스템에 채용 공고를 등록하면 자동으로 여러 채용 플랫폼과 회사 SNS에 동시 게시합니다.",
"Webhook, HTTP Request, Twitter/LinkedIn/Slack",
"""Webhook (HR 시스템: 채용 공고 등록)
  │ 데이터: 포지션, 부서, 요건, 마감일
  ▼
Code 노드 (플랫폼별 포맷 변환)
  │
  ▼
[병렬 발행]
  ├─ HTTP Request (원티드 API)
  ├─ HTTP Request (사람인 API)
  ├─ LinkedIn (Company Page 포스팅)
  ├─ Slack (#채용채널: 내부 추천 요청)
  └─ Gmail (전 직원: 지인 추천 요청)
  ▼
Google Sheets (게시 현황 기록)
  └─ 각 플랫폼 게시 URL, 게시일 기록
  ▼
[완료]""",
["내부 추천제 활성화 시 Slack 공지에 추천 보상 내용 포함", "각 플랫폼 API 한도와 포스팅 형식을 사전 확인 필수", "채용 마감일 D-3에 자동 리마인더 재게시 설정 권장"])

# ─── CH4: 마케팅 자동화 (031~042) ─────────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 4: 마케팅 자동화", 1)
par(doc, "콘텐츠 발행, 광고 관리, 리드 너처링, 분석 자동화로 마케팅 팀의 생산성을 높이고 캠페인 성과를 향상시킵니다.")

case(doc, "031", "블로그 포스트 → 소셜 미디어 자동 발행",
"블로그 새 포스트가 게시되면 Twitter, LinkedIn, Facebook에 맞춤화된 형식으로 자동 공유합니다.",
"RSS Feed Trigger, OpenAI, HTTP Request (Social APIs)",
"""RSS Feed Trigger (블로그 새 포스트)
  │
  ▼
OpenAI (플랫폼별 게시글 생성)
  │ - Twitter: 280자 요약 + 해시태그
  │ - LinkedIn: 전문적 요약 + CTA
  │ - Facebook: 캐주얼한 소개글
  ▼
[병렬 발행]
  ├─ Twitter API (트윗 게시)
  ├─ LinkedIn API (포스팅)
  └─ Facebook Graph API (페이지 포스팅)
  ▼
Google Sheets (발행 현황 기록)
  └─ 포스트 URL, 각 플랫폼 링크, 발행 시각
  ▼
[완료]""",
["각 플랫폼의 최적 게시 시간(Twitter 오전 9시, LinkedIn 오전 7~8시)에 예약 발행", "이미지는 블로그 썸네일을 재사용하거나 Canva API로 플랫폼별 리사이징", "해시태그는 업계 트렌드 분석으로 주기적 업데이트 권장"])

case(doc, "032", "광고 성과 자동 리포팅",
"Google Ads, Meta Ads 등 광고 플랫폼의 일일 성과를 자동 수집하여 대시보드와 Slack에 리포트합니다.",
"Schedule Trigger, Google Ads API, Meta API, Google Sheets, Slack",
"""Schedule Trigger (매일 오전 8시)
  │
  ▼
[병렬 데이터 수집]
  ├─ Google Ads API (전일 성과)
  │   - 노출, 클릭, 전환, 비용
  ├─ Meta Ads API (전일 성과)
  └─ Naver 검색광고 API (전일 성과)
  ▼
Code 노드 (통합 집계 + KPI 계산)
  │ - 총 광고비, CPC, CPA, ROAS
  │ - 전일/전주 대비 변화율
  ▼
Google Sheets (데이터 기록)
  │
  ▼
Code 노드 (Slack 블록 포매팅)
  │
  ▼
Slack (#마케팅채널)
  └─ "📊 전일 광고 현황: ROAS {{roas}}, CPA {{cpa}}"
     이상값 감지 시 🚨 경고 포함""",
["ROAS < 목표치 시 자동으로 마케팅 매니저에게 별도 알림 발송", "월간 예산 소진율 계산하여 과소비/과절감 감지", "Google Data Studio 연동으로 실시간 대시보드 자동화"])

case(doc, "033", "이벤트 참가자 자동 관리",
"행사 등록 시 확인 이메일, 리마인더, 입장 QR코드 발송까지 전체 참가자 관리를 자동화합니다.",
"Webhook (이벤투브라이트/자체폼), Gmail, Airtable",
"""Webhook (이벤트 등록)
  │
  ▼
Airtable (참가자 정보 저장)
  │
  ▼
Code 노드 (QR 코드 생성)
  │ - 참가자별 고유 QR 코드
  │ - 이벤트 ID + 참가자 ID 인코딩
  ▼
Gmail (등록 확인 이메일)
  │ - QR 코드 첨부
  │ - 행사 상세 안내
  ▼
Schedule (D-7: 사전 안내 이메일)
  ▼
Schedule (D-1: 최종 리마인더)
  └─ 주차, 준비물, 당일 일정 안내
  ▼
Schedule (D+1: 감사 이메일 + 자료 공유)""",
["이벤트 정원 도달 시 자동 대기자 명단 등록 처리", "QR 코드 스캔으로 현장 체크인 연동 가능", "참가율 추적하여 다음 이벤트 참가 독려 이메일 분기 처리"])

case(doc, "034", "리드 스코어링 자동화",
"리드의 행동 데이터(이메일 오픈, 웹 방문, 콘텐츠 다운로드)를 수집하여 자동으로 점수를 계산합니다.",
"Webhook, HubSpot, Code, Slack",
"""Webhook (다양한 행동 이벤트)
  │ - 이메일 오픈, 클릭
  │ - 웹페이지 방문
  │ - 콘텐츠 다운로드
  │ - 무료 체험 신청
  ▼
HubSpot (리드 조회)
  │
  ▼
Code 노드 (행동별 점수 가산)
  │ 이메일 오픈: +2점
  │ 이메일 클릭: +5점
  │ 가격 페이지 방문: +10점
  │ 콘텐츠 다운로드: +7점
  │ 무료체험 신청: +25점
  ▼
HubSpot (리드 점수 업데이트)
  │
  ▼
[IF] 점수 >= 50? (MQL 임계값)
  │ YES
  ▼
Slack (영업팀 알림)
  └─ "🔥 MQL 전환: {{name}} (점수: {{score}})"
     즉시 연락 권장""",
["점수 감소 로직도 추가: 30일 무활동 시 -10점", "직책/회사 규모 등 인구통계 점수와 행동 점수를 분리 관리", "MQL 임계값은 실제 전환율 데이터로 지속 보정 필요"])

case(doc, "035", "경쟁사 가격 모니터링",
"경쟁사 웹사이트에서 가격 정보를 주기적으로 수집하여 변동 시 마케팅/영업팀에 알립니다.",
"Schedule Trigger, HTTP Request, Code, Google Sheets, Slack",
"""Schedule Trigger (매일 오전 7시)
  │
  ▼
[Loop] 경쟁사 목록
  │
  ▼
HTTP Request (경쟁사 가격 페이지 스크래핑)
  │ (또는 가격 비교 API 활용)
  ▼
Code 노드 (가격 데이터 파싱)
  │
  ▼
Google Sheets (전일 가격과 비교)
  │
  ▼
[IF] 가격 변동 감지?
  │ YES
  ▼
Slack (#마케팅채널)
  └─ "⚠️ 경쟁사 가격 변동: {{company}}
     {{product}}: {{old_price}} → {{new_price}}"
  ▼
Google Sheets (변동 이력 기록)""",
["법적 문제 방지를 위해 공개된 정보만 수집 (로그인 불필요 페이지)", "가격 변동이 빈번한 시즌에는 수집 주기를 증가", "경쟁사 가격 대비 자사 포지셔닝 자동 계산 추가 가능"])

case(doc, "036", "웹사이트 이탈 방문자 재타겟팅",
"특정 페이지를 방문하고 전환하지 않은 방문자를 감지하여 개인화된 이메일을 발송합니다.",
"Webhook (웹사이트 이벤트), HubSpot, Schedule, Gmail",
"""Webhook (웹사이트 이벤트)
  │ - 가격 페이지 방문 이벤트
  │ - 데모 페이지 방문 이벤트
  ▼
HubSpot (방문자 이메일 주소 확인)
  │ (이메일로 식별된 방문자만)
  ▼
HubSpot (방문 기록 저장)
  │
  ▼
Schedule (24시간 후: 전환 여부 확인)
  │
  ▼
[IF] 24시간 내 전환 없음?
  │ YES
  ▼
Gmail (재타겟팅 이메일)
  └─ 방문 페이지 맞춤 내용
     가격 페이지 → 특별 할인 제안
     데모 페이지 → 데모 예약 링크
  ▼
HubSpot (재타겟팅 활동 기록)""",
["이메일 식별이 안 된 익명 방문자는 쿠키 기반 광고 재타겟팅 활용", "재타겟팅 이메일 발송 빈도 조절 (최대 주 2회)로 opt-out 방지", "이메일 오픈/클릭 후 전환 여부 추적으로 효과 측정"])

case(doc, "037", "A/B 테스트 자동화",
"이메일 캠페인의 A/B 테스트를 자동화하여 성과가 좋은 버전을 전체 발송에 자동 적용합니다.",
"Schedule Trigger, Google Sheets, Gmail, Code",
"""[캠페인 등록 단계]
Google Sheets에 A/B 버전 정보 입력
  - 제목 A, 제목 B, 발송 비율 (10%씩)
  │
  ▼
Schedule Trigger (캠페인 시작일)
  │
  ▼
전체 수신자 목록 → 랜덤 10% 선택 × 2그룹
  │
  ▼
[병렬 발송]
  ├─ Group A: Gmail (제목 A 발송)
  └─ Group B: Gmail (제목 B 발송)
  │
  ▼
Wait (24시간)
  │
  ▼
Code 노드 (A/B 오픈율 비교)
  │
  ▼
[IF] 통계적 유의미한 차이 있음?
  │ YES
  ▼
Gmail (승리 버전으로 나머지 80% 발송)""",
["샘플 크기가 작으면 통계적 신뢰도가 낮으므로 최소 200명 이상 필요", "오픈율 외 클릭율, 전환율도 함께 비교하는 것이 바람직", "테스트 변수는 한 번에 하나씩 (제목 OR 발신자 OR 본문)"])

case(doc, "038", "고객 리뷰 수집 및 SNS 공유",
"구매 완료 후 자동으로 리뷰 요청을 발송하고 좋은 리뷰를 소셜 미디어에 자동 공유합니다.",
"Webhook (구매 완료), Gmail, Schedule, OpenAI, Twitter",
"""Webhook (구매 완료 이벤트)
  │
  ▼
Schedule (D+7: 리뷰 요청 이메일)
  │
  ▼
[Wait] 리뷰 작성 Webhook
  │
  ▼
OpenAI (리뷰 감정 분석 + 별점 확인)
  │
  ├─ 별점 4~5점 → 소셜 공유 진행
  │   │
  │   ▼
  │  OpenAI (SNS 게시글 생성)
  │   └─ 고객 이름, 리뷰 발췌, 제품명
  │   ▼
  │  Twitter/LinkedIn (자동 게시)
  │
  └─ 별점 1~3점 → CS팀 즉시 알림
      "⚠️ 낮은 평점 리뷰 접수: {{customer}}"
      즉각 대응 요청""",
["리뷰 게시 전 고객 동의 확인 (약관에 포함하거나 별도 동의 받기)", "낮은 평점 즉각 대응이 고객 이탈 방지의 핵심", "월간 리뷰 통계 (평균 별점 트렌드) 자동 리포트 생성"])

case(doc, "039", "뉴스레터 성과 분석 자동화",
"발송한 뉴스레터의 오픈율, 클릭율, 구독 취소율 등을 자동으로 분석하여 인사이트를 제공합니다.",
"Schedule Trigger, Mailchimp API, Code, Google Sheets, Slack",
"""Schedule Trigger (매주 월요일 오전 9시)
  │
  ▼
Mailchimp API (지난 주 캠페인 성과 조회)
  │ - 발송수, 오픈율, 클릭율
  │ - 구독취소율, 반송율
  │ - 클릭 링크별 성과
  ▼
Code 노드 (업계 평균 대비 분석)
  │ 업계 평균: 오픈율 21%, 클릭율 2.5%
  │ - 평균 이상/이하 판단
  │ - 트렌드 분석 (4주 이동 평균)
  ▼
Google Sheets (이력 기록)
  │
  ▼
OpenAI (인사이트 및 개선점 제안)
  │ "오픈율이 하락 추세입니다. 제목 A/B 테스트를 고려하세요"
  ▼
Slack (#마케팅채널)
  └─ 주간 뉴스레터 성과 요약""",
["구독 취소율이 0.5% 초과 시 즉시 콘텐츠 전략 재검토 알림", "가장 클릭 많은 콘텐츠 유형 파악하여 다음 호에 활용", "A/B 테스트 이력과 비교하여 최적 전략 도출"])

case(doc, "040", "SEO 순위 모니터링 자동화",
"주요 키워드의 Google 검색 순위를 주기적으로 추적하고 순위 변동 시 마케팅팀에 알립니다.",
"Schedule Trigger, HTTP Request (SEO API), Google Sheets, Slack",
"""Schedule Trigger (매주 월요일)
  │
  ▼
HTTP Request (Google Search Console API)
  │ 또는 SEMrush/Ahrefs API
  │ - 주요 키워드 10~50개 순위 조회
  ▼
Code 노드 (전주 대비 순위 변동 계산)
  │
  ▼
Google Sheets (순위 이력 기록)
  │
  ▼
[IF] 주요 키워드 순위 변동 감지?
  │ 상승                   │ 하락
  ▼                        ▼
Slack (긍정 알림)         Slack (경고 알림)
"📈 {{keyword}}: 3위↑"  "📉 {{keyword}}: 5위↓"
  ▼
월간 SEO 트렌드 리포트 자동 생성 (매월 1일)""",
["순위 변동이 큰 키워드는 알고리즘 업데이트 여부 별도 확인", "경쟁사의 같은 키워드 순위도 함께 추적하면 상대적 포지션 파악 가능", "하락 키워드에 대한 SEO 액션 플랜 자동 체크리스트 생성"])

case(doc, "041", "인플루언서 캠페인 추적 자동화",
"인플루언서 협업 캠페인의 게시물, 도달, 참여율을 자동 추적하여 ROI를 측정합니다.",
"Schedule Trigger, Instagram/YouTube API, Google Sheets, Slack",
"""Schedule Trigger (매일 오전 8시)
  │
  ▼
Google Sheets (진행 중 캠페인 + 인플루언서 목록 조회)
  │
  ▼
[Loop] 각 인플루언서 처리
  │
  ▼
Instagram/YouTube API (게시물 성과 조회)
  │ - 좋아요, 댓글, 조회수
  │ - 저장수, 공유수
  │ - 클릭 및 전환 (UTM 추적)
  ▼
Code 노드 (CPE, ROAS 계산)
  │
  ▼
Google Sheets (일별 성과 기록)
  │
  ▼
[IF] 성과 목표 80% 미달?
  │ YES
  ▼
Slack (캠페인 매니저 알림)
  └─ 추가 콘텐츠 협의 필요""",
["UTM 파라미터로 인플루언서별 전환 추적 구현 필수", "캠페인 종료 후 최종 ROI 자동 산출 및 파트너사 공유", "성과 우수 인플루언서 자동 태그로 재협업 대상 관리"])

case(doc, "042", "콘텐츠 캘린더 자동 관리",
"Google Sheets 콘텐츠 캘린더를 기반으로 제작 리마인더, 검토 요청, 발행 알림을 자동화합니다.",
"Schedule Trigger, Google Sheets, Slack, Gmail",
"""Schedule Trigger (매일 오전 9시)
  │
  ▼
Google Sheets (이번 주 콘텐츠 캘린더 조회)
  │
  ▼
Code 노드 (D+7, D+3, D+1 항목 분류)
  │
  ▼
[Loop] 각 콘텐츠 항목 처리
  │
  ├─ D+7 (7일 후 발행): Slack (담당자)
  │   "📝 {{제목}} 초안 작성 시작하세요"
  ├─ D+3 (3일 후 발행): Slack (에디터)
  │   "✏️ {{제목}} 검토 요청 (마감: D+2)"
  └─ D+0 (오늘 발행): Slack (담당자)
      "🚀 {{제목}} 오늘 발행 예정입니다"
  ▼
Google Sheets (알림 발송 기록 업데이트)""",
["콘텐츠 발행 확인 후 체크박스 자동 업데이트로 현황 실시간 파악", "발행 지연 3일 이상 시 콘텐츠 매니저에게 에스컬레이션", "월별 발행 수, 유형별 분포 자동 통계 생성"])

doc.save(PATH)
print("[OK] gen_vol1_revised_p2.py 완료 - Ch3(HR)+Ch4(마케팅) 케이스 021-042 추가")
