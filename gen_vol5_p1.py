# -*- coding: utf-8 -*-
"""Vol 5 Part 1: 표지 + Ch1 n8n AI 노드 완전 가이드 + Ch2 LLM 모델별 가이드"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = 'C:/Users/장우경/oomni/n8n_cases_vol5.docx'
doc = Document()

for section in doc.sections:
    section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5)
    section.left_margin=Cm(3.0); section.right_margin=Cm(2.5)

def sc(cell, hx):
    tc=cell._tc; pr=tc.get_or_add_tcPr()
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hx); pr.append(s)

def h(doc,text,lv=1):
    p=doc.add_heading(text,level=lv)
    C={1:RGBColor(0x1F,0x49,0x7D),2:RGBColor(0x2E,0x75,0xB6),3:RGBColor(0x17,0x5E,0x40),4:RGBColor(0x40,0x40,0x40)}
    S={1:Pt(22),2:Pt(16),3:Pt(13),4:Pt(11)}
    if p.runs: p.runs[0].font.name='맑은 고딕'; p.runs[0].font.size=S.get(lv,Pt(11)); p.runs[0].font.color.rgb=C.get(lv,RGBColor(0,0,0))
    return p

def par(doc,text,bold=False,color=None,size=Pt(10),indent=None):
    p=doc.add_paragraph()
    if indent: p.paragraph_format.left_indent=indent
    r=p.add_run(text); r.font.name='맑은 고딕'; r.font.size=size; r.bold=bold
    if color: r.font.color.rgb=color
    return p

def bul(doc,text,lv=0):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent=Cm(0.5+lv*0.5)
    r=p.add_run(text); r.font.name='맑은 고딕'; r.font.size=Pt(10); return p

def code(doc,text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run(text); r.font.name='Courier New'; r.font.size=Pt(8)
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),'F2F2F2'); pr.append(s)
    return p

def tbl(doc, headers, rows, col_widths, header_color='1F497D'):
    t=doc.add_table(rows=len(rows)+1, cols=len(headers)); t.style='Table Grid'; t.autofit=False
    for i,w in enumerate(col_widths): t.columns[i].width=w
    for ci,hdr in enumerate(headers):
        sc(t.rows[0].cells[ci], header_color)
        t.rows[0].cells[ci].text=hdr
        for p in t.rows[0].cells[ci].paragraphs:
            for r in p.runs: r.font.name='맑은 고딕'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            sc(t.rows[ri+1].cells[ci], 'EBF3FB' if ci==0 else 'FFFFFF')
            t.rows[ri+1].cells[ci].text=str(val)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.name='맑은 고딕'; r.font.size=Pt(9)
                if ci==0:
                    for r in p.runs: r.bold=True
    return t

def divider(doc):
    p=doc.add_paragraph(); pr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr')
    bt=OxmlElement('w:bottom'); bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'4')
    bt.set(qn('w:space'),'1'); bt.set(qn('w:color'),'DDDDDD'); pb.append(bt); pr.append(pb)

# ══════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(72)
r=p.add_run('n8n AI 에이전트 완전 가이드')
r.font.name='맑은 고딕'; r.font.size=Pt(30); r.bold=True; r.font.color.rgb=RGBColor(0x1F,0x49,0x7D)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('6,380개 패턴 분석 | LLM 모델별 최적 전략 | 멀티에이전트 설계')
r.font.name='맑은 고딕'; r.font.size=Pt(16); r.font.color.rgb=RGBColor(0x2E,0x75,0xB6)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Volume 5 — n8n 마스터 시리즈')
r.font.name='맑은 고딕'; r.font.size=Pt(13); r.font.color.rgb=RGBColor(0x70,0xAD,0x47)

doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('n8n AI Agent 노드 완전 분해 | RAG 심화 | 프롬프트 엔지니어링\n'
            'Human-in-the-loop | 프로덕션 운영 | 비용 최적화')
r.font.name='맑은 고딕'; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x40,0x40,0x40)

doc.add_page_break()

# 목차
h(doc,'목차',1)
toc=[
    ('Chapter 1','n8n AI 노드 완전 가이드','AI Agent, LangChain 노드 25종 전체 분해'),
    ('Chapter 2','LLM 모델별 최적 전략','GPT-4o/Claude/Gemini/Llama/Mistral 비교 매트릭스'),
    ('Chapter 3','멀티에이전트 10종 설계 패턴','SupervisorAgent, ParallelAgent 등 고급 패턴'),
    ('Chapter 4','RAG 심화 완전 가이드','벡터DB 5종, 청킹 전략, 임베딩 최적화'),
    ('Chapter 5','6,380개 AI 워크플로우 패턴 분석','커뮤니티 패턴 심층 분류 및 트렌드'),
    ('Chapter 6','프롬프트 엔지니어링 완전 가이드','시스템 프롬프트, Chain-of-Thought, Few-shot'),
    ('Chapter 7','Human-in-the-loop 설계 패턴','승인 게이트, 피드백 루프, RLHF 연계'),
    ('Chapter 8','프로덕션 AI 운영 가이드','비용 관리, 모니터링, 보안, 규제 대응'),
    ('부록','참조 자료 및 체크리스트','노드 연결 치트시트, 디버깅 가이드'),
]
t=doc.add_table(rows=len(toc)+1, cols=3); t.style='Table Grid'; t.autofit=False
for w,ci in zip([Cm(2.5),Cm(5),Cm(10)],range(3)): t.columns[ci].width=w
for ci,hdr in enumerate(['챕터','제목','내용 요약']):
    sc(t.rows[0].cells[ci],'1F497D')
    t.rows[0].cells[ci].text=hdr
    for p in t.rows[0].cells[ci].paragraphs:
        for r in p.runs: r.font.name='맑은 고딕'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for i,(ch,title,desc) in enumerate(toc):
    sc(t.rows[i+1].cells[0],'EBF3FB')
    for ci,val in enumerate([ch,title,desc]):
        t.rows[i+1].cells[ci].text=val
        for p in t.rows[i+1].cells[ci].paragraphs:
            for r in p.runs: r.font.name='맑은 고딕'; r.font.size=Pt(9)
            if ci==0:
                for r in p.runs: r.bold=True

doc.add_page_break()

# ══════════════════════════════════════════════════════
# CHAPTER 1: n8n AI 노드 완전 가이드
# ══════════════════════════════════════════════════════
h(doc,'Chapter 1. n8n AI 노드 완전 가이드',1)
par(doc,'n8n v1.22+ 기준으로 사용 가능한 AI/LangChain 노드 25종을 완전 분해합니다. '
    '각 노드의 역할, 연결 포트, 필수 파라미터, 실전 활용법을 체계적으로 정리했습니다.', size=Pt(10))
doc.add_paragraph()

h(doc,'1.1 n8n AI 아키텍처 개요',2)
par(doc,'n8n의 AI 시스템은 LangChain Expression Language(LCEL) 기반으로 구축되어 있으며, '
    '모든 AI 노드는 특수 연결 포트(ai_*)를 통해 계층적으로 연결됩니다.', size=Pt(10))
doc.add_paragraph()

code(doc,
'n8n AI 노드 연결 포트 체계\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'  ai_languageModel   ──▶  AI Agent / Chain 노드에 LLM 공급\n'
'  ai_tool            ──▶  AI Agent에 도구 기능 추가\n'
'  ai_memory          ──▶  AI Agent에 대화 메모리 추가\n'
'  ai_vectorStore     ──▶  Vector Store Tool에 DB 연결\n'
'  ai_embedding       ──▶  Vector Store에 임베딩 모델 공급\n'
'  ai_document        ──▶  Document Loader에서 문서 공급\n'
'  ai_textSplitter    ──▶  Document에 청킹 전략 적용\n'
'  ai_outputParser    ──▶  LLM 출력 구조화 파싱\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'1.2 AI Agent 노드 (핵심)',2)
par(doc,'AI Agent는 n8n AI 시스템의 중심 노드입니다. LLM + Tool + Memory를 조합하여 '
    '자율적으로 작업을 수행하는 에이전트를 구성합니다.', size=Pt(10))
doc.add_paragraph()

h(doc,'1.2.1 Agent 타입별 특성',3)
agent_types=[
    ('toolsAgent (기본)','도구 호출 최적화. Function Calling 지원 모델에서 사용. 대부분의 사용 케이스에 적합.','GPT-4o, Claude 3.5+, Gemini 1.5+'),
    ('conversationalAgent','대화 이력 유지 최적화. Buffer Memory 필수 연결. 고객 서비스 챗봇에 적합.','모든 Chat 모델'),
    ('openAiFunctionsAgent','OpenAI Function Calling 전용. toolsAgent보다 이전 방식. GPT-3.5/4에서 사용.','GPT-3.5-turbo, GPT-4'),
    ('reActAgent','Reasoning+Acting 반복 패턴. 복잡한 다단계 추론에 적합. 느리지만 정확.','Claude 3, GPT-4o'),
    ('planAndExecuteAgent','계획 수립 후 실행. 긴 태스크 분해에 적합. 실험적 기능.','GPT-4o, Claude 3 Opus'),
    ('sqlAgent','SQL 쿼리 자동 생성 및 실행 전문. 데이터베이스 Q&A에 특화.','GPT-4o, Claude 3.5'),
]
tbl(doc,['Agent 타입','특성 및 적합 용도','지원 모델'],agent_types,
    [Cm(4.5),Cm(8.5),Cm(4.5)],header_color='2E75B6')
doc.add_paragraph()

h(doc,'1.2.2 AI Agent 핵심 파라미터',3)
params=[
    ('systemMessage','에이전트의 역할, 페르소나, 행동 규칙 정의. 가장 중요한 파라미터.','string','당신은 고객지원 전문가입니다...'),
    ('maxIterations','최대 Tool 호출 횟수. 무한 루프 방지. 기본값 10.','number','10 (단순), 25 (복잡)'),
    ('returnIntermediateSteps','중간 사고 과정 반환. 디버깅에 유용. 프로덕션에서는 false 권장.','boolean','false'),
    ('humanMessageTemplate','사용자 메시지 형식 커스터마이즈. {input} 변수 사용.','string','질문: {input}\n컨텍스트: {context}'),
    ('passthroughBinaryImages','이미지 전달 허용. Vision 모델 필요.','boolean','false'),
    ('promptType','chat(기본) 또는 define(커스텀 프롬프트).','enum','chat'),
]
tbl(doc,['파라미터','설명','타입','권장값'],params,[Cm(4),Cm(7),Cm(2),Cm(4.5)],header_color='2E75B6')
doc.add_paragraph()

h(doc,'1.3 LLM 모델 노드 (ai_languageModel)',2)
par(doc,'AI Agent에 ai_languageModel 포트로 연결하는 언어 모델 노드들입니다.', size=Pt(10))
doc.add_paragraph()

llm_nodes=[
    ('lmChatOpenAi','OpenAI Chat (GPT-4o, GPT-4o-mini, GPT-3.5-turbo)','openAiApi','model, temperature, maxTokens, responseFormat'),
    ('lmChatAnthropic','Claude 3.5 Sonnet, Claude 3 Opus/Haiku','anthropicApi','model, temperature, maxTokens, system'),
    ('lmChatGoogleGemini','Gemini 1.5 Pro/Flash, Gemini 2.0','googlePalmApi','model, temperature, maxOutputTokens'),
    ('lmChatOllama','Llama 3, Mistral, Phi-3 등 로컬 모델','ollamaApi (URL)','model, temperature, numCtx (컨텍스트 길이)'),
    ('lmChatAzureOpenAi','Azure OpenAI 서비스 (기업 전용)','azureOpenAiApi','deploymentName, temperature, maxTokens'),
    ('lmChatAwsBedrock','AWS Bedrock (Claude, Llama, Titan)','aws','model, region, temperature'),
    ('lmChatMistralCloud','Mistral AI 클라우드 (mistral-large 등)','mistralCloudApi','model, temperature'),
    ('lmChatGroq','Groq 고속 추론 (Llama 3, Mixtral)','groqApi','model, temperature, maxTokens'),
    ('lmChatHuggingFaceInferenceApi','HuggingFace 모델 (Zephyr, Yi 등)','huggingFaceApi','model, temperature'),
    ('lmChatCohere','Cohere Command R+ (RAG 특화)','cohereApi','model, temperature'),
]
tbl(doc,['노드명','지원 모델','자격증명','주요 파라미터'],llm_nodes,
    [Cm(4),Cm(4.5),Cm(3),Cm(6)],header_color='175E40')
doc.add_paragraph()

h(doc,'1.3.1 모델 선택 빠른 가이드',3)
code(doc,
'용도별 LLM 선택 가이드\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'일반 텍스트 생성/요약   → gpt-4o-mini          ($0.15/1M input)\n'
'복잡한 추론/코드 생성   → gpt-4o / claude-3-5-sonnet  ($2.5/1M)\n'
'한국어 특화             → claude-3-5-sonnet / gpt-4o\n'
'이미지 분석 (Vision)    → gpt-4o / claude-3-5-sonnet\n'
'초고속 응답 (실시간)    → groq + llama-3.1-70b  (무료 플랜 있음)\n'
'비용 제로 (로컬)        → ollama + llama-3.1-8b (GPU 필요)\n'
'엔터프라이즈 보안       → azure-openai / aws-bedrock\n'
'RAG/검색 특화           → cohere command-r-plus\n'
'긴 문서 처리 (200K ctx) → claude-3-5-sonnet / gemini-1.5-pro\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'1.4 Tool 노드 (ai_tool)',2)
par(doc,'AI Agent가 자율적으로 호출할 수 있는 도구 노드들입니다. '
    'ai_tool 포트로 Agent에 연결하면 에이전트가 필요 시 자동으로 도구를 선택하여 실행합니다.', size=Pt(10))
doc.add_paragraph()

tool_nodes=[
    ('toolHttpRequest','HTTP API 호출 도구. 모든 REST API를 Tool로 변환.','name, description, url, method, headers'),
    ('toolCode','JavaScript 코드 실행 도구. 커스텀 로직 구현.','name, description, jsCode'),
    ('toolCalculator','수식 계산 도구. 수치 연산이 필요한 에이전트에 유용.','자동 설정'),
    ('toolWikipedia','Wikipedia 검색 도구. 백과사전 정보 조회.','자동 설정'),
    ('toolGoogleCalendar','Google Calendar 조회/생성 도구.','calendarId, 시간 범위'),
    ('toolGmail','Gmail 검색/발송 도구.','검색 쿼리, 발신인'),
    ('toolVectorStoreRetriever','벡터 DB 검색 도구. RAG 구현의 핵심.','description, topK, vectorStore 연결'),
    ('toolWorkflow','다른 n8n 워크플로우를 Tool로 호출. 모듈화에 유용.','workflowId, description'),
    ('toolMcp','MCP(Model Context Protocol) 서버 도구.','mcpServerUrl'),
    ('toolSerpApi','Google 검색 도구 (SerpAPI).','apiKey'),
    ('toolThinkTool','에이전트 내부 사고 도구. Anthropic "Think" 패턴.','자동 설정'),
]
tbl(doc,['Tool 노드','설명','주요 설정'],tool_nodes,[Cm(4.5),Cm(7),Cm(6)],header_color='175E40')
doc.add_paragraph()

h(doc,'1.4.1 커스텀 Tool 만들기 (toolHttpRequest)',3)
par(doc,'toolHttpRequest 노드를 사용하면 어떤 REST API든 AI 에이전트의 Tool로 만들 수 있습니다.', size=Pt(10))
code(doc,
'toolHttpRequest 설정 예시 - 날씨 조회 Tool\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'name: get_weather\n'
'description: 특정 도시의 현재 날씨를 조회합니다.\n'
'             필요 입력: city (도시명, 영어)\n'
'url: https://api.openweathermap.org/data/2.5/weather\n'
'method: GET\n'
'qs:\n'
'  q: ={{ $fromAI("city", "조회할 도시명") }}\n'
'  appid: YOUR_API_KEY\n'
'  units: metric\n'
'  lang: kr\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'$fromAI("파라미터명", "설명") 구문으로 에이전트가\n'
'실행 시 적절한 값을 자동으로 채워 넣습니다.'
)
doc.add_paragraph()

h(doc,'1.5 Memory 노드 (ai_memory)',2)
par(doc,'대화 이력을 저장하여 멀티턴 대화를 가능하게 하는 메모리 노드들입니다.', size=Pt(10))
doc.add_paragraph()

memory_nodes=[
    ('memoryBufferWindow','최근 N개 메시지만 보관. 가장 간단하고 빠름.','windowSize: 10 (권장)','토큰 제한 없는 단순 챗봇'),
    ('memoryTokenBuffer','토큰 수 기준으로 메모리 관리. 더 정밀한 컨텍스트 제어.','maxTokenLimit: 4000','긴 대화 챗봇'),
    ('memorySummaryBuffer','오래된 대화를 요약으로 압축. 장기 대화에 적합.','maxTokenLimit: 2000\nsummarizationModel 연결','장기 CS 대화봇'),
    ('memoryPostgresChat','PostgreSQL에 영구 저장. 재시작 후에도 이력 유지.','sessionId, tableName','프로덕션 챗봇'),
    ('memoryRedisChat','Redis에 TTL 기반 저장. 세션 만료 자동 처리.','sessionId, sessionTTL','고성능 챗봇'),
    ('memoryMotorhead','Motorhead 서버 기반 분산 메모리.','sessionId, baseUrl','엔터프라이즈'),
    ('memoryZep','Zep 벡터 기반 장기 메모리. 관련 이력 자동 검색.','sessionId, baseUrl','개인화 AI'),
]
tbl(doc,['Memory 노드','특성','주요 설정','적합 케이스'],memory_nodes,
    [Cm(4.5),Cm(5),Cm(4),Cm(4)],header_color='2E75B6')
doc.add_paragraph()

code(doc,
'메모리 세션 ID 설정 패턴\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'사용자별 대화: sessionId = {{ $json.userId }}\n'
'채널별 대화:   sessionId = {{ $json.channelId }}\n'
'쓰레드별:      sessionId = {{ $json.threadId }}\n'
'임시 세션:     sessionId = {{ $executionId }}\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'중요: sessionId를 사용자/채널별로 분리해야\n'
'다른 사람의 대화가 섞이지 않음!'
)
doc.add_paragraph()

h(doc,'1.6 Vector Store 노드 (ai_vectorStore)',2)
par(doc,'문서를 벡터로 변환하여 저장하고 의미론적 유사도 검색을 제공하는 노드들입니다. '
    'RAG(Retrieval Augmented Generation) 구현의 핵심입니다.', size=Pt(10))
doc.add_paragraph()

vs_nodes=[
    ('vectorStorePinecone','Pinecone 관리형 벡터 DB. 가장 널리 사용.','pineconeApi','인덱스 생성 필요. 무료 1개 인덱스'),
    ('vectorStoreQdrant','Qdrant 오픈소스 자체 호스팅. 무료.','qdrantApi (URL)','Docker로 로컬 운영 가능'),
    ('vectorStoreSupabase','Supabase pgvector. 기존 Supabase 사용자에게 최적.','supabaseApi','Extension 활성화 필요: vector'),
    ('vectorStoreInMemory','메모리 내 임시 저장. 테스트용.','없음 (인스턴스 내)','프로덕션 사용 금지'),
    ('vectorStoreChroma','Chroma 오픈소스. Python 생태계와 호환.','chromaApi (URL)','로컬/Self-hosted'),
    ('vectorStoreWeaviate','Weaviate 엔터프라이즈 벡터 DB.','weaviateApi','대용량 엔터프라이즈'),
    ('vectorStoreMilvus','Milvus 대규모 벡터 DB. 10억+ 벡터 처리.','milvusApi','초대용량 시스템'),
    ('vectorStoreMongoDbAtlas','MongoDB Atlas Vector Search.','mongoDbAtlasApi','기존 MongoDB 사용자'),
]
tbl(doc,['Vector Store','특성','자격증명','권장 사용 환경'],vs_nodes,
    [Cm(4.5),Cm(5.5),Cm(3.5),Cm(4)],header_color='175E40')
doc.add_paragraph()

h(doc,'1.7 Embedding 모델 노드 (ai_embedding)',2)
par(doc,'텍스트를 벡터로 변환하는 임베딩 모델 노드들입니다. '
    'Vector Store 노드에 ai_embedding 포트로 연결합니다.', size=Pt(10))
doc.add_paragraph()

emb_nodes=[
    ('embeddingsOpenAi','text-embedding-3-small (1536d), text-embedding-3-large (3072d)','openAiApi','소규모~중규모. 가장 보편적.','$0.02/1M tokens'),
    ('embeddingsGoogleGemini','embedding-001 (768d)','googlePalmApi','Google 스택 사용 시.','무료 (한도 있음)'),
    ('embeddingsHuggingFaceInferenceApi','sentence-transformers 등 오픈소스','huggingFaceApi','한국어 특화 모델 사용 시.','무료~저렴'),
    ('embeddingsOllama','nomic-embed-text (로컬)','ollamaApi','완전 로컬 처리.','무료 (GPU 필요)'),
    ('embeddingsCohere','embed-multilingual-v3.0 (다국어 특화)','cohereApi','다국어 문서 처리.','$0.10/1M tokens'),
    ('embeddingsMistralCloud','mistral-embed','mistralCloudApi','유럽 GDPR 준수.','$0.10/1M tokens'),
]
tbl(doc,['Embedding 노드','모델','자격증명','권장 사용','비용'],emb_nodes,
    [Cm(4.5),Cm(4),Cm(3),Cm(3.5),Cm(2.5)],header_color='2E75B6')
doc.add_paragraph()

h(doc,'1.8 Document Loader 노드 (ai_document)',2)
par(doc,'다양한 소스에서 문서를 불러와 벡터 DB에 인덱싱할 때 사용하는 노드들입니다.', size=Pt(10))
doc.add_paragraph()

doc_nodes=[
    ('documentDefaultDataLoader','n8n 워크플로우 내 데이터를 문서로 변환. 가장 범용적.','data, dataType (json/text/binary)'),
    ('documentBinaryInputDataLoader','PDF, Word, 이미지 등 Binary 파일 로드.','binaryDataKey, mimeType'),
    ('documentGithubLoader','GitHub 저장소 문서 자동 로드. 코드베이스 인덱싱.','repository, branch, filePath'),
    ('documentConfluenceLoader','Confluence 페이지 로드.','spaceKey, pageTitle'),
    ('documentNotionLoader','Notion 페이지/데이터베이스 로드.','notionPageId'),
]
tbl(doc,['Document Loader','설명','주요 파라미터'],doc_nodes,[Cm(5),Cm(6),Cm(6.5)],header_color='175E40')
doc.add_paragraph()

h(doc,'1.9 Text Splitter 노드 (ai_textSplitter)',2)
par(doc,'긴 문서를 벡터 DB에 저장하기 위해 청크로 분할하는 노드들입니다. '
    '청킹 전략은 RAG 품질에 직접적인 영향을 미칩니다.', size=Pt(10))
doc.add_paragraph()

splitter_nodes=[
    ('textSplitterRecursiveCharacterTextSplitter',
     '계층적 구분자로 분할 (단락→문장→단어 순). 가장 권장.',
     'chunkSize: 1000\nchunkOverlap: 200'),
    ('textSplitterCharacterTextSplitter',
     '단일 구분자로 분할. 구조가 명확한 문서에 적합.',
     'separator: "\\n\\n"\nchunkSize: 1000'),
    ('textSplitterTokenSplitter',
     '토큰 단위 분할. LLM 컨텍스트 윈도우에 정확히 맞출 때.',
     'chunkSize: 512 (tokens)\nchunkOverlap: 50'),
    ('textSplitterMarkdownHeaderTextSplitter',
     'Markdown H1/H2/H3 헤딩 기준 분할. 문서 구조 보존.',
     '헤딩 레벨 설정'),
]
tbl(doc,['Text Splitter','설명','권장 파라미터'],splitter_nodes,[Cm(5.5),Cm(6),Cm(6)],header_color='2E75B6')
doc.add_paragraph()

code(doc,
'청킹 전략 권장 설정\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'일반 텍스트 문서:   chunkSize=1000, overlap=200 (20%)\n'
'기술 문서/코드:     chunkSize=500,  overlap=100\n'
'긴 계약서/법률:     chunkSize=2000, overlap=400\n'
'구조화 문서(MD):    Markdown Header Splitter 사용\n'
'\n'
'⚠️ overlap은 chunkSize의 15~20%가 최적\n'
'⚠️ 너무 작은 chunk: 컨텍스트 손실\n'
'⚠️ 너무 큰 chunk: 검색 정밀도 저하\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'1.10 Output Parser 노드 (ai_outputParser)',2)
par(doc,'LLM의 자유 형식 텍스트 출력을 구조화된 데이터로 파싱하는 노드들입니다.', size=Pt(10))
doc.add_paragraph()

parser_nodes=[
    ('outputParserStructured','JSON Schema 기반 구조화 출력. 가장 안정적.','schema (JSON Schema 정의)','{"category":"...","score":0}'),
    ('outputParserAutofixing','파싱 실패 시 LLM이 자동 수정 재시도. 안정성 향상.','parser + LLM 연결','오류 자동 복구'),
    ('outputParserItemList','쉼표 구분 목록 파싱.','없음','tag1, tag2, tag3 → 배열'),
    ('outputParserListOutput','번호/불릿 목록 파싱.','없음','1. item\n2. item → 배열'),
]
tbl(doc,['Output Parser','특성','설정','출력 예시'],parser_nodes,[Cm(4.5),Cm(5),Cm(3.5),Cm(4.5)],header_color='175E40')
doc.add_paragraph()

code(doc,
'Structured Output Parser 사용 예시\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'// JSON Schema 정의:\n'
'{\n'
'  "type": "object",\n'
'  "properties": {\n'
'    "sentiment": { "type": "string", "enum": ["positive","neutral","negative"] },\n'
'    "score": { "type": "number", "minimum": 0, "maximum": 10 },\n'
'    "keywords": { "type": "array", "items": { "type": "string" } },\n'
'    "summary": { "type": "string", "maxLength": 200 }\n'
'  },\n'
'  "required": ["sentiment", "score", "keywords", "summary"]\n'
'}\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'→ LLM이 반드시 위 스키마에 맞는 JSON을 반환하게 강제'
)
doc.add_paragraph()

h(doc,'1.11 Chain 노드',2)
par(doc,'에이전트 없이 LLM을 직접 사용하는 체인 노드들입니다. '
    'Agent보다 결정론적이고 빠르지만, 도구 사용이 필요 없는 단순 텍스트 처리에 적합합니다.', size=Pt(10))
doc.add_paragraph()

chain_nodes=[
    ('chainLlm','기본 LLM Chain. 프롬프트 → LLM → 출력.','systemPromptTemplate, humanMessageTemplate'),
    ('chainSummarization','긴 문서 요약 특화. Map-Reduce 방식 지원.','summarizationType (stuff/map_reduce/refine)'),
    ('chainRetrievalQa','Vector Store 기반 Q&A. RAG 단순 버전.','vectorStore + retriever 연결'),
    ('chainMultiRetrievalQa','여러 Vector Store를 라우팅하여 검색.','여러 retriever + router LLM'),
    ('chainConversational','대화 이력 포함 기본 대화 체인.','memory 연결'),
]
tbl(doc,['Chain 노드','용도','주요 파라미터'],chain_nodes,[Cm(5),Cm(6),Cm(6.5)],header_color='2E75B6')
doc.add_paragraph()

# 노드 연결 아키텍처 다이어그램
h(doc,'1.12 전체 AI 시스템 연결 아키텍처',2)
code(doc,
'완전한 RAG AI Agent 시스템 아키텍처\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'[사용자 입력]                                [외부 Tools]\n'
'    │                                            │\n'
'    ▼                                            │\n'
'[Webhook / Chat]                          ┌──────┴──────────┐\n'
'    │                                     │ Tool: Web Search │\n'
'    ▼                           ai_tool──▶│ Tool: Calendar   │\n'
'┌──────────────────────────────────┐      │ Tool: Gmail      │\n'
'│         AI Agent 노드             │      │ Tool: VectorDB   │\n'
'│  (toolsAgent / conversational)   │      └─────────────────-┘\n'
'└──────┬────────┬────────┬─────────┘\n'
'       │        │        │\n'
'  ai_language  ai_     ai_\n'
'    Model    memory   tool\n'
'       │        │\n'
'       ▼        ▼\n'
'  ┌─────────┐ ┌──────────────┐\n'
'  │ GPT-4o  │ │ Buffer Window│\n'
'  │ Claude  │ │ PostgreSQL   │\n'
'  │ Gemini  │ │ Redis        │\n'
'  └─────────┘ └──────────────┘\n'
'\n'
'[Vector Store Tool 내부]\n'
'  Tool → ai_vectorStore → Pinecone/Qdrant\n'
'                └→ ai_embedding → OpenAI Embeddings\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════
# CHAPTER 2: LLM 모델별 최적 전략
# ══════════════════════════════════════════════════════
h(doc,'Chapter 2. LLM 모델별 최적 전략',1)
par(doc,'n8n에서 사용 가능한 주요 LLM 모델들의 특성, 강점, 약점, 최적 사용 케이스를 '
    '심층 분석합니다. 모델 선택은 품질, 비용, 속도의 트레이드오프를 고려해야 합니다.', size=Pt(10))
doc.add_paragraph()

h(doc,'2.1 모델 성능 종합 비교 매트릭스',2)
comparison=[
    ('GPT-4o','OpenAI','128K','$2.50','$10.00','⭐⭐⭐⭐⭐','⭐⭐⭐⭐','⭐⭐⭐','텍스트+이미지+함수호출','일반 에이전트, 코드 생성, 멀티모달'),
    ('GPT-4o-mini','OpenAI','128K','$0.15','$0.60','⭐⭐⭐⭐','⭐⭐⭐⭐⭐','⭐⭐⭐⭐','텍스트+이미지','분류, 요약, 간단한 에이전트'),
    ('Claude 3.5 Sonnet','Anthropic','200K','$3.00','$15.00','⭐⭐⭐⭐⭐','⭐⭐⭐⭐','⭐⭐⭐⭐','텍스트+이미지+Tool Use','추론, 긴 문서, 보안 민감'),
    ('Claude 3 Haiku','Anthropic','200K','$0.25','$1.25','⭐⭐⭐⭐','⭐⭐⭐⭐⭐','⭐⭐⭐⭐⭐','텍스트+이미지','초저비용, 빠른 응답'),
    ('Gemini 1.5 Pro','Google','1M','$1.25','$5.00','⭐⭐⭐⭐⭐','⭐⭐⭐⭐','⭐⭐⭐','텍스트+이미지+영상','초장문 문서, Google 스택'),
    ('Gemini 1.5 Flash','Google','1M','$0.075','$0.30','⭐⭐⭐⭐','⭐⭐⭐⭐⭐','⭐⭐⭐⭐⭐','텍스트+이미지','초저비용 대량 처리'),
    ('Llama 3.1 70B','Meta/Groq','128K','무료*','무료*','⭐⭐⭐⭐','⭐⭐⭐⭐⭐','⭐⭐⭐⭐','텍스트','로컬/프라이버시, 무제한'),
    ('Llama 3.1 8B','Meta/Ollama','128K','무료','무료','⭐⭐⭐','⭐⭐⭐⭐⭐','⭐⭐⭐⭐⭐','텍스트','엣지/로컬 배포'),
    ('Mistral Large','Mistral','128K','$2.00','$6.00','⭐⭐⭐⭐','⭐⭐⭐⭐','⭐⭐⭐⭐','텍스트','EU GDPR, 유럽 기업'),
    ('Command R+','Cohere','128K','$2.50','$10.00','⭐⭐⭐⭐','⭐⭐⭐⭐','⭐⭐⭐','텍스트+검색','RAG 특화, 인용 포함'),
]
t=doc.add_table(rows=len(comparison)+1, cols=10); t.style='Table Grid'; t.autofit=False
widths=[Cm(3.5),Cm(2.5),Cm(1.5),Cm(1.8),Cm(1.8),Cm(1.8),Cm(1.5),Cm(1.8),Cm(3),Cm(3)]
for i,w in enumerate(widths): t.columns[i].width=w
headers=['모델','제공사','컨텍스트','입력비용\n/1M','출력비용\n/1M','추론\n품질','속도','비용\n효율','지원 형식','최적 사용 케이스']
for ci,hdr in enumerate(headers):
    sc(t.rows[0].cells[ci],'1F497D')
    t.rows[0].cells[ci].text=hdr
    for p in t.rows[0].cells[ci].paragraphs:
        for r in p.runs: r.font.name='맑은 고딕'; r.font.size=Pt(7.5); r.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for ri,row in enumerate(comparison):
    for ci,val in enumerate(row):
        sc(t.rows[ri+1].cells[ci],'EBF3FB' if ci==0 else 'FFFFFF')
        t.rows[ri+1].cells[ci].text=val
        for p in t.rows[ri+1].cells[ci].paragraphs:
            for r in p.runs: r.font.name='맑은 고딕'; r.font.size=Pt(7.5)
            if ci==0:
                for r in p.runs: r.bold=True

par(doc,'*Groq 무료 플랜: 분당 30회, 일 14,400회 제한. 상용은 유료. Ollama 로컬 실행 시 완전 무료(GPU 필요).',
    size=Pt(8), color=RGBColor(0x80,0x80,0x80))
doc.add_paragraph()

h(doc,'2.2 모델별 심층 분석',2)

h(doc,'2.2.1 GPT-4o 시리즈 (OpenAI)',3)
par(doc,'n8n에서 가장 널리 사용되는 모델 시리즈입니다. Function Calling 지원이 안정적이며 '
    'n8n AI Agent와 최고의 호환성을 보입니다.', size=Pt(10))
code(doc,
'GPT-4o 최적 n8n 설정\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'모델:         gpt-4o (복잡한 추론)\n'
'              gpt-4o-mini (빠름/저렴, 대부분의 케이스)\n'
'Temperature:  0 (결정론적, 분류/추출)\n'
'              0.3~0.7 (창의적 텍스트 생성)\n'
'              1.0+ (매우 창의적, 브레인스토밍)\n'
'MaxTokens:    1000 (단답형)\n'
'              4000 (중간 길이 생성)\n'
'              16000 (긴 문서 생성)\n'
'ResponseFormat: json_object (구조화 출력 강제)\n'
'\n'
'⚡ n8n toolsAgent와 최고 호환성\n'
'⚡ Function Calling이 가장 안정적\n'
'⚠️ gpt-4o는 비용 주의 ($10/1M output tokens)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'2.2.2 Claude 3.5 시리즈 (Anthropic)',3)
par(doc,'복잡한 추론, 긴 문서 처리, 한국어 품질에서 GPT-4o와 동등하거나 우수합니다. '
    '200K 컨텍스트 윈도우로 초장문 문서 처리에 탁월합니다.', size=Pt(10))
code(doc,
'Claude 최적 n8n 설정\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'모델:         claude-3-5-sonnet-20241022 (최고 성능)\n'
'              claude-3-haiku-20240307 (초저비용)\n'
'MaxTokens:    최대 8192 (현재 n8n 노드 기준)\n'
'\n'
'Claude 특화 기능:\n'
'  1. Extended Thinking: 복잡한 수학/추론 문제\n'
'     → additionalFields.thinking.type = "enabled"\n'
'  2. Tool Use: toolsAgent와 호환\n'
'  3. Vision: 이미지 분석 (GPT-4o와 동급)\n'
'\n'
'⚡ 200K 컨텍스트: 책 한 권 전체를 한 번에 처리\n'
'⚡ 긴 문서 요약, 계약서 분석, 코드베이스 리뷰\n'
'⚠️ n8n에서 streaming 시 일부 제한 있음\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'2.2.3 Gemini 시리즈 (Google)',3)
par(doc,'1M 토큰 컨텍스트 윈도우가 최대 강점입니다. 동영상, 오디오를 포함한 멀티모달 처리와 '
    'Google Workspace 연동에 최적화되어 있습니다.', size=Pt(10))
code(doc,
'Gemini 최적 n8n 설정\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'모델:         gemini-1.5-pro (1M 컨텍스트, 복잡한 작업)\n'
'              gemini-1.5-flash (초고속/초저가)\n'
'              gemini-2.0-flash-exp (최신, 실험적)\n'
'\n'
'특화 케이스:\n'
'  - Google Drive 문서 전체 분석\n'
'  - YouTube 영상 스크립트 처리\n'
'  - Google Sheets 대용량 데이터 분석\n'
'  - 매우 긴 법률/계약 문서 처리\n'
'\n'
'⚡ 1M 토큰: PDF 1000페이지 전체 처리 가능\n'
'⚡ Gemini 1.5 Flash: 가장 저렴한 고성능 모델\n'
'⚠️ Tool Use(Function Calling): GPT-4o 대비 안정성 낮음\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'2.2.4 Ollama 로컬 모델',3)
par(doc,'Ollama를 통해 로컬에서 실행하는 오픈소스 모델들입니다. '
    '데이터 프라이버시가 중요하거나 API 비용을 없애고 싶을 때 사용합니다.', size=Pt(10))
code(doc,
'Ollama n8n 연동 설정\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'사전 작업:\n'
'  1. ollama.ai에서 Ollama 설치\n'
'  2. ollama pull llama3.1:8b\n'
'  3. ollama serve (기본 포트: 11434)\n'
'\n'
'n8n Credentials 설정:\n'
'  Type: Ollama API\n'
'  Base URL: http://localhost:11434 (n8n과 같은 서버)\n'
'             http://ollama:11434 (Docker Compose)\n'
'\n'
'권장 모델:\n'
'  llama3.1:8b   → 일반 텍스트 (8GB RAM 필요)\n'
'  llama3.1:70b  → 고품질 (40GB+ RAM 또는 GPU)\n'
'  mistral:7b    → 유럽 언어 특화\n'
'  nomic-embed   → 임베딩 전용\n'
'  llava:7b      → 이미지 분석 (Vision)\n'
'\n'
'⚡ 완전 로컬: 데이터 외부 전송 없음\n'
'⚡ 무제한 호출: API 비용 없음\n'
'⚠️ GPU 없으면 속도 매우 느림 (CPU 추론)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'2.3 모델 선택 결정 트리',2)
code(doc,
'n8n 모델 선택 결정 트리\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'시작\n'
'  │\n'
'  ├─ 데이터가 외부로 나가면 안 됨?\n'
'  │    └─ YES → Ollama (로컬) 또는 Azure/AWS (기업 전용)\n'
'  │\n'
'  ├─ EU GDPR 준수 필수?\n'
'  │    └─ YES → Mistral AI (유럽 서버) 또는 Azure OpenAI\n'
'  │\n'
'  ├─ 컨텍스트 길이 > 128K 필요?\n'
'  │    └─ YES → Gemini 1.5 Pro (1M) 또는 Claude 3.5 (200K)\n'
'  │\n'
'  ├─ 비용이 최우선?\n'
'  │    ├─ 빠름 필요: Groq + Llama 3 (무료 플랜)\n'
'  │    └─ 품질도 중요: Gemini 1.5 Flash ($0.075/1M)\n'
'  │\n'
'  ├─ 도구 호출(Tool Use) 안정성 최우선?\n'
'  │    └─ YES → GPT-4o 또는 GPT-4o-mini\n'
'  │\n'
'  ├─ 코드 생성/분석 특화?\n'
'  │    └─ YES → GPT-4o 또는 Claude 3.5 Sonnet\n'
'  │\n'
'  ├─ RAG/검색 특화?\n'
'  │    └─ YES → Cohere Command R+\n'
'  │\n'
'  └─ 일반 목적 (균형적)?\n'
'       └─ gpt-4o-mini (기본 추천)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'2.4 비용 최적화 전략',2)
par(doc,'LLM 비용은 프로덕션 운영에서 중요한 요소입니다. 아래 전략으로 품질을 유지하면서 비용을 50-90% 절감할 수 있습니다.', size=Pt(10))
doc.add_paragraph()

cost_strategies=[
    ('모델 티어링','간단한 분류→gpt-4o-mini, 복잡한 추론→gpt-4o. 혼합 사용으로 60% 절감 가능.','IF 노드로 복잡도 분류 후 모델 선택'),
    ('캐싱','동일한 질문/프롬프트는 결과를 Redis에 캐싱. 반복 API 호출 방지.','Code 노드: Redis GET/SET으로 결과 캐싱'),
    ('Prompt 최적화','시스템 프롬프트 압축. 불필요한 예시 제거. 프롬프트 길이 30% 단축 = 비용 30% 절감.','Claude/GPT 프롬프트 토큰 분석'),
    ('배치 처리','개별 API 호출 대신 여러 항목을 한 번에 처리. OpenAI Batch API 활용.','Batch API: 50% 할인'),
    ('로컬 하이브리드','민감도 낮은 작업은 Ollama 로컬 처리, 중요 작업만 클라우드 API 사용.','IF 노드로 민감도 분류'),
    ('Groq 무료 플랜','Llama 3.1 70B를 Groq에서 무료로 사용. 분당 30회, 일 14,400회 제한.','Groq API: 개발/테스트에 적극 활용'),
    ('Gemini Flash','저렴한 고성능 모델. gemini-1.5-flash = $0.075/1M input. GPT-4o-mini의 절반 가격.','간단한 작업 전부 Flash로 이전'),
    ('maxTokens 제한','출력 토큰이 비용의 주범. 필요한 길이만 생성하도록 maxTokens 엄격 설정.','요약: 500, 분류: 100, 분석: 2000'),
]
tbl(doc,['전략','설명','n8n 구현 방법'],cost_strategies,[Cm(3.5),Cm(8),Cm(6)],header_color='175E40')
doc.add_paragraph()

code(doc,
'월간 AI 비용 추정 계산기 (1000회/일 실행 기준)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'가정: 평균 입력 1,000 tokens + 출력 500 tokens\n'
'\n'
'GPT-4o:          ($2.50×1K + $10.00×0.5K) / 1M × 30K회\n'
'                 = ($2.50 + $5.00) / 1,000 × 30,000 = $225/월\n'
'\n'
'GPT-4o-mini:     ($0.15×1K + $0.60×0.5K) / 1M × 30K회\n'
'                 = $0.45 / 1,000 × 30,000 = $13.5/월 → 94% 절감!\n'
'\n'
'Gemini 1.5 Flash:($0.075×1K + $0.30×0.5K) / 1M × 30K회\n'
'                 = $0.225 / 1,000 × 30,000 = $6.75/월\n'
'\n'
'Groq(Llama 3.1): 무료 플랜 = $0/월 (14,400회/일 한도)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

doc.add_page_break()
doc.save(OUT)
print('Vol5 Part1 done:', OUT)
