# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "C:/Users/장우경/oomni/n8n_cases_full.docx"
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

def sc(doc, text, level=1):
    colors = {1: RGBColor(0x1F,0x49,0x7D), 2: RGBColor(0x2E,0x74,0xB5), 3: RGBColor(0x1F,0x49,0x7D)}
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(sizes[level]); run.font.bold = True; run.font.color.rgb = colors[level]
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after = Pt(6)

def par(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

def bul(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run("  • " + item)
        run.font.size = Pt(10); run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)

def flow(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(9); run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    p._element.pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)

def case(doc, num, title, desc, nodes, flowchart, tips):
    sc(doc, f"[{num}] {title}", 2)
    par(doc, desc)
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run("▶ 주요 노드: ")
    run.font.bold = True; run.font.size = Pt(10); run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run2 = p.add_run(nodes)
    run2.font.size = Pt(10); run2.font.name = '맑은 고딕'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    p2 = doc.add_paragraph()
    p2.style = doc.styles['Normal']
    run3 = p2.add_run("▶ 워크플로우 순서도")
    run3.font.bold = True; run3.font.size = Pt(10); run3.font.name = '맑은 고딕'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p2.paragraph_format.space_before = Pt(4); p2.paragraph_format.space_after = Pt(2)
    flow(doc, flowchart)
    p3 = doc.add_paragraph()
    p3.style = doc.styles['Normal']
    run4 = p3.add_run("▶ 핵심 포인트")
    run4.font.bold = True; run4.font.size = Pt(10); run4.font.name = '맑은 고딕'
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p3.paragraph_format.space_before = Pt(4); p3.paragraph_format.space_after = Pt(2)
    bul(doc, tips)
    doc.add_paragraph()

# ─── 표지 ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("n8n 자동화 마스터 레퍼런스")
run.font.size = Pt(18); run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Vol.1: 실전 자동화 84 케이스")
run.font.size = Pt(28); run.font.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n이메일 | CRM | HR | 마케팅 | 고객서비스 | 재무 | DevOps | 데이터 | AI\n워크플로우 순서도 포함 개정판 (v2)")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\nOOMNI Research Team | 2025")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_page_break()

# ─── CH1: 이메일 자동화 (001~010) ────────────────────────────────────────────
sc(doc, "Chapter 1: 이메일 자동화", 1)
par(doc, "이메일은 가장 널리 사용되는 비즈니스 커뮤니케이션 채널이며 n8n으로 강력하게 자동화할 수 있습니다.")

case(doc, "001", "신규 이메일 → Slack 알림",
"중요한 이메일이 수신될 때 즉시 Slack에 알림을 보냅니다. 발신자, 제목, 미리보기를 포함하여 빠른 확인이 가능합니다.",
"Gmail Trigger, IF, Slack",
"""Gmail Trigger (새 이메일 수신)
  │
  ▼
[IF] 발신자 필터 조건 충족?
  │ YES                    │ NO
  ▼                        ▼
Slack (메시지 전송)       [종료]
  │
  └─ 채널: #이메일알림
     메시지: "📧 {{sender}}: {{subject}}"
  ▼
[완료]""",
["Gmail Trigger의 'Poll Time'을 5분으로 설정하면 비용 효율적", "중요 발신자만 필터링하여 Slack 노이즈 최소화", "Slack 메시지에 Gmail 링크 포함으로 빠른 접근"])

case(doc, "002", "이메일 첨부파일 → Google Drive 저장",
"특정 이메일의 첨부파일을 자동으로 Google Drive의 지정 폴더에 저장합니다. 계약서, 청구서 등 중요 문서 관리에 활용합니다.",
"Gmail Trigger, Google Drive",
"""Gmail Trigger (첨부파일 있는 이메일)
  │
  ▼
[IF] 첨부파일 존재?
  │ YES                    │ NO
  ▼                        ▼
Google Drive              [종료]
(파일 업로드)
  │
  └─ 폴더: /이메일첨부파일/{{year}}/{{month}}
     파일명: {{date}}_{{sender}}_{{filename}}
  ▼
Gmail (읽음 표시 + 라벨 추가)""",
["파일명에 날짜와 발신자를 포함하여 추후 검색 용이", "폴더를 연/월로 자동 분류하여 체계적 관리", "대용량 파일은 Google Drive 용량 확인 필요"])

case(doc, "003", "뉴스레터 구독 → CRM 자동 등록",
"웹사이트 뉴스레터 구독 폼 제출 시 자동으로 CRM에 연락처를 생성하고 환영 이메일을 발송합니다.",
"Webhook, HubSpot, Gmail",
"""Webhook (구독 폼 제출)
  │
  ▼
[데이터 정제] 이메일 형식 검증
  │
  ▼
HubSpot (연락처 검색)
  │
  ├─ 기존 연락처 존재? → HubSpot (속성 업데이트)
  │
  └─ 신규? → HubSpot (연락처 생성)
                │
                ▼
             Gmail (환영 이메일 발송)
                │
                ▼
             [완료]""",
["중복 구독 방지를 위해 이메일 중복 체크 필수", "환영 이메일에 구독 확인 링크(double opt-in) 포함 권장", "GDPR 준수를 위해 구독 동의 기록 보관"])

case(doc, "004", "이메일 내용 → Notion 데이터베이스 저장",
"특정 라벨/폴더의 이메일을 파싱하여 Notion 데이터베이스에 자동 저장합니다. 업무 이메일 트래킹에 유용합니다.",
"Gmail Trigger, Code, Notion",
"""Gmail Trigger (특정 라벨 이메일)
  │
  ▼
Code 노드 (이메일 파싱)
  │ - 발신자 추출
  │ - 날짜 포매팅
  │ - 본문 요약 (처음 500자)
  ▼
Notion (페이지 생성)
  │
  └─ DB: 이메일 트래킹
     속성: 발신자, 제목, 날짜, 요약, 상태
  ▼
[완료]""",
["Code 노드에서 이메일 HTML을 텍스트로 변환 처리", "Notion 상태 속성으로 처리 현황 관리 (신규/처리중/완료)", "정기적 Notion DB 클린업 워크플로우 병행 권장"])

case(doc, "005", "일일 이메일 다이제스트",
"매일 아침 전날 수신된 중요 이메일을 요약하여 Slack 또는 이메일로 전송합니다.",
"Schedule Trigger, Gmail, Code, Slack",
"""Schedule Trigger (매일 오전 9시)
  │
  ▼
Gmail (어제 수신 이메일 목록 조회)
  │
  ▼
[Loop] 각 이메일 처리
  │
  ▼
Code 노드 (다이제스트 텍스트 생성)
  │ - 발신자별 그룹화
  │ - 제목 목록 생성
  │ - 총 건수 집계
  ▼
Slack (다이제스트 발송)
  └─ #업무채널
     "📬 어제 이메일 {{count}}건: ..."
  ▼
[완료]""",
["주말/공휴일 스킵 로직 추가 (IF 노드로 요일 확인)", "발신자 우선순위 설정으로 VIP 이메일 상단 표시", "이메일 수가 많을 때 Loop 노드 배치 제한 설정"])

case(doc, "006", "이메일 자동 분류 및 라벨링",
"AI를 활용하여 수신 이메일을 내용에 따라 자동으로 분류하고 Gmail 라벨을 적용합니다.",
"Gmail Trigger, OpenAI, Gmail",
"""Gmail Trigger (새 이메일)
  │
  ▼
OpenAI (이메일 분류)
  │ System: "이메일을 다음 중 하나로 분류:
  │          영업/CS/청구서/내부/스팸"
  │ User: {{subject}} + {{body_preview}}
  ▼
Switch 노드 (분류 결과)
  ├─ 영업 → Gmail (라벨: 🟢 영업)
  ├─ CS → Gmail (라벨: 🔴 CS) + Slack 알림
  ├─ 청구서 → Gmail (라벨: 💛 청구서) + Drive 저장
  ├─ 내부 → Gmail (라벨: 🔵 내부)
  └─ 스팸 → Gmail (스팸 이동)
  ▼
[완료]""",
["OpenAI 비용 절감을 위해 gpt-4o-mini 사용 권장 (분류 작업)", "Few-shot 예시를 시스템 프롬프트에 추가하면 정확도 향상", "분류 결과 로그 저장으로 정확도 모니터링 가능"])

case(doc, "007", "이메일 → Jira 티켓 자동 생성",
"CS 이메일이 수신될 때 자동으로 Jira 이슈를 생성하여 개발팀이 즉시 처리할 수 있도록 합니다.",
"Gmail Trigger, OpenAI, Jira",
"""Gmail Trigger (CS@ 이메일)
  │
  ▼
OpenAI (이슈 정보 추출)
  │ - 이슈 요약 (한줄)
  │ - 우선순위 (Critical/High/Medium/Low)
  │ - 카테고리 (버그/기능요청/질문)
  ▼
Jira (이슈 생성)
  │ - 프로젝트: CS
  │ - 제목: [AI 요약]
  │ - 설명: 원본 이메일 내용
  │ - 우선순위: [AI 판단]
  ▼
Gmail (접수 확인 자동 회신)
  └─ "안녕하세요, 문의 접수되었습니다. (티켓 #{{jiraKey}})"
  ▼
[완료]""",
["Jira 이슈 키를 Gmail 라벨로 추가하여 추적 편리성 향상", "우선순위 Critical 시 Slack PagerDuty 알림 추가 권장", "자동 회신에 예상 처리 시간을 포함하면 고객 만족도 향상"])

case(doc, "008", "대량 개인화 이메일 발송",
"Google Sheets의 수신자 목록을 기반으로 개인화된 이메일을 일괄 발송합니다. 이름, 회사, 맞춤 내용 포함.",
"Google Sheets, Loop, Gmail",
"""Google Sheets (수신자 목록 읽기)
  │ 컬럼: 이름, 이메일, 회사, 맞춤내용
  ▼
[Loop] 각 수신자 처리
  │
  ▼
[IF] 발송 여부 확인 (이미 발송됨?)
  │ NO                     │ YES
  ▼                        ▼
Gmail (개인화 이메일 발송) [스킵]
  │ 제목: "{{이름}}님께 드리는..."
  │ 본문: 개인화 내용 포함
  ▼
Google Sheets (발송 완료 표시)
  └─ 상태 컬럼 = "발송완료", 일시 기록
  ▼
Wait 노드 (1초 지연 — Rate Limit 방지)""",
["Gmail API 일일 한도(500건) 확인, 초과 시 SendGrid/Mailchimp 연동", "Wait 노드로 발송 간격 조절하여 스팸 필터 우회", "발송 실패 시 에러 이메일 별도 시트에 기록하여 재발송"])

case(doc, "009", "이메일 수신 → Google Calendar 일정 추가",
"이메일 내용에서 날짜/시간 정보를 추출하여 자동으로 Google Calendar에 일정을 등록합니다.",
"Gmail Trigger, OpenAI, Google Calendar",
"""Gmail Trigger (일정 관련 이메일)
  │ 필터: 제목에 "미팅", "회의", "meeting" 포함
  ▼
OpenAI (날짜/시간 정보 추출)
  │ JSON 출력: {
  │   "title": "회의명",
  │   "date": "2025-03-15",
  │   "time": "14:00",
  │   "duration": 60,
  │   "location": "장소 또는 링크"
  │ }
  ▼
[IF] 유효한 날짜 정보 추출됨?
  │ YES                    │ NO
  ▼                        ▼
Google Calendar           Gmail (수동 확인 요청)
(이벤트 생성)
  ▼
Gmail (확인 이메일 발송)""",
["타임존 설정을 명확히 해야 시간 오류 방지 (Asia/Seoul)", "반복 일정인 경우 recurrence 필드 처리 로직 추가", "이미 등록된 일정 중복 방지를 위한 체크 로직 권장"])

case(doc, "010", "이메일 서명에서 연락처 자동 추출",
"수신 이메일의 서명에서 이름, 회사, 전화번호, 직책 등을 추출하여 CRM에 자동 저장합니다.",
"Gmail Trigger, OpenAI, HubSpot",
"""Gmail Trigger (외부 이메일)
  │
  ▼
Code 노드 (서명 부분 추출)
  │ - 이메일 본문 하단 서명 섹션 분리
  ▼
OpenAI (연락처 정보 추출)
  │ JSON: { name, company, title,
  │         phone, email, website }
  ▼
[IF] 유효한 연락처 정보?
  │ YES
  ▼
HubSpot (연락처 검색 → 생성 or 업데이트)
  │
  ▼
[완료] (로그 기록)""",
["이미 알고 있는 내부 도메인 이메일은 필터링하여 처리 제외", "OpenAI 추출 정확도가 낮은 경우 정규식 보조 처리 추가", "추출된 전화번호는 국가 코드 포함 형식으로 정규화"])

# ─── CH2: CRM/영업 자동화 (011~020) ──────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 2: CRM/영업 자동화", 1)
par(doc, "영업 파이프라인을 자동화하여 리드 손실을 방지하고 영업 담당자가 핵심 업무에 집중할 수 있도록 지원합니다.")

case(doc, "011", "웹사이트 리드 → CRM 자동 등록",
"웹사이트 문의 폼 제출 시 CRM에 자동으로 리드를 생성하고 담당자에게 알림을 발송합니다.",
"Webhook, HubSpot, Slack, Gmail",
"""Webhook (문의 폼 제출)
  │
  ▼
Code 노드 (데이터 유효성 검증)
  │ - 이메일 형식 확인
  │ - 필수 필드 확인
  ▼
HubSpot (기존 연락처 검색)
  │
  ├─ 기존 → Deal 업데이트 + 활동 기록
  │
  └─ 신규 → 연락처 생성 → Deal 생성 (단계: 신규)
              │
              ▼
           Slack 알림 (담당자에게)
           Gmail (자동 답장)
  ▼
[완료]""",
["리드 소스(UTM 파라미터)를 CRM에 함께 저장하여 마케팅 효과 측정", "응답 시간이 ROI에 직결되므로 알림은 즉시 Slack으로 발송", "Webhook 응답 시간 < 1초 유지를 위해 HubSpot 처리는 비동기로"])

case(doc, "012", "LinkedIn 리드 → CRM 연동",
"LinkedIn Sales Navigator에서 저장한 리드를 자동으로 CRM으로 전송하고 초기 연락 이메일을 발송합니다.",
"Schedule Trigger, HTTP Request, HubSpot, Gmail",
"""Schedule Trigger (4시간마다)
  │
  ▼
LinkedIn API (새 저장 리드 조회)
  │ (또는 LinkedIn CSV Export 파싱)
  ▼
[Loop] 각 리드 처리
  │
  ▼
HubSpot (연락처 중복 확인)
  │ YES: 스킵  │ NO: 계속
              ▼
           HubSpot (연락처 생성)
              │ - 이름, 회사, 직책
              │ - LinkedIn URL
              │ - 리드 소스: LinkedIn
              ▼
           Gmail (초기 연락 이메일)
              └─ 개인화 메시지 발송""",
["LinkedIn API 제한으로 인해 공식 Sales Nav API 또는 웹훅 활용", "초기 연락 이메일은 24시간 이내 발송이 응답률 3배 높음", "연락처 enrichment: Clearbit/Hunter.io API로 추가 정보 보완"])

case(doc, "013", "Deal 단계 변경 → 팀 알림",
"CRM에서 Deal이 새 단계로 이동할 때 관련 팀원에게 자동으로 Slack 알림을 발송합니다.",
"HubSpot Trigger, Slack",
"""HubSpot Trigger (Deal 단계 변경)
  │
  ▼
Switch 노드 (새 단계)
  ├─ 제안서 발송 → Slack (영업팀장)
  │                "📋 {{company}} 제안서 단계 진입"
  ├─ 계약 검토 → Slack (법무팀)
  │               "⚖️ {{company}} 계약 검토 필요"
  ├─ 계약 완료 → Slack (#전체채널)
  │               "🎉 {{company}} 계약 성사! {{amount}}원"
  └─ 잃음 → Slack (영업팀장)
              "❌ {{company}} 딜 종료. 사유: {{reason}}"
  ▼
HubSpot (활동 기록 추가)""",
["계약 완료 시 전체 채널 공유로 팀 사기 진작 효과", "잃음 단계 시 30일 후 재접촉 자동 알림 설정 권장", "금액 기준으로 대형 딜은 CEO에게 추가 알림 설정"])

case(doc, "014", "견적서 자동 생성 및 발송",
"CRM의 Deal 정보를 기반으로 Google Docs 템플릿에서 견적서를 자동 생성하고 PDF로 변환하여 이메일 발송합니다.",
"HubSpot Trigger, Google Docs, Google Drive, Gmail",
"""HubSpot Trigger (Deal: 견적 요청)
  │
  ▼
HubSpot (Deal + Contact 상세 정보 조회)
  │
  ▼
Google Docs (템플릿 복사)
  │ - {{company_name}} 치환
  │ - {{products}} 테이블 채우기
  │ - {{total_amount}} 계산
  ▼
Google Drive (PDF 변환)
  │
  ▼
Gmail (견적서 첨부 발송)
  └─ 수신: 고객 이메일
     참조: 담당 영업사원
  ▼
HubSpot (활동 기록: 견적서 발송)""",
["Google Docs 템플릿에 미리 스타일 정의하여 일관된 브랜딩 유지", "견적 유효기간 자동 계산 (발송일 + 30일) 포함", "PDF 파일명: 견적서_{{company}}_{{yyyymmdd}}.pdf 형식 통일"])

case(doc, "015", "거래처 생일/기념일 자동 축하",
"CRM에 등록된 주요 고객의 생일이나 계약 기념일에 자동으로 축하 이메일이나 메시지를 발송합니다.",
"Schedule Trigger, HubSpot, Gmail",
"""Schedule Trigger (매일 오전 8시)
  │
  ▼
HubSpot (오늘 생일/기념일인 연락처 조회)
  │ 필터: birthday = today OR
  │       contract_date anniversary = today
  ▼
[IF] 해당 연락처 존재?
  │ YES
  ▼
[Loop] 각 연락처 처리
  │
  ▼
Switch (생일 vs 기념일)
  ├─ 생일 → Gmail (생일 축하 이메일)
  └─ 기념일 → Gmail (거래 기념일 감사 이메일)
  ▼
HubSpot (활동 기록)""",
["생일 이메일에 특별 할인 쿠폰 포함 시 전환율 크게 향상", "템플릿을 개인화하여 단순 자동 메일처럼 보이지 않도록", "VIP 고객은 이메일 대신 담당자 직접 연락 알림으로 처리"])

case(doc, "016", "영업 활동 주간 리포트 자동화",
"매주 금요일 오후 각 영업사원의 주간 활동 지표를 CRM에서 집계하여 리포트를 Slack으로 발송합니다.",
"Schedule Trigger, HubSpot, Code, Slack",
"""Schedule Trigger (매주 금요일 17:00)
  │
  ▼
HubSpot (이번 주 영업 활동 조회)
  │ - 생성 Deal 수, 금액
  │ - 완료 통화/미팅 수
  │ - 이메일 발송 수
  │ - 전환율
  ▼
Code 노드 (영업사원별 집계 + 순위 계산)
  │
  ▼
Code 노드 (Slack 블록 포매팅)
  │ - 순위표 생성
  │ - 이번 주 MVP 선정
  ▼
Slack (주간 리포트 발송)
  └─ #영업채널 + 개인 DM (본인 지표)""",
["영업팀장에게는 상세 데이터, 팀원에게는 본인 데이터만 발송", "전주 대비 증감율 표시로 트렌드 파악 용이", "월간 누적 달성률과 목표 대비 현황 함께 표시"])

case(doc, "017", "고객 이탈 위험 감지 알림",
"CRM 데이터를 분석하여 로그인 없음, 응답 없음, 지원 티켓 급증 등 이탈 신호를 감지하고 CS팀에 알립니다.",
"Schedule Trigger, HubSpot, Slack",
"""Schedule Trigger (매일 오전 10시)
  │
  ▼
HubSpot (고객 건강 지표 조회)
  │ - 최근 30일 로그인 횟수
  │ - 미응답 이메일 수
  │ - 지원 티켓 수
  │ - 사용량 데이터
  ▼
Code 노드 (이탈 위험 점수 계산)
  │ 점수 = 로그인감소*0.3 + 미응답*0.3
  │        + 티켓급증*0.4
  ▼
[IF] 위험 점수 > 70?
  │ YES
  ▼
Slack (CS 매니저에게 긴급 알림)
  └─ "⚠️ {{company}} 이탈 위험 (점수: {{score}})"
     "마지막 로그인: {{days}}일 전"
  ▼
HubSpot (고객 상태 태그 업데이트)""",
["이탈 예방은 신규 획득보다 5~7배 비용 효율적", "위험 점수 기준은 실제 이탈 데이터로 캘리브레이션 필요", "자동 알림 후 72시간 내 미조치 시 에스컬레이션 추가"])

case(doc, "018", "신규 계약 → 온보딩 시퀀스 시작",
"CRM에서 Deal이 '계약 완료'로 변경될 때 자동으로 온보딩 이메일 시퀀스를 시작합니다.",
"HubSpot Trigger, Gmail, Schedule Trigger",
"""HubSpot Trigger (Deal 단계: 계약완료)
  │
  ▼
HubSpot (고객 상세 정보 조회)
  │
  ▼
Gmail (D+0: 환영 이메일 즉시 발송)
  │ - 담당 CS 매니저 소개
  │ - 온보딩 일정 안내
  ▼
Schedule (D+3: 킥오프 미팅 리마인더)
  ▼
Schedule (D+7: 제품 사용 팁 이메일)
  ▼
Schedule (D+14: 피드백 설문)
  ▼
Schedule (D+30: 성공 사례 공유)
  ▼
HubSpot (온보딩 완료 태그 추가)""",
["이메일 시퀀스 각 단계에서 고객 반응(오픈/클릭) 추적", "설문 응답에 따라 다음 단계 콘텐츠 분기 처리 가능", "D+30 시점 갱신 가능성이 높으므로 업셀 기회 탐지"])

case(doc, "019", "영업 미팅 후 자동 Follow-up",
"Google Calendar에서 미팅이 종료되면 자동으로 감사 이메일과 다음 액션 아이템 이메일을 발송합니다.",
"Google Calendar Trigger, OpenAI, Gmail, HubSpot",
"""Google Calendar Trigger (미팅 종료)
  │ 조건: 미팅 제목에 "영업" or "sales" 포함
  ▼
[Wait] 30분 (미팅 정리 시간)
  │
  ▼
OpenAI (Follow-up 이메일 초안 생성)
  │ 입력: 미팅 제목, 참석자, 메모(있는 경우)
  │ 출력: 감사 인사 + 다음 단계 이메일
  ▼
Gmail (Draft 생성)
  │ (즉시 발송 아닌 초안으로 저장)
  ▼
Slack (담당자에게 알림)
  └─ "미팅 follow-up 초안 생성됨. 검토 후 발송하세요."
  ▼
HubSpot (미팅 활동 기록)""",
["자동 발송보다 초안(Draft) 생성 후 검토 방식이 품질 보장", "캘린더 미팅 메모에 태그(#followup) 추가로 트리거 정밀화", "Follow-up 지연 시간을 업종/고객 특성에 맞게 조정"])

case(doc, "020", "경쟁사 언급 모니터링 → 영업 알림",
"소셜 미디어나 뉴스에서 경쟁사 언급을 감지하여 영업팀에 알림을 보냅니다.",
"Schedule Trigger, HTTP Request (Search API), OpenAI, Slack",
"""Schedule Trigger (2시간마다)
  │
  ▼
HTTP Request (Google Alerts/Twitter API)
  │ 키워드: [경쟁사명, 업계 키워드]
  ▼
[IF] 신규 언급 존재?
  │ YES
  ▼
[Loop] 각 언급 처리
  │
  ▼
OpenAI (감정 분석 + 영업 기회 판단)
  │ "이 언급이 영업 기회가 될 수 있나요?"
  ▼
[IF] 영업 기회 (부정적 언급/불만)?
  │ YES
  ▼
Slack (영업팀 알림)
  └─ "🎯 영업 기회: {{source}} - {{content}}"
  ▼
HubSpot (태스크 생성: 후속 연락)""",
["Google Alerts를 RSS 피드로 구독하면 별도 API 없이 구현 가능", "경쟁사 고객 불만은 전환 성공률 30% 이상 높은 황금 기회", "감정 분석으로 단순 언급과 실제 기회를 구분하여 노이즈 감소"])

doc.save(PATH)
print("[OK] gen_vol1_revised_p1.py 완료 - n8n_cases_full.docx 생성 (표지+Ch1+Ch2, 케이스001-020)")
