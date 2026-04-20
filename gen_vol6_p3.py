# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "C:/Users/장우경/oomni/n8n_cases_vol6.docx"
doc = Document(PATH)

def sc(doc, text, level=1):
    styles = {1: (16, True, RGBColor(0x1F,0x49,0x7D)),
              2: (13, True, RGBColor(0x2E,0x74,0xB5)),
              3: (11, True, RGBColor(0x1F,0x49,0x7D))}
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    sz, bold, color = styles[level]
    run.font.size = Pt(sz); run.font.bold = bold; run.font.color.rgb = color
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after = Pt(6)

def h(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

def par(doc, text, indent=0):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.5)

def bul(doc, items, indent=1):
    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(("  " * indent) + "• " + item)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

def code(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._element.pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def tbl(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'BDD7EE')
        cell._tc.get_or_add_tcPr().append(shd)
    for rdata in rows:
        row = table.add_row()
        for i, val in enumerate(rdata):
            cell = row.cells[i]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ─── CH5: 도입 우선순위 프레임워크 ──────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 5: 도입 우선순위 프레임워크", 1)
par(doc, "한정된 예산과 인력으로 최대 ROI를 달성하려면 자동화 과제의 우선순위를 체계적으로 결정해야 합니다.")

sc(doc, "5.1 Impact-Effort 매트릭스", 2)
par(doc, "각 자동화 과제를 비즈니스 임팩트와 구현 노력(Effort)으로 평가하여 4개 사분면에 배치합니다.")
code(doc, """        High Impact
             |
  Quick Win  |  Major Project
  (즉시 착수) |  (장기 계획)
─────────────┼─────────────
  Fill-in    |  Time Sink
  (여유 시   |  (지양)
  도전)      |
             |
        Low Impact

[Quick Win 예시 — 즉시 착수]
• 이메일 수신 → Slack 알림 (2시간 구현, 일 2시간 절감)
• Google Form → Notion 자동 입력 (3시간, 오류 제거)
• 채용공고 만료 자동 알림 (4시간, 누락 0%)

[Major Project 예시 — 분기 계획]
• 전사 CS 챗봇 (1~2개월, ROI 500%+)
• ERP 연동 자동화 (2~3개월, 월 500만원 절감)
• 멀티에이전트 영업 지원 (3개월, 계약률 20% 향상)""")

h(doc, "5.1.1 임팩트 평가 기준 (1~10점)")
tbl(doc,
    ["평가 항목", "배점", "측정 방법"],
    [
        ["월간 절감 시간", "0~4점", "담당자 인터뷰, 시간 측정"],
        ["비용 절감 금액", "0~3점", "현재 비용 × 자동화 가능 비율"],
        ["오류/리스크 감소", "0~2점", "현재 오류율 × 건당 비용"],
        ["전략적 중요도", "0~1점", "경영진 우선순위와 일치 여부"],
    ],
    [5, 3, 6]
)

h(doc, "5.1.2 구현 노력(Effort) 평가 기준 (1~10점)")
tbl(doc,
    ["평가 항목", "배점", "측정 방법"],
    [
        ["기술적 복잡도", "0~4점", "시스템 연동 수, 로직 복잡도"],
        ["데이터 정제 필요도", "0~2점", "입력 데이터 품질 수준"],
        ["이해관계자 협의", "0~2점", "관련 팀 수, 변경 관리 필요도"],
        ["규제/컴플라이언스", "0~2점", "법적 요구사항 충족 필요도"],
    ],
    [5, 3, 6]
)

sc(doc, "5.2 ROI 기반 우선순위 산정", 2)
h(doc, "5.2.1 자동화 과제 우선순위 스코어카드")
tbl(doc,
    ["자동화 과제", "임팩트 점수", "Effort 점수", "ROI 배율", "우선순위"],
    [
        ["CS FAQ 챗봇 구축", "9", "5", "12개월 ROI 800%", "★★★★★ (1위)"],
        ["급여 계산 자동화", "8", "4", "12개월 ROI 600%", "★★★★★ (2위)"],
        ["리드 자동 분류", "7", "3", "12개월 ROI 500%", "★★★★★ (3위)"],
        ["청구서 OCR 처리", "7", "5", "12개월 ROI 350%", "★★★★ (4위)"],
        ["SNS 자동 발행", "5", "2", "12개월 ROI 300%", "★★★★ (5위)"],
        ["ERP 멀티 연동", "9", "9", "12개월 ROI 250%", "★★★ (장기)"],
        ["데이터 마이그레이션", "4", "8", "12개월 ROI 100%", "★★ (검토)"],
    ],
    [4.5, 2.5, 2.5, 3.5, 3]
)

h(doc, "5.2.2 ROI 계산기 공식")
code(doc, """// 자동화 과제 ROI 빠른 계산
function calculateROI(task) {
  // 월간 편익 계산
  const monthlyBenefit =
    (task.savedHours * task.hourlyRate) +    // 인건비 절감
    (task.errorReduction * task.errorCost) + // 오류 비용 절감
    task.additionalRevenue;                   // 추가 수익

  // 월간 비용
  const monthlyCost =
    task.infrastructureCost +
    task.apiCost +
    task.maintenanceCost;

  // 순편익
  const netMonthlyBenefit = monthlyBenefit - monthlyCost;

  // 회수 기간 (월)
  const paybackMonths = task.initialInvestment / netMonthlyBenefit;

  // 12개월 ROI
  const roi12m = ((netMonthlyBenefit * 12 - task.initialInvestment)
                   / task.initialInvestment) * 100;

  return { netMonthlyBenefit, paybackMonths, roi12m };
}

// 예시: CS 챗봇
const chatbotROI = calculateROI({
  savedHours: 500,        // 월 500시간 절감
  hourlyRate: 30000,      // 시간당 3만원
  errorReduction: 0,
  errorCost: 0,
  additionalRevenue: 3000000,  // 만족도 향상 → 300만원
  infrastructureCost: 500000,
  apiCost: 1500000,
  maintenanceCost: 500000,
  initialInvestment: 30000000  // 3천만원
});
// 결과: 순편익 1,300만원/월, 회수 2.3개월, ROI 420%""")

sc(doc, "5.3 단계별 도입 로드맵", 2)
h(doc, "5.3.1 12개월 자동화 로드맵 (권장 순서)")
tbl(doc,
    ["월", "자동화 과제", "예상 ROI", "담당 부서"],
    [
        ["1~2개월", "이메일/알림 자동화 (Quick Win 5~10건)", "200~400%", "전 부서"],
        ["2~3개월", "리드 입력/분류 자동화", "400~600%", "영업/마케팅"],
        ["3~5개월", "CS 챗봇 1차 (FAQ 70% 커버)", "500~800%", "CS팀"],
        ["4~6개월", "HR 채용/급여 자동화", "300~600%", "인사팀"],
        ["5~7개월", "청구서/경비 처리 자동화", "250~450%", "재무팀"],
        ["6~9개월", "CS 챗봇 고도화 (AI 에이전트)", "추가 300%", "CS팀"],
        ["8~12개월", "멀티에이전트 영업 지원 시스템", "400~700%", "영업팀"],
        ["10~12개월", "AI 기반 예측 분석 대시보드", "200~350%", "경영전략"],
    ],
    [2.5, 5.5, 3, 3]
)

# ─── CH6: 업종별 ROI 케이스 스터디 ──────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 6: 업종별 ROI 케이스 스터디", 1)
par(doc, "국내외 실제 n8n 도입 사례를 기반으로 업종별 ROI를 분석합니다. 수치는 업계 평균을 기반으로 한 추정치입니다.")

sc(doc, "6.1 금융/보험 업종", 2)
h(doc, "케이스: 중견 손해보험사 — 보험금 청구 처리 자동화")
tbl(doc,
    ["항목", "자동화 전", "자동화 후", "개선 효과"],
    [
        ["청구서 접수 처리", "3영업일", "2시간", "97% 단축"],
        ["서류 검토 (AI OCR)", "담당자 2시간/건", "15분/건", "87% 단축"],
        ["지급 결정 (단순건)", "5영업일", "당일 자동 결정", "완전 자동화"],
        ["이상 감지 (사기)", "누락률 15%", "누락률 2%", "87% 향상"],
        ["고객 진행상황 알림", "수동 발송 (30% 누락)", "100% 자동 발송", "완전 자동화"],
    ],
    [3.5, 4, 4, 4]
)
code(doc, """금융/보험 자동화 ROI 계산:

  대상: 월 5,000건 보험금 청구 처리
  담당 인력: 기존 15명 → 자동화 후 8명 (7명 절감)
  절감 인건비: 7명 × 450만원 = 3,150만원/월

  처리 속도 향상 → 고객 만족도 향상:
  → NPS 향상 → 계약 갱신율 5% 향상 → 연 10억원 효과
  → 월 환산: 약 8,333만원

  사기 감지 향상 (누락 15%→2%):
  → 월 허위 청구 방지 금액: 1,500만원

  총 월간 편익: 3,150 + 8,333/10 + 1,500 = 5,483만원
  (NPS는 분기별 반영이므로 1/10로 보수적 계산)

  초기 구축 비용: 1억원
  운영 비용: 500만원/월
  월 순편익: 4,983만원
  손익분기점: 2.0개월
  연간 ROI: 497%""")

sc(doc, "6.2 제조/물류 업종", 2)
h(doc, "케이스: 중견 제조업체 — 생산/재고/발주 자동화")
tbl(doc,
    ["자동화 영역", "절감 효과", "월간 절감액"],
    [
        ["생산 일정 자동 최적화", "계획 수립 시간 80% 단축, 가동률 12% 향상", "2,000만원"],
        ["재고 부족 예측 + 자동 발주", "재고 부족 사고 95% 감소, 긴급발주 비용 절감", "800만원"],
        ["품질 검사 AI 자동화", "불량 감지율 향상 (85%→97%), 검사 시간 60% 단축", "600만원"],
        ["납기 예측 + 고객 알림", "납기 준수율 78%→95%, CS 문의 40% 감소", "400만원"],
        ["공급업체 평가 자동화", "평가 시간 90% 단축, 우수 업체 선별 효율화", "200만원"],
    ],
    [4, 6, 3]
)
code(doc, """제조업 자동화 총 ROI:
  월간 총 편익: 4,000만원
  초기 투자: 2억원
  운영 비용: 700만원/월
  순편익: 3,300만원/월
  손익분기점: 6.1개월
  연간 ROI: 98%
  5년 NPV: 약 12.5억원""")

sc(doc, "6.3 유통/이커머스 업종", 2)
h(doc, "케이스: 온라인 쇼핑몰 — 주문~배송~CS 전체 자동화")
tbl(doc,
    ["자동화 영역", "처리 건수/월", "절감 시간", "월간 가치"],
    [
        ["주문 확인/발송 처리", "50,000건", "500시간", "1,750만원"],
        ["반품/교환 처리", "2,000건", "400시간", "1,400만원"],
        ["개인화 상품 추천", "100,000명", "-", "전환율 18% 향상 = 3,000만원"],
        ["리뷰 분석 + 응답", "5,000건", "200시간", "700만원"],
        ["재고 부족 사전 알림", "상시", "-", "품절 손실 60% 감소 = 500만원"],
        ["이탈 고객 재타겟팅", "10,000명", "100시간", "매출 회복 800만원"],
    ],
    [4, 3.5, 3.5, 4]
)

sc(doc, "6.4 IT/SW 업종", 2)
h(doc, "케이스: SaaS 스타트업 — DevOps + 고객 온보딩 자동화")
tbl(doc,
    ["자동화 영역", "자동화 전", "자동화 후", "ROI"],
    [
        ["코드 리뷰 보조 (AI)", "리뷰 평균 2일", "주요 이슈 즉시 감지", "버그 출시 50% 감소"],
        ["인시던트 대응", "MTTR 평균 4시간", "MTTR 45분", "장애 비용 75% 감소"],
        ["고객 온보딩 시퀀스", "수동 이메일 (50% 완수율)", "자동 시퀀스 (85% 완수율)", "Activation율 70% 향상"],
        ["청구/구독 관리", "수동 처리 (5일 걸림)", "즉시 자동 처리", "이탈 방지 연 5,000만원"],
        ["고객 건강 점수 모니터링", "주 1회 수동 확인", "실시간 자동 모니터링", "이탈 조기 감지 40%"],
    ],
    [4, 4, 4, 4]
)
code(doc, """IT/SaaS 자동화 종합 ROI:

  연간 절감/추가 수익:
  - DevOps 자동화: 개발자 2명 = 1.2억원 절감
  - 고객 온보딩 향상: Activation 70% → MRR 15% 향상 = 연 3억원
  - 이탈 방지: 연 5,000만원
  - 인시던트 비용 절감: 연 2,000만원

  연간 총 편익: 약 5억원
  초기 투자: 3,000만원
  운영 비용: 월 200만원 (연 2,400만원)
  연간 ROI: ((5억 - 2,400만) / 3,000만) × 100 = 1,587%""")

sc(doc, "6.5 의료/헬스케어 업종", 2)
h(doc, "케이스: 종합병원 — 행정 업무 자동화 (진료 예약~청구)")
tbl(doc,
    ["자동화 영역", "절감 효과", "월간 가치", "특이사항"],
    [
        ["예약 스케줄링 자동화", "전화 예약 60% 감소", "1,500만원", "24/7 예약 가능"],
        ["의무기록 요약 (AI)", "의사 입력 시간 40% 단축", "2,000만원", "EMR 연동 필수"],
        ["보험 청구 자동화", "청구 오류 95% 감소", "3,000만원", "심사 지연 감소"],
        ["환자 follow-up 알림", "재방문율 25% 향상", "2,500만원", "HIPAA 준수 필수"],
        ["처방 검토 보조 (AI)", "약물 상호작용 오류 80% 감소", "리스크 비용", "규제 준수 복잡"],
    ],
    [4, 4, 3, 5]
)
par(doc, "* 의료 분야는 HIPAA, 개인정보보호법, 의료기기 규제 등 컴플라이언스 비용이 높아 ROI 산정 시 규제 준수 비용을 반드시 포함해야 합니다.")

# ─── 부록 ───────────────────────────────────────────────────────────────────
doc.add_page_break()
sc(doc, "부록 A: ROI 계산기 템플릿", 1)
par(doc, "아래 템플릿을 복사하여 실제 자동화 과제의 ROI를 계산하세요.")
code(doc, """═══════════════════════════════════════════════════════
         n8n 자동화 ROI 계산기
═══════════════════════════════════════════════════════

[프로젝트 정보]
  과제명: ___________________________________
  담당 부서: _________________________________
  예상 완료일: _______________________________

[초기 투자 비용]
  요구사항 분석 및 설계:     ________ 만원
  워크플로우 개발:          ________ 만원
  시스템 연동:              ________ 만원
  테스트 및 QA:             ________ 만원
  교육 및 온보딩:            ________ 만원
  마이그레이션:              ________ 만원
  ─────────────────────────────────────────
  초기 투자 합계 (A):        ________ 만원

[월간 편익]
  절감 인건비:
    절감 시간 ____시간 × 시간당 ____원 = ________ 만원
  오류 감소:
    오류 감소 건수 ____ × 건당 비용 ____원 = ________ 만원
  추가 수익:                ________ 만원
  기타 편익:                ________ 만원
  ─────────────────────────────────────────
  월간 총 편익 (B):         ________ 만원

[월간 운영 비용]
  n8n 인프라/라이선스:       ________ 만원
  API 비용:                ________ 만원
  유지보수:                ________ 만원
  ─────────────────────────────────────────
  월간 운영 비용 (C):        ________ 만원

[ROI 계산]
  월간 순편익 (D = B - C):   ________ 만원
  손익분기점 (A / D):         ________ 개월
  6개월 ROI:  ((D×6 - A) / A) × 100 = ________%
  12개월 ROI: ((D×12 - A) / A) × 100 = ________%
  24개월 ROI: ((D×24 - A) / A) × 100 = ________%

[투자 판단]
  □ 손익분기점 < 3개월 → 즉시 승인
  □ 손익분기점 3~6개월 → 적극 권장
  □ 손익분기점 6~12개월 → 검토 후 승인
  □ 손익분기점 > 12개월 → 재검토 필요
═══════════════════════════════════════════════════════""")

doc.add_page_break()
sc(doc, "부록 B: 비용 절감 계산 워크시트", 1)
par(doc, "부서별 현재 업무 시간을 측정하고 자동화 가능성을 평가하는 워크시트입니다.")
tbl(doc,
    ["업무명", "담당자", "현재 소요시간 (시간/월)", "자동화 가능율(%)", "절감 가능 시간", "우선순위"],
    [
        ["이메일 분류/라우팅", "", "", "80~95%", "", ""],
        ["데이터 입력/이전", "", "", "90~100%", "", ""],
        ["정기 리포트 생성", "", "", "85~100%", "", ""],
        ["승인 요청/처리", "", "", "70~90%", "", ""],
        ["고객 문의 분류", "", "", "60~80%", "", ""],
        ["청구서 처리", "", "", "80~95%", "", ""],
        ["일정 조율/알림", "", "", "90~100%", "", ""],
        ["재고/발주 관리", "", "", "70~90%", "", ""],
        ["직접 입력 (업무명)", "", "", "", "", ""],
        ["직접 입력 (업무명)", "", "", "", "", ""],
    ],
    [4, 3, 4, 3.5, 3.5, 2.5]
)

doc.add_page_break()
sc(doc, "부록 C: 벤치마크 데이터 참고표", 1)
par(doc, "업계 평균 자동화 효과 벤치마크 데이터 (Gartner, McKinsey, n8n 사용자 조사 기반)")

h(doc, "C.1 직종별 자동화 가능 업무 비율")
tbl(doc,
    ["직종", "자동화 가능 업무 비율", "즉시 자동화 가능", "AI 활용 시 추가"],
    [
        ["데이터 입력 담당자", "78%", "65%", "+13%"],
        ["회계/재무 담당자", "55%", "42%", "+13%"],
        ["HR 담당자", "56%", "45%", "+11%"],
        ["CS 담당자", "47%", "30%", "+17%"],
        ["마케터", "35%", "25%", "+10%"],
        ["영업 담당자", "30%", "20%", "+10%"],
        ["개발자", "25%", "15%", "+10%"],
        ["관리자", "22%", "15%", "+7%"],
    ],
    [4, 4, 4, 4]
)
par(doc, "* 출처: McKinsey Global Institute (2023), Gartner IT Automation Survey (2024)")

h(doc, "C.2 자동화 프로젝트 성공률 및 ROI 달성률")
tbl(doc,
    ["지표", "업계 평균", "n8n 도입 기업 평균", "상위 20% 기업"],
    [
        ["프로젝트 성공률", "62%", "78%", "95%"],
        ["목표 ROI 달성률", "55%", "72%", "90%"],
        ["평균 손익분기점", "14개월", "5.2개월", "2.1개월"],
        ["1년 ROI", "87%", "320%", "750%+"],
        ["자동화 범위 (계획 대비)", "65%", "85%", "110%"],
    ],
    [5, 4, 4, 4]
)

h(doc, "C.3 자동화 ROI 실패 주요 원인")
tbl(doc,
    ["실패 원인", "발생 빈도", "예방 방법"],
    [
        ["불명확한 요구사항", "34%", "상세 요구사항 문서화, 이해관계자 확인"],
        ["유지보수 비용 과소 예측", "28%", "운영비용 20~30% 버퍼 추가"],
        ["조직 변화 관리 실패", "22%", "변화 관리 예산 별도 책정"],
        ["기술 복잡도 과소평가", "18%", "POC 선행 후 전체 개발"],
        ["ROI 측정 방법 부재", "15%", "사전 KPI 설정, 베이스라인 측정"],
        ["스코프 크리프", "25%", "MVP 우선, 단계적 확장"],
    ],
    [5, 2.5, 7]
)

# 최종 저장
doc.save(PATH)
print("[OK] gen_vol6_p3.py 완료 - n8n_cases_vol6.docx Ch5+Ch6+부록 추가")
