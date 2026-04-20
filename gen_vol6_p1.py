# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "C:/Users/장우경/oomni/n8n_cases_vol6.docx"

doc = Document()

# 페이지 마진
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

def sc(doc, text, level=1):
    styles = {1: (16, True, RGBColor(0x1F,0x49,0x7D)),
              2: (13, True, RGBColor(0x2E,0x74,0xB5)),
              3: (11, True, RGBColor(0x1F,0x49,0x7D))}
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    sz, bold, color = styles[level]
    run.font.size = Pt(sz); run.font.bold = bold; run.font.color.rgb = color
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p

def h(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p

def par(doc, text, indent=0):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.5)
    return p

def bul(doc, items, indent=1):
    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(("  " * indent) + "• " + item)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

def code(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._element.pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def tbl(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'BDD7EE')
        cell._tc.get_or_add_tcPr().append(shd)
    for rdata in rows:
        row = table.add_row()
        for i, val in enumerate(rdata):
            cell = row.cells[i]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ─── 표지 ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("n8n 자동화 마스터 레퍼런스")
run.font.size = Pt(18); run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Vol.6: ROI & 비용 분석 완전 가이드")
run.font.size = Pt(28); run.font.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\nn8n vs Zapier vs Make 비교 | 케이스별 절감 계산 | ROI 계산기\n우선순위 프레임워크 | 업종별 ROI | 비용 최적화 전략")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\nOOMNI Research Team | 2025")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
run.font.name = '맑은 고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_page_break()

# ─── 목차 ────────────────────────────────────────────────────────────────────
sc(doc, "목차", 1)
toc_items = [
    ("Chapter 1", "ROI 분석 개론 — 자동화 투자 수익률 기초"),
    ("  1.1", "ROI 정의 및 계산 방법론"),
    ("  1.2", "자동화 비용 항목 분류"),
    ("  1.3", "가치 측정의 정량적/정성적 접근"),
    ("  1.4", "ROI 목표 설정 프레임워크"),
    ("Chapter 2", "케이스별 절감 효과 계산"),
    ("  2.1", "HR 자동화 ROI 계산"),
    ("  2.2", "CRM/영업 자동화 ROI"),
    ("  2.3", "마케팅 자동화 ROI"),
    ("  2.4", "고객 서비스 자동화 ROI"),
    ("  2.5", "재무/회계 자동화 ROI"),
    ("  2.6", "DevOps 자동화 ROI"),
    ("  2.7", "AI 에이전트 추가 가치"),
    ("Chapter 3", "플랫폼 비교 — n8n vs Zapier vs Make"),
    ("  3.1", "기능 비교 매트릭스"),
    ("  3.2", "가격 모델 비교"),
    ("  3.3", "총소유비용(TCO) 분석"),
    ("  3.4", "마이그레이션 비용 고려사항"),
    ("Chapter 4", "비용 최적화 전략"),
    ("  4.1", "n8n 비용 구조 최적화"),
    ("  4.2", "실행 횟수 최적화"),
    ("  4.3", "인프라 비용 최적화"),
    ("Chapter 5", "도입 우선순위 프레임워크"),
    ("  5.1", "Impact-Effort 매트릭스"),
    ("  5.2", "ROI 기반 우선순위 산정"),
    ("  5.3", "단계별 도입 로드맵"),
    ("Chapter 6", "업종별 ROI 케이스 스터디"),
    ("  6.1", "금융/보험 업종"),
    ("  6.2", "제조/물류 업종"),
    ("  6.3", "유통/이커머스 업종"),
    ("  6.4", "IT/SW 업종"),
    ("  6.5", "의료/헬스케어 업종"),
    ("부록 A", "ROI 계산기 템플릿"),
    ("부록 B", "비용 절감 계산 워크시트"),
    ("부록 C", "벤치마크 데이터 참고표"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if not num.startswith("  "):
        run.font.bold = True

doc.add_page_break()

# ─── CH1: ROI 분석 개론 ──────────────────────────────────────────────────────
sc(doc, "Chapter 1: ROI 분석 개론 — 자동화 투자 수익률 기초", 1)
par(doc, "n8n 자동화 도입을 경영진에게 설득하고 예산을 확보하려면 명확한 ROI 분석이 필수입니다. 이 장에서는 자동화 프로젝트의 ROI를 측정하는 표준 방법론과 프레임워크를 제시합니다.")

sc(doc, "1.1 ROI 정의 및 계산 방법론", 2)

h(doc, "1.1.1 기본 ROI 공식")
code(doc, """ROI (%) = (순편익 / 총투자비용) × 100

순편익 = 총절감액 + 추가수익 - 운영비용

예시:
  총절감액    = 월 1,500만원 (인건비 절감 + 오류 감소)
  추가수익    = 월 300만원 (처리속도 향상으로 추가 거래)
  운영비용    = 월 200만원 (n8n 인프라 + API 비용)
  순편익      = 1,500 + 300 - 200 = 1,600만원/월
  총투자비용  = 5,000만원 (초기 개발 + 구축)

월 ROI = (1,600 / 5,000) × 100 = 32%
손익분기점 = 5,000 / 1,600 ≈ 3.1개월""")

h(doc, "1.1.2 순현재가치(NPV) 분석")
par(doc, "장기 투자 효과를 비교할 때는 NPV를 활용하여 미래 현금흐름을 현재가치로 환산합니다.")
code(doc, """NPV = Σ [CFt / (1 + r)^t] - 초기투자비용

CF = 연간 순현금흐름 (절감액 + 추가수익 - 운영비용)
r  = 할인율 (일반적으로 기업 자본비용: 8~15%)
t  = 연도

예시 (5년, 할인율 10%):
  초기투자   = 5,000만원
  연간 CF   = 1,600만원/월 × 12 = 1억 9,200만원

  NPV = 1억9,200만/(1.1)¹ + 1억9,200만/(1.1)² + ...
      - 5,000만원 = 약 6억 8,000만원

IRR (내부수익률) ≈ 385% → 투자 매우 타당""")

h(doc, "1.1.3 회수 기간 (Payback Period)")
tbl(doc,
    ["회수 기간", "투자 판단", "일반적 케이스"],
    [
        ["< 3개월", "즉시 승인", "단순 반복 업무 자동화 (데이터 입력, 리포트 생성)"],
        ["3~6개월", "적극 권장", "CS 챗봇, 이메일 자동화, 승인 워크플로우"],
        ["6~12개월", "검토 권장", "복잡한 멀티시스템 연동, AI 에이전트 구축"],
        ["12~24개월", "신중 검토", "대규모 ERP 연동, 커스텀 개발 포함"],
        ["> 24개월", "재검토 필요", "과도한 커스터마이징, 불명확한 요구사항"],
    ],
    [3, 3, 8]
)

sc(doc, "1.2 자동화 비용 항목 분류", 2)
h(doc, "1.2.1 초기 투자 비용 (One-time Cost)")
tbl(doc,
    ["비용 항목", "세부 내용", "소규모 (SMB)", "중규모", "대규모"],
    [
        ["요구사항 분석", "현업 인터뷰, 프로세스 매핑", "100~300만원", "300~800만원", "800~2,000만원"],
        ["n8n 구축/설정", "서버 설치, 초기 설정", "50~200만원", "200~500만원", "500~1,500만원"],
        ["워크플로우 개발", "노드 구성, 로직 개발", "300~1,000만원", "1,000~3,000만원", "3,000~1억원"],
        ["시스템 연동", "API 연결, 인증 설정", "100~500만원", "500~2,000만원", "2,000~8,000만원"],
        ["테스트/QA", "기능/성능/보안 테스트", "100~300만원", "300~1,000만원", "1,000~3,000만원"],
        ["교육/온보딩", "담당자 교육", "50~200만원", "200~500만원", "500~2,000만원"],
        ["마이그레이션", "기존 시스템 데이터 이전", "0~500만원", "500~2,000만원", "2,000~1억원"],
    ],
    [4, 4, 3, 3, 3]
)

h(doc, "1.2.2 운영 비용 (Recurring Cost)")
tbl(doc,
    ["비용 항목", "세부 내용", "월 비용 범위", "절감 방법"],
    [
        ["n8n 라이선스", "Cloud/Enterprise 플랜", "$0~2,000/월", "셀프호스팅으로 절감"],
        ["서버/인프라", "AWS/GCP/Azure 또는 온프레미스", "5~500만원/월", "적절한 사이즈 선택"],
        ["API 사용료", "외부 API 호출 비용", "10~1,000만원/월", "캐싱, 배치 처리"],
        ["LLM API", "OpenAI, Claude 등", "5~500만원/월", "모델 라우팅, 캐싱"],
        ["유지보수", "워크플로우 수정, 오류 대응", "50~500만원/월", "문서화, 표준화"],
        ["모니터링", "APM, 로그 관리 도구", "10~100만원/월", "오픈소스 활용"],
    ],
    [3.5, 4, 3.5, 4]
)

h(doc, "1.2.3 숨겨진 비용 (Hidden Cost)")
par(doc, "ROI 분석 시 종종 누락되는 숨겨진 비용을 반드시 포함해야 정확한 분석이 가능합니다.")
bul(doc, [
    "조직 변화 관리: 기존 담당자 재교육, 저항 관리, 변화 커뮤니케이션",
    "예외 케이스 처리: 자동화가 못 처리하는 케이스를 위한 수동 프로세스 유지",
    "오류 복구 비용: 자동화 오류 발생 시 데이터 복원, 재처리",
    "기술 부채: 빠른 구축으로 인한 향후 재개발 비용",
    "통합 유지보수: 연동 시스템 업데이트 시 워크플로우 수정",
    "보안 감사: 연간 보안 점검 및 취약점 패치",
])

sc(doc, "1.3 가치 측정의 정량적/정성적 접근", 2)
h(doc, "1.3.1 정량적 가치 측정")
tbl(doc,
    ["가치 유형", "측정 방법", "일반적 규모"],
    [
        ["인건비 절감", "(절감 시간 × 시간당 인건비)", "전체 ROI의 60~70%"],
        ["오류 감소", "(오류 건수 감소 × 건당 처리비용)", "전체 ROI의 10~20%"],
        ["처리 속도 향상", "(추가 처리건수 × 건당 수익)", "전체 ROI의 5~15%"],
        ["24/7 운영", "(야간/주말 처리건수 × 인건비)", "전체 ROI의 10~20%"],
        ["스케일링 비용 절감", "(증가 처리량 × 한계비용 절감)", "전체 ROI의 5~10%"],
    ],
    [4, 6, 4]
)

h(doc, "1.3.2 정성적 가치 (화폐화 가이드)")
par(doc, "정성적 효과도 가능한 한 화폐 가치로 환산하여 경영진을 설득하는 데 활용합니다.")
tbl(doc,
    ["정성적 효과", "화폐화 방법", "예시 수치"],
    [
        ["직원 만족도 향상", "이직률 감소 × 채용비용", "이직률 10% 감소 → 연 2,000만원"],
        ["고객 만족도 향상", "CSAT 향상 × LTV 증가율", "CSAT 10p 향상 → 매출 3% 증가"],
        ["브랜드 이미지 개선", "고객 획득 비용 감소", "CAC 15% 감소"],
        ["컴플라이언스 준수", "규제 위반 리스크 × 벌금 확률", "연간 리스크 비용 80% 감소"],
        ["의사결정 속도 향상", "기회 손실 비용 감소", "제품 출시 2주 단축 → 매출 5% 증가"],
    ],
    [4, 5, 5]
)

sc(doc, "1.4 ROI 목표 설정 프레임워크", 2)
par(doc, "성공적인 자동화 프로젝트를 위해 명확하고 달성 가능한 ROI 목표를 설정합니다.")
h(doc, "SMART ROI 목표 설정")
tbl(doc,
    ["원칙", "의미", "자동화 적용 예시"],
    [
        ["Specific (구체적)", "명확한 지표와 단위 정의", "\"CS 응대 시간 35% 단축\""],
        ["Measurable (측정 가능)", "정량적 측정 방법 사전 정의", "\"월간 티켓 처리 건수 측정\""],
        ["Achievable (달성 가능)", "현실적인 목표 수치 설정", "\"6개월 내 ROI 150% 달성\""],
        ["Relevant (관련성)", "비즈니스 목표와 연계", "\"고객 만족도 향상 전략과 연계\""],
        ["Time-bound (기한)", "구체적인 달성 기한 명시", "\"2025년 Q3까지 완료\""],
    ],
    [3, 4, 7]
)

# ─── CH2: 케이스별 절감 효과 계산 ────────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 2: 케이스별 절감 효과 계산", 1)
par(doc, "n8n으로 자동화 가능한 주요 업무 카테고리별로 구체적인 시간 및 비용 절감 효과를 산정합니다. 아래 수치는 국내외 실제 도입 사례를 기반으로 한 업계 평균치입니다.")

sc(doc, "2.1 HR 자동화 ROI 계산", 2)
par(doc, "HR 부서는 반복적인 행정 업무가 많아 자동화 효과가 크게 나타나는 분야입니다.")
h(doc, "2.1.1 채용 프로세스 자동화")
tbl(doc,
    ["자동화 항목", "절감 전 (시간/건)", "절감 후 (시간/건)", "월간 처리량", "월간 절감 시간"],
    [
        ["이력서 1차 스크리닝", "3시간", "0.1시간", "200건", "580시간"],
        ["면접 일정 조율", "1.5시간", "0.1시간", "50건", "70시간"],
        ["레퍼런스 체크 메일", "0.5시간", "0.05시간", "30건", "13.5시간"],
        ["입사 서류 처리", "2시간", "0.3시간", "20건", "34시간"],
        ["온보딩 안내 발송", "1시간", "0.05시간", "20건", "19시간"],
    ],
    [4.5, 3, 3, 3, 3]
)
code(doc, """채용 자동화 ROI 계산:
  월간 총 절감 시간 = 580 + 70 + 13.5 + 34 + 19 = 716.5시간
  HR 담당자 시간당 비용 = 3.5만원 (연봉 5,040만원 기준)
  월간 절감액 = 716.5 × 3.5만원 = 2,508만원

  n8n 운영 비용 = 100만원/월
  순 절감액 = 2,508 - 100 = 2,408만원/월

  초기 구축 비용: 1,500만원
  손익분기점: 1,500 / 2,408 ≈ 0.6개월 (18일!)
  연간 ROI = (2,408 × 12 - 1,500) / 1,500 × 100 = 1,826%""")

h(doc, "2.1.2 급여/복리후생 자동화")
tbl(doc,
    ["자동화 항목", "절감 시간/월", "오류 감소율", "절감 비용/월"],
    [
        ["급여 계산 및 대장 작성", "80시간", "98%", "280만원"],
        ["4대보험 신고/처리", "40시간", "95%", "140만원"],
        ["연차/휴가 관리", "30시간", "100%", "105만원"],
        ["경비 정산 처리", "60시간", "90%", "210만원"],
        ["세금 신고 준비", "50시간 (분기)", "85%", "175만원 (분기)"],
    ],
    [5, 3, 3, 3]
)

sc(doc, "2.2 CRM/영업 자동화 ROI", 2)
tbl(doc,
    ["자동화 항목", "자동화 전", "자동화 후", "개선율", "월간 절감/수익"],
    [
        ["리드 입력/분류", "30분/건, 오류 15%", "즉시, 오류 1%", "98% 시간절감", "450만원"],
        ["후속 이메일 발송", "15분/건, 발송률 60%", "즉시, 발송률 100%", "100% 자동화", "계약 15% 증가"],
        ["견적서 생성", "2시간/건", "10분/건", "83% 단축", "280만원"],
        ["계약 갱신 알림", "누락률 20%", "누락률 0%", "완전 자동화", "이탈 방지 연 5억"],
        ["파이프라인 업데이트", "30분/일", "실시간", "100% 자동화", "120만원"],
    ],
    [4, 4, 4, 3, 4]
)

code(doc, """CRM 자동화 종합 ROI:
  직접 비용 절감 = 850만원/월
  추가 수익 (계약 증가 15%) = 영업 월매출 2억 × 15% = 3,000만원

  초기 구축비용: 2,000만원
  운영비용: 150만원/월

  월 순편익 = 850 + 3,000 - 150 = 3,700만원
  손익분기점 = 2,000 / 3,700 ≈ 0.5개월
  연간 ROI ≈ 2,100%""")

sc(doc, "2.3 마케팅 자동화 ROI", 2)
tbl(doc,
    ["자동화 항목", "절감 시간/월", "효과 개선", "월간 가치"],
    [
        ["콘텐츠 발행 스케줄링", "40시간", "발행 일관성 100%", "140만원"],
        ["소셜미디어 모니터링", "60시간", "반응 시간 90% 단축", "210만원"],
        ["이메일 캠페인 세분화", "30시간", "오픈율 35% 향상", "350만원 (전환 증가)"],
        ["광고 성과 리포트", "20시간", "일간 자동 생성", "70만원"],
        ["리드 너처링 시퀀스", "50시간", "전환율 28% 향상", "560만원"],
        ["A/B 테스트 자동화", "25시간", "테스트 속도 5배", "150만원"],
    ],
    [4.5, 3.5, 4.5, 4]
)

sc(doc, "2.4 고객 서비스 자동화 ROI", 2)
par(doc, "CS 자동화는 비용 절감과 고객 만족도 향상을 동시에 달성할 수 있는 최고의 ROI 영역 중 하나입니다.")
tbl(doc,
    ["자동화 항목", "자동화 전 KPI", "자동화 후 KPI", "개선 폭", "월간 절감"],
    [
        ["FAQ 챗봇 응대", "첫 응답 4시간, CSAT 72%", "즉시, CSAT 85%", "응답 100배 빠름", "600만원"],
        ["티켓 분류/라우팅", "수동 15분/티켓", "즉시 자동 분류", "100% 자동화", "200만원"],
        ["주문 조회 자동화", "2분/건, 상담원 필요", "30초, 셀프서비스", "상담원 30% 감소", "450만원"],
        ["이슈 에스컬레이션", "누락 15%", "누락 0%", "완전 자동화", "SLA 위반 방지"],
        ["고객 피드백 분석", "2일/배치", "실시간", "즉각 인사이트", "150만원"],
    ],
    [3.5, 4, 4, 3, 4]
)

code(doc, """CS 챗봇 도입 상세 ROI (월 티켓 10,000건 기준):

  자동 해결 비율: 65% = 6,500건 자동 처리

  절감 비용:
  - 상담원 처리 절감 = 6,500건 × 5분 × (3만원/시간 ÷ 60) = 1,625만원
  - 야간/주말 인건비 절감 = 300만원
  - 채용 불필요 = 상담원 2명 × 300만원 = 600만원

  추가 수익:
  - CSAT 향상 (72→85%) → 이탈 감소 → 월 300만원 LTV 증가

  총 월간 편익 = 1,625 + 300 + 600 + 300 = 2,825만원
  운영 비용 (n8n + LLM API) = 200만원/월
  순편익 = 2,625만원/월
  초기 구축비: 3,000만원
  손익분기점: 3,000 / 2,625 = 1.1개월""")

sc(doc, "2.5 재무/회계 자동화 ROI", 2)
tbl(doc,
    ["자동화 항목", "절감 시간/월", "오류 감소율", "월간 절감액"],
    [
        ["청구서 처리 (OCR+AI)", "150시간", "97%", "525만원"],
        ["경비 승인 워크플로우", "80시간", "100%", "280만원"],
        ["월말 마감 리포트", "60시간", "99%", "210만원"],
        ["세금 계산서 처리", "100시간", "99%", "350만원"],
        ["예산 대비 실적 분석", "40시간", "-", "140만원"],
        ["미수금 알림 발송", "30시간", "100%", "105만원 + 회수율 15% 향상"],
    ],
    [5, 3, 3, 4]
)

sc(doc, "2.6 DevOps 자동화 ROI", 2)
tbl(doc,
    ["자동화 항목", "절감 시간/월", "인시던트 감소", "비용 절감"],
    [
        ["CI/CD 파이프라인", "120시간", "배포 오류 80% 감소", "420만원 + 장애 비용 절감"],
        ["인프라 모니터링/알림", "80시간", "MTTR 70% 단축", "장애 비용 연 5,000만원 절감"],
        ["보안 취약점 스캔", "40시간", "보안 사고 90% 예방", "리스크 비용 절감"],
        ["로그 분석 자동화", "60시간", "이상 감지 3배 빠름", "210만원"],
        ["백업/복구 자동화", "30시간", "데이터 손실 0%", "105만원 + 리스크"],
    ],
    [4.5, 3.5, 4, 4]
)

sc(doc, "2.7 AI 에이전트 추가 가치", 2)
par(doc, "일반 n8n 자동화에 AI 에이전트를 추가하면 기존 자동화 대비 추가적인 가치가 창출됩니다.")
tbl(doc,
    ["AI 에이전트 기능", "기존 자동화 한계", "AI 추가 가치", "추가 ROI"],
    [
        ["비정형 데이터 처리", "정형 데이터만 처리", "자연어/이미지/PDF 처리", "+40~60% 처리 범위"],
        ["판단/추론 필요 업무", "규칙 기반 분기만 가능", "맥락 이해 후 판단", "복잡 케이스 자동화"],
        ["자연어 인터페이스", "정해진 입력 형식 필요", "자유형식 대화 처리", "사용자 편의성 대폭 향상"],
        ["지식 기반 응답", "데이터 조회만 가능", "RAG로 전문 지식 제공", "CS 해결률 30% 향상"],
        ["멀티태스킹", "단일 워크플로우 순차 실행", "병렬 도구 호출, 계획 수립", "처리 시간 50% 단축"],
    ],
    [3.5, 4, 4, 4]
)

doc.save(PATH)
print("[OK] gen_vol6_p1.py 완료 - n8n_cases_vol6.docx 생성 (표지+TOC+Ch1+Ch2)")
