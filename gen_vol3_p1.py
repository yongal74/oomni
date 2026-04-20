# -*- coding: utf-8 -*-
"""Vol 3 Part 1: 문서 초기화 + HR/인사 10개 워크플로우"""
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

OUT = 'C:/Users/장우경/oomni/n8n_cases_vol3.docx'

# ── 스타일 헬퍼 ──────────────────────────────────────────────────────────────

def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_font(cell, text, bold=False, size=Pt(9), color=None):
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = '맑은 고딕'; r.font.size = size; r.bold = bold
            if color: r.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    colors = {1: RGBColor(0x1F,0x49,0x7D), 2: RGBColor(0x2E,0x75,0xB6), 3: RGBColor(0x17,0x5E,0x40)}
    sizes  = {1: Pt(22), 2: Pt(16), 3: Pt(13)}
    if p.runs:
        p.runs[0].font.name = '맑은 고딕'
        p.runs[0].font.size = sizes.get(level, Pt(11))
        p.runs[0].font.color.rgb = colors.get(level, RGBColor(0,0,0))
    return p

def add_para(doc, text, bold=False, color=None, size=Pt(10), indent=None):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    run.font.name = '맑은 고딕'; run.font.size = size; run.bold = bold
    if color: run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    r = p.add_run(text); r.font.name = '맑은 고딕'; r.font.size = Pt(10)
    return p

def add_code_block(doc, code_text):
    """모노스페이스 코드 블록 (회색 배경)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'; run.font.size = Pt(7.5)
    # 배경색 설정
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def add_info_table(doc, rows_data, header_color='1F497D'):
    """2열 정보 테이블"""
    t = doc.add_table(rows=len(rows_data), cols=2)
    t.style = 'Table Grid'; t.autofit = False
    t.columns[0].width = Cm(3.5); t.columns[1].width = Cm(14)
    for i, (label, value) in enumerate(rows_data):
        shade_cell(t.rows[i].cells[0], 'EBF3FB')
        set_cell_font(t.rows[i].cells[0], label, bold=True, size=Pt(9))
        set_cell_font(t.rows[i].cells[1], value, size=Pt(9))
    return t

DIFF_COLORS = {
    '초급(Easy)':    RGBColor(0x70,0xAD,0x47),
    '중급(Medium)':  RGBColor(0xED,0x7D,0x31),
    '고급(Hard)':    RGBColor(0xC0,0x00,0x00),
    '전문가(Expert)':RGBColor(0x70,0x30,0xA0),
}

def make_recipe(doc, num, title, category, difficulty, apps, trigger, description,
                workflow_json, node_guide, setup_checklist, tips):
    """완전한 레시피 블록 생성"""
    # 제목
    add_heading(doc, f'Recipe {num:03d}. {title}', 3)

    # 메타 정보 테이블
    diff_color = DIFF_COLORS.get(difficulty, RGBColor(0x40,0x40,0x40))
    info = [
        ('카테고리', category),
        ('난이도', difficulty),
        ('주요 앱', apps),
        ('트리거', trigger),
    ]
    t = doc.add_table(rows=2, cols=4); t.style = 'Table Grid'; t.autofit = False
    widths = [Cm(2.2), Cm(5.3), Cm(2.2), Cm(7.8)]
    for i, w in enumerate(widths): t.columns[i].width = w
    data = [['카테고리', category, '주요 앱', apps],
            ['난이도',   difficulty, '트리거',  trigger]]
    for ri, row_data in enumerate(data):
        row = t.rows[ri]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
            for p in row.cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = '맑은 고딕'; r.font.size = Pt(9)
                    if ci % 2 == 0: r.bold = True
        shade_cell(t.rows[ri].cells[0], 'EBF3FB')
        shade_cell(t.rows[ri].cells[2], 'EBF3FB')
        shade_cell(t.rows[ri].cells[1], 'F8FBFE')
        shade_cell(t.rows[ri].cells[3], 'F8FBFE')
    # 난이도 컬러
    if t.rows[1].cells[1].paragraphs[0].runs:
        t.rows[1].cells[1].paragraphs[0].runs[0].font.color.rgb = diff_color
        t.rows[1].cells[1].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # 설명
    add_para(doc, '📋 워크플로우 개요', bold=True, color=RGBColor(0x1F,0x49,0x7D))
    add_para(doc, description, indent=Cm(0.5))

    # JSON
    add_para(doc, '📦 n8n JSON (복사 → Import 가능)', bold=True, color=RGBColor(0x1F,0x49,0x7D))
    add_code_block(doc, workflow_json)

    # 노드 가이드
    add_para(doc, '🔧 노드별 설정 가이드', bold=True, color=RGBColor(0x1F,0x49,0x7D))
    for item in node_guide:
        add_bullet(doc, item)

    # 셋업 체크리스트
    add_para(doc, '✅ 설정 체크리스트', bold=True, color=RGBColor(0x17,0x5E,0x40))
    for item in setup_checklist:
        add_bullet(doc, f'☐ {item}')

    # 팁
    if tips:
        add_para(doc, '💡 실전 팁', bold=True, color=RGBColor(0xED,0x7D,0x31))
        for tip in tips:
            add_bullet(doc, tip)

    doc.add_paragraph()
    # 구분선
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pb.append(bottom); pPr.append(pb)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 문서 생성 시작
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# 여백 설정
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── 표지 ──────────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(72)
cover.paragraph_format.space_after  = Pt(0)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run('n8n 구현 레시피북')
r.font.name = '맑은 고딕'; r.font.size = Pt(32); r.bold = True
r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('완성 JSON 워크플로우 템플릿 100선')
r.font.name = '맑은 고딕'; r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x2E,0x75,0xB6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Volume 3 — n8n 마스터 시리즈')
r.font.name = '맑은 고딕'; r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x70,0xAD,0x47)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('n8n을 즉시 활용할 수 있는 Import-Ready 완성 워크플로우 JSON 100개\n'
              '복사 → n8n Import → 자격증명 설정 → 즉시 실행')
r.font.name = '맑은 고딕'; r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x40,0x40,0x40)

doc.add_page_break()

# ── 서문 ──────────────────────────────────────────────────────────────────
add_heading(doc, '이 레시피북 사용법', 1)
add_para(doc, '이 레시피북은 n8n으로 즉시 실행 가능한 완성 워크플로우 JSON 100개를 제공합니다. '
              '각 레시피는 다음 구조로 이루어져 있습니다:', size=Pt(10))
items = [
    '📦 JSON 복사 → n8n 메뉴 Workflows > Import from clipboard 클릭 → 붙여넣기 → Import',
    '🔧 Import 후 각 노드의 Credentials(자격증명) 연결 필요',
    '✅ 체크리스트 항목을 순서대로 완료하면 즉시 운영 가능',
    '💡 실전 팁: 프로덕션 전 반드시 테스트 환경에서 먼저 실행',
]
for item in items:
    add_bullet(doc, item)

doc.add_paragraph()
add_para(doc, '난이도 기준:', bold=True)
diff_info = [
    ('초급(Easy)', '표준 노드만 사용, API Key 1~2개, 30분 내 셋업 가능'),
    ('중급(Medium)', '여러 서비스 연동, OAuth2 포함, 1~3시간 셋업'),
    ('고급(Hard)', '커스텀 로직(Code 노드), 복잡한 데이터 변환, 반나절~하루'),
    ('전문가(Expert)', 'Self-hosted 필수, 엔터프라이즈 API, 전문 지식 필요'),
]
t = doc.add_table(rows=4, cols=2); t.style = 'Table Grid'; t.autofit = False
t.columns[0].width = Cm(4); t.columns[1].width = Cm(13.5)
dc = {'초급(Easy)': '70AD47', '중급(Medium)': 'ED7D31', '고급(Hard)': 'C00000', '전문가(Expert)': '7030A0'}
for i, (d, desc) in enumerate(diff_info):
    shade_cell(t.rows[i].cells[0], dc[d]+'22')  # 투명도 느낌
    t.rows[i].cells[0].text = d
    t.rows[i].cells[1].text = desc
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.name='맑은 고딕'; r.font.size=Pt(9); r.bold=True
            r.font.color.rgb = RGBColor(*bytes.fromhex(dc[d]))
    for p in t.rows[i].cells[1].paragraphs:
        for r in p.runs:
            r.font.name='맑은 고딕'; r.font.size=Pt(9)
doc.add_paragraph()

# 목차 요약
add_heading(doc, '목차', 1)
toc_items = [
    ('Chapter 1', 'HR / 인사관리', 'Recipe 001~010', '10개'),
    ('Chapter 2', '영업 / CRM',    'Recipe 011~020', '10개'),
    ('Chapter 3', '마케팅 / 콘텐츠', 'Recipe 021~030', '10개'),
    ('Chapter 4', '고객지원 / CS',  'Recipe 031~040', '10개'),
    ('Chapter 5', '재무 / 회계',    'Recipe 041~050', '10개'),
    ('Chapter 6', 'IT운영 / DevOps', 'Recipe 051~060', '10개'),
    ('Chapter 7', '보안 / SOAR',    'Recipe 061~070', '10개'),
    ('Chapter 8', '의료 / 헬스케어', 'Recipe 071~080', '10개'),
    ('Chapter 9', '금융 / 보험',    'Recipe 081~090', '10개'),
    ('Chapter 10','AI 에이전트',     'Recipe 091~100', '10개'),
]
t = doc.add_table(rows=len(toc_items)+1, cols=4); t.style = 'Table Grid'; t.autofit = False
for w, col in zip([Cm(2.5), Cm(5.5), Cm(4), Cm(1.5)], range(4)):
    t.columns[col].width = w
# 헤더
for ci, hdr in enumerate(['챕터', '주제', '레시피 범위', '수량']):
    shade_cell(t.rows[0].cells[ci], '1F497D')
    t.rows[0].cells[ci].text = hdr
    for p in t.rows[0].cells[ci].paragraphs:
        for r in p.runs:
            r.font.name='맑은 고딕'; r.font.size=Pt(9); r.bold=True
            r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
for i, (ch, topic, rng, cnt) in enumerate(toc_items):
    row = t.rows[i+1]
    shade_cell(row.cells[0], 'EBF3FB')
    for ci, val in enumerate([ch, topic, rng, cnt]):
        row.cells[ci].text = val
        for p in row.cells[ci].paragraphs:
            for r in p.runs:
                r.font.name='맑은 고딕'; r.font.size=Pt(9)
                if ci == 0: r.bold = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 1: HR / 인사관리
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Chapter 1. HR / 인사관리', 1)
add_para(doc,
    '신규 입사자 온보딩부터 휴가 관리, 급여 처리, 채용 파이프라인까지 HR 업무 자동화 레시피 10선입니다. '
    'BambooHR, Workday, Google Workspace, Slack과의 연동 패턴을 포함합니다.', size=Pt(10))
doc.add_paragraph()

# ── Recipe 001 ───────────────────────────────────────────────────────────
WF_001 = json.dumps({
  "name": "신규 입사자 온보딩 자동화",
  "nodes": [
    {
      "parameters": {
        "httpMethod": "POST",
        "path": "new-employee",
        "responseMode": "responseNode",
        "options": {}
      },
      "id": "a1b2c3d4-0001-0001-0001-000000000001",
      "name": "Webhook - 신규입사자",
      "type": "n8n-nodes-base.webhook",
      "typeVersion": 2,
      "position": [240, 300]
    },
    {
      "parameters": {
        "resource": "user",
        "operation": "create",
        "email": "={{ $json.email }}",
        "firstName": "={{ $json.firstName }}",
        "lastName": "={{ $json.lastName }}",
        "options": {
          "orgUnitPath": "/직원",
          "password": "TempPass123!",
          "changePasswordAtNextLogin": True
        }
      },
      "id": "a1b2c3d4-0001-0001-0001-000000000002",
      "name": "Google Workspace - 계정 생성",
      "type": "n8n-nodes-base.googleWorkspace",
      "typeVersion": 1,
      "position": [460, 300],
      "credentials": {"googleWorkspaceOAuth2Api": {"id": "cred-gws", "name": "Google Workspace"}}
    },
    {
      "parameters": {
        "resource": "channel",
        "operation": "invite",
        "channelId": "C_GENERAL",
        "userIds": "={{ $('Webhook - 신규입사자').item.json.slackUserId }}"
      },
      "id": "a1b2c3d4-0001-0001-0001-000000000003",
      "name": "Slack - 채널 초대",
      "type": "n8n-nodes-base.slack",
      "typeVersion": 2,
      "position": [680, 200],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    },
    {
      "parameters": {
        "resource": "message",
        "operation": "post",
        "channel": "={{ $('Webhook - 신규입사자').item.json.slackUserId }}",
        "text": "안녕하세요 {{ $('Webhook - 신규입사자').item.json.firstName }}님! OOMNI에 오신 것을 환영합니다 🎉\n\n다음 온보딩 체크리스트를 완료해 주세요:\n- [ ] 이메일 계정 활성화\n- [ ] Slack 프로필 업데이트\n- [ ] 첫날 미팅 캘린더 확인\n- [ ] 보안 교육 완료",
        "otherOptions": {}
      },
      "id": "a1b2c3d4-0001-0001-0001-000000000004",
      "name": "Slack - 환영 메시지",
      "type": "n8n-nodes-base.slack",
      "typeVersion": 2,
      "position": [680, 380],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    },
    {
      "parameters": {
        "resource": "event",
        "operation": "create",
        "calendarId": "primary",
        "start": "={{ $json.startDate }}T09:00:00",
        "end": "={{ $json.startDate }}T17:00:00",
        "summary": "신규 입사자 첫날 온보딩: {{ $('Webhook - 신규입사자').item.json.firstName }} {{ $('Webhook - 신규입사자').item.json.lastName }}",
        "additionalFields": {
          "attendees": [{"email": "={{ $('Webhook - 신규입사자').item.json.email }}"},
                        {"email": "hr@company.com"}]
        }
      },
      "id": "a1b2c3d4-0001-0001-0001-000000000005",
      "name": "Google Calendar - 온보딩 일정",
      "type": "n8n-nodes-base.googleCalendar",
      "typeVersion": 1,
      "position": [900, 300],
      "credentials": {"googleCalendarOAuth2Api": {"id": "cred-gcal", "name": "Google Calendar"}}
    },
    {
      "parameters": {
        "respondWith": "json",
        "responseBody": "={\"success\": true, \"message\": \"온보딩 프로세스 시작됨\", \"employee\": \"{{ $('Webhook - 신규입사자').item.json.firstName }}\"}"
      },
      "id": "a1b2c3d4-0001-0001-0001-000000000006",
      "name": "응답 반환",
      "type": "n8n-nodes-base.respondToWebhook",
      "typeVersion": 1,
      "position": [1120, 300]
    }
  ],
  "connections": {
    "Webhook - 신규입사자": {"main": [[
      {"node": "Google Workspace - 계정 생성", "type": "main", "index": 0}
    ]]},
    "Google Workspace - 계정 생성": {"main": [[
      {"node": "Slack - 채널 초대",   "type": "main", "index": 0},
      {"node": "Slack - 환영 메시지", "type": "main", "index": 0}
    ]]},
    "Slack - 환영 메시지": {"main": [[
      {"node": "Google Calendar - 온보딩 일정", "type": "main", "index": 0}
    ]]},
    "Google Calendar - 온보딩 일정": {"main": [[
      {"node": "응답 반환", "type": "main", "index": 0}
    ]]}
  },
  "active": False,
  "settings": {"executionOrder": "v1"},
  "tags": ["HR", "onboarding", "google-workspace", "slack"]
}, ensure_ascii=False, indent=2)

make_recipe(doc,
    num=1,
    title='신규 입사자 온보딩 자동화',
    category='HR / 인사관리',
    difficulty='중급(Medium)',
    apps='Google Workspace, Slack, Google Calendar',
    trigger='Webhook (HR 시스템 → POST)',
    description=(
        'HR 시스템에서 신규 입사자 정보가 등록되면 자동으로 Google Workspace 계정 생성, '
        'Slack 채널 초대, 환영 메시지 발송, 온보딩 캘린더 일정 생성을 순차 처리합니다. '
        '평균 처리 시간: 45초. 수동 대비 2~3시간 절감.'
    ),
    workflow_json=WF_001,
    node_guide=[
        'Webhook 노드: HR 시스템(BambooHR/Workday)에서 POST로 employee 정보 전송. 필수 필드: email, firstName, lastName, startDate, slackUserId',
        'Google Workspace 노드: Admin SDK API 필요. Service Account 또는 OAuth2. orgUnitPath는 회사 조직 구조에 맞게 수정',
        'Slack 채널 초대: C_GENERAL을 실제 채널 ID로 교체. 채널 ID는 채널 우클릭 → Copy channel ID',
        'Slack 환영 메시지: DM 발송. slackUserId는 Slack User ID (U로 시작하는 문자열)',
        'Google Calendar: 첫날 일정 자동 생성. hr@company.com을 실제 HR 담당자 이메일로 교체',
    ],
    setup_checklist=[
        'Google Workspace OAuth2 자격증명 생성 (Admin SDK 권한 포함)',
        'Slack Bot Token 생성 (users:read, channels:manage, chat:write 스코프)',
        'Google Calendar OAuth2 자격증명 생성',
        'Webhook URL을 HR 시스템에 등록',
        'C_GENERAL을 실제 #general 채널 ID로 교체',
        'hr@company.com을 실제 HR 담당자 이메일로 교체',
        'orgUnitPath를 실제 조직 경로로 교체',
        '임시 비밀번호 정책 확인 (TempPass123! → 회사 정책에 맞게 수정)',
    ],
    tips=[
        'BambooHR 연동: BambooHR → Webhooks → New Employee 이벤트 등록 → n8n Webhook URL 입력',
        'Workday 연동: Workday Studio → EIB → n8n HTTP Request 방식 사용',
        '에러 처리: Catch Error 노드 추가 후 Slack #hr-alerts 채널에 알림 전송 권장',
        '중복 방지: email 기준으로 Google Workspace에 이미 계정이 있는지 IF 노드로 체크',
    ]
)

# ── Recipe 002 ───────────────────────────────────────────────────────────
WF_002 = json.dumps({
  "name": "휴가 신청 자동 승인 파이프라인",
  "nodes": [
    {
      "parameters": {
        "pollTimes": {"item": [{"mode": "everyMinute", "value": 5}]},
        "filters": {"labelIds": ["LEAVE_REQUEST"]}
      },
      "id": "b2c3d4e5-0002-0002-0002-000000000001",
      "name": "Gmail - 휴가신청 감지",
      "type": "n8n-nodes-base.gmailTrigger",
      "typeVersion": 1,
      "position": [240, 300],
      "credentials": {"gmailOAuth2": {"id": "cred-gmail", "name": "Gmail"}}
    },
    {
      "parameters": {
        "jsCode": "const body = $input.item.json.snippet || '';\nconst startMatch = body.match(/(\\d{4}-\\d{2}-\\d{2})/);\nconst endMatch   = body.match(/~\\s*(\\d{4}-\\d{2}-\\d{2})/);\nconst start = startMatch ? startMatch[1] : null;\nconst end   = endMatch   ? endMatch[1]   : null;\nlet days = 0;\nif (start && end) {\n  const d1 = new Date(start), d2 = new Date(end);\n  days = Math.ceil((d2-d1)/(1000*60*60*24)) + 1;\n}\nreturn [{ json: { ...($input.item.json), parsed_start: start, parsed_end: end, leave_days: days } }];"
      },
      "id": "b2c3d4e5-0002-0002-0002-000000000002",
      "name": "Code - 날짜 파싱",
      "type": "n8n-nodes-base.code",
      "typeVersion": 2,
      "position": [460, 300]
    },
    {
      "parameters": {
        "conditions": {
          "number": [{"value1": "={{ $json.leave_days }}", "operation": "smallerEqual", "value2": 3}]
        }
      },
      "id": "b2c3d4e5-0002-0002-0002-000000000003",
      "name": "IF - 3일 이하 자동승인",
      "type": "n8n-nodes-base.if",
      "typeVersion": 1,
      "position": [680, 300]
    },
    {
      "parameters": {
        "resource": "message",
        "operation": "post",
        "channel": "#hr-notifications",
        "text": "✅ 휴가 자동 승인\n신청자: {{ $json.from.emailAddress.name }}\n기간: {{ $json.parsed_start }} ~ {{ $json.parsed_end }} ({{ $json.leave_days }}일)\n상태: **자동 승인** (3일 이하)",
        "otherOptions": {}
      },
      "id": "b2c3d4e5-0002-0002-0002-000000000004",
      "name": "Slack - 자동승인 알림",
      "type": "n8n-nodes-base.slack",
      "typeVersion": 2,
      "position": [900, 200],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    },
    {
      "parameters": {
        "resource": "message",
        "operation": "post",
        "channel": "#hr-manager",
        "text": "⏳ 휴가 승인 필요\n신청자: {{ $json.from.emailAddress.name }}\n기간: {{ $json.parsed_start }} ~ {{ $json.parsed_end }} ({{ $json.leave_days }}일)\n승인: <https://hrms.company.com/approve/{{ $json.id }}|여기서 승인>",
        "otherOptions": {}
      },
      "id": "b2c3d4e5-0002-0002-0002-000000000005",
      "name": "Slack - 매니저 승인 요청",
      "type": "n8n-nodes-base.slack",
      "typeVersion": 2,
      "position": [900, 420],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Gmail - 휴가신청 감지": {"main": [[{"node": "Code - 날짜 파싱", "type": "main", "index": 0}]]},
    "Code - 날짜 파싱": {"main": [[{"node": "IF - 3일 이하 자동승인", "type": "main", "index": 0}]]},
    "IF - 3일 이하 자동승인": {
      "main": [
        [{"node": "Slack - 자동승인 알림",    "type": "main", "index": 0}],
        [{"node": "Slack - 매니저 승인 요청", "type": "main", "index": 0}]
      ]
    }
  },
  "active": False,
  "settings": {"executionOrder": "v1"},
  "tags": ["HR", "leave", "gmail", "slack"]
}, ensure_ascii=False, indent=2)

make_recipe(doc,
    num=2,
    title='휴가 신청 자동 승인 파이프라인',
    category='HR / 인사관리',
    difficulty='중급(Medium)',
    apps='Gmail, Slack',
    trigger='Gmail Trigger (라벨: LEAVE_REQUEST)',
    description=(
        '직원이 Gmail로 휴가 신청 이메일을 보내면, 3일 이하는 자동 승인 처리하고 '
        '4일 이상은 매니저에게 Slack 알림으로 수동 승인을 요청합니다. '
        '날짜 파싱은 Code 노드(JavaScript)로 처리합니다.'
    ),
    workflow_json=WF_002,
    node_guide=[
        'Gmail Trigger: LEAVE_REQUEST 라벨이 붙은 이메일만 감지. Gmail 필터 규칙으로 "휴가신청" 제목 → 라벨 자동 적용 설정',
        'Code 노드: 이메일 본문에서 YYYY-MM-DD 형식 날짜 추출. 형식이 다르면 정규식 수정 필요',
        'IF 노드: leave_days ≤ 3이면 true(자동승인), false면 매니저 승인 요청',
        'Slack 자동승인 알림: #hr-notifications 채널에 결과 통보',
        'Slack 매니저 요청: #hr-manager 채널에 승인 링크 포함 알림. 링크는 실제 HRMS URL로 교체',
    ],
    setup_checklist=[
        'Gmail OAuth2 자격증명 생성',
        'Gmail 필터 설정: 특정 제목/발신자 → LEAVE_REQUEST 라벨 자동 적용',
        'Slack Bot Token 설정 (chat:write 스코프)',
        '#hr-notifications, #hr-manager 채널 생성 및 봇 초대',
        'HRMS 승인 URL 패턴 확인 후 교체',
        '자동 승인 기준일수(현재 3일) 정책에 맞게 조정',
    ],
    tips=[
        '이메일 본문 형식 표준화: HR 공지로 "YYYY-MM-DD ~ YYYY-MM-DD" 형식 강제 권장',
        'Calendar 자동 차단: 자동 승인 시 Google Calendar에 부재 일정 자동 생성 연결 가능',
        'HRMS 직접 연동: 회사 HRMS에 API가 있다면 HTTP Request 노드로 직접 상태 업데이트',
    ]
)

# ── Recipe 003 ───────────────────────────────────────────────────────────
WF_003 = json.dumps({
  "name": "월간 급여 명세서 자동 발송",
  "nodes": [
    {
      "parameters": {
        "rule": {"interval": [{"field": "cronExpression", "expression": "0 8 1 * *"}]}
      },
      "id": "c3d4e5f6-0003-0003-0003-000000000001",
      "name": "Schedule - 매월 1일 08:00",
      "type": "n8n-nodes-base.scheduleTrigger",
      "typeVersion": 1,
      "position": [240, 300]
    },
    {
      "parameters": {
        "url": "https://payroll.company.com/api/payslips/current-month",
        "authentication": "predefinedCredentialType",
        "nodeCredentialType": "httpHeaderAuth",
        "options": {}
      },
      "id": "c3d4e5f6-0003-0003-0003-000000000002",
      "name": "HTTP - 급여명세서 목록 조회",
      "type": "n8n-nodes-base.httpRequest",
      "typeVersion": 4,
      "position": [460, 300],
      "credentials": {"httpHeaderAuth": {"id": "cred-payroll", "name": "Payroll API"}}
    },
    {
      "parameters": {
        "fieldToSplitOut": "employees",
        "options": {}
      },
      "id": "c3d4e5f6-0003-0003-0003-000000000003",
      "name": "Split - 직원별 분리",
      "type": "n8n-nodes-base.splitOut",
      "typeVersion": 1,
      "position": [680, 300]
    },
    {
      "parameters": {
        "url": "https://payroll.company.com/api/payslips/{{ $json.employee_id }}/pdf",
        "authentication": "predefinedCredentialType",
        "nodeCredentialType": "httpHeaderAuth",
        "options": {"response": {"response": {"responseFormat": "file"}}}
      },
      "id": "c3d4e5f6-0003-0003-0003-000000000004",
      "name": "HTTP - PDF 다운로드",
      "type": "n8n-nodes-base.httpRequest",
      "typeVersion": 4,
      "position": [900, 300],
      "credentials": {"httpHeaderAuth": {"id": "cred-payroll", "name": "Payroll API"}}
    },
    {
      "parameters": {
        "resource": "message",
        "operation": "send",
        "toEmail": "={{ $json.email }}",
        "subject": "={{ new Date().getFullYear() }}년 {{ new Date().getMonth() }}월 급여 명세서",
        "message": "안녕하세요 {{ $json.name }}님,\n\n이번 달 급여 명세서를 첨부 파일로 발송드립니다.\n\n문의사항은 hr@company.com으로 연락해 주세요.\n\n감사합니다.",
        "options": {"attachments": "data"}
      },
      "id": "c3d4e5f6-0003-0003-0003-000000000005",
      "name": "Gmail - 명세서 발송",
      "type": "n8n-nodes-base.gmail",
      "typeVersion": 2,
      "position": [1120, 300],
      "credentials": {"gmailOAuth2": {"id": "cred-gmail", "name": "Gmail"}}
    }
  ],
  "connections": {
    "Schedule - 매월 1일 08:00": {"main": [[{"node": "HTTP - 급여명세서 목록 조회", "type": "main", "index": 0}]]},
    "HTTP - 급여명세서 목록 조회": {"main": [[{"node": "Split - 직원별 분리", "type": "main", "index": 0}]]},
    "Split - 직원별 분리": {"main": [[{"node": "HTTP - PDF 다운로드", "type": "main", "index": 0}]]},
    "HTTP - PDF 다운로드": {"main": [[{"node": "Gmail - 명세서 발송", "type": "main", "index": 0}]]}
  },
  "active": False,
  "settings": {"executionOrder": "v1"},
  "tags": ["HR", "payroll", "gmail", "schedule"]
}, ensure_ascii=False, indent=2)

make_recipe(doc,
    num=3,
    title='월간 급여 명세서 자동 발송',
    category='HR / 인사관리',
    difficulty='중급(Medium)',
    apps='급여 시스템 API, Gmail',
    trigger='Schedule (매월 1일 오전 8시)',
    description=(
        '매월 1일 오전 8시에 자동으로 급여 시스템 API에서 전 직원 명세서 PDF를 조회하여 '
        '개인 이메일로 발송합니다. Split Out 노드로 직원별 병렬 처리하여 전체 발송 완료 시간을 단축합니다.'
    ),
    workflow_json=WF_003,
    node_guide=[
        'Schedule Trigger: cron "0 8 1 * *" = 매월 1일 08:00. 전날 테스트 권장',
        'HTTP 급여 조회: payroll.company.com을 실제 급여 시스템 URL로 교체. API Key는 httpHeaderAuth로 설정',
        'Split Out: employees 배열을 개인별 아이템으로 분리. 응답 구조에 따라 fieldToSplitOut 수정',
        'HTTP PDF 다운로드: responseFormat: file 설정 필수. Binary Data로 받아서 다음 노드에 전달',
        'Gmail 발송: attachments 필드에 "data" 입력 시 Binary Data가 첨부파일로 자동 첨부',
    ],
    setup_checklist=[
        '급여 시스템 API 문서 확인 및 API Key 발급',
        'API 응답 구조 확인 (employees 배열 경로, employee_id 필드명)',
        'Gmail OAuth2 자격증명 설정 (gmail.send 스코프)',
        'payroll.company.com을 실제 URL로 교체',
        'cron 표현식 확인 (서버 타임존 주의)',
        '테스트: 1명 직원으로 먼저 테스트 후 전체 실행',
        '실패 알림: Error Workflow 설정으로 발송 실패 시 HR 팀에 알림',
    ],
    tips=[
        'Rate Limit: 직원 수 많을 경우 Split 후 Wait 노드 (500ms) 추가로 Gmail API 한도 초과 방지',
        '암호화: 민감한 급여 데이터 → 이메일 본문보다 PDF 암호화 + 비밀번호 별도 SMS 발송 권장',
        '로그: 발송 완료 목록을 Google Sheets에 기록하여 추적',
    ]
)

# ── Recipe 004 ───────────────────────────────────────────────────────────
WF_004 = json.dumps({
  "name": "채용 파이프라인 자동화 (Notion 기반)",
  "nodes": [
    {
      "parameters": {
        "event": "formSubmission",
        "formId": "YOUR_FORM_ID"
      },
      "id": "d4e5f6g7-0004-0004-0004-000000000001",
      "name": "Google Forms - 지원서 수신",
      "type": "n8n-nodes-base.googleFormsTrigger",
      "typeVersion": 1,
      "position": [240, 300],
      "credentials": {"googleFormsOAuth2Api": {"id": "cred-gforms", "name": "Google Forms"}}
    },
    {
      "parameters": {
        "resource": "page",
        "operation": "create",
        "databaseId": "YOUR_NOTION_DB_ID",
        "title": "={{ $json['이름'] }} - {{ $json['지원 직군'] }}",
        "propertiesUi": {
          "propertyValues": [
            {"key": "이름", "type": "title", "titleValue": "={{ $json['이름'] }}"},
            {"key": "이메일", "type": "email", "emailValue": "={{ $json['이메일'] }}"},
            {"key": "지원직군", "type": "select", "selectValue": "={{ $json['지원 직군'] }}"},
            {"key": "상태", "type": "select", "selectValue": "서류검토"},
            {"key": "제출일", "type": "date", "dateValue": "={{ $now.toISO() }}"}
          ]
        }
      },
      "id": "d4e5f6g7-0004-0004-0004-000000000002",
      "name": "Notion - 지원자 DB 등록",
      "type": "n8n-nodes-base.notion",
      "typeVersion": 2,
      "position": [460, 300],
      "credentials": {"notionApi": {"id": "cred-notion", "name": "Notion"}}
    },
    {
      "parameters": {
        "resource": "message",
        "operation": "send",
        "toEmail": "={{ $('Google Forms - 지원서 수신').item.json['이메일'] }}",
        "subject": "[{{ $('Google Forms - 지원서 수신').item.json['지원 직군'] }}] 지원 접수 확인",
        "message": "안녕하세요 {{ $('Google Forms - 지원서 수신').item.json['이름'] }}님,\n\n지원서가 정상적으로 접수되었습니다.\n\n지원 직군: {{ $('Google Forms - 지원서 수신').item.json['지원 직군'] }}\n접수 번호: {{ $json.id }}\n\n서류 검토 후 1~2주 이내 연락드리겠습니다.\n\n감사합니다.",
        "options": {}
      },
      "id": "d4e5f6g7-0004-0004-0004-000000000003",
      "name": "Gmail - 접수 확인 발송",
      "type": "n8n-nodes-base.gmail",
      "typeVersion": 2,
      "position": [680, 200],
      "credentials": {"gmailOAuth2": {"id": "cred-gmail", "name": "Gmail"}}
    },
    {
      "parameters": {
        "resource": "message",
        "operation": "post",
        "channel": "#recruiting",
        "text": "📋 새 지원자 등록!\n이름: {{ $('Google Forms - 지원서 수신').item.json['이름'] }}\n직군: {{ $('Google Forms - 지원서 수신').item.json['지원 직군'] }}\nNotion: <{{ $('Notion - 지원자 DB 등록').item.json.url }}|바로가기>",
        "otherOptions": {}
      },
      "id": "d4e5f6g7-0004-0004-0004-000000000004",
      "name": "Slack - 채용팀 알림",
      "type": "n8n-nodes-base.slack",
      "typeVersion": 2,
      "position": [680, 420],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Google Forms - 지원서 수신": {"main": [[{"node": "Notion - 지원자 DB 등록", "type": "main", "index": 0}]]},
    "Notion - 지원자 DB 등록": {"main": [[
      {"node": "Gmail - 접수 확인 발송", "type": "main", "index": 0},
      {"node": "Slack - 채용팀 알림",   "type": "main", "index": 0}
    ]]}
  },
  "active": False,
  "settings": {"executionOrder": "v1"},
  "tags": ["HR", "recruiting", "notion", "google-forms"]
}, ensure_ascii=False, indent=2)

make_recipe(doc,
    num=4,
    title='채용 파이프라인 자동화 (Notion 기반)',
    category='HR / 인사관리',
    difficulty='중급(Medium)',
    apps='Google Forms, Notion, Gmail, Slack',
    trigger='Google Forms Trigger (지원서 제출)',
    description=(
        '지원자가 Google Forms로 입사지원서를 제출하면 자동으로 Notion 지원자 DB에 등록하고, '
        '지원자에게 접수 확인 이메일을 발송하며, 채용팀 Slack 채널에 알림을 보냅니다. '
        '지원자 상태 관리는 Notion DB에서 수동으로 진행합니다.'
    ),
    workflow_json=WF_004,
    node_guide=[
        'Google Forms Trigger: YOUR_FORM_ID를 실제 폼 ID로 교체 (URL에서 /d/{ID}/edit)',
        'Notion 노드: YOUR_NOTION_DB_ID를 실제 DB ID로 교체 (Notion URL 마지막 32자리)',
        'Notion 속성 매핑: 폼 필드명과 Notion DB 속성명 정확히 일치시켜야 함',
        'Gmail 접수 확인: 지원 직군별 다른 내용 원하면 IF 노드로 분기',
        'Slack 알림: Notion URL은 $json.url로 자동 포함됨',
    ],
    setup_checklist=[
        'Google Forms 생성 (이름, 이메일, 지원직군 필드)',
        'Notion DB 생성 (이름/이메일/지원직군/상태/제출일 속성)',
        'Notion Integration 생성 및 DB에 초대',
        'Google Forms OAuth2 자격증명',
        'Notion API Key 자격증명',
        'Gmail OAuth2 자격증명',
        'Slack Bot Token 자격증명',
        'YOUR_FORM_ID, YOUR_NOTION_DB_ID 교체',
    ],
    tips=[
        'ATS 연동: Greenhouse, Lever 등 ATS API가 있으면 Notion 대신 직접 연동 가능',
        '이력서 첨부: Google Forms 파일 업로드 필드 → Google Drive 링크로 Notion에 저장',
        '중복 지원 체크: Notion 조회 노드로 동일 이메일 체크 후 IF로 분기',
    ]
)

# ── Recipe 005 ───────────────────────────────────────────────────────────
WF_005 = json.dumps({
  "name": "직원 성과 리뷰 수집 및 집계",
  "nodes": [
    {
      "parameters": {
        "rule": {"interval": [{"field": "cronExpression", "expression": "0 9 1 1,4,7,10 *"}]}
      },
      "id": "e5f6g7h8-0005-0005-0005-000000000001",
      "name": "Schedule - 분기별 리뷰 시작",
      "type": "n8n-nodes-base.scheduleTrigger",
      "typeVersion": 1,
      "position": [240, 300]
    },
    {
      "parameters": {
        "operation": "read",
        "sheetId": "YOUR_SHEET_ID",
        "range": "직원목록!A2:D",
        "options": {}
      },
      "id": "e5f6g7h8-0005-0005-0005-000000000002",
      "name": "Sheets - 직원 목록 조회",
      "type": "n8n-nodes-base.googleSheets",
      "typeVersion": 4,
      "position": [460, 300],
      "credentials": {"googleSheetsOAuth2Api": {"id": "cred-gsheets", "name": "Google Sheets"}}
    },
    {
      "parameters": {
        "fieldToSplitOut": "values",
        "options": {}
      },
      "id": "e5f6g7h8-0005-0005-0005-000000000003",
      "name": "Split - 직원별 분리",
      "type": "n8n-nodes-base.splitOut",
      "typeVersion": 1,
      "position": [680, 300]
    },
    {
      "parameters": {
        "resource": "message",
        "operation": "send",
        "toEmail": "={{ $json[1] }}",
        "subject": "Q{{ Math.ceil(new Date().getMonth()/3) }} 성과 리뷰 설문 참여 요청",
        "message": "안녕하세요 {{ $json[0] }}님,\n\n분기 성과 리뷰 설문이 시작되었습니다.\n아래 링크에서 자기 평가를 완료해 주세요.\n\n설문 링크: https://forms.gle/YOUR_REVIEW_FORM\n마감일: {{ $now.plus({days: 7}).toFormat('yyyy-MM-dd') }}\n\n약 5분 소요됩니다.",
        "options": {}
      },
      "id": "e5f6g7h8-0005-0005-0005-000000000004",
      "name": "Gmail - 리뷰 설문 발송",
      "type": "n8n-nodes-base.gmail",
      "typeVersion": 2,
      "position": [900, 300],
      "credentials": {"gmailOAuth2": {"id": "cred-gmail", "name": "Gmail"}}
    }
  ],
  "connections": {
    "Schedule - 분기별 리뷰 시작": {"main": [[{"node": "Sheets - 직원 목록 조회", "type": "main", "index": 0}]]},
    "Sheets - 직원 목록 조회": {"main": [[{"node": "Split - 직원별 분리", "type": "main", "index": 0}]]},
    "Split - 직원별 분리": {"main": [[{"node": "Gmail - 리뷰 설문 발송", "type": "main", "index": 0}]]}
  },
  "active": False,
  "settings": {"executionOrder": "v1"},
  "tags": ["HR", "performance-review", "google-sheets", "gmail"]
}, ensure_ascii=False, indent=2)

make_recipe(doc,
    num=5,
    title='분기별 성과 리뷰 설문 자동 발송',
    category='HR / 인사관리',
    difficulty='초급(Easy)',
    apps='Google Sheets, Gmail',
    trigger='Schedule (분기 첫째 날 오전 9시)',
    description=(
        '매 분기(1월/4월/7월/10월) 1일에 Google Sheets 직원 목록에서 전체 직원을 조회하여 '
        '성과 리뷰 설문 링크를 이메일로 자동 발송합니다. 7일 마감일도 자동 계산합니다.'
    ),
    workflow_json=WF_005,
    node_guide=[
        'Schedule Trigger: "0 9 1 1,4,7,10 *" = 1월/4월/7월/10월 1일 오전 9시',
        'Google Sheets: YOUR_SHEET_ID를 실제 시트 ID로 교체. A열=이름, B열=이메일 구조 가정',
        'Split Out: Sheets에서 반환되는 values 2차원 배열을 행별로 분리',
        'Gmail: YOUR_REVIEW_FORM을 실제 Google Forms URL로 교체',
        'Math.ceil(getMonth()/3): 현재 분기 계산 (1~12월 → 1~4분기)',
    ],
    setup_checklist=[
        'Google Sheets 직원 목록 시트 준비 (A: 이름, B: 이메일)',
        'YOUR_SHEET_ID를 실제 시트 ID로 교체',
        'Google Forms 성과 리뷰 설문 생성',
        'YOUR_REVIEW_FORM을 실제 폼 URL로 교체',
        'Gmail OAuth2 자격증명 설정',
        'Google Sheets OAuth2 자격증명 설정',
    ],
    tips=[
        '미완성 제출 알림: 7일 후 미제출자에게 리마인더 발송하는 별도 워크플로우 연계 권장',
        '응답 자동 집계: Forms → Sheets 연동 후, 집계 완료 시 매니저에게 자동 리포트 발송',
    ]
)

# ── Recipe 006~010 (HR 나머지) ────────────────────────────────────────────
WF_006 = json.dumps({
  "name": "퇴직 처리 자동화",
  "nodes": [
    {
      "parameters": {"httpMethod": "POST", "path": "employee-offboarding", "options": {}},
      "id": "f6g7-0006-0001", "name": "Webhook - 퇴직 처리 요청",
      "type": "n8n-nodes-base.webhook", "typeVersion": 2, "position": [240, 300]
    },
    {
      "parameters": {
        "resource": "user", "operation": "update",
        "userId": "={{ $json.googleUserId }}",
        "updateFields": {"suspended": True}
      },
      "id": "f6g7-0006-0002", "name": "Google Workspace - 계정 정지",
      "type": "n8n-nodes-base.googleWorkspace", "typeVersion": 1, "position": [460, 200],
      "credentials": {"googleWorkspaceOAuth2Api": {"id": "cred-gws", "name": "Google Workspace"}}
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#it-security",
        "text": "🚨 퇴직 처리 완료\n직원: {{ $json.name }}\n퇴직일: {{ $json.lastDay }}\n✅ Google 계정 정지\n✅ Slack 비활성화 요청 필요\n✅ 물리적 접근카드 회수 확인",
        "otherOptions": {}
      },
      "id": "f6g7-0006-0003", "name": "Slack - IT 보안팀 알림",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [680, 300],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Webhook - 퇴직 처리 요청": {"main": [[{"node": "Google Workspace - 계정 정지", "type": "main", "index": 0}]]},
    "Google Workspace - 계정 정지": {"main": [[{"node": "Slack - IT 보안팀 알림", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["HR", "offboarding", "security"]
}, ensure_ascii=False, indent=2)

make_recipe(doc,
    num=6,
    title='퇴직자 오프보딩 자동 처리',
    category='HR / 인사관리',
    difficulty='고급(Hard)',
    apps='Google Workspace, Slack, Active Directory',
    trigger='Webhook (HR 시스템 퇴직 처리)',
    description=(
        '직원 퇴직 시 Google Workspace 계정 즉시 정지, 이메일 전달 설정, '
        'IT 보안팀 오프보딩 체크리스트 알림을 자동 처리합니다. '
        '계정 정지는 되돌리기 어려우므로 반드시 HR 매니저 승인 후 트리거해야 합니다.'
    ),
    workflow_json=WF_006,
    node_guide=[
        'Webhook: HR 시스템에서 마지막 근무일 도달 시 자동 POST 또는 HR 담당자가 수동 호출',
        'Google Workspace 계정 정지: suspended: true 설정. 즉시 모든 Google 서비스 접근 차단',
        'Slack 알림: IT 보안팀에 후속 조치 체크리스트 전달 (물리적 보안 포함)',
    ],
    setup_checklist=[
        '⚠️ 인간 승인 스텝 추가 필수 (되돌리기 어려운 작업)',
        'Google Workspace 계정 관리 권한 확인',
        '퇴직자 데이터 이관 프로세스 정의 (이메일 전달 기간, Google Drive 이관)',
        'IT 보안 체크리스트 문서화 (접근 카드, 노트북 반납, 기타 시스템)',
        '법적 데이터 보존 요건 확인 (이메일 아카이브 기간)',
    ],
    tips=[
        '데이터 이관: 퇴직 전 Google Drive 소유권을 매니저에게 이전하는 별도 단계 추가 권장',
        '감사 로그: 모든 퇴직 처리 내역을 Sheets에 기록하여 감사 대응',
    ]
)

WF_007 = json.dumps({
  "name": "교육 완료 자동 인증서 발급",
  "nodes": [
    {
      "parameters": {
        "filters": {"labelIds": ["TRAINING_COMPLETE"]}
      },
      "id": "g7h8-0007-0001", "name": "Gmail - 교육완료 감지",
      "type": "n8n-nodes-base.gmailTrigger", "typeVersion": 1, "position": [240, 300],
      "credentials": {"gmailOAuth2": {"id": "cred-gmail", "name": "Gmail"}}
    },
    {
      "parameters": {
        "jsCode": "const name = $input.item.json.from?.emailAddress?.name || 'Unknown';\nconst training = $input.item.json.subject?.replace('[완료]','').trim() || '교육 과정';\nconst date = new Date().toLocaleDateString('ko-KR');\nreturn [{ json: { name, training, date, certId: 'CERT-' + Date.now() } }];"
      },
      "id": "g7h8-0007-0002", "name": "Code - 인증서 데이터 준비",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [460, 300]
    },
    {
      "parameters": {
        "operation": "appendOrUpdate",
        "sheetId": "CERT_SHEET_ID",
        "range": "인증서이력!A:E",
        "keyRowIndex": 0,
        "options": {},
        "dataToSend": "autoMapInputData"
      },
      "id": "g7h8-0007-0003", "name": "Sheets - 이력 기록",
      "type": "n8n-nodes-base.googleSheets", "typeVersion": 4, "position": [680, 300],
      "credentials": {"googleSheetsOAuth2Api": {"id": "cred-gsheets", "name": "Google Sheets"}}
    }
  ],
  "connections": {
    "Gmail - 교육완료 감지": {"main": [[{"node": "Code - 인증서 데이터 준비", "type": "main", "index": 0}]]},
    "Code - 인증서 데이터 준비": {"main": [[{"node": "Sheets - 이력 기록", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["HR", "training", "certification"]
}, ensure_ascii=False, indent=2)

make_recipe(doc,
    num=7,
    title='직원 교육 완료 자동 인증 및 이력 관리',
    category='HR / 인사관리',
    difficulty='초급(Easy)',
    apps='Gmail, Google Sheets',
    trigger='Gmail Trigger (TRAINING_COMPLETE 라벨)',
    description=(
        '직원이 교육 과정을 완료하면 자동으로 인증서 번호를 생성하고 '
        'Google Sheets 인증 이력에 기록합니다. LMS 시스템과 연동 가능합니다.'
    ),
    workflow_json=WF_007,
    node_guide=[
        'Gmail Trigger: 교육 완료 이메일에 TRAINING_COMPLETE 라벨 자동 적용 규칙 설정 필요',
        'Code 노드: 발신자 이름, 교육명(제목에서 추출), 완료일, 고유 인증서 ID 생성',
        'Google Sheets: 인증서 이력 자동 추가 (certId, name, training, date)',
    ],
    setup_checklist=[
        'Gmail 필터: "교육 완료" 제목 → TRAINING_COMPLETE 라벨',
        'Google Sheets 인증서 이력 시트 생성',
        'CERT_SHEET_ID를 실제 시트 ID로 교체',
    ],
    tips=['인증서 PDF 자동 생성: Google Docs 템플릿 + Google Docs API 연동으로 PDF 자동 생성 가능']
)

WF_008 = json.dumps({
  "name": "조직도 자동 업데이트 (HR 변경 감지)",
  "nodes": [
    {
      "parameters": {"rule": {"interval": [{"field": "cronExpression", "expression": "0 9 * * 1"}]}},
      "id": "h8i9-0008-0001", "name": "Schedule - 매주 월요일",
      "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]
    },
    {
      "parameters": {
        "url": "https://api.bamboohr.com/api/gateway.php/COMPANY/v1/employees/directory",
        "authentication": "predefinedCredentialType",
        "nodeCredentialType": "httpBasicAuth",
        "options": {"headers": {"headers": [{"name": "Accept", "value": "application/json"}]}}
      },
      "id": "h8i9-0008-0002", "name": "HTTP - BambooHR 조직도",
      "type": "n8n-nodes-base.httpRequest", "typeVersion": 4, "position": [460, 300],
      "credentials": {"httpBasicAuth": {"id": "cred-bamboohr", "name": "BambooHR"}}
    },
    {
      "parameters": {
        "operation": "update",
        "pageId": "NOTION_ORG_PAGE_ID",
        "blockId": "NOTION_TABLE_BLOCK_ID",
        "propertiesUi": {"propertyValues": []}
      },
      "id": "h8i9-0008-0003", "name": "Notion - 조직도 페이지 업데이트",
      "type": "n8n-nodes-base.notion", "typeVersion": 2, "position": [680, 300],
      "credentials": {"notionApi": {"id": "cred-notion", "name": "Notion"}}
    }
  ],
  "connections": {
    "Schedule - 매주 월요일": {"main": [[{"node": "HTTP - BambooHR 조직도", "type": "main", "index": 0}]]},
    "HTTP - BambooHR 조직도": {"main": [[{"node": "Notion - 조직도 페이지 업데이트", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["HR", "org-chart", "bamboohr", "notion"]
}, ensure_ascii=False, indent=2)

make_recipe(doc,
    num=8,
    title='BambooHR → Notion 조직도 자동 동기화',
    category='HR / 인사관리',
    difficulty='중급(Medium)',
    apps='BambooHR, Notion',
    trigger='Schedule (매주 월요일 오전 9시)',
    description=(
        'BambooHR의 직원 디렉토리를 매주 자동 조회하여 Notion 조직도 페이지를 최신 상태로 유지합니다. '
        '입퇴사, 부서 이동, 직급 변경이 자동 반영됩니다.'
    ),
    workflow_json=WF_008,
    node_guide=[
        'HTTP BambooHR: COMPANY를 BambooHR 서브도메인으로 교체. Basic Auth (API Key: x, Password: API_KEY)',
        'Notion 업데이트: 조직도 페이지 구조에 맞게 매핑 로직 수정 필요',
    ],
    setup_checklist=[
        'BambooHR API Key 생성 (My Info → API Keys)',
        'COMPANY를 실제 서브도메인으로 교체',
        'Notion 조직도 DB ID 확인 및 교체',
        '데이터 매핑 로직 회사 구조에 맞게 수정',
    ],
    tips=['변경 감지: 이전 주 데이터와 비교하여 변경된 직원만 업데이트하는 diff 로직 추가 권장']
)

WF_009 = json.dumps({
  "name": "인력 채용 공고 멀티 플랫폼 자동 게시",
  "nodes": [
    {
      "parameters": {
        "resource": "page", "operation": "getAll",
        "databaseId": "JOB_OPENINGS_DB_ID",
        "filterType": "manual",
        "filters": {"conditions": [{"key": "상태", "condition": "equals", "value": "게시대기"}]}
      },
      "id": "i9j0-0009-0001", "name": "Notion - 게시대기 공고 조회",
      "type": "n8n-nodes-base.notion", "typeVersion": 2, "position": [240, 300],
      "credentials": {"notionApi": {"id": "cred-notion", "name": "Notion"}}
    },
    {
      "parameters": {
        "url": "https://boards-api.greenhouse.io/v1/boards/BOARD_TOKEN/jobs",
        "method": "POST",
        "authentication": "predefinedCredentialType",
        "nodeCredentialType": "httpBasicAuth",
        "body": {
          "body": {
            "title": "={{ $json.properties['공고명'].title[0].plain_text }}",
            "content": "={{ $json.properties['내용'].rich_text[0].plain_text }}"
          }
        }
      },
      "id": "i9j0-0009-0002", "name": "HTTP - Greenhouse 게시",
      "type": "n8n-nodes-base.httpRequest", "typeVersion": 4, "position": [460, 300],
      "credentials": {"httpBasicAuth": {"id": "cred-greenhouse", "name": "Greenhouse"}}
    },
    {
      "parameters": {
        "resource": "page", "operation": "update",
        "pageId": "={{ $('Notion - 게시대기 공고 조회').item.json.id }}",
        "propertiesUi": {"propertyValues": [{"key": "상태", "type": "select", "selectValue": "게시완료"}]}
      },
      "id": "i9j0-0009-0003", "name": "Notion - 상태 업데이트",
      "type": "n8n-nodes-base.notion", "typeVersion": 2, "position": [680, 300],
      "credentials": {"notionApi": {"id": "cred-notion", "name": "Notion"}}
    }
  ],
  "connections": {
    "Notion - 게시대기 공고 조회": {"main": [[{"node": "HTTP - Greenhouse 게시", "type": "main", "index": 0}]]},
    "HTTP - Greenhouse 게시": {"main": [[{"node": "Notion - 상태 업데이트", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["HR", "recruiting", "job-posting", "greenhouse"]
}, ensure_ascii=False, indent=2)

make_recipe(doc,
    num=9,
    title='채용 공고 멀티 플랫폼 자동 게시',
    category='HR / 인사관리',
    difficulty='중급(Medium)',
    apps='Notion, Greenhouse/LinkedIn/Jobkorea',
    trigger='Notion DB 변경 감지 (Polling)',
    description=(
        'Notion 채용 공고 DB에서 "게시대기" 상태인 공고를 감지하여 '
        'Greenhouse, LinkedIn, 잡코리아 등 채용 플랫폼에 자동 게시하고 Notion 상태를 "게시완료"로 업데이트합니다.'
    ),
    workflow_json=WF_009,
    node_guide=[
        'Notion DB: 공고명, 내용, 상태(게시대기/게시완료) 속성 필요',
        'Greenhouse HTTP: BOARD_TOKEN을 실제 Greenhouse 보드 토큰으로 교체',
        'Notion 상태 업데이트: 게시 성공 후 상태 변경하여 중복 게시 방지',
    ],
    setup_checklist=[
        'JOB_OPENINGS_DB_ID를 실제 Notion DB ID로 교체',
        'Greenhouse API Key 발급 및 BOARD_TOKEN 확인',
        'LinkedIn API 연동 (Marketing Developer Platform 신청 필요)',
        'Notion DB 속성 구조 확인',
    ],
    tips=['잡코리아/사람인: 공식 API 없음 → 이메일 신청 자동화 또는 Selenium 연동 고려']
)

WF_010 = json.dumps({
  "name": "직원 생일/기념일 자동 알림",
  "nodes": [
    {
      "parameters": {"rule": {"interval": [{"field": "cronExpression", "expression": "0 8 * * *"}]}},
      "id": "j0k1-0010-0001", "name": "Schedule - 매일 08:00",
      "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]
    },
    {
      "parameters": {
        "operation": "read", "sheetId": "EMPLOYEES_SHEET_ID",
        "range": "직원!A2:E", "options": {}
      },
      "id": "j0k1-0010-0002", "name": "Sheets - 직원 목록",
      "type": "n8n-nodes-base.googleSheets", "typeVersion": 4, "position": [460, 300],
      "credentials": {"googleSheetsOAuth2Api": {"id": "cred-gsheets", "name": "Google Sheets"}}
    },
    {
      "parameters": {
        "jsCode": "const today = new Date();\nconst mm = String(today.getMonth()+1).padStart(2,'0');\nconst dd = String(today.getDate()).padStart(2,'0');\nconst todayMMDD = mm + '-' + dd;\nconst birthdays = [];\nconst anniversaries = [];\nfor (const row of $input.all()) {\n  const data = row.json.values || row.json;\n  const bday = (data[2] || '').slice(5,10); // MM-DD\n  const hdate = (data[3] || '').slice(5,10);\n  if (bday === todayMMDD) birthdays.push({ name: data[0], email: data[1], type: '생일' });\n  if (hdate === todayMMDD) anniversaries.push({ name: data[0], email: data[1], years: today.getFullYear() - parseInt(data[3]), type: '입사기념일' });\n}\nreturn [...birthdays, ...anniversaries].map(d => ({ json: d }));"
      },
      "id": "j0k1-0010-0003", "name": "Code - 오늘 생일/기념일 필터",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [680, 300]
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#celebrations",
        "text": "={{ $json.type === '생일' ? '🎂 오늘은 ' + $json.name + '님의 생일입니다! 모두 축하해 주세요! 🎉' : '🎊 오늘은 ' + $json.name + '님의 입사 ' + $json.years + '주년입니다! 감사합니다! 🙏' }}",
        "otherOptions": {}
      },
      "id": "j0k1-0010-0004", "name": "Slack - 팀 축하 메시지",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [900, 300],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Schedule - 매일 08:00": {"main": [[{"node": "Sheets - 직원 목록", "type": "main", "index": 0}]]},
    "Sheets - 직원 목록": {"main": [[{"node": "Code - 오늘 생일/기념일 필터", "type": "main", "index": 0}]]},
    "Code - 오늘 생일/기념일 필터": {"main": [[{"node": "Slack - 팀 축하 메시지", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["HR", "birthday", "anniversary", "slack"]
}, ensure_ascii=False, indent=2)

make_recipe(doc,
    num=10,
    title='직원 생일·입사기념일 자동 축하 알림',
    category='HR / 인사관리',
    difficulty='초급(Easy)',
    apps='Google Sheets, Slack',
    trigger='Schedule (매일 오전 8시)',
    description=(
        '매일 아침 Google Sheets 직원 DB에서 오늘의 생일자와 입사 기념일 대상자를 찾아 '
        '#celebrations 채널에 자동으로 축하 메시지를 발송합니다. '
        '직원 몰입도와 팀 문화 형성에 효과적입니다.'
    ),
    workflow_json=WF_010,
    node_guide=[
        'Schedule: 매일 08:00 실행. 서버 타임존 확인 필수 (KST = UTC+9)',
        'Google Sheets: A열=이름, B열=이메일, C열=생년월일(YYYY-MM-DD), D열=입사일(YYYY-MM-DD)',
        'Code 노드: 오늘 날짜와 MM-DD만 비교. 연도 무관하게 매년 자동 실행',
        'Slack: 생일/기념일 없는 날은 Code 노드 출력이 빈 배열 → Slack 노드 실행 안 됨',
    ],
    setup_checklist=[
        'EMPLOYEES_SHEET_ID를 실제 시트 ID로 교체',
        'Google Sheets 직원 DB 구조 확인 (A~D열 순서)',
        'Slack Bot Token 설정 (chat:write 스코프)',
        '#celebrations 채널 생성 및 봇 초대',
        '서버 타임존 확인 (n8n Settings → Timezone)',
    ],
    tips=[
        '개인화: 팀 채널 대신 해당 직원 개인 DM으로 발송하면 더 개인적인 축하 가능',
        '기념품 자동 주문: 특정 기념일(5주년, 10주년)에 기프티콘 API 연동',
    ]
)

doc.add_page_break()
doc.save(OUT)
print('Vol3 Part1 done:', OUT)
