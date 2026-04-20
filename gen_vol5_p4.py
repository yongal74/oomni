# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

PATH = "C:/Users/장우경/oomni/n8n_cases_vol5.docx"

doc = Document(PATH)

def sc(doc, text, level=1):
    styles = {1: ('Heading 1', 16, True, RGBColor(0x1F,0x49,0x7D)),
              2: ('Heading 2', 13, True, RGBColor(0x2E,0x74,0xB5)),
              3: ('Heading 3', 11, True, RGBColor(0x1F,0x49,0x7D))}
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    sz, bold, color = styles[level][1], styles[level][2], styles[level][3]
    run.font.size = Pt(sz); run.font.bold = bold; run.font.color.rgb = color
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p

def h(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p

def par(doc, text, indent=0):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.5)
    return p

def bul(doc, items, indent=1):
    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(("  " * indent) + "• " + item)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

def code(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._element.pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def tbl(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'BDD7EE')
        cell._tc.get_or_add_tcPr().append(shd)
    for rdata in rows:
        row = table.add_row()
        for i, val in enumerate(rdata):
            cell = row.cells[i]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ─── CH8: 프로덕션 AI 운영 가이드 ───────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 8: 프로덕션 AI 운영 가이드", 1)
par(doc, "AI 에이전트를 개발 환경에서 프로덕션으로 이전할 때 발생하는 운영상의 도전과제를 해결하는 종합 가이드입니다. 비용 관리, 모니터링, 보안, 규제 준수의 4대 축을 중심으로 실무 운영 전략을 제시합니다.")

# 8.1 비용 관리
sc(doc, "8.1 LLM API 비용 관리 전략", 2)
par(doc, "AI 에이전트 운영에서 LLM API 비용은 주요 변수입니다. 체계적인 비용 관리 없이는 예산 초과가 발생할 수 있습니다.")

h(doc, "8.1.1 비용 구조 이해")
tbl(doc,
    ["비용 항목", "산정 방식", "일반적 비중", "최적화 레버"],
    [
        ["Input tokens", "프롬프트 + 컨텍스트 길이", "40~60%", "프롬프트 최적화, 컨텍스트 압축"],
        ["Output tokens", "생성 텍스트 길이", "20~40%", "max_tokens 제한, 간결 출력 유도"],
        ["Embedding 호출", "텍스트 → 벡터 변환", "5~15%", "임베딩 캐싱, 배치 처리"],
        ["Tool call 횟수", "함수 호출 + 응답 파싱", "10~20%", "도구 사용 제한, 결과 캐싱"],
        ["Fine-tuning", "학습 데이터 + GPU 시간", "일회성", "소량 데이터 효율화"],
        ["Storage (벡터DB)", "벡터 저장 + 쿼리 비용", "5~10%", "인덱스 최적화, TTL 설정"],
    ],
    [4, 4, 3, 5]
)

h(doc, "8.1.2 모델 라우팅 전략")
par(doc, "작업 복잡도에 따라 적절한 모델을 자동 선택하면 비용을 50~80% 절감할 수 있습니다.")
code(doc, """// 모델 라우팅 로직 (n8n Code 노드)
function selectModel(task) {
  // 간단한 분류/추출 → 저비용 모델
  if (task.type === 'classification' || task.complexity < 3) {
    return { model: 'gpt-4o-mini', maxTokens: 512 };
  }
  // 표준 생성 작업 → 중간 모델
  if (task.complexity < 7) {
    return { model: 'gpt-4o', maxTokens: 2048 };
  }
  // 복잡한 추론/코딩 → 고성능 모델
  return { model: 'claude-opus-4', maxTokens: 4096 };
}

// 비용 예측 함수
function estimateCost(model, inputTokens, outputTokens) {
  const pricing = {
    'gpt-4o-mini':  { input: 0.15, output: 0.60 },   // per 1M tokens
    'gpt-4o':       { input: 5.00, output: 15.00 },
    'claude-opus-4': { input: 15.00, output: 75.00 },
    'claude-sonnet-4': { input: 3.00, output: 15.00 }
  };
  const p = pricing[model];
  return (inputTokens * p.input + outputTokens * p.output) / 1_000_000;
}""")

h(doc, "8.1.3 캐싱 전략")
tbl(doc,
    ["캐싱 유형", "적용 대상", "히트율 목표", "구현 방법"],
    [
        ["Prompt 캐싱", "동일/유사 프롬프트", "30~60%", "Redis, 해시 기반 키"],
        ["Semantic 캐싱", "의미적으로 유사한 쿼리", "20~40%", "임베딩 유사도 검색"],
        ["Embedding 캐싱", "자주 쿼리되는 텍스트", "70~90%", "Redis TTL, 배치 처리"],
        ["결과 캐싱", "동일 입력의 출력", "40~70%", "캐시 DB + 만료 정책"],
        ["Tool 결과 캐싱", "API 호출 결과", "50~80%", "TTL 기반 캐싱"],
    ],
    [4, 4, 3, 5]
)

code(doc, """// n8n Redis 캐싱 패턴
// 1. 캐시 확인 노드 (Redis Get)
{
  "key": "{{ 'ai_cache_' + $crypto.MD5($json.query).toString('hex') }}"
}

// 2. IF 노드: 캐시 히트 여부
// True → 캐시 결과 반환
// False → LLM 호출 → Redis Set (TTL: 3600s)

// Semantic 캐싱 구현
const queryEmbedding = await embed(query);
const cached = await vectorDB.search(queryEmbedding, { topK: 1, threshold: 0.95 });
if (cached.score > 0.95) return cached.content;""")

h(doc, "8.1.4 토큰 최적화 기법")
bul(doc, [
    "컨텍스트 윈도우 압축: 오래된 대화 요약 후 원본 삭제 (토큰 70% 절감)",
    "청킹 최적화: 불필요한 텍스트 제거, 핵심만 RAG 문서에 포함",
    "구조화된 출력: JSON Schema 사용으로 불필요한 설명 제거",
    "배치 처리: 여러 요청을 묶어서 API 호출 횟수 감소",
    "프롬프트 최적화: A/B 테스트로 최소 토큰 프롬프트 선별",
    "스트리밍 취소: 충분한 답변 생성 시 조기 종료",
])

h(doc, "8.1.5 월별 비용 예산 프레임워크")
tbl(doc,
    ["규모", "월간 요청", "예상 비용 범위", "권장 전략"],
    [
        ["스타트업", "1,000~10,000건", "$50~500", "GPT-4o-mini 중심, 캐싱 집중"],
        ["SMB", "10,000~100,000건", "$500~5,000", "모델 라우팅, 임베딩 캐싱"],
        ["중견기업", "100K~1M건", "$5,000~50,000", "전체 캐싱, 로컬 모델 하이브리드"],
        ["대기업", "1M+ 건", "$50,000+", "전용 인프라, Fine-tuning 검토"],
    ],
    [3, 4, 4, 5]
)

# 8.2 모니터링
sc(doc, "8.2 AI 에이전트 모니터링 시스템", 2)
par(doc, "AI 에이전트의 비결정적(non-deterministic) 특성으로 인해 전통적인 소프트웨어 모니터링과 다른 접근이 필요합니다.")

h(doc, "8.2.1 모니터링 4대 차원")
tbl(doc,
    ["차원", "핵심 지표", "임계값 예시", "대응 액션"],
    [
        ["품질", "답변 정확도, 관련성 점수, Hallucination율", "정확도 < 85%", "프롬프트 수정, 모델 교체"],
        ["성능", "응답 시간, TTFT, 토큰/초", "P95 > 10초", "스케일아웃, 캐싱 강화"],
        ["비용", "일일/월간 API 비용, 토큰 효율", "일비용 > 예산 120%", "라우팅 조정, 캐싱 확대"],
        ["신뢰성", "에러율, 재시도율, 타임아웃율", "에러율 > 5%", "Fallback 활성화, 알림"],
    ],
    [3, 5, 4, 4]
)

h(doc, "8.2.2 LLM Observability 구현")
par(doc, "LangSmith, LangFuse, Helicone, Arize AI 등 전문 LLM 관찰성 플랫폼을 n8n과 연동합니다.")
code(doc, """// LangFuse 연동 (n8n HTTP Request 노드)
// 워크플로우 실행 시작 시 trace 생성
POST https://cloud.langfuse.com/api/public/traces
{
  "name": "customer-support-agent",
  "metadata": {
    "userId": "{{ $json.userId }}",
    "sessionId": "{{ $json.sessionId }}",
    "workflowId": "{{ $workflow.id }}"
  }
}

// 각 LLM 호출 후 span 기록
POST https://cloud.langfuse.com/api/public/generations
{
  "traceId": "{{ $json.traceId }}",
  "name": "main-llm-call",
  "model": "gpt-4o",
  "input": { "messages": "{{ $json.messages }}" },
  "output": "{{ $json.response }}",
  "usage": {
    "input": "{{ $json.usage.prompt_tokens }}",
    "output": "{{ $json.usage.completion_tokens }}"
  },
  "startTime": "{{ $json.startTime }}",
  "endTime": "{{ $json.endTime }}"
}""")

h(doc, "8.2.3 Hallucination 감지 시스템")
par(doc, "AI 에이전트가 사실이 아닌 정보를 생성하는 Hallucination은 프로덕션에서 가장 위험한 문제입니다.")
tbl(doc,
    ["감지 방법", "원리", "정확도", "처리 속도"],
    [
        ["Self-consistency", "동일 질문 N회 → 일관성 확인", "70~80%", "느림 (N배 비용)"],
        ["RAG 사실 검증", "응답 vs 소스 문서 비교", "80~90%", "중간"],
        ["NLI 모델", "자연어 추론으로 사실 검증", "75~85%", "빠름"],
        ["외부 검색 교차검증", "웹/DB 검색으로 확인", "85~95%", "느림"],
        ["신뢰도 점수", "모델 출력 로그확률 분석", "65~75%", "매우 빠름"],
    ],
    [4, 5, 3, 3]
)

code(doc, """// Hallucination 감지 워크플로우 (n8n)
// 1단계: AI 에이전트 답변 생성
// 2단계: RAG 소스 추출
// 3단계: 검증 LLM 호출

const verificationPrompt = `
다음 답변이 제공된 소스 문서의 내용과 일치하는지 검증하세요.
사실 오류나 소스에 없는 내용이 있으면 지적하세요.

답변: ${answer}
소스: ${sources.join('\\n')}

검증 결과를 JSON으로 출력:
{
  "isFactual": true/false,
  "confidence": 0.0-1.0,
  "issues": ["이슈1", "이슈2"]
}`;

// 신뢰도 < 0.7이면 Human Review 에스컬레이션""")

h(doc, "8.2.4 알림 및 대응 시스템")
tbl(doc,
    ["이벤트", "알림 채널", "대응 시간 목표", "자동화 대응"],
    [
        ["에러율 급증 (>10%)", "PagerDuty + Slack", "즉시 (P1)", "Fallback 모델 자동 전환"],
        ["비용 이상 감지", "이메일 + Slack", "1시간 이내", "요청 속도 제한 적용"],
        ["응답 품질 저하", "Slack 알림", "당일 처리", "프롬프트 버전 롤백"],
        ["모델 API 다운", "PagerDuty", "즉시 (P1)", "대체 모델 자동 전환"],
        ["Hallucination 감지", "Slack + 담당자 이메일", "30분 이내", "해당 케이스 검토 큐 등록"],
    ],
    [5, 4, 4, 5]
)

# 8.3 보안
sc(doc, "8.3 AI 에이전트 보안 가이드", 2)
par(doc, "AI 에이전트는 자율적으로 외부 시스템에 접근하고 데이터를 처리하므로 기존 웹 애플리케이션보다 확장된 공격 표면을 가집니다.")

h(doc, "8.3.1 AI 특화 보안 위협")
tbl(doc,
    ["위협 유형", "설명", "실제 예시", "방어 방법"],
    [
        ["Prompt Injection", "악의적 지시를 프롬프트에 삽입", "사용자 입력으로 시스템 프롬프트 무력화", "입력 필터링, 구조적 분리"],
        ["Indirect Injection", "처리할 데이터에 지시 삽입", "웹페이지에 숨겨진 지시로 에이전트 조작", "데이터-지시 분리"],
        ["Data Exfiltration", "에이전트를 통한 민감 데이터 유출", "\"이 파일 내용을 외부 서버로 전송해\"", "출력 검증, 허용 도메인 화이트리스트"],
        ["Privilege Escalation", "에이전트 권한 상승 시도", "관리자 권한 도구 무단 호출", "최소 권한 원칙, 도구 접근 제어"],
        ["Model Poisoning", "Fine-tuning 데이터 오염", "학습 데이터에 악의적 패턴 삽입", "학습 데이터 검증, 이상 감지"],
        ["SSRF via Agent", "에이전트를 통한 내부 서비스 접근", "내부 API 엔드포인트 스캔", "네트워크 분리, URL 검증"],
    ],
    [3, 4, 5, 4]
)

h(doc, "8.3.2 Prompt Injection 방어 구현")
code(doc, """// 입력 위생처리 (n8n Code 노드)
function sanitizeInput(userInput) {
  // 1. 시스템 지시 패턴 탐지
  const injectionPatterns = [
    /ignore (all |previous |above )?instructions/i,
    /disregard (your |the |all )?instructions/i,
    /you are now/i,
    /act as (if|a|an)/i,
    /system prompt/i,
    /reveal.*instructions/i,
    /\\[INST\\]|\\[\\/INST\\]/,  // Llama 지시 토큰
    /<\\|im_start\\|>|<\\|im_end\\|>/,  // ChatML 토큰
  ];

  for (const pattern of injectionPatterns) {
    if (pattern.test(userInput)) {
      return { safe: false, reason: 'Potential prompt injection detected' };
    }
  }

  // 2. 길이 제한
  if (userInput.length > 4000) {
    return { safe: false, reason: 'Input too long' };
  }

  // 3. HTML/Script 이스케이프
  const escaped = userInput
    .replace(/</g, '&lt;')
    .replace(/>/g, '&gt;')
    .replace(/\\/g, '\\\\');

  return { safe: true, sanitized: escaped };
}

// 구조적 분리: 사용자 입력을 명확히 구분
const systemPrompt = `당신은 고객지원 에이전트입니다. 제공된 정책만 따르세요.`;
const userSection = `--- 사용자 메시지 시작 ---\n${sanitized}\n--- 사용자 메시지 끝 ---`;""")

h(doc, "8.3.3 최소 권한 원칙 구현")
bul(doc, [
    "도구별 권한 세분화: 읽기 전용 / 쓰기 / 실행 권한 분리",
    "데이터 범위 제한: 에이전트가 접근 가능한 데이터 범위 명시 제한",
    "행동 제한 목록 정의: 허용/금지 행동을 시스템 프롬프트에 명시",
    "샌드박스 실행: 코드 실행 에이전트는 격리된 환경에서만 실행",
    "감사 로그: 모든 도구 호출 기록 및 이상 행동 감지",
])

code(doc, """// n8n 도구 접근 제어 설정
{
  "agent_config": {
    "allowed_tools": ["search_kb", "send_slack", "create_ticket"],
    "forbidden_actions": [
      "delete_*",
      "admin_*",
      "export_all_*",
      "modify_system_*"
    ],
    "data_scope": {
      "customer_data": "read_own_only",
      "internal_docs": "read_department_only",
      "financial_data": "none"
    },
    "output_validation": {
      "block_pii": true,
      "block_secrets": true,
      "allowed_domains": ["company.com", "trusted-partner.com"]
    }
  }
}""")

h(doc, "8.3.4 PII(개인정보) 보호")
tbl(doc,
    ["PII 유형", "감지 방법", "처리 방식", "보관 정책"],
    [
        ["이름/연락처", "정규식 + NER 모델", "마스킹 (**** 처리)", "30일 후 삭제"],
        ["주민등록번호", "패턴 매칭", "즉시 거부/삭제", "보관 금지"],
        ["신용카드번호", "Luhn 알고리즘", "마스킹, 토큰화", "PCI DSS 준수"],
        ["의료정보", "의료 NER 모델", "암호화 저장", "HIPAA 준수"],
        ["이메일/IP", "정규식", "익명화/해시화", "90일 후 삭제"],
    ],
    [3, 4, 4, 4]
)

# 8.4 규제 준수
sc(doc, "8.4 AI 규제 및 컴플라이언스", 2)
par(doc, "EU AI Act, GDPR, 국내 AI 기본법 등 AI 관련 규제가 강화되고 있습니다. 프로덕션 AI 에이전트는 관련 규제를 준수해야 합니다.")

h(doc, "8.4.1 주요 AI 규제 현황")
tbl(doc,
    ["규제", "적용 대상", "핵심 요구사항", "미준수 시 제재"],
    [
        ["EU AI Act (2024)", "EU 시장 AI 시스템", "고위험 AI 등록, 투명성, 인간 감독", "최대 €3천만 또는 매출 6%"],
        ["GDPR", "EU 시민 데이터 처리", "목적 제한, 데이터 최소화, 권리 보장", "최대 €2천만 또는 매출 4%"],
        ["국내 AI 기본법 (2024)", "국내 AI 사업자", "고영향 AI 안전성 평가, 설명 의무", "과태료 + 사업 중단"],
        ["미국 EO 14110", "연방 AI 시스템", "안전 테스트, 투명성 보고", "조달 자격 박탈"],
        ["PCI DSS v4.0", "결제 데이터 처리", "AI 모델의 카드 데이터 처리 제한", "가맹점 자격 박탈"],
    ],
    [3, 4, 5, 4]
)

h(doc, "8.4.2 AI 거버넌스 프레임워크")
par(doc, "프로덕션 AI 에이전트 운영을 위한 거버넌스 체계를 수립합니다.")
bul(doc, [
    "AI 인벤토리 관리: 운영 중인 모든 AI 에이전트 등록 및 위험 분류",
    "설명 가능성(XAI): 주요 결정에 대한 근거 제공 메커니즘 구현",
    "인간 감독(Human Oversight): 고위험 결정에 대한 검토/승인 프로세스",
    "편향성(Bias) 테스트: 정기적인 공정성 평가 및 편향 모니터링",
    "모델 카드 작성: 각 AI 모델의 사용 목적, 한계, 성능 문서화",
    "감사 추적: 모든 AI 의사결정 이력 보관 (최소 3년)",
    "인시던트 대응: AI 오동작 시 신속 대응 프로세스 정의",
])

h(doc, "8.4.3 데이터 주권 및 처리 위치")
tbl(doc,
    ["산업", "데이터 주권 요구사항", "권장 아키텍처"],
    [
        ["금융", "국내 데이터 처리 의무, 해외 전송 제한", "온프레미스 or 국내 클라우드 + 로컬 LLM"],
        ["의료", "의료정보 국외 이전 금지 (일부 예외)", "폐쇄망 환경, Self-hosted 모델"],
        ["공공", "공공데이터 국내 처리 원칙", "정부 클라우드 (G-Cloud), 오픈소스 LLM"],
        ["일반기업", "GDPR 대상 시 EU 처리 or SCCs", "Azure/AWS EU Region, BCR 체결"],
    ],
    [3, 5, 6]
)

# 8.5 성능 최적화
sc(doc, "8.5 응답 성능 최적화", 2)

h(doc, "8.5.1 응답 시간 최적화 기법")
tbl(doc,
    ["기법", "효과", "구현 난이도", "적용 시나리오"],
    [
        ["스트리밍 응답", "체감 속도 70% 향상", "낮음", "챗봇, 실시간 대화"],
        ["병렬 도구 호출", "총 시간 40~60% 감소", "중간", "멀티 소스 검색"],
        ["프리페치(Prefetch)", "예측 쿼리 미리 처리", "높음", "패턴이 명확한 워크플로우"],
        ["모델 양자화", "추론 속도 2~4배 향상", "높음", "셀프호스팅 모델"],
        ["KV Cache 활용", "반복 프롬프트 50% 빠름", "낮음", "시스템 프롬프트 고정"],
        ["Speculative Decoding", "생성 속도 2~3배 향상", "매우 높음", "대용량 모델 운영"],
    ],
    [3.5, 3, 3, 5]
)

h(doc, "8.5.2 확장성 설계")
code(doc, """// n8n 워커 스케일링 설정 (docker-compose.yml 참고)
// 큐 모드 활성화로 워커 수평 확장
{
  "EXECUTIONS_MODE": "queue",
  "QUEUE_BULL_REDIS_HOST": "redis",
  "N8N_CONCURRENCY_PRODUCTION_LIMIT": 10,
  // AI 워크플로우용 전용 워커 설정
  "QUEUE_WORKER_TIMEOUT": 300000,  // 5분 타임아웃
  "N8N_PAYLOAD_SIZE_MAX": 64  // 64MB 최대 페이로드
}

// 로드 밸런싱: AI 에이전트 별 우선순위 큐 분리
// High Priority: 실시간 고객 대화
// Medium Priority: 문서 분석
// Low Priority: 배치 처리, 리포트 생성""")

h(doc, "8.5.3 오류 처리 및 복원력")
tbl(doc,
    ["오류 유형", "감지 방법", "즉각 대응", "장기 대응"],
    [
        ["API Rate Limit", "429 응답 감지", "지수 백오프 재시도", "요청 분산, 다중 API 키"],
        ["API 다운타임", "5xx 응답, 타임아웃", "대체 모델 자동 전환", "멀티 프로바이더 아키텍처"],
        ["컨텍스트 초과", "토큰 한계 오류", "컨텍스트 압축 후 재시도", "청킹 전략 최적화"],
        ["품질 저하", "평가 점수 임계값 하회", "프롬프트 버전 롤백", "프롬프트 재최적화"],
        ["무한 루프", "실행 횟수 카운터", "강제 중단 + 알림", "루프 탈출 조건 강화"],
    ],
    [3.5, 4, 4.5, 4]
)

# ─── 부록 ───────────────────────────────────────────────────────────────────
doc.add_page_break()
sc(doc, "부록 A: n8n AI 노드 연결 치트시트", 1)
par(doc, "n8n의 AI 관련 노드들의 연결 포트와 호환 관계를 빠르게 참조할 수 있는 가이드입니다.")

sc(doc, "A.1 AI 연결 포트 종류", 2)
tbl(doc,
    ["포트 색상", "포트 이름", "연결 방향", "연결 가능 노드"],
    [
        ["보라색", "ai_languageModel", "입력", "Agent, Chain 노드"],
        ["주황색", "ai_tool", "입력", "Agent 노드만"],
        ["파란색", "ai_memory", "입력", "Agent, Chain 노드"],
        ["초록색", "ai_vectorStore", "입력", "Agent, Retriever 노드"],
        ["분홍색", "ai_document", "입력", "Vector Store 노드"],
        ["회색", "ai_textSplitter", "입력", "Document Loader 노드"],
        ["노란색", "ai_embedding", "입력", "Vector Store 노드"],
        ["청록색", "ai_outputParser", "입력", "Chain, Agent 노드"],
        ["빨간색", "ai_retriever", "입력", "Chain 노드"],
    ],
    [3, 4, 3, 5]
)

sc(doc, "A.2 자주 쓰는 AI 워크플로우 구성", 2)
h(doc, "A.2.1 기본 챗봇")
code(doc, """Chat Trigger
  ↓ (main)
AI Agent ←─ [ai_languageModel] ─ OpenAI Chat Model (gpt-4o)
           ←─ [ai_memory] ─────── Window Buffer Memory (10턴)
           ←─ [ai_tool] ────────── Calculator Tool
           ←─ [ai_tool] ────────── SerpAPI Tool
  ↓ (main)
Respond to Webhook""")

h(doc, "A.2.2 RAG 에이전트")
code(doc, """Webhook Trigger
  ↓
AI Agent ←─ [ai_languageModel] ─ Anthropic Claude
          ←─ [ai_memory] ─────── Redis Chat Memory
          ←─ [ai_tool] ────────── Vector Store Tool
                                    ← [ai_vectorStore] ─ Pinecone
                                    ← [ai_embedding] ─── OpenAI Embeddings
  ↓
HTTP Response""")

h(doc, "A.2.3 문서 임베딩 파이프라인")
code(doc, """File Upload Trigger
  ↓
Default Data Loader
  ← [ai_textSplitter] ─ Recursive Character Text Splitter
  ↓ (ai_document)
Pinecone Vector Store (Insert)
  ← [ai_embedding] ─── OpenAI Embeddings (text-embedding-3-small)""")

h(doc, "A.2.4 멀티에이전트 (Supervisor 패턴)")
code(doc, """Webhook
  ↓
Supervisor Agent (gpt-4o)
  ├─→ [sub1] Research Agent ←─ [tool] Tavily Search
  ├─→ [sub2] Analysis Agent ←─ [tool] Code Interpreter
  └─→ [sub3] Writer Agent   ←─ [tool] Google Docs
  ↓
Merge & Format
  ↓
Response""")

doc.add_page_break()
sc(doc, "부록 B: AI 에이전트 디버깅 가이드", 1)

sc(doc, "B.1 일반적인 문제와 해결법", 2)
tbl(doc,
    ["증상", "가능한 원인", "진단 방법", "해결책"],
    [
        ["에이전트가 루프에 빠짐", "최대 반복 횟수 미설정, 탈출 조건 부재", "실행 로그의 반복 패턴 확인", "maxIterations 설정, 명시적 종료 조건 추가"],
        ["도구 호출 실패 반복", "도구 입력 스키마 불일치", "도구 에러 메시지 확인", "도구 설명과 파라미터 스키마 명확화"],
        ["답변 품질 저하", "프롬프트 모호성, 컨텍스트 부족", "프롬프트 로그 분석", "시스템 프롬프트 개선, Few-shot 예시 추가"],
        ["RAG 검색 실패", "임베딩 모델 불일치, 청킹 문제", "검색 결과 점수 확인", "동일 임베딩 모델 사용, 청킹 크기 조정"],
        ["응답 시간 과다", "도구 체인 직렬화, 큰 컨텍스트", "각 단계 시간 측정", "병렬 도구 호출, 컨텍스트 압축"],
        ["메모리 컨텍스트 오염", "이전 세션 데이터 잔존", "메모리 내용 덤프 확인", "세션별 메모리 격리, TTL 설정"],
    ],
    [4, 4, 4, 4]
)

sc(doc, "B.2 n8n AI 에이전트 디버깅 체크리스트", 2)
h(doc, "개발 단계")
bul(doc, [
    "[ ] 시스템 프롬프트가 명확하고 역할이 정의되어 있는가?",
    "[ ] 각 도구의 description이 충분히 상세한가?",
    "[ ] 도구의 입력/출력 스키마가 정확히 정의되어 있는가?",
    "[ ] maxIterations 값이 적절히 설정되어 있는가? (권장: 5~15)",
    "[ ] 에이전트 타입이 작업에 적합한가? (ReAct vs Plan & Execute vs OpenAI Functions)",
    "[ ] 메모리 타입과 크기가 적절한가?",
    "[ ] 오류 처리 노드(Try/Catch)가 구현되어 있는가?",
])

h(doc, "RAG 설정")
bul(doc, [
    "[ ] 임베딩 모델이 저장 시와 검색 시 동일한가?",
    "[ ] 청킹 크기가 컨텍스트 윈도우 대비 적절한가?",
    "[ ] topK 값이 적절한가? (일반적으로 3~5)",
    "[ ] 유사도 임계값이 설정되어 있는가?",
    "[ ] 메타데이터 필터가 올바르게 적용되는가?",
])

h(doc, "프로덕션 전환")
bul(doc, [
    "[ ] API 키가 n8n Credential로 안전하게 저장되었는가?",
    "[ ] Rate limit 처리(백오프)가 구현되어 있는가?",
    "[ ] 비용 모니터링이 설정되어 있는가?",
    "[ ] 알림 시스템이 연결되어 있는가?",
    "[ ] 로그 보관 정책이 수립되어 있는가?",
    "[ ] Fallback 모델/시나리오가 준비되어 있는가?",
    "[ ] 보안 취약점 점검(Prompt Injection 방어 등)이 완료되었는가?",
])

doc.add_page_break()
sc(doc, "부록 C: AI 에이전트 성능 벤치마크", 1)

sc(doc, "C.1 모델별 실무 성능 (2024년 기준)", 2)
tbl(doc,
    ["모델", "MMLU", "HumanEval", "GSM8K", "응답속도(TTFT)", "비용/1M tok"],
    [
        ["GPT-4o", "88.7%", "90.2%", "92.0%", "0.8초", "$5/$15"],
        ["GPT-4o-mini", "82.0%", "87.2%", "87.0%", "0.5초", "$0.15/$0.60"],
        ["Claude 3.5 Sonnet", "88.7%", "92.0%", "95.0%", "1.2초", "$3/$15"],
        ["Claude 3 Haiku", "75.2%", "75.9%", "88.9%", "0.6초", "$0.25/$1.25"],
        ["Gemini 1.5 Pro", "85.9%", "84.1%", "90.8%", "1.5초", "$1.25/$5"],
        ["Llama 3.1 70B", "82.0%", "80.5%", "87.3%", "0.4초 (셀프호스팅)", "인프라비용만"],
    ],
    [4, 3, 3, 4, 4]
)
par(doc, "* 입력 토큰 비용 / 출력 토큰 비용 (USD per 1M tokens)")

sc(doc, "C.2 사용 사례별 권장 모델", 2)
tbl(doc,
    ["사용 사례", "1순위", "2순위 (대안)", "선택 이유"],
    [
        ["고객 지원 챗봇", "GPT-4o-mini", "Claude 3 Haiku", "비용 효율, 빠른 응답"],
        ["코드 생성/리뷰", "Claude 3.5 Sonnet", "GPT-4o", "높은 HumanEval 성능"],
        ["문서 분석/요약", "Gemini 1.5 Pro", "Claude 3.5 Sonnet", "긴 컨텍스트 처리"],
        ["수학/논리 추론", "Claude 3.5 Sonnet", "GPT-4o", "높은 GSM8K 성능"],
        ["다국어 처리", "GPT-4o", "Claude 3.5 Sonnet", "한국어 포함 다국어 강점"],
        ["비용 민감 대량 처리", "Llama 3.1 70B (셀프)", "GPT-4o-mini", "API 비용 제로"],
        ["최고 정확도 요구", "Claude 3.5 Sonnet", "GPT-4o", "SOTA 성능"],
    ],
    [4, 4, 4, 5]
)

doc.add_page_break()
sc(doc, "부록 D: 산업별 AI 에이전트 도입 로드맵", 1)

tbl(doc,
    ["단계", "기간", "주요 마일스톤", "성공 지표"],
    [
        ["1단계: 탐색", "1~2개월", "파일럿 유스케이스 선정, POC 구축, 내부 데모", "기술 검증 완료, 경영진 승인"],
        ["2단계: 구축", "2~4개월", "핵심 워크플로우 개발, 시스템 연동, 테스트", "유닛/통합 테스트 통과"],
        ["3단계: 배포", "1~2개월", "스테이징 검증, 보안 감사, 점진적 롤아웃", "에러율 < 5%, 사용자 승인"],
        ["4단계: 운영", "지속", "모니터링 최적화, 성능 개선, 확장", "ROI 목표 달성, 만족도 > 80%"],
        ["5단계: 확장", "3~6개월", "추가 유스케이스 적용, 조직 문화 내재화", "자동화율 목표 달성"],
    ],
    [2.5, 2.5, 6, 4]
)

sc(doc, "D.1 산업별 우선 도입 권장 영역", 2)
tbl(doc,
    ["산업", "1순위 도입 영역", "예상 ROI", "도입 난이도"],
    [
        ["금융/보험", "AML 이상거래 감지, 고객 상담 자동화", "300~500%", "높음 (규제 준수 필요)"],
        ["제조", "품질 검사 AI, 예측 유지보수", "200~400%", "중간"],
        ["유통/이커머스", "개인화 추천, CS 챗봇", "150~300%", "낮음"],
        ["의료", "의료기록 요약, 예약 자동화", "200~350%", "매우 높음 (규제)"],
        ["IT/SW", "코드 리뷰 자동화, 인시던트 대응", "400~700%", "낮음"],
        ["물류", "경로 최적화, 수요 예측", "250~450%", "중간"],
        ["교육", "맞춤형 튜터링, 채점 자동화", "150~250%", "낮음"],
    ],
    [3, 5, 3, 3]
)

# 최종 저장
doc.save(PATH)
print("[OK] gen_vol5_p4.py 완료 - n8n_cases_vol5.docx Ch8 + 부록 추가")
