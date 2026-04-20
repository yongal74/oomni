# -*- coding: utf-8 -*-
"""Vol 5 Part 2: Ch3 멀티에이전트 패턴 + Ch4 RAG 심화"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'C:/Users/장우경/oomni/n8n_cases_vol5.docx'
doc = Document(OUT)

def sc(cell,hx):
    tc=cell._tc;pr=tc.get_or_add_tcPr()
    s=OxmlElement('w:shd');s.set(qn('w:val'),'clear');s.set(qn('w:color'),'auto');s.set(qn('w:fill'),hx);pr.append(s)

def h(doc,text,lv=1):
    p=doc.add_heading(text,level=lv)
    C={1:RGBColor(0x1F,0x49,0x7D),2:RGBColor(0x2E,0x75,0xB6),3:RGBColor(0x17,0x5E,0x40),4:RGBColor(0x40,0x40,0x40)}
    S={1:Pt(22),2:Pt(16),3:Pt(13),4:Pt(11)}
    if p.runs:p.runs[0].font.name='맑은 고딕';p.runs[0].font.size=S.get(lv,Pt(11));p.runs[0].font.color.rgb=C.get(lv,RGBColor(0,0,0))
    return p

def par(doc,text,bold=False,color=None,size=Pt(10),indent=None):
    p=doc.add_paragraph()
    if indent:p.paragraph_format.left_indent=indent
    r=p.add_run(text);r.font.name='맑은 고딕';r.font.size=size;r.bold=bold
    if color:r.font.color.rgb=color
    return p

def bul(doc,text,lv=0):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent=Cm(0.5+lv*0.5)
    r=p.add_run(text);r.font.name='맑은 고딕';r.font.size=Pt(10);return p

def code(doc,text):
    p=doc.add_paragraph();p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run(text);r.font.name='Courier New';r.font.size=Pt(8)
    pr=p._p.get_or_add_pPr();s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear');s.set(qn('w:color'),'auto');s.set(qn('w:fill'),'F2F2F2');pr.append(s)
    return p

def tbl(doc,headers,rows,col_widths,hc='1F497D'):
    t=doc.add_table(rows=len(rows)+1,cols=len(headers));t.style='Table Grid';t.autofit=False
    for i,w in enumerate(col_widths):t.columns[i].width=w
    for ci,hdr in enumerate(headers):
        sc(t.rows[0].cells[ci],hc)
        t.rows[0].cells[ci].text=hdr
        for p in t.rows[0].cells[ci].paragraphs:
            for r in p.runs:r.font.name='맑은 고딕';r.font.size=Pt(9);r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            sc(t.rows[ri+1].cells[ci],'EBF3FB' if ci==0 else 'FFFFFF')
            t.rows[ri+1].cells[ci].text=str(val)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:r.font.name='맑은 고딕';r.font.size=Pt(9)
                if ci==0:
                    for r in p.runs:r.bold=True
    return t

# ════════════════════════════════════════════════════
# CHAPTER 3: 멀티에이전트 10종 설계 패턴
# ════════════════════════════════════════════════════
h(doc,'Chapter 3. 멀티에이전트 10종 설계 패턴',1)
par(doc,'단일 AI 에이전트의 한계를 극복하는 멀티에이전트 아키텍처를 설계하는 10가지 핵심 패턴을 '
    '상세히 분석합니다. 각 패턴의 구조, 적합한 케이스, n8n 구현 방법을 제공합니다.', size=Pt(10))
doc.add_paragraph()

h(doc,'3.1 멀티에이전트가 필요한 이유',2)
reasons=[
    ('컨텍스트 한계 극복','단일 에이전트의 컨텍스트 윈도우 제한을 분산 처리로 해결'),
    ('전문화','각 에이전트가 특정 영역에 특화되어 전체 품질 향상'),
    ('병렬 처리','독립적인 태스크를 동시 처리하여 속도 향상'),
    ('체크-앤-밸런스','한 에이전트의 출력을 다른 에이전트가 검증'),
    ('비용 최적화','간단한 태스크는 저렴한 모델, 복잡한 태스크는 고급 모델 배정'),
    ('오류 격리','한 에이전트 실패가 전체 시스템에 영향 최소화'),
]
tbl(doc,['이유','설명'],reasons,[Cm(5),Cm(12.5)],hc='2E75B6')
doc.add_paragraph()

# 패턴 1
h(doc,'패턴 1: Supervisor-Worker 패턴',2)
par(doc,'가장 일반적인 멀티에이전트 패턴입니다. 상위 Supervisor 에이전트가 태스크를 분해하고 '
    '적절한 Worker 에이전트에게 위임합니다. Worker 결과를 수집하여 최종 응답을 생성합니다.', size=Pt(10))
code(doc,
'Supervisor-Worker 패턴 구조\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'  [사용자 요청]\n'
'       │\n'
'       ▼\n'
'  ┌─────────────────────────────────────┐\n'
'  │     Supervisor Agent (GPT-4o)        │\n'
'  │  "태스크를 분해하고 Worker 지정"     │\n'
'  └──────┬─────────┬──────────┬─────────┘\n'
'         │         │          │\n'
'         ▼         ▼          ▼\n'
'   ┌──────────┐ ┌──────────┐ ┌──────────┐\n'
'   │Research  │ │Writing   │ │Review    │\n'
'   │Worker    │ │Worker    │ │Worker    │\n'
'   │(Claude)  │ │(GPT-4o)  │ │(GPT-4o-m)│\n'
'   └────┬─────┘ └────┬─────┘ └────┬─────┘\n'
'        │             │             │\n'
'        └─────────────┼─────────────┘\n'
'                      ▼\n'
'               [최종 결과 통합]\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()
par(doc,'n8n 구현 방법:', bold=True, color=RGBColor(0x1F,0x49,0x7D))
bul(doc,'Supervisor AI Agent → toolWorkflow 도구로 각 Worker 워크플로우 호출')
bul(doc,'각 Worker는 별도 n8n 워크플로우로 구성 (전문화된 시스템 프롬프트)')
bul(doc,'Supervisor가 Worker 결과를 수집하여 최종 통합')
bul(doc,'실제 구현: Supervisor의 systemMessage에 Worker 목록과 호출 기준 명시')
doc.add_paragraph()
par(doc,'적합한 케이스:', bold=True)
bul(doc,'복잡한 리서치 리포트 자동 작성 (리서치→분석→작성→검토 분리)')
bul(doc,'멀티모달 콘텐츠 제작 (텍스트+이미지+SEO 각각 전담)')
bul(doc,'고객 문의 처리 (분류→FAQ 검색→답변 생성→품질 검토)')
doc.add_paragraph()

# 패턴 2
h(doc,'패턴 2: Parallel 에이전트 패턴',2)
par(doc,'여러 에이전트가 독립적인 태스크를 동시에 처리하고 결과를 병합합니다. '
    'n8n의 병렬 실행 특성을 최대로 활용하는 패턴입니다.', size=Pt(10))
code(doc,
'Parallel 에이전트 패턴\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'[입력 데이터]\n'
'      │\n'
'      ├──────────────────────────────────┐\n'
'      │                                  │\n'
'      ▼ (동시 실행)                      ▼\n'
'┌──────────────────┐          ┌──────────────────┐\n'
'│  감성 분석       │          │  키워드 추출      │\n'
'│  Agent A         │          │  Agent B          │\n'
'└────────┬─────────┘          └────────┬──────────┘\n'
'         │                             │\n'
'         └──────────┬──────────────────┘\n'
'                    ▼\n'
'            [결과 병합 노드]\n'
'                    │\n'
'                    ▼\n'
'             [최종 출력]\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()
par(doc,'n8n 구현 방법:', bold=True, color=RGBColor(0x1F,0x49,0x7D))
bul(doc,'하나의 노드에서 여러 AI Agent 노드로 동시 연결 (n8n 자동 병렬 처리)')
bul(doc,'각 Agent 완료 후 Merge 노드(Combine All)로 결과 합치기')
bul(doc,'Code 노드에서 여러 Agent 결과를 하나의 객체로 통합')
doc.add_paragraph()

# 패턴 3
h(doc,'패턴 3: Pipeline 체인 패턴',2)
par(doc,'에이전트들이 순서대로 실행되며 이전 에이전트의 출력이 다음 에이전트의 입력이 됩니다. '
    '단계별 점진적 정제에 적합합니다.', size=Pt(10))
code(doc,
'Pipeline 체인 패턴\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'[원본 데이터]\n'
'     │\n'
'     ▼\n'
'[Agent 1: 번역]   → 한국어 → 영어 변환\n'
'     │\n'
'     ▼\n'
'[Agent 2: 요약]   → 영어 원문 → 핵심 요약 (200자)\n'
'     │\n'
'     ▼\n'
'[Agent 3: SEO]    → 요약 → SEO 키워드 최적화\n'
'     │\n'
'     ▼\n'
'[Agent 4: 리뷰]   → 최종 검토 및 품질 점수 부여\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()
par(doc,'적합한 케이스: 콘텐츠 생성 파이프라인, 번역→요약→발행, 코드 생성→리뷰→테스트 자동화', size=Pt(10))
doc.add_paragraph()

# 패턴 4
h(doc,'패턴 4: Debate 패턴 (비평-수정)',2)
par(doc,'하나의 에이전트가 답변을 생성하고, 다른 에이전트가 비판하며, '
    '원래 에이전트가 피드백을 반영하여 수정합니다. 품질이 중요한 작업에 적합합니다.', size=Pt(10))
code(doc,
'Debate 패턴 (자기 수정 루프)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'[초안 생성 에이전트]\n'
'    │\n'
'    ▼ (초안)\n'
'[비평 에이전트] → "논리적 오류, 개선 제안"\n'
'    │\n'
'    ▼ (비평)\n'
'[수정 에이전트] → 비평 반영하여 개선\n'
'    │\n'
'    ├─ 품질 점수 < 8/10 → [비평 에이전트] 루프\n'
'    └─ 품질 점수 ≥ 8/10 → [최종 출력]\n'
'\n'
'n8n 구현: Loop 노드 + IF 노드로 품질 임계값 달성까지 반복\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

# 패턴 5
h(doc,'패턴 5: Specialist Router 패턴',2)
par(doc,'입력을 분석하여 가장 적합한 전문가 에이전트에게 라우팅합니다. '
    '다양한 유형의 요청을 처리하는 범용 AI 시스템에 적합합니다.', size=Pt(10))
code(doc,
'Specialist Router 패턴\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'[사용자 질문]\n'
'      │\n'
'      ▼\n'
'[Router Agent / IF 노드]\n'
'"이 질문의 카테고리는?"\n'
'      │\n'
'      ├─ 법률 질문 ──────▶ [법률 전문 Agent]\n'
'      │                    (법률 문서 RAG + 법조문 DB)\n'
'      ├─ 기술 지원 ─────▶ [Tech 전문 Agent]\n'
'      │                    (매뉴얼 RAG + 에러 코드 DB)\n'
'      ├─ 결제/환불 ─────▶ [결제 전문 Agent]\n'
'      │                    (주문 DB + 환불 정책)\n'
'      └─ 일반 문의 ─────▶ [General Agent]\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
par(doc,'n8n 구현 방법:', bold=True, color=RGBColor(0x1F,0x49,0x7D))
bul(doc,'방법 1: GPT-4o-mini로 카테고리 분류 → Switch 노드로 라우팅')
bul(doc,'방법 2: AI Agent에 toolWorkflow 여러 개 연결 → 에이전트가 자율 선택')
bul(doc,'방법 3: 키워드 기반 IF 노드 (빠름, 저렴, 단순)')
doc.add_paragraph()

# 패턴 6-10 요약
patterns_56_10=[
    ('패턴 6: Map-Reduce 패턴',
     '대용량 데이터를 청크로 분할하여 병렬 처리(Map) 후 결과를 통합(Reduce).\n'
     '예: 1000개 리뷰 동시 감성분석 → 결과 집계\n'
     'n8n: Split In Batches → 병렬 AI Agent → Aggregate 노드'),
    ('패턴 7: Human-in-the-Loop 패턴',
     '에이전트가 불확실한 경우 인간에게 승인/피드백을 요청하고 재개.\n'
     '예: AI 초안 생성 → 담당자 Slack 승인 → 승인 시 자동 발행\n'
     'n8n: AI Agent → Wait 노드(Webhook) → 승인 후 재개'),
    ('패턴 8: 에이전트 메모리 공유 패턴',
     '여러 에이전트가 공유 벡터 DB/Redis를 통해 지식을 공유.\n'
     '에이전트 A가 학습한 내용을 에이전트 B가 활용.\n'
     'n8n: 공유 Pinecone/Redis → 모든 Agent에 동일 VectorStore 연결'),
    ('패턴 9: Reflection 패턴',
     '에이전트가 자신의 출력을 스스로 평가하고 개선하는 자기성찰 루프.\n'
     '"이 답변의 문제점은? 어떻게 개선할 수 있나?" 질문을 자동 생성\n'
     'n8n: AI Agent → Code(자기평가 프롬프트 생성) → 루프'),
    ('패턴 10: 이벤트 주도 에이전트 패턴',
     '외부 이벤트(이메일, 알럿, 웹훅)에 자율적으로 반응하는 에이전트.\n'
     '"항상 깨어있는" 자율 AI 시스템 구현 패턴.\n'
     'n8n: 여러 트리거 → 이벤트 분류 → 적절한 에이전트 실행'),
]
for title, desc in patterns_56_10:
    h(doc, title, 2)
    code(doc, desc)
    doc.add_paragraph()

h(doc,'3.2 멀티에이전트 패턴 선택 가이드',2)
pattern_guide=[
    ('단순 자동화','단일 Agent','빠름/저렴','코드 리뷰, 번역, 분류'),
    ('순차 처리','Pipeline Chain','중간','번역→요약→발행'),
    ('병렬 처리 필요','Parallel','빠름','동시 다중 분석'),
    ('품질 중시','Debate','느림/비쌈','계약서 검토, 의학 진단'),
    ('다양한 요청 유형','Specialist Router','중간','범용 고객지원'),
    ('대용량 처리','Map-Reduce','매우 빠름','1000+ 문서 처리'),
    ('인간 감독 필요','Human-in-the-Loop','중간','금융/의료 자동화'),
    ('자율 학습','Memory Sharing','장기적 향상','개인화 AI 비서'),
    ('고품질 출력','Reflection','느림','광고 카피, 법적 문서'),
    ('실시간 반응','Event-Driven','즉각','SOC 자동화, 알럿 처리'),
]
tbl(doc,['필요 기능','추천 패턴','특성','예시 케이스'],pattern_guide,
    [Cm(4),Cm(4.5),Cm(3),Cm(6)],hc='175E40')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════
# CHAPTER 4: RAG 심화 완전 가이드
# ════════════════════════════════════════════════════
h(doc,'Chapter 4. RAG 심화 완전 가이드',1)
par(doc,'RAG(Retrieval Augmented Generation)는 LLM의 한계를 극복하는 핵심 기술입니다. '
    '기본 RAG부터 고급 Advanced RAG, Agentic RAG까지 n8n 구현 방법을 완전 가이드합니다.', size=Pt(10))
doc.add_paragraph()

h(doc,'4.1 RAG 아키텍처 이해',2)
code(doc,
'RAG 기본 아키텍처\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'[인덱싱 단계 - 사전 처리]\n'
'문서 수집 → 청킹 → 임베딩 → 벡터 DB 저장\n'
'(Notion/Drive/Web) → (Splitter) → (Embeddings) → (Pinecone)\n'
'\n'
'[검색 및 생성 단계 - 실시간]\n'
'사용자 질문\n'
'    │\n'
'    ▼\n'
'질문 임베딩 → 벡터 DB 유사도 검색 → Top-K 문서 조각 검색\n'
'    │\n'
'    ▼\n'
'[프롬프트 조립]\n'
'"다음 컨텍스트를 바탕으로 답변하세요:\n'
'컨텍스트: {검색된 문서}\n'
'질문: {사용자 질문}"\n'
'    │\n'
'    ▼\n'
'LLM 답변 생성 (컨텍스트 기반)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'4.2 기본 RAG vs Advanced RAG vs Agentic RAG',2)
rag_types=[
    ('기본 RAG (Naive RAG)',
     '단순 유사도 검색 → 컨텍스트 추가 → 답변',
     'chainRetrievalQa 노드 사용',
     '저비용, 구현 단순',
     'FAQ 챗봇, 단순 Q&A',
     '복잡한 다단계 질문 약함'),
    ('Advanced RAG',
     'Query 재작성, 하이브리드 검색, 재순위화, Self-Query',
     'AI Agent + 여러 Vector Store Tool',
     '높은 정확도',
     '기업 지식베이스, 법률/의료 Q&A',
     '구현 복잡, 비용 증가'),
    ('Agentic RAG',
     '에이전트가 검색 전략을 자율 결정. 반복 검색. 여러 소스 통합.',
     'AI Agent + 다중 Tool + 메모리',
     '최고 정확도',
     '복잡한 리서치, 멀티 DB 통합 Q&A',
     '가장 느림, 가장 비쌈'),
]
t=doc.add_table(rows=len(rag_types)+1,cols=7);t.style='Table Grid';t.autofit=False
widths=[Cm(3),Cm(4),Cm(3),Cm(2.5),Cm(3),Cm(2)]
# 헤더
hdrs=['RAG 유형','동작 방식','n8n 구현','장점','적합 케이스','단점']
for ci,hdr in enumerate(hdrs[:6]):
    sc(t.rows[0].cells[ci],'1F497D')
    t.rows[0].cells[ci].text=hdr
    for p in t.rows[0].cells[ci].paragraphs:
        for r in p.runs:r.font.name='맑은 고딕';r.font.size=Pt(8.5);r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for ri,row in enumerate(rag_types):
    for ci,val in enumerate(row[:6]):
        sc(t.rows[ri+1].cells[ci],'EBF3FB' if ci==0 else 'FFFFFF')
        t.rows[ri+1].cells[ci].text=val
        for p in t.rows[ri+1].cells[ci].paragraphs:
            for r in p.runs:r.font.name='맑은 고딕';r.font.size=Pt(8.5)
            if ci==0:
                for r in p.runs:r.bold=True
doc.add_paragraph()

h(doc,'4.3 청킹(Chunking) 전략 완전 가이드',2)
par(doc,'청킹은 RAG 품질에 가장 큰 영향을 미치는 요소 중 하나입니다. '
    '문서 유형과 사용 케이스에 따라 최적의 청킹 전략을 선택해야 합니다.', size=Pt(10))
doc.add_paragraph()

chunking_strategies=[
    ('고정 크기 청킹',
     'chunkSize=1000, overlap=200\n(RecursiveCharacterSplitter)',
     '가장 간단, 빠름',
     '균일한 일반 텍스트',
     '문장 중간 끊김 가능'),
    ('의미론적 청킹',
     '문장 임베딩 유사도 기반 경계 결정\n(LangChain SemanticChunker)',
     '의미 단위 보존',
     '구조가 다양한 문서',
     'n8n Code 노드 커스텀 구현 필요'),
    ('문서 구조 기반',
     'H1/H2/H3 헤딩 기준 분할\n(MarkdownHeaderSplitter)',
     '섹션 컨텍스트 보존',
     'Markdown/HTML 문서',
     '헤딩 없는 문서에 부적합'),
    ('Parent-Child 청킹',
     '큰 Parent 청크 + 작은 Child 청크 이중 구조\n검색은 Child, 컨텍스트는 Parent 사용',
     '검색 정밀도 + 컨텍스트 균형',
     '길고 복잡한 기술 문서',
     '구현 복잡, 저장 공간 2배'),
    ('RAPTOR 계층 청킹',
     '청크를 재귀적으로 요약하여 트리 구조 생성\n다양한 추상화 레벨에서 검색',
     '전체적+세부적 동시 검색',
     '책, 긴 보고서, 매뉴얼',
     '인덱싱 비용 높음'),
]
tbl(doc,['청킹 전략','설정 방법','장점','적합 문서','단점'],
    chunking_strategies,[Cm(3.5),Cm(5),Cm(3),Cm(3),Cm(3)],hc='2E75B6')
doc.add_paragraph()

code(doc,
'청킹 전략 n8n 구현 예시\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'// Parent-Child 청킹 (Code 노드)\n'
'const text = $input.item.json.content;\n'
'\n'
'// Parent: 2000자 청크\n'
'const parentChunks = [];\nconst parentSize = 2000;\nconst parentOverlap = 200;\nfor (let i = 0; i < text.length; i += parentSize - parentOverlap) {\n  parentChunks.push(text.slice(i, i + parentSize));\n}\n'
'\n'
'// Child: 500자 청크 (Parent 내에서 분할)\n'
'const childChunks = [];\nparentChunks.forEach((parent, pIdx) => {\n  const childSize = 500, childOverlap = 50;\n  for (let i = 0; i < parent.length; i += childSize - childOverlap) {\n    childChunks.push({\n      content: parent.slice(i, i + childSize),\n      parentId: pIdx,\n      parentContent: parent\n    });\n  }\n});\n'
'\n'
'return childChunks.map(c => ({ json: c }));\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'4.4 임베딩 모델 선택 및 최적화',2)

emb_guide=[
    ('소규모 (< 10만 문서)','text-embedding-3-small','1536차원, $0.02/1M tokens','가성비 최고, 영어/한국어 우수'),
    ('대규모 (> 10만 문서)','text-embedding-3-large','3072차원, $0.13/1M tokens','정확도 우선시, 대용량'),
    ('다국어 특화','Cohere embed-multilingual-v3','1024차원, $0.10/1M','한국어+영어 혼합 문서'),
    ('완전 로컬','nomic-embed-text (Ollama)','768차원, 무료','데이터 프라이버시 필수'),
    ('Google 스택','text-embedding-004 (Gemini)','768차원, 무료 한도 있음','Google Cloud 사용 시'),
    ('Fine-tuned 임베딩','OpenAI Fine-tuning API','커스텀','도메인 특화 전문용어 많을 때'),
]
tbl(doc,['사용 규모','추천 모델','스펙','권장 이유'],emb_guide,[Cm(4),Cm(5),Cm(4),Cm(4.5)],hc='175E40')
doc.add_paragraph()

code(doc,
'임베딩 차원 수와 검색 성능의 관계\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'차원수 ↑  → 표현력 ↑, 저장 공간 ↑, 검색 속도 ↓\n'
'차원수 ↓  → 표현력 ↓, 저장 공간 ↓, 검색 속도 ↑\n'
'\n'
'text-embedding-3-small의 차원 축소 기능:\n'
'  기본: 1536차원\n'
'  dimensions=256으로 설정 시: 256차원 (속도 2배, 품질 소폭 하락)\n'
'\n'
'→ 대용량 인덱스: dimensions=256~512로 절충\n'
'→ 정확도 우선: dimensions=1536 (기본값)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'4.5 검색 전략 최적화',2)
par(doc,'기본 유사도 검색만으로는 복잡한 질문에 정확히 답하기 어렵습니다. '
    '여러 고급 검색 전략으로 정확도를 크게 향상시킬 수 있습니다.', size=Pt(10))
doc.add_paragraph()

search_strategies=[
    ('기본 유사도 검색','코사인 유사도 기반 Top-K 반환','간단, 빠름','유사 표현 다양할 때 미스'),
    ('하이브리드 검색','키워드(BM25) + 벡터 검색 결합','키워드+의미 모두 매칭','두 스코어 가중치 조정 필요'),
    ('MMR(최대 주변 관련성)','관련성 + 다양성 균형. 중복 청크 방지','다양한 관점 포함','n8n: retriever의 searchType=mmr'),
    ('Self-Query Retrieval','LLM이 메타데이터 필터 자동 생성\n예: "2024년 법률 문서" → filter:year=2024','정밀한 필터링','메타데이터 구조 사전 설계 필요'),
    ('Multi-Query Retrieval','동일 질문을 다양한 표현으로 재작성 후 통합 검색','검색 커버리지 향상','API 비용 증가 (3~5배 호출)'),
    ('HyDE (가상 문서 임베딩)','LLM으로 가상 답변 생성 → 그 답변으로 검색','정확도 대폭 향상','LLM 추가 호출 필요'),
    ('Re-Ranking','검색된 Top-N 청크를 Cross-Encoder로 재순위화','정밀도 향상','Cohere Rerank API 활용'),
]
tbl(doc,['검색 전략','동작 방식','장점','고려사항'],search_strategies,
    [Cm(3.5),Cm(5.5),Cm(3),Cm(5.5)],hc='2E75B6')
doc.add_paragraph()

code(doc,
'Multi-Query RAG n8n 구현 (Code 노드)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'// Step 1: 원본 질문을 3가지 다른 표현으로 재작성 (OpenAI)\n'
'// 시스템 프롬프트:\n'
'"다음 질문을 3가지 다른 방식으로 재작성하세요.\n'
'각 버전을 새 줄로 구분하세요.\n'
'질문: {input}"\n'
'\n'
'// Step 2: 재작성된 질문들로 병렬 검색 (각각 Vector Store)\n'
'const queries = openAIResult.split("\\n").filter(q => q.trim());\n'
'// → Split 노드로 각 쿼리별 검색\n'
'\n'
'// Step 3: 검색 결과 중복 제거 및 통합 (Code 노드)\n'
'const allResults = $input.all().map(i => i.json);\n'
'const seen = new Set();\n'
'const unique = allResults.filter(r => {\n'
'  const key = r.pageContent?.slice(0,100);\n'
'  if (seen.has(key)) return false;\n'
'  seen.add(key); return true;\n'
'});\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'4.6 벡터 DB 비교 및 선택 가이드',2)

vdb_comparison=[
    ('Pinecone','완전 관리형 SaaS. 설정 최소화.','pineconeApi','무료 1인덱스\n$0.096/100만 벡터/월','가장 쉬운 시작. 스타트업 권장.'),
    ('Qdrant','오픈소스 Self-hosted 또는 Cloud.','qdrantApi (URL)','Self-hosted 무료\nCloud $25/월~','프라이버시/비용 우선. Docker 운영.'),
    ('Supabase pgvector','기존 PostgreSQL에 벡터 확장.','supabaseApi','기존 Supabase 사용 시 무료','이미 Supabase 사용 중인 팀.'),
    ('Chroma','Python 생태계 오픈소스.','chromaApi (URL)','Self-hosted 무료','기존 Python ML 파이프라인.'),
    ('Weaviate','엔터프라이즈 기능 풍부.','weaviateApi','Self-hosted 무료\nCloud $25/월~','복잡한 메타데이터 필터링 필요 시.'),
    ('Milvus','초대규모 (10억+ 벡터) 특화.','milvusApi','Self-hosted 무료','대기업 초대용량 시스템.'),
    ('MongoDB Atlas','기존 MongoDB 확장. JSON 친화적.','mongoDbAtlasApi','M10+ 클러스터 필요','이미 MongoDB 사용 중인 팀.'),
]
tbl(doc,['벡터 DB','특성','n8n 자격증명','비용','권장 케이스'],vdb_comparison,
    [Cm(3),Cm(4),Cm(3.5),Cm(3.5),Cm(3.5)],hc='175E40')
doc.add_paragraph()

h(doc,'4.7 RAG 성능 평가 및 개선',2)
par(doc,'RAG 시스템의 품질을 객관적으로 측정하는 지표와 개선 방법입니다.', size=Pt(10))
doc.add_paragraph()

rag_metrics=[
    ('Context Precision','검색된 청크 중 실제 관련 청크의 비율. 높을수록 좋음.','목표: > 0.8','청킹 크기 줄이기, 메타데이터 필터 개선'),
    ('Context Recall','정답에 필요한 정보가 검색된 비율. 높을수록 좋음.','목표: > 0.8','Top-K 늘리기, 멀티 쿼리 사용'),
    ('Faithfulness','답변이 제공된 컨텍스트에 충실한 비율. 환각 측정.','목표: > 0.9','시스템 프롬프트 강화: "컨텍스트에만 기반"'),
    ('Answer Relevancy','답변이 질문과 관련성 있는 정도.','목표: > 0.8','프롬프트 개선, 더 정밀한 쿼리'),
    ('Latency','응답 시간.','목표: < 3초','더 작은 Top-K, 빠른 임베딩 모델'),
]
tbl(doc,['평가 지표','설명','목표값','개선 방법'],rag_metrics,[Cm(4),Cm(5),Cm(2.5),Cm(6)],hc='2E75B6')
doc.add_paragraph()

h(doc,'4.8 RAG 파이프라인 완성 예시 (n8n)',2)
code(doc,
'완전한 RAG 파이프라인 n8n 워크플로우 구조\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'[인덱싱 워크플로우 - 1회 실행]\n'
'Notion 페이지 로드       → Notion Document Loader\n'
'    │\n'
'    ▼\n'
'문서 청킹                → Recursive Text Splitter\n'
'    │                      (chunkSize=1000, overlap=200)\n'
'    ▼\n'
'임베딩 생성              → OpenAI Embeddings\n'
'    │                      (text-embedding-3-small)\n'
'    ▼\n'
'벡터 DB 저장             → Pinecone Vector Store (Insert)\n'
'\n'
'[검색 워크플로우 - 실시간]\n'
'Webhook (질문 수신)\n'
'    │\n'
'    ▼\n'
'Redis 캐시 확인          → [캐시 히트] → 즉시 반환\n'
'    │ [캐시 미스]\n'
'    ▼\n'
'AI Agent\n'
'  ├─ Tool: VectorStore Retriever (Pinecone, topK=5)\n'
'  │       └─ Embeddings: OpenAI\n'
'  ├─ Tool: Web Search (최신 정보 보완)\n'
'  └─ LLM: gpt-4o-mini (temperature=0)\n'
'    │\n'
'    ▼\n'
'Redis 캐시 저장 (TTL: 1시간)\n'
'    │\n'
'    ▼\n'
'Respond to Webhook\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

doc.add_page_break()
doc.save(OUT)
print('Vol5 Part2 done:', OUT)
