# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "C:/Users/장우경/oomni/n8n_cases_vol6.docx"
doc = Document(PATH)

def sc(doc, text, level=1):
    styles = {1: (16, True, RGBColor(0x1F,0x49,0x7D)),
              2: (13, True, RGBColor(0x2E,0x74,0xB5)),
              3: (11, True, RGBColor(0x1F,0x49,0x7D))}
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    sz, bold, color = styles[level]
    run.font.size = Pt(sz); run.font.bold = bold; run.font.color.rgb = color
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p

def h(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

def par(doc, text, indent=0):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.5)

def bul(doc, items, indent=1):
    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(("  " * indent) + "• " + item)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

def code(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._element.pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def tbl(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'BDD7EE')
        cell._tc.get_or_add_tcPr().append(shd)
    for rdata in rows:
        row = table.add_row()
        for i, val in enumerate(rdata):
            cell = row.cells[i]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ─── CH3: 플랫폼 비교 ─────────────────────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 3: 플랫폼 비교 — n8n vs Zapier vs Make", 1)
par(doc, "3대 자동화 플랫폼의 기능, 가격, TCO를 종합 비교하여 최적의 플랫폼 선택을 지원합니다.")

sc(doc, "3.1 기능 비교 매트릭스", 2)
tbl(doc,
    ["기능 항목", "n8n", "Zapier", "Make (Integromat)"],
    [
        ["비즈니스 로직", "복잡한 로직, 반복문, 조건문", "제한적 (기본 필터)", "중간 수준"],
        ["코드 실행", "JavaScript/Python 네이티브", "불가 (Code by Zapier 제한적)", "불가"],
        ["AI/LLM 연동", "25종+ LangChain 노드, AI Agent", "OpenAI 기본 연동", "OpenAI 기본 연동"],
        ["셀프호스팅", "완전 지원 (무료 오픈소스)", "불가", "불가"],
        ["워크플로우 복잡도", "무제한 복잡도", "단순~중간", "중간~복잡"],
        ["실시간 처리", "지원 (Webhook 즉시)", "지원", "지원"],
        ["API 커스텀 연동", "완전 자유 (HTTP 노드)", "제한적", "HTTP 모듈 지원"],
        ["데이터 변환", "JavaScript로 무제한", "기본 데이터 포매터", "중간 수준"],
        ["오류 처리", "Try/Catch, 세밀한 제어", "기본 오류 알림", "오류 핸들러 지원"],
        ["버전 관리", "Git 연동 가능", "없음", "없음"],
        ["디버깅", "실시간 실행 로그", "기본 히스토리", "실행 히스토리"],
        ["멀티테넌시", "지원 (Enterprise)", "지원 (Teams)", "지원 (Teams)"],
        ["화이트레이블링", "가능 (셀프호스팅)", "불가", "불가"],
        ["사전 구축 템플릿", "700+ 템플릿", "6,000+ Zaps", "500+ 템플릿"],
        ["통합 앱 수", "400+ (커스텀 무제한)", "6,000+", "1,500+"],
    ],
    [4.5, 3.5, 3, 3.5]
)

sc(doc, "3.2 가격 모델 비교", 2)
h(doc, "3.2.1 n8n 가격 체계")
tbl(doc,
    ["플랜", "가격", "실행 횟수", "주요 특징"],
    [
        ["Community (셀프호스팅)", "무료", "무제한", "오픈소스, 직접 서버 운영 필요"],
        ["Starter (Cloud)", "$20/월", "2,500 실행/월", "클라우드 호스팅, 2명 사용자"],
        ["Pro (Cloud)", "$50/월", "10,000 실행/월", "5명 사용자, 우선 지원"],
        ["Enterprise", "문의", "무제한", "SSO, 감사로그, 전담 지원, SLA"],
    ],
    [3, 2.5, 3.5, 5.5]
)

h(doc, "3.2.2 Zapier 가격 체계")
tbl(doc,
    ["플랜", "가격", "태스크/월", "주요 특징"],
    [
        ["Free", "$0", "100", "5 Zaps, 15분 업데이트 주기"],
        ["Starter", "$19.99/월", "750", "20 Zaps, 15분 주기"],
        ["Professional", "$49/월", "2,000", "무제한 Zaps, 2분 주기"],
        ["Team", "$69/월", "2,000", "25명, 공유 앱"],
        ["Company", "$103/월~", "50,000~", "SAML SSO, 고급 관리자"],
    ],
    [3, 2.5, 3.5, 5.5]
)

h(doc, "3.2.3 Make (Integromat) 가격 체계")
tbl(doc,
    ["플랜", "가격", "작업/월", "주요 특징"],
    [
        ["Free", "$0", "1,000", "2 활성 시나리오"],
        ["Core", "$9/월", "10,000", "무제한 시나리오, 1분 주기"],
        ["Pro", "$16/월", "10,000", "고급 도구, 전체 실행 이력"],
        ["Teams", "$29/월", "10,000", "팀 협업 기능"],
        ["Enterprise", "문의", "맞춤", "고급 보안, 전담 지원"],
    ],
    [3, 2.5, 3.5, 5.5]
)

sc(doc, "3.3 총소유비용(TCO) 분석", 2)
par(doc, "단순 라이선스 비용이 아닌 구축, 운영, 유지보수를 포함한 5년 TCO를 비교합니다.")
h(doc, "시나리오: 중규모 기업 (월 50,000 실행, 10명 사용자)")
tbl(doc,
    ["비용 항목", "n8n 셀프호스팅", "n8n Cloud", "Zapier", "Make"],
    [
        ["라이선스 (연)", "0원", "$600 (Pro)", "$4,788 (Company)", "$348 (Pro)"],
        ["인프라 (연)", "500만원 (AWS t3.medium)", "포함", "-", "-"],
        ["초기 구축 (일회성)", "500만원", "300만원", "200만원", "200만원"],
        ["유지보수 (연)", "200만원", "100만원", "50만원", "50만원"],
        ["인력 교육 (일회성)", "200만원", "100만원", "50만원", "70만원"],
        ["5년 TCO", "1,900만원", "860만원 + $3,000", "$28,940 + 한화 비용", "$2,740 + 한화 비용"],
    ],
    [3.5, 3, 3, 3, 3]
)

par(doc, "* 실행 횟수가 많을수록 n8n의 TCO 우위가 커집니다. 월 50,000+ 실행 시 n8n 셀프호스팅이 5년 기준 Zapier 대비 80% 이상 절감됩니다.")

h(doc, "3.3.1 실행 횟수별 최적 플랫폼 추천")
tbl(doc,
    ["월 실행 횟수", "권장 플랫폼", "예상 월 비용", "선택 이유"],
    [
        ["< 100건", "Zapier Free 또는 Make Free", "$0", "소규모, 단순 자동화"],
        ["100~2,000건", "n8n Cloud Starter 또는 Make Core", "$9~20", "비용 효율, 쉬운 시작"],
        ["2,000~10,000건", "n8n Cloud Pro 또는 Make Pro", "$16~50", "기능 충분, 적정 비용"],
        ["10,000~50,000건", "n8n 셀프호스팅 or Cloud Pro", "30~150만원", "n8n 기능 우위"],
        ["50,000+ 건", "n8n 셀프호스팅 (필수)", "50~200만원", "Zapier 대비 80%+ 절감"],
        ["엔터프라이즈 복잡 로직", "n8n Enterprise", "협의", "커스텀 로직, AI 에이전트"],
    ],
    [3.5, 3.5, 3, 5]
)

sc(doc, "3.4 마이그레이션 비용 고려사항", 2)
par(doc, "기존 Zapier/Make에서 n8n으로 전환 시 고려해야 할 비용과 리스크를 분석합니다.")
tbl(doc,
    ["마이그레이션 항목", "비용 추정", "기간", "리스크"],
    [
        ["워크플로우 재구축", "Zap당 30분~3시간", "1~3개월", "로직 손실 가능성"],
        ["크레덴셜 재설정", "앱당 30분", "1~2일", "인증 오류"],
        ["테스트/검증", "원본의 20%~50% 시간", "2~4주", "엣지 케이스 누락"],
        ["다운타임 관리", "롤링 마이그레이션", "0~2일", "서비스 중단"],
        ["팀 교육", "인당 4~16시간", "1~2주", "생산성 저하"],
    ],
    [4, 4, 3, 5]
)

bul(doc, [
    "마이그레이션 절감 효과가 1년 내 마이그레이션 비용을 커버하는 경우에만 전환 권장",
    "Zapier 100 Zaps → n8n: 약 2~4주 작업, 비용 40~80만원",
    "점진적 전환 (새 자동화는 n8n, 기존은 유지) 방식 추천",
    "n8n-node-types GitHub에서 커뮤니티 마이그레이션 도구 활용",
])

# ─── CH4: 비용 최적화 전략 ───────────────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 4: 비용 최적화 전략", 1)
par(doc, "n8n 운영 비용을 최소화하면서 성능을 극대화하는 전략적 접근 방법을 제시합니다.")

sc(doc, "4.1 n8n 비용 구조 최적화", 2)
h(doc, "4.1.1 Cloud vs 셀프호스팅 전환 시점")
tbl(doc,
    ["기준", "Cloud 유지", "셀프호스팅 전환"],
    [
        ["기술 역량", "DevOps 부재", "서버 관리 가능"],
        ["월간 실행량", "< 10,000건", "> 10,000건"],
        ["예산", "월 $50 이하 Cloud가 편리", "월 $100+ 될 때 전환 고려"],
        ["보안/컴플라이언스", "일반 수준", "데이터 주권, 규제 요구"],
        ["커스터마이징 요구", "표준 기능으로 충분", "화이트레이블, 고급 설정 필요"],
    ],
    [5, 4.5, 4.5]
)

h(doc, "4.1.2 셀프호스팅 최적 인프라 구성")
tbl(doc,
    ["사용 규모", "권장 서버 사양", "월 예상 비용", "추가 구성"],
    [
        ["소규모 (< 1만 실행/월)", "t3.small (2vCPU, 2GB)", "AWS $15~20/월", "Redis t3.micro + RDS db.t3.micro"],
        ["중규모 (1~10만/월)", "t3.medium (2vCPU, 4GB)", "AWS $35~50/월", "Redis + RDS t3.small"],
        ["대규모 (10~100만/월)", "t3.large 또는 t3.xlarge", "AWS $80~200/월", "ElastiCache + RDS m5.large"],
        ["엔터프라이즈 (100만+/월)", "c5.2xlarge 또는 EKS", "AWS $300~2,000/월", "클러스터 구성, 로드밸런서"],
    ],
    [4, 4, 3.5, 5]
)

code(doc, """// 비용 최적화 docker-compose 설정
services:
  n8n:
    image: n8nio/n8n:latest
    environment:
      - EXECUTIONS_MODE=queue          # 큐 모드로 효율적 처리
      - N8N_CONCURRENCY_PRODUCTION_LIMIT=5  # 동시 실행 수 제한
      - EXECUTIONS_DATA_PRUNE=true     # 오래된 실행 데이터 자동 삭제
      - EXECUTIONS_DATA_MAX_AGE=168    # 7일 보관 (기본 336시간)
      - N8N_LOG_LEVEL=warn             # 로그 최소화
    resources:
      limits:
        memory: 2G                      # 메모리 제한으로 비용 절감
        cpus: '1.0'""")

sc(doc, "4.2 실행 횟수 최적화", 2)
par(doc, "n8n Cloud 사용 시 실행 횟수가 과금 기준이므로 불필요한 실행을 줄이는 것이 핵심입니다.")

h(doc, "4.2.1 불필요한 실행 제거 방법")
bul(doc, [
    "폴링 간격 최적화: 1분 폴링을 15분으로 변경 → 실행 횟수 93% 감소",
    "Webhook 전환: 폴링 대신 Webhook으로 이벤트 발생 시만 실행",
    "배치 처리: 개별 처리 대신 묶음 처리 (100건 개별 = 100회 → 1회 배치)",
    "조건부 실행: IF 노드로 조건 미충족 시 이후 노드 실행 스킵",
    "실행 트리거 최적화: 필요한 이벤트에만 반응하도록 필터 설정",
])

h(doc, "4.2.2 실행 횟수 절감 사례")
tbl(doc,
    ["최적화 전", "최적화 후", "절감률", "월 절감 비용 (Pro 기준)"],
    [
        ["Gmail 1분 폴링: 43,200회/월", "Gmail Webhook: 실제 수신 시만 (300회/월)", "99.3%", "플랜 다운그레이드 가능"],
        ["개별 Slack 알림 500건", "배치 Slack 알림 5건", "99%", "$30/월 절감"],
        ["Shopify 신규주문 확인 5분폴링", "Shopify Webhook 연동", "95%", "플랜 유지 + 여유 실행"],
        ["DB 변경감지 1분폴링", "PostgreSQL Trigger + Listen", "98%", "서버 부하 감소"],
    ],
    [5, 5, 2, 4]
)

h(doc, "4.2.3 실행 횟수 모니터링 설정")
code(doc, """// n8n 실행 횟수 모니터링 워크플로우
// 1. 매일 자정: n8n API로 실행 통계 조회
GET /api/v1/executions?limit=100&status=success

// 2. 워크플로우별 실행 횟수 집계
// 3. 예산 초과 임박 시 알림
if (monthlyExecutions > threshold * 0.8) {
  // Slack 알림: "이번 달 실행 횟수 80% 도달"
  // 과도한 폴링 워크플로우 자동 감지 및 보고
}

// 실행 데이터 정리로 스토리지 절감
// 성공 실행: 7일 보관 / 실패 실행: 30일 보관""")

sc(doc, "4.3 인프라 비용 최적화", 2)
h(doc, "4.3.1 AWS 비용 최적화")
bul(doc, [
    "Reserved Instance: 1년 예약 구매로 EC2 비용 40% 절감",
    "Spot Instance: 비실시간 배치 작업에 Spot 사용 (70% 절감)",
    "Auto Scaling: 야간/주말 인스턴스 수 자동 축소",
    "S3 Intelligent-Tiering: 첨부파일/로그 자동 저비용 스토리지 이동",
    "CloudFront: 정적 자산 캐싱으로 데이터 전송 비용 절감",
    "RDS Proxy: 연결 풀링으로 소형 DB 인스턴스 사용 가능",
])

h(doc, "4.3.2 API 비용 최적화")
tbl(doc,
    ["API 유형", "최적화 방법", "절감률", "구현 방법"],
    [
        ["OpenAI/LLM API", "캐싱 + 모델 라우팅", "50~80%", "Redis 시맨틱 캐시"],
        ["외부 SaaS API", "결과 캐싱 (TTL 적용)", "30~60%", "n8n Redis 캐시 노드"],
        ["이메일 발송 API", "배치 발송", "70~90%", "큐 + 배치 처리"],
        ["SMS API", "카카오 알림톡 전환", "40~60%", "카카오 비즈메시지"],
        ["지도/주소 API", "결과 DB 저장 재사용", "60~80%", "PostgreSQL 캐시 테이블"],
    ],
    [3.5, 4, 3, 4]
)

h(doc, "4.3.3 비용 절감 로드맵")
tbl(doc,
    ["단계", "기간", "주요 액션", "예상 절감률"],
    [
        ["즉시 (0~1개월)", "당장", "폴링→Webhook 전환, 배치 처리 도입", "30~50%"],
        ["단기 (1~3개월)", "1개월 후", "LLM 캐싱 구현, 모델 라우팅 설정", "추가 20~40%"],
        ["중기 (3~6개월)", "3개월 후", "셀프호스팅 전환 검토, Reserved Instance", "추가 20~30%"],
        ["장기 (6개월+)", "6개월 후", "자체 Fine-tuning, 온프레미스 LLM", "추가 30~50%"],
    ],
    [3, 3, 6, 3]
)

doc.save(PATH)
print("[OK] gen_vol6_p2.py 완료 - Ch3+Ch4 추가")
