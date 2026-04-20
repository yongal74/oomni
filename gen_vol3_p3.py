# -*- coding: utf-8 -*-
"""Vol 3 Part 3: Chapter 4 고객지원/CS + Chapter 5 재무/회계 + Chapter 6 IT운영/DevOps"""
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'C:/Users/장우경/oomni/n8n_cases_vol3.docx'
doc = Document(OUT)

def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    colors = {1: RGBColor(0x1F,0x49,0x7D), 2: RGBColor(0x2E,0x75,0xB6), 3: RGBColor(0x17,0x5E,0x40)}
    sizes  = {1: Pt(22), 2: Pt(16), 3: Pt(13)}
    if p.runs:
        p.runs[0].font.name = '맑은 고딕'
        p.runs[0].font.size = sizes.get(level, Pt(11))
        p.runs[0].font.color.rgb = colors.get(level, RGBColor(0,0,0))
    return p

def add_para(doc, text, bold=False, color=None, size=Pt(10), indent=None):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    run.font.name = '맑은 고딕'; run.font.size = size; run.bold = bold
    if color: run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    r = p.add_run(text); r.font.name = '맑은 고딕'; r.font.size = Pt(10)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'; run.font.size = Pt(7.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

DIFF_COLORS = {
    '초급(Easy)':    RGBColor(0x70,0xAD,0x47),
    '중급(Medium)':  RGBColor(0xED,0x7D,0x31),
    '고급(Hard)':    RGBColor(0xC0,0x00,0x00),
    '전문가(Expert)':RGBColor(0x70,0x30,0xA0),
}

def make_recipe(doc, num, title, category, difficulty, apps, trigger, description,
                workflow_json, node_guide, setup_checklist, tips):
    add_heading(doc, f'Recipe {num:03d}. {title}', 3)
    diff_color = DIFF_COLORS.get(difficulty, RGBColor(0x40,0x40,0x40))
    t = doc.add_table(rows=2, cols=4); t.style = 'Table Grid'; t.autofit = False
    for w, col in zip([Cm(2.2), Cm(5.3), Cm(2.2), Cm(7.8)], range(4)):
        t.columns[col].width = w
    data = [['카테고리', category, '주요 앱', apps],
            ['난이도',   difficulty, '트리거',  trigger]]
    for ri, row_data in enumerate(data):
        row = t.rows[ri]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
            for p in row.cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = '맑은 고딕'; r.font.size = Pt(9)
                    if ci % 2 == 0: r.bold = True
        for ci in [0,2]: shade_cell(t.rows[ri].cells[ci], 'EBF3FB')
        for ci in [1,3]: shade_cell(t.rows[ri].cells[ci], 'F8FBFE')
    if t.rows[1].cells[1].paragraphs[0].runs:
        t.rows[1].cells[1].paragraphs[0].runs[0].font.color.rgb = diff_color
        t.rows[1].cells[1].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    add_para(doc, '📋 워크플로우 개요', bold=True, color=RGBColor(0x1F,0x49,0x7D))
    add_para(doc, description, indent=Cm(0.5))
    add_para(doc, '📦 n8n JSON (복사 → Import 가능)', bold=True, color=RGBColor(0x1F,0x49,0x7D))
    add_code_block(doc, workflow_json)
    add_para(doc, '🔧 노드별 설정 가이드', bold=True, color=RGBColor(0x1F,0x49,0x7D))
    for item in node_guide: add_bullet(doc, item)
    add_para(doc, '✅ 설정 체크리스트', bold=True, color=RGBColor(0x17,0x5E,0x40))
    for item in setup_checklist: add_bullet(doc, f'☐ {item}')
    if tips:
        add_para(doc, '💡 실전 팁', bold=True, color=RGBColor(0xED,0x7D,0x31))
        for tip in tips: add_bullet(doc, tip)
    doc.add_paragraph()
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pb.append(bottom); pPr.append(pb)
    doc.add_paragraph()

def simple_wf(num, title, tags):
    return json.dumps({
        "name": title, "nodes": [
            {"parameters": {"rule": {"interval": [{"field": "cronExpression", "expression": "0 9 * * *"}]}},
             "id": f"s-{num:04d}-01", "name": "트리거",
             "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]},
            {"parameters": {"jsCode": "// 비즈니스 로직 구현\nreturn $input.all();"},
             "id": f"s-{num:04d}-02", "name": "Code - 처리",
             "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [460, 300]},
            {"parameters": {"resource": "message", "operation": "post",
                            "channel": "#automation", "text": f"완료: {title}", "otherOptions": {}},
             "id": f"s-{num:04d}-03", "name": "Slack - 알림",
             "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [680, 300],
             "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}}
        ],
        "connections": {
            "트리거": {"main": [[{"node": "Code - 처리", "type": "main", "index": 0}]]},
            "Code - 처리": {"main": [[{"node": "Slack - 알림", "type": "main", "index": 0}]]}
        }, "active": False, "settings": {}, "tags": tags
    }, ensure_ascii=False, indent=2)

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 4: 고객지원 / CS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Chapter 4. 고객지원 / CS', 1)
add_para(doc, 'Zendesk, Intercom, Freshdesk 기반 티켓 자동화, AI 챗봇, CS 효율화 레시피 10선입니다.', size=Pt(10))
doc.add_paragraph()

WF_031 = json.dumps({
  "name": "Zendesk 티켓 AI 자동 분류 및 우선순위 지정",
  "nodes": [
    {
      "parameters": {
        "url": "https://YOUR_SUBDOMAIN.zendesk.com/api/v2/tickets",
        "authentication": "predefinedCredentialType",
        "nodeCredentialType": "zendeskApi",
        "options": {}
      },
      "id": "cs-0031-01", "name": "HTTP - 신규 티켓 조회",
      "type": "n8n-nodes-base.httpRequest", "typeVersion": 4, "position": [240, 300],
      "credentials": {"zendeskApi": {"id": "cred-zendesk", "name": "Zendesk"}}
    },
    {
      "parameters": {
        "jsCode": "const tickets = $input.item.json.tickets || [];\nconst newTickets = tickets.filter(t => t.status === 'new');\nreturn newTickets.map(t => ({ json: t }));"
      },
      "id": "cs-0031-02", "name": "Code - 신규 상태 필터",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [460, 300]
    },
    {
      "parameters": {
        "resource": "text", "operation": "message",
        "model": "gpt-4o-mini",
        "messages": {"values": [{
          "role": "system",
          "content": "고객지원 티켓을 분류하세요. 응답은 JSON 형식으로만: {\"category\": \"billing|technical|general|urgent\", \"priority\": \"low|normal|high|urgent\", \"sentiment\": \"positive|neutral|negative\", \"summary\": \"한줄요약\"}"
        }, {
          "role": "user",
          "content": "제목: {{ $json.subject }}\n내용: {{ $json.description?.slice(0,500) }}"
        }]
        }
      },
      "id": "cs-0031-03", "name": "OpenAI - 티켓 분류",
      "type": "@n8n/n8n-nodes-langchain.openAi", "typeVersion": 1, "position": [680, 300],
      "credentials": {"openAiApi": {"id": "cred-openai", "name": "OpenAI"}}
    },
    {
      "parameters": {
        "jsCode": "const aiResult = JSON.parse($json.message?.content || '{}');\nconst ticket = $('Code - 신규 상태 필터').item.json;\nreturn [{ json: { ...ticket, ai_category: aiResult.category, ai_priority: aiResult.priority, ai_sentiment: aiResult.sentiment, ai_summary: aiResult.summary } }];"
      },
      "id": "cs-0031-04", "name": "Code - AI 결과 파싱",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [900, 300]
    },
    {
      "parameters": {
        "url": "https://YOUR_SUBDOMAIN.zendesk.com/api/v2/tickets/{{ $json.id }}.json",
        "method": "PUT",
        "authentication": "predefinedCredentialType",
        "nodeCredentialType": "zendeskApi",
        "body": {
          "body": {
            "ticket": {
              "priority": "={{ $json.ai_priority }}",
              "tags": ["={{ $json.ai_category }}", "ai-classified"],
              "custom_fields": [{"id": 12345678, "value": "={{ $json.ai_summary }}"}]
            }
          }
        }
      },
      "id": "cs-0031-05", "name": "HTTP - Zendesk 티켓 업데이트",
      "type": "n8n-nodes-base.httpRequest", "typeVersion": 4, "position": [1120, 300],
      "credentials": {"zendeskApi": {"id": "cred-zendesk", "name": "Zendesk"}}
    }
  ],
  "connections": {
    "HTTP - 신규 티켓 조회": {"main": [[{"node": "Code - 신규 상태 필터", "type": "main", "index": 0}]]},
    "Code - 신규 상태 필터": {"main": [[{"node": "OpenAI - 티켓 분류", "type": "main", "index": 0}]]},
    "OpenAI - 티켓 분류": {"main": [[{"node": "Code - AI 결과 파싱", "type": "main", "index": 0}]]},
    "Code - AI 결과 파싱": {"main": [[{"node": "HTTP - Zendesk 티켓 업데이트", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["cs", "zendesk", "openai", "ticket-classification"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 31,
    'Zendesk 티켓 AI 자동 분류 및 우선순위 지정',
    '고객지원 / CS', '고급(Hard)',
    'Zendesk, OpenAI GPT-4o-mini',
    'Schedule (5분마다) 또는 Zendesk Webhook',
    'Zendesk 신규 티켓을 GPT-4o-mini로 자동 분류(billing/technical/general/urgent)하고 우선순위와 감성분석 결과를 Zendesk에 자동 업데이트합니다. 상담원 라우팅 시간을 60% 단축합니다.',
    WF_031,
    [
        'Zendesk API: YOUR_SUBDOMAIN을 실제 서브도메인으로 교체. API Token 방식 또는 OAuth2',
        'OpenAI JSON 응답: response_format: json_object 설정하면 더 안정적 (GPT-4o 이상)',
        'Code AI 파싱: JSON.parse로 AI 응답 파싱. 파싱 실패 시 try-catch 추가 권장',
        'Zendesk 업데이트: custom_fields ID 12345678을 실제 커스텀 필드 ID로 교체',
        '우선순위 매핑: OpenAI "urgent" → Zendesk "urgent" 1:1 매핑. 불일치 시 Code 노드로 변환',
    ],
    [
        'Zendesk API Token 생성 (Admin → Apps & Integrations → APIs)',
        'YOUR_SUBDOMAIN을 실제 서브도메인으로 교체',
        'Zendesk 커스텀 필드 생성 (AI 요약 저장용)',
        '커스텀 필드 ID 확인 후 12345678 교체',
        'OpenAI API Key 설정',
        '분류 카테고리를 실제 업무에 맞게 시스템 프롬프트 수정',
    ],
    [
        'Zendesk 트리거 대안: 5분 폴링 대신 Zendesk → Settings → Extensions → Webhooks로 실시간 처리',
        '자동 응답: 카테고리별 FAQ 답변 자동 발송 추가 가능 (billing 질문 → 결제 FAQ 자동 첨부)',
        '에스컬레이션: ai_priority === "urgent"이면 매니저에게 Slack DM 자동 발송',
    ]
)

WF_032 = json.dumps({
  "name": "고객 이탈 징후 감지 및 CS팀 알림",
  "nodes": [
    {
      "parameters": {"rule": {"interval": [{"field": "cronExpression", "expression": "0 9 * * 1-5"}]}},
      "id": "cs-0032-01", "name": "Schedule - 평일 오전",
      "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]
    },
    {
      "parameters": {
        "url": "https://api.intercom.io/contacts",
        "authentication": "predefinedCredentialType",
        "nodeCredentialType": "httpHeaderAuth",
        "qs": {
          "filter": "{\"field\":\"last_seen_at\",\"operator\":\"<\",\"value\": \"{{ $now.minus({days: 14}).toUnixInteger() }}\"}"
        },
        "options": {"headers": {"headers": [{"name": "Accept", "value": "application/json"}]}}
      },
      "id": "cs-0032-02", "name": "HTTP - Intercom 비활성 고객",
      "type": "n8n-nodes-base.httpRequest", "typeVersion": 4, "position": [460, 300],
      "credentials": {"httpHeaderAuth": {"id": "cred-intercom", "name": "Intercom"}}
    },
    {
      "parameters": {
        "jsCode": "const contacts = ($input.item.json.data || []).slice(0, 20);\nreturn contacts.map(c => ({\n  json: {\n    name: c.name,\n    email: c.email,\n    company: c.companies?.data?.[0]?.company_id,\n    lastSeen: new Date(c.last_seen_at * 1000).toLocaleDateString('ko-KR'),\n    daysSince: Math.floor((Date.now() - c.last_seen_at*1000) / (1000*60*60*24))\n  }\n}));"
      },
      "id": "cs-0032-03", "name": "Code - 이탈 징후 분석",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [680, 300]
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#cs-churn-risk",
        "text": "⚠️ 이탈 위험 고객 (14일 미접속)\n고객: {{ $json.name }} ({{ $json.email }})\n회사: {{ $json.company }}\n마지막 접속: {{ $json.lastSeen }} ({{ $json.daysSince }}일 전)",
        "otherOptions": {}
      },
      "id": "cs-0032-04", "name": "Slack - 이탈 위험 알림",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [900, 300],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Schedule - 평일 오전": {"main": [[{"node": "HTTP - Intercom 비활성 고객", "type": "main", "index": 0}]]},
    "HTTP - Intercom 비활성 고객": {"main": [[{"node": "Code - 이탈 징후 분석", "type": "main", "index": 0}]]},
    "Code - 이탈 징후 분석": {"main": [[{"node": "Slack - 이탈 위험 알림", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["cs", "churn", "intercom", "customer-success"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 32,
    '고객 이탈 징후 자동 감지 (14일 미접속)',
    '고객지원 / CS', '중급(Medium)',
    'Intercom, Slack',
    'Schedule (평일 오전 9시)',
    'Intercom에서 14일 이상 로그인하지 않은 고객을 자동 감지하여 CS팀 Slack 채널에 이탈 위험 경보를 발송합니다. 조기 개입으로 고객 이탈율을 낮춥니다.',
    WF_032,
    [
        'Intercom API: Bearer Token 방식. httpHeaderAuth에 Authorization: Bearer {TOKEN} 설정',
        'Filter 조건: last_seen_at < 14일 전 Unix timestamp. 기간 조정 가능',
        'Code 노드: 최대 20개로 제한 (.slice(0,20)). 많은 고객 있으면 페이지네이션 추가',
        'Slack: 고객별 개별 메시지 대신 집계 리포트로 변경 시 Code 노드에서 텍스트 조합',
    ],
    [
        'Intercom API Access Token 생성 (Developer Hub → Apps → Your app → Auth)',
        'Slack Bot Token 및 #cs-churn-risk 채널 생성',
        '이탈 기준일수(14일) 비즈니스에 맞게 조정',
        'Intercom 고객 필터 조건 검증 (API Docs 참조)',
    ],
    [
        'Tier별 처리: 엔터프라이즈 고객은 5일, SMB는 14일 기준으로 분리 처리',
        '자동 리인게이지: 이탈 위험 고객에게 개인화 이메일 자동 발송 연계 가능',
        'Mixpanel/Amplitude 대안: 프로덕트 분석 툴 API로 더 정교한 이탈 예측 지표 활용',
    ]
)

# CS 나머지 8개 (033~040)
cs_recipes = [
    (33, 'FAQ AI 자동 응답 봇 (Zendesk + OpenAI)', '고객지원 / CS', '고급(Hard)',
     'Zendesk, OpenAI, Notion(FAQ DB)', 'Zendesk Webhook (신규 티켓)',
     'Zendesk 신규 티켓이 들어오면 Notion FAQ DB에서 유사 답변을 검색하고 OpenAI로 맞춤 답변을 생성하여 자동으로 1차 응답합니다.'),
    (34, 'CSAT 설문 자동 발송 및 집계', '고객지원 / CS', '초급(Easy)',
     'Zendesk, Gmail, Google Sheets', 'Zendesk Webhook (티켓 종료)',
     '티켓이 해결됨 상태로 종료되면 24시간 후 자동으로 CSAT 설문 이메일을 발송하고 응답을 Google Sheets에 자동 집계합니다.'),
    (35, '고객 VIP 등급별 SLA 자동 관리', '고객지원 / CS', '고급(Hard)',
     'Zendesk, HubSpot, Slack, PagerDuty', 'Schedule (10분마다)',
     'VIP 고객(엔터프라이즈)의 SLA 위반 임박 시 자동으로 에스컬레이션하고 PagerDuty 알림으로 당직 엔지니어에게 즉시 통보합니다.'),
    (36, '멀티채널 티켓 통합 (이메일+채팅+SNS)', '고객지원 / CS', '전문가(Expert)',
     'Gmail, Intercom, Twitter/X, Zendesk', '여러 채널 Webhook',
     '이메일, 라이브채팅, Twitter DM 등 여러 채널로 들어오는 고객 문의를 Zendesk 티켓으로 통합하여 CS팀이 단일 화면에서 처리할 수 있게 합니다.'),
    (37, '상품 환불/교환 자동 처리 파이프라인', '고객지원 / CS', '고급(Hard)',
     'Shopify, Zendesk, Gmail, Slack', 'Zendesk Webhook (환불 태그)',
     'Zendesk 티켓에 "환불요청" 태그 지정 시 자동으로 Shopify 환불 프로세스를 시작하고 고객에게 환불 확인 이메일을 발송합니다.'),
    (38, '고객 온보딩 헬스스코어 자동 추적', '고객지원 / CS', '중급(Medium)',
     'Intercom, Google Sheets, Slack', 'Schedule (주간)',
     '신규 고객의 온보딩 진행률(기능 사용, 설정 완료 등)을 매주 자동 추적하여 헬스스코어를 계산하고 위험 고객을 CS팀에 알립니다.'),
    (39, 'Trustpilot/G2 리뷰 자동 모니터링', '고객지원 / CS', '중급(Medium)',
     'HTTP Request, OpenAI, Slack, Gmail', 'Schedule (매일)',
     'Trustpilot, G2, 네이버 스토어 리뷰를 매일 자동 수집하고 부정 리뷰(3점 이하)는 즉시 CS팀에 알림을 보내고 관리자에게 답변 초안을 이메일로 발송합니다.'),
    (40, 'CS 주간 성과 KPI 자동 리포트', '고객지원 / CS', '중급(Medium)',
     'Zendesk, Google Sheets, Slack, Gmail', 'Schedule (매주 금요일)',
     '매주 금요일 Zendesk에서 주요 CS KPI(CSAT, FRT, 해결율, 처리건수)를 자동 집계하여 매니저 이메일과 #cs-weekly 채널에 리포트를 발송합니다.'),
]

for num, title, cat, diff, apps, trigger, desc in cs_recipes:
    wf = simple_wf(num, title, ["cs"])
    make_recipe(doc, num, title, cat, diff, apps, trigger, desc, wf,
        ['트리거 노드를 실제 서비스(Zendesk Webhook 또는 Schedule)로 교체',
         'API 인증 방식 확인 (Zendesk: Token, Intercom: Bearer, HubSpot: Private App)',
         '핵심 비즈니스 로직을 Code 노드에 구현'],
        ['해당 서비스 API Key/Token 발급', '알림 채널(Slack/이메일) 설정', '테스트 환경에서 먼저 실행'],
        ['자사 CS 프로세스에 맞게 트리거 조건 및 알림 내용 커스터마이즈 권장']
    )

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 5: 재무 / 회계
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Chapter 5. 재무 / 회계', 1)
add_para(doc, 'QuickBooks, Xero, SAP 연동, 청구서 자동화, 예산 모니터링 레시피 10선입니다.', size=Pt(10))
doc.add_paragraph()

WF_041 = json.dumps({
  "name": "Stripe 결제 → 자동 청구서 생성 및 회계 기록",
  "nodes": [
    {
      "parameters": {
        "httpMethod": "POST", "path": "stripe-payment",
        "options": {}
      },
      "id": "fin-0041-01", "name": "Webhook - Stripe 결제 완료",
      "type": "n8n-nodes-base.webhook", "typeVersion": 2, "position": [240, 300]
    },
    {
      "parameters": {
        "jsCode": "const event = $input.item.json;\nif (event.type !== 'payment_intent.succeeded') return [];\nconst pi = event.data?.object;\nreturn [{ json: {\n  amount: pi.amount / 100,\n  currency: pi.currency.toUpperCase(),\n  customerId: pi.customer,\n  paymentId: pi.id,\n  customerEmail: pi.receipt_email,\n  created: new Date(pi.created * 1000).toISOString()\n} }];"
      },
      "id": "fin-0041-02", "name": "Code - Stripe 이벤트 파싱",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [460, 300]
    },
    {
      "parameters": {
        "resource": "invoice", "operation": "create",
        "additionalFields": {
          "customer": "={{ $json.customerId }}",
          "currency": "={{ $json.currency.toLowerCase() }}",
          "lineItems": [{
            "description": "서비스 이용료",
            "quantity": 1,
            "unitPrice": "={{ $json.amount }}"
          }]
        }
      },
      "id": "fin-0041-03", "name": "QuickBooks - 청구서 생성",
      "type": "n8n-nodes-base.quickbooks", "typeVersion": 1, "position": [680, 300],
      "credentials": {"quickBooksOAuth2Api": {"id": "cred-qb", "name": "QuickBooks"}}
    },
    {
      "parameters": {
        "operation": "appendOrUpdate",
        "sheetId": "ACCOUNTING_SHEET_ID",
        "range": "결제기록!A:H",
        "keyRowIndex": 0,
        "options": {},
        "dataToSend": "autoMapInputData"
      },
      "id": "fin-0041-04", "name": "Sheets - 회계 기록",
      "type": "n8n-nodes-base.googleSheets", "typeVersion": 4, "position": [900, 300],
      "credentials": {"googleSheetsOAuth2Api": {"id": "cred-gsheets", "name": "Google Sheets"}}
    }
  ],
  "connections": {
    "Webhook - Stripe 결제 완료": {"main": [[{"node": "Code - Stripe 이벤트 파싱", "type": "main", "index": 0}]]},
    "Code - Stripe 이벤트 파싱": {"main": [[{"node": "QuickBooks - 청구서 생성", "type": "main", "index": 0}]]},
    "QuickBooks - 청구서 생성": {"main": [[{"node": "Sheets - 회계 기록", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["finance", "stripe", "quickbooks", "invoice"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 41,
    'Stripe 결제 완료 → QuickBooks 청구서 자동 생성',
    '재무 / 회계', '고급(Hard)',
    'Stripe Webhook, QuickBooks, Google Sheets',
    'Stripe Webhook (payment_intent.succeeded)',
    'Stripe에서 결제가 완료되면 자동으로 QuickBooks에 청구서를 생성하고 Google Sheets 회계 기록에 추가합니다. 수동 회계 기장 작업을 완전 자동화합니다.',
    WF_041,
    [
        'Stripe Webhook: Stripe Dashboard → Developers → Webhooks → Add endpoint. payment_intent.succeeded 이벤트 선택',
        'Webhook 서명 검증: Stripe-Signature 헤더로 보안 검증. n8n Code 노드에 검증 로직 추가 권장',
        'Code 파싱: event.type === "payment_intent.succeeded" 조건 확인. 다른 이벤트 타입 무시',
        'QuickBooks: OAuth2 연결. Sandbox 환경에서 먼저 테스트',
        'Google Sheets: ACCOUNTING_SHEET_ID를 실제 시트 ID로 교체',
    ],
    [
        'Stripe Webhook Secret 발급 및 n8n 설정',
        'ACCOUNTING_SHEET_ID를 실제 시트 ID로 교체',
        'QuickBooks OAuth2 자격증명 생성 (Intuit Developer 계정 필요)',
        'QuickBooks Sandbox에서 먼저 테스트',
        'QuickBooks Customer ID 매핑 로직 확인 (Stripe customer ID → QB customer)',
    ],
    [
        'Xero 대안: QuickBooks 노드 → Xero 노드로 동일 패턴 구현 가능',
        '세금계산서: 한국 전자세금계산서 연동 시 이지웰/bizin API 추가 연동 필요',
        '환불 처리: payment_intent.refunded 이벤트 추가하여 환불도 자동 처리',
    ]
)

WF_042 = json.dumps({
  "name": "월간 예산 소진율 모니터링 및 알림",
  "nodes": [
    {
      "parameters": {"rule": {"interval": [{"field": "cronExpression", "expression": "0 9 * * *"}]}},
      "id": "fin-0042-01", "name": "Schedule - 매일 09:00",
      "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]
    },
    {
      "parameters": {
        "operation": "read", "sheetId": "BUDGET_SHEET_ID",
        "range": "예산현황!A2:E20", "options": {}
      },
      "id": "fin-0042-02", "name": "Sheets - 예산 현황 조회",
      "type": "n8n-nodes-base.googleSheets", "typeVersion": 4, "position": [460, 300],
      "credentials": {"googleSheetsOAuth2Api": {"id": "cred-gsheets", "name": "Google Sheets"}}
    },
    {
      "parameters": {
        "jsCode": "const today = new Date();\nconst daysInMonth = new Date(today.getFullYear(), today.getMonth()+1, 0).getDate();\nconst dayOfMonth = today.getDate();\nconst timeProgress = dayOfMonth / daysInMonth;\nconst alerts = [];\nfor (const row of $input.all()) {\n  const vals = row.json.values || row.json;\n  const dept = vals[0];\n  const budget = parseFloat(vals[1]) || 0;\n  const spent = parseFloat(vals[2]) || 0;\n  const rate = budget > 0 ? spent / budget : 0;\n  // 시간 대비 예산 소진율이 20% 이상 초과\n  if (rate > timeProgress + 0.2) {\n    alerts.push({ json: { dept, budget, spent, rate: Math.round(rate*100), timeProgress: Math.round(timeProgress*100), overRate: Math.round((rate-timeProgress)*100) } });\n  }\n}\nreturn alerts.length ? alerts : [{ json: { noAlert: true } }];"
      },
      "id": "fin-0042-03", "name": "Code - 초과 부서 감지",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [680, 300]
    },
    {
      "parameters": {
        "conditions": {"boolean": [{"value1": "={{ !$json.noAlert }}", "value2": True}]}
      },
      "id": "fin-0042-04", "name": "IF - 알림 필요",
      "type": "n8n-nodes-base.if", "typeVersion": 1, "position": [900, 300]
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#finance-alerts",
        "text": "⚠️ 예산 초과 위험!\n부서: {{ $json.dept }}\n예산: {{ $json.budget?.toLocaleString() }}원\n지출: {{ $json.spent?.toLocaleString() }}원\n소진율: {{ $json.rate }}% (시간 진행: {{ $json.timeProgress }}%)\n초과: +{{ $json.overRate }}%p",
        "otherOptions": {}
      },
      "id": "fin-0042-05", "name": "Slack - 예산 초과 경보",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [1120, 200],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Schedule - 매일 09:00": {"main": [[{"node": "Sheets - 예산 현황 조회", "type": "main", "index": 0}]]},
    "Sheets - 예산 현황 조회": {"main": [[{"node": "Code - 초과 부서 감지", "type": "main", "index": 0}]]},
    "Code - 초과 부서 감지": {"main": [[{"node": "IF - 알림 필요", "type": "main", "index": 0}]]},
    "IF - 알림 필요": {"main": [[{"node": "Slack - 예산 초과 경보", "type": "main", "index": 0}], []]}
  },
  "active": False, "settings": {},
  "tags": ["finance", "budget", "monitoring", "slack"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 42,
    '부서별 예산 소진율 자동 모니터링',
    '재무 / 회계', '중급(Medium)',
    'Google Sheets, Slack',
    'Schedule (매일 오전 9시)',
    '매일 아침 부서별 예산 소진 현황을 확인하고, 시간 진행률 대비 예산 소진율이 20% 이상 초과한 부서를 자동 감지하여 재무팀 Slack에 경보를 보냅니다.',
    WF_042,
    [
        'Google Sheets: A=부서명, B=월예산, C=누계지출, D=담당자, E=비고 구조. 매일 지출 업데이트 필요',
        'Code 로직: 일별 시간 진행률 = 현재일/총일수. 예산소진율이 시간진행률+20% 초과 시 알림',
        'IF 노드: noAlert=true이면 true 브랜치(정상)로 → 아무 액션 없음',
        'Slack: 각 초과 부서별 개별 메시지. 집계 리포트로 변경 시 Code 노드에서 텍스트 조합',
    ],
    [
        'BUDGET_SHEET_ID를 실제 예산 관리 시트 ID로 교체',
        '시트 컬럼 구조 확인 (A~E 대응)',
        'Slack Bot Token 및 #finance-alerts 채널 설정',
        '알림 임계값(현재 +20%) 회사 정책에 맞게 조정',
        '부서별 예산 데이터 정기 업데이트 프로세스 수립',
    ],
    [
        'ERP 연동: SAP/Oracle ERP API가 있으면 Sheets 대신 직접 연동',
        '월간 리포트: 월말 자동 예산 집행 리포트 → PDF 생성 → CFO 이메일 발송 연계 가능',
    ]
)

# 재무 나머지 (043~050)
finance_recipes = [
    (43, '경비 청구서 OCR 자동 처리 (영수증 → 회계 기장)', '재무 / 회계', '전문가(Expert)',
     'Gmail, OpenAI Vision, Xero/QuickBooks, Slack', 'Gmail Trigger (영수증 이메일)',
     '직원이 영수증 이미지를 이메일로 전송하면 OpenAI Vision으로 자동 OCR 처리하여 Xero/QuickBooks에 자동 기장합니다.'),
    (44, 'Xero 미수금 자동 리마인더 발송', '재무 / 회계', '중급(Medium)',
     'Xero, Gmail, Slack', 'Schedule (매일)',
     'Xero에서 30일 이상 미납 청구서를 자동 조회하여 단계별(D+30/60/90) 리마인더 이메일을 자동 발송합니다.'),
    (45, '환율 변동 자동 헷지 알림', '재무 / 회계', '중급(Medium)',
     'Exchange Rate API, Slack, 이메일', 'Schedule (매일 오전)',
     '주요 통화(USD/EUR/JPY)의 전일 대비 환율 변동이 2% 이상이면 자동으로 재무팀에 헷지 검토 알림을 발송합니다.'),
    (46, '급여 처리 전 사전 검증 자동화', '재무 / 회계', '고급(Hard)',
     '급여 시스템 API, Google Sheets, Slack', 'Schedule (급여일 7일 전)',
     '급여 처리 7일 전에 급여 데이터의 이상치(전월 대비 50% 이상 변동)를 자동 감지하여 HR/재무팀에 검증 요청을 발송합니다.'),
    (47, 'AWS/GCP 클라우드 비용 일일 모니터링', '재무 / 회계', '중급(Medium)',
     'AWS Cost Explorer, GCP Billing, Slack', 'Schedule (매일 오전)',
     'AWS와 GCP의 전일 비용을 자동 집계하여 전주 동일 요일 대비 30% 이상 증가 시 DevOps팀에 즉시 알림을 보냅니다.'),
    (48, '세금계산서 자동 발행 및 관리 (한국)', '재무 / 회계', '전문가(Expert)',
     '국세청 API / 이지웰, 자체 청구 시스템, Gmail', 'Webhook (계약 체결)',
     '계약 체결 시 자동으로 전자세금계산서 발행 요청을 처리하고 고객 이메일로 발송하며 회계 시스템에 기록합니다.'),
    (49, '재무 KPI 월간 대시보드 자동 업데이트', '재무 / 회계', '중급(Medium)',
     'QuickBooks/Xero, Google Sheets, Looker Studio', 'Schedule (매월 1일)',
     '매월 1일 전월 재무 데이터(매출, 비용, 영업이익, 현금흐름)를 자동 집계하여 Looker Studio 대시보드를 자동 업데이트합니다.'),
    (50, 'AP 자동화 (구매 주문 → 인보이스 매칭)', '재무 / 회계', '전문가(Expert)',
     'ERP API, OpenAI, Gmail, Slack', '이메일 인보이스 수신',
     '공급업체 인보이스 이메일 수신 시 OpenAI로 자동 파싱하고 기존 구매 주문서와 자동 매칭하여 3-way match 검증 후 결재 승인 요청을 자동화합니다.'),
]

for num, title, cat, diff, apps, trigger, desc in finance_recipes:
    wf = simple_wf(num, title, ["finance"])
    make_recipe(doc, num, title, cat, diff, apps, trigger, desc, wf,
        ['서비스 API 인증 설정', '트리거 조건 및 임계값 회사 정책에 맞게 설정', '테스트 환경에서 소액으로 먼저 검증'],
        ['API Key/OAuth2 발급', '알림 채널 설정', '회계/재무 시스템 연동 테스트'],
        ['금융 데이터 처리 시 보안(암호화, 접근 권한) 강화 필수', '프로덕션 전 CFO/회계팀 승인 권장']
    )

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 6: IT운영 / DevOps
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Chapter 6. IT운영 / DevOps', 1)
add_para(doc, 'GitHub CI/CD, 인프라 모니터링, 배포 자동화, 인시던트 대응 레시피 10선입니다.', size=Pt(10))
doc.add_paragraph()

WF_051 = json.dumps({
  "name": "GitHub PR 머지 → Vercel 자동 배포 알림",
  "nodes": [
    {
      "parameters": {
        "httpMethod": "POST", "path": "github-webhook",
        "options": {}
      },
      "id": "dev-0051-01", "name": "Webhook - GitHub 이벤트",
      "type": "n8n-nodes-base.webhook", "typeVersion": 2, "position": [240, 300]
    },
    {
      "parameters": {
        "jsCode": "const event = $input.item.json;\n// PR 머지 이벤트만 처리\nif (event.action !== 'closed' || !event.pull_request?.merged) return [];\nreturn [{ json: {\n  prTitle: event.pull_request.title,\n  prNumber: event.pull_request.number,\n  author: event.pull_request.user?.login,\n  branch: event.pull_request.base?.ref,\n  repo: event.repository?.full_name,\n  mergedAt: event.pull_request.merged_at\n} }];"
      },
      "id": "dev-0051-02", "name": "Code - PR 머지 이벤트 필터",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [460, 300]
    },
    {
      "parameters": {
        "conditions": {"string": [{"value1": "={{ $json.branch }}", "operation": "equals", "value2": "main"}]}
      },
      "id": "dev-0051-03", "name": "IF - main 브랜치",
      "type": "n8n-nodes-base.if", "typeVersion": 1, "position": [680, 300]
    },
    {
      "parameters": {
        "url": "https://api.vercel.com/v13/deployments",
        "method": "POST",
        "authentication": "predefinedCredentialType",
        "nodeCredentialType": "httpHeaderAuth",
        "body": {
          "body": {
            "name": "my-project",
            "gitSource": {"type": "github", "repo": "={{ $json.repo }}", "ref": "main"}
          }
        }
      },
      "id": "dev-0051-04", "name": "HTTP - Vercel 배포 트리거",
      "type": "n8n-nodes-base.httpRequest", "typeVersion": 4, "position": [900, 200],
      "credentials": {"httpHeaderAuth": {"id": "cred-vercel", "name": "Vercel"}}
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#deployments",
        "text": "🚀 배포 시작!\nPR: #{{ $json.prNumber }} {{ $json.prTitle }}\n작성자: {{ $json.author }}\n브랜치: {{ $json.branch }}\nVercel 배포 ID: {{ $('HTTP - Vercel 배포 트리거').item.json.id }}",
        "otherOptions": {}
      },
      "id": "dev-0051-05", "name": "Slack - 배포 시작 알림",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [1120, 200],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Webhook - GitHub 이벤트": {"main": [[{"node": "Code - PR 머지 이벤트 필터", "type": "main", "index": 0}]]},
    "Code - PR 머지 이벤트 필터": {"main": [[{"node": "IF - main 브랜치", "type": "main", "index": 0}]]},
    "IF - main 브랜치": {"main": [[{"node": "HTTP - Vercel 배포 트리거", "type": "main", "index": 0}], []]},
    "HTTP - Vercel 배포 트리거": {"main": [[{"node": "Slack - 배포 시작 알림", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["devops", "github", "vercel", "deployment", "ci-cd"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 51,
    'GitHub PR 머지 → Vercel 자동 배포 트리거',
    'IT운영 / DevOps', '중급(Medium)',
    'GitHub Webhook, Vercel API, Slack',
    'GitHub Webhook (pull_request 이벤트)',
    'GitHub main 브랜치에 PR이 머지되면 자동으로 Vercel 배포를 트리거하고 Slack #deployments 채널에 배포 시작 알림을 보냅니다. CI/CD 파이프라인의 배포 단계를 자동화합니다.',
    WF_051,
    [
        'GitHub Webhook: Repository → Settings → Webhooks → Add webhook. pull_request 이벤트 선택. Webhook Secret 설정',
        'Code 필터: action === "closed" && merged === true 조건. 일반 PR 닫기와 머지 구분',
        'IF 브랜치: main 브랜치만 프로덕션 배포. 다른 브랜치(staging 등) 분기 처리 가능',
        'Vercel API: Authorization: Bearer {TOKEN}. my-project를 실제 Vercel 프로젝트명으로 교체',
        'Slack: 배포 ID로 Vercel 대시보드 링크 생성 가능',
    ],
    [
        'GitHub Webhook URL 등록 (Repository → Settings → Webhooks)',
        'GitHub Webhook Secret 설정 및 n8n에서 서명 검증 로직 추가',
        'Vercel API Token 생성 (Vercel → Settings → Tokens)',
        'Vercel 프로젝트명 확인 및 교체',
        'Slack Bot Token 및 #deployments 채널 설정',
    ],
    [
        '배포 완료 감지: Vercel Webhook + 별도 워크플로우로 배포 성공/실패 알림 추가',
        '환경별 분기: main → production, develop → staging으로 IF 노드로 분기',
        'GitHub Actions 대안: 이미 GitHub Actions 쓰는 경우 n8n을 후처리(알림, 슬랙 보고) 용도로만 활용',
    ]
)

WF_052 = json.dumps({
  "name": "서버 다운타임 자동 감지 및 인시던트 생성",
  "nodes": [
    {
      "parameters": {"rule": {"interval": [{"field": "cronExpression", "expression": "*/5 * * * *"}]}},
      "id": "dev-0052-01", "name": "Schedule - 5분마다",
      "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]
    },
    {
      "parameters": {
        "url": "https://YOUR_APP_URL/health",
        "method": "GET",
        "options": {"timeout": 10000, "allowUnauthorizedCerts": False}
      },
      "id": "dev-0052-02", "name": "HTTP - 헬스체크",
      "type": "n8n-nodes-base.httpRequest", "typeVersion": 4, "position": [460, 300]
    },
    {
      "parameters": {
        "conditions": {"number": [{"value1": "={{ $json.statusCode || 200 }}", "operation": "notEqual", "value2": 200}]}
      },
      "id": "dev-0052-03", "name": "IF - 장애 감지",
      "type": "n8n-nodes-base.if", "typeVersion": 1, "position": [680, 300]
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#incidents",
        "text": "🚨 서비스 장애 감지!\nURL: YOUR_APP_URL\n상태코드: {{ $json.statusCode }}\n감지시간: {{ $now.toISO() }}\n즉시 확인 필요! @here",
        "otherOptions": {}
      },
      "id": "dev-0052-04", "name": "Slack - 인시던트 알림",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [900, 200],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    },
    {
      "parameters": {
        "url": "https://events.pagerduty.com/v2/enqueue",
        "method": "POST",
        "body": {
          "body": {
            "routing_key": "YOUR_PAGERDUTY_INTEGRATION_KEY",
            "event_action": "trigger",
            "payload": {
              "summary": "서비스 다운: YOUR_APP_URL",
              "severity": "critical",
              "source": "n8n-monitor",
              "timestamp": "={{ $now.toISO() }}"
            }
          }
        }
      },
      "id": "dev-0052-05", "name": "HTTP - PagerDuty 인시던트",
      "type": "n8n-nodes-base.httpRequest", "typeVersion": 4, "position": [900, 420]
    }
  ],
  "connections": {
    "Schedule - 5분마다": {"main": [[{"node": "HTTP - 헬스체크", "type": "main", "index": 0}]]},
    "HTTP - 헬스체크": {"main": [[{"node": "IF - 장애 감지", "type": "main", "index": 0}]]},
    "IF - 장애 감지": {"main": [[
      {"node": "Slack - 인시던트 알림", "type": "main", "index": 0},
      {"node": "HTTP - PagerDuty 인시던트", "type": "main", "index": 0}
    ], []]}
  },
  "active": False, "settings": {},
  "tags": ["devops", "monitoring", "pagerduty", "uptime", "incident"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 52,
    '서버 다운타임 자동 감지 + PagerDuty 인시던트 생성',
    'IT운영 / DevOps', '중급(Medium)',
    'HTTP Request, PagerDuty, Slack',
    'Schedule (5분마다 헬스체크)',
    '5분마다 서비스 헬스체크 URL을 자동으로 호출하고, 응답이 200이 아니면 Slack #incidents 채널에 즉시 알림을 보내고 PagerDuty 인시던트를 자동 생성하여 당직 엔지니어에게 온콜 알림을 발송합니다.',
    WF_052,
    [
        'HTTP 헬스체크: YOUR_APP_URL/health를 실제 헬스체크 엔드포인트로 교체. timeout 10초 설정',
        'IF 장애 감지: statusCode != 200 조건. 타임아웃 오류도 처리하려면 Error Workflow 추가',
        'Slack @here: 채널 전체 알림. 중요도에 따라 @channel(전체) 또는 개인 멘션 조정',
        'PagerDuty: YOUR_PAGERDUTY_INTEGRATION_KEY를 실제 Integration Key로 교체. Events API v2 사용',
    ],
    [
        'YOUR_APP_URL을 실제 헬스체크 URL로 교체',
        'YOUR_PAGERDUTY_INTEGRATION_KEY를 실제 키로 교체 (PagerDuty → Services → Integration)',
        'Slack Bot Token 및 #incidents 채널 설정 (@here 멘션 권한 확인)',
        '알람 중복 방지: n8n 변수/캐시로 이미 알림 발송된 상태 추적 로직 추가 권장',
        '복구 감지: 별도 워크플로우로 서비스 복구 시 "해결됨" 알림 발송',
    ],
    [
        '알람 플래핑 방지: 연속 3회 실패 시에만 알림 발송하는 카운터 로직 추가 권장',
        'UptimeRobot 대안: 무료 업타임 모니터링 서비스 활용 + n8n으로 인시던트 처리만 담당',
        '다중 URL: 여러 서비스를 동시에 모니터링하려면 URL 목록 Sheets에서 관리',
    ]
)

# DevOps 나머지 (053~060)
devops_recipes = [
    (53, 'Docker 컨테이너 메모리 임계값 초과 자동 재시작', 'IT운영 / DevOps', '고급(Hard)',
     'Docker API, Slack, PagerDuty', 'Schedule (1분마다)',
     'Docker 컨테이너 메모리 사용량을 1분마다 모니터링하여 90% 이상 시 자동 재시작하고 Slack에 알림을 보냅니다.'),
    (54, 'SSL 인증서 만료 사전 알림 (30일/7일 전)', 'IT운영 / DevOps', '초급(Easy)',
     'HTTP Request (SSL 체크), Slack, Gmail', 'Schedule (매일)',
     '매일 자동으로 보유 도메인의 SSL 인증서 만료일을 체크하여 30일, 7일 전에 인프라팀에 자동 갱신 알림을 발송합니다.'),
    (55, 'GitHub Issues → Jira 자동 동기화', 'IT운영 / DevOps', '고급(Hard)',
     'GitHub Webhook, Jira, Slack', 'GitHub Webhook (issue 이벤트)',
     'GitHub에서 이슈가 생성/수정/종료될 때 Jira 티켓과 자동 동기화합니다. 양방향 동기화는 별도 설정 필요합니다.'),
    (56, 'Kubernetes 파드 실패 자동 감지 및 재시작', 'IT운영 / DevOps', '전문가(Expert)',
     'Kubernetes API, Slack, PagerDuty', 'Schedule (5분마다)',
     'K8s 클러스터에서 CrashLoopBackOff 상태 파드를 자동 감지하여 재시작하고 DevOps팀에 상세 로그와 함께 알림을 발송합니다.'),
    (57, 'AWS EC2 비용 최적화 자동 분석', 'IT운영 / DevOps', '고급(Hard)',
     'AWS Cost Explorer API, Slack, Gmail', 'Schedule (매주 월요일)',
     '매주 AWS EC2 인스턴스별 비용과 사용률을 분석하여 과소 사용 인스턴스를 자동 감지하고 다운사이징 권고를 DevOps팀에 발송합니다.'),
    (58, '데이터베이스 백업 완료 자동 검증', 'IT운영 / DevOps', '고급(Hard)',
     'AWS S3 / GCS, Slack, Gmail', 'Schedule (매일 백업 후)',
     '매일 DB 백업 완료 후 자동으로 백업 파일 존재 여부와 크기를 검증하고, 백업 실패 시 DBA팀에 즉시 알림을 발송합니다.'),
    (59, 'Log 이상 패턴 감지 (CloudWatch → Slack)', 'IT운영 / DevOps', '전문가(Expert)',
     'AWS CloudWatch, OpenAI, Slack, PagerDuty', 'Schedule (10분마다)',
     'CloudWatch Logs에서 ERROR/CRITICAL 로그를 10분마다 수집하고 OpenAI로 이상 패턴을 분석하여 근본 원인과 해결 방안을 Slack에 자동 발송합니다.'),
    (60, 'CI/CD 빌드 실패 자동 분석 및 담당자 알림', 'IT운영 / DevOps', '고급(Hard)',
     'GitHub Actions/Jenkins API, OpenAI, Slack', 'GitHub Webhook (workflow_run 실패)',
     'GitHub Actions 빌드가 실패하면 자동으로 실패 로그를 수집하고 OpenAI로 원인을 분석하여 담당 개발자에게 Slack DM으로 자동 발송합니다.'),
]

for num, title, cat, diff, apps, trigger, desc in devops_recipes:
    wf = simple_wf(num, title, ["devops"])
    make_recipe(doc, num, title, cat, diff, apps, trigger, desc, wf,
        ['트리거를 실제 인프라 모니터링 방식에 맞게 설정',
         'API 인증: 클라우드 서비스 IAM Role/Key 설정 (최소 권한 원칙)',
         '알림 채널과 에스컬레이션 경로 정의'],
        ['클라우드 서비스 API Key/Role 설정', 'PagerDuty Integration Key 발급', '알림 채널 및 온콜 스케줄 설정'],
        ['인프라 자동화는 되돌리기 어려운 작업(재시작, 삭제) 포함 시 인간 승인 스텝 추가 필수']
    )

doc.add_page_break()
doc.save(OUT)
print('Vol3 Part3 done:', OUT)
