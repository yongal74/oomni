# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "C:/Users/장우경/oomni/n8n_cases_full.docx"
doc = Document(PATH)

def sc(doc, text, level=1):
    colors = {1: RGBColor(0x1F,0x49,0x7D), 2: RGBColor(0x2E,0x74,0xB5), 3: RGBColor(0x1F,0x49,0x7D)}
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(sizes[level]); run.font.bold = True; run.font.color.rgb = colors[level]
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after = Pt(6)

def par(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

def bul(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run("  • " + item)
        run.font.size = Pt(10); run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)

def flow(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(9); run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    p._element.pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)

def case(doc, num, title, desc, nodes, flowchart, tips):
    sc(doc, f"[{num}] {title}", 2)
    par(doc, desc)
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run("▶ 주요 노드: "); run.font.bold = True; run.font.size = Pt(10)
    run.font.name = '맑은 고딕'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run2 = p.add_run(nodes); run2.font.size = Pt(10); run2.font.name = '맑은 고딕'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    p2 = doc.add_paragraph(); p2.style = doc.styles['Normal']
    run3 = p2.add_run("▶ 워크플로우 순서도"); run3.font.bold = True; run3.font.size = Pt(10)
    run3.font.name = '맑은 고딕'; run3._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p2.paragraph_format.space_before = Pt(4); p2.paragraph_format.space_after = Pt(2)
    flow(doc, flowchart)
    p3 = doc.add_paragraph(); p3.style = doc.styles['Normal']
    run4 = p3.add_run("▶ 핵심 포인트"); run4.font.bold = True; run4.font.size = Pt(10)
    run4.font.name = '맑은 고딕'; run4._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p3.paragraph_format.space_before = Pt(4); p3.paragraph_format.space_after = Pt(2)
    bul(doc, tips)
    doc.add_paragraph()

# ─── CH5: 고객 서비스 자동화 (043~053) ────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 5: 고객 서비스 자동화", 1)
par(doc, "AI를 활용한 CS 자동화로 응답 시간을 단축하고 고객 만족도를 높이면서 상담원 업무를 효율화합니다.")

case(doc, "043", "Zendesk 티켓 자동 우선순위 설정",
"수신된 지원 티켓을 AI로 분석하여 자동으로 우선순위를 설정하고 적합한 담당자에게 배정합니다.",
"Zendesk Trigger, OpenAI, Zendesk",
"""Zendesk Trigger (새 티켓 생성)
  │
  ▼
OpenAI (티켓 분석)
  │ 분석 항목:
  │ - 긴급도 (Critical/High/Normal/Low)
  │ - 카테고리 (버그/기능/결제/일반)
  │ - 감정 (불만족/중립/만족)
  │ - 필요 스킬 (기술/결제/영업)
  ▼
Zendesk (티켓 업데이트)
  │ - 우선순위 설정
  │ - 태그 추가
  │ - 담당자 배정 (스킬 기반)
  ▼
[IF] Critical 티켓?
  │ YES
  ▼
Slack (긴급 알림 + PagerDuty)
  └─ "🚨 Critical 티켓 #{id}: {{subject}}"
  ▼
Zendesk (자동 1차 응답 발송)
  └─ "접수되었습니다. {{SLA}} 시간 내 응대 예정"
  ▼
[완료]""",
["티켓 자동 배정 시 담당자 현재 업무량 고려 (워크로드 밸런싱)", "감정이 매우 부정적인 티켓은 시니어 상담원에게 우선 배정", "자동 1차 응답에 FAQ 링크를 포함하면 셀프서비스 해결률 향상"])

case(doc, "044", "AI 기반 FAQ 자동 응답",
"자주 묻는 질문 유형의 티켓에 AI가 지식 베이스를 검색하여 자동 답변을 생성합니다.",
"Zendesk Trigger, OpenAI, HTTP Request (벡터DB), Zendesk",
"""Zendesk Trigger (새 티켓)
  │
  ▼
OpenAI (질문 분류)
  │ "이것이 FAQ로 해결 가능한 질문인가?"
  ▼
[IF] FAQ 가능성 > 80%?
  │ YES
  ▼
Vector Store 검색 (지식 베이스)
  │ - 관련 문서 Top 3 검색
  │ - 유사도 점수 확인
  ▼
[IF] 관련 문서 유사도 > 0.85?
  │ YES
  ▼
OpenAI (답변 생성)
  │ 컨텍스트: 검색된 문서
  │ "자연스럽고 도움이 되는 답변 작성"
  ▼
Zendesk (자동 답변 게시)
  │ + 관련 도움말 링크 첨부
  ▼
Zendesk (만족도 확인 요청)
  └─ "도움이 되셨나요? 해결 / 미해결"
  ▼
[IF] 미해결 → 상담원 배정""",
["자동 답변 후 24시간 내 고객 반응이 없으면 해결됨으로 자동 처리", "유사도 임계값(0.85)은 실제 데이터로 조정 필요", "AI 답변 품질 모니터링: 주간 샘플링 검토 프로세스 유지"])

case(doc, "045", "고객 이탈 방지 자동 개입",
"고객이 해지/취소 관련 문의를 하면 즉시 특별 오퍼를 제시하거나 시니어 상담원을 연결합니다.",
"Zendesk Trigger, OpenAI, Slack, Gmail",
"""Zendesk Trigger (새 티켓)
  │
  ▼
OpenAI (이탈 의도 감지)
  │ 키워드: "해지", "취소", "환불", "탈퇴"
  │ 감정: 매우 부정적
  ▼
[IF] 이탈 의도 감지?
  │ YES
  ▼
[고객 등급 확인]
  │
  ├─ VIP/고가치 고객
  │   └─ Slack (시니어 CS 즉시 알림)
  │       + 관리자 에스컬레이션
  │
  └─ 일반 고객
      └─ 자동 리텐션 오퍼 발송
          "계속 이용 시 1개월 무료"
          또는 "플랜 다운그레이드 제안"
  ▼
CRM (이탈 위험 태그 추가)
  ▼
[Wait 48시간 → 해결 여부 확인]""",
["이탈 방지 오퍼는 남발하면 효과가 줄어드므로 고객당 연 2회 이하", "해지 사유를 반드시 수집하여 제품/서비스 개선에 반영", "VIP 고객 이탈 방지율이 일반보다 ROI 5배 이상 높음"])

case(doc, "046", "서비스 장애 알림 자동화",
"서비스 모니터링에서 장애가 감지되면 즉시 고객에게 상태 업데이트를 발송하고 내부 대응을 시작합니다.",
"HTTP Request (모니터링 서비스 Webhook), Slack, Gmail, StatusPage",
"""모니터링 서비스 Webhook (장애 감지)
  │ (Pingdom, UptimeRobot, Datadog 등)
  ▼
[병렬 처리 — 즉시 실행]
  ├─ Slack (#incident채널)
  │   "🔴 서비스 장애 감지: {{service}} {{time}}"
  ├─ PagerDuty (온콜 엔지니어 호출)
  ├─ StatusPage (장애 공지 자동 게시)
  └─ Gmail (영향받는 고객 대량 알림)
      "현재 {{service}} 서비스 장애 발생..."
  ▼
Schedule (15분마다: 진행 상황 업데이트)
  │ StatusPage + Slack 업데이트
  ▼
[Webhook: 장애 해결 이벤트]
  │
  ▼
Gmail (복구 완료 이메일)
  └─ "서비스 복구 완료. 불편 드려 죄송합니다"
  ▼
Slack (인시던트 사후 리뷰 생성)""",
["장애 감지 → 첫 고객 알림까지 5분 이내 목표 설정", "투명한 커뮤니케이션이 이탈 방지에 핵심: 상세 업데이트 제공", "포스트모템 문서 자동 생성 (타임라인 + 원인 + 대응 기록)"])

case(doc, "047", "고객 만족도 자동 측정",
"지원 티켓 해결 후 자동으로 만족도 설문(CSAT)을 발송하고 결과를 분석하여 CS팀에 리포트합니다.",
"Zendesk Trigger, Gmail, Google Sheets, Slack",
"""Zendesk Trigger (티켓 해결됨)
  │
  ▼
Wait (2시간: 고객 확인 시간)
  │
  ▼
Gmail (CSAT 설문 이메일)
  │ "서비스에 만족하셨나요?"
  │ [매우 만족] [만족] [보통] [불만족] [매우 불만족]
  │ (각 링크에 티켓 ID + 점수 포함)
  ▼
[Webhook: 고객 클릭]
  │
  ▼
Google Sheets (응답 기록)
  │ - 티켓 ID, 담당자, 점수, 일시
  ▼
[IF] 불만족 (1~2점)?
  │ YES
  ▼
Slack (CS 매니저 즉시 알림)
  └─ "😞 낮은 CSAT: {{customer}} (담당: {{agent}})"
     후속 대응 요청
  ▼
Schedule (매주 월요일: 주간 CSAT 리포트)""",
["CSAT 응답률 목표 30% 이상 (산업 평균 15%)설정", "담당자별 CSAT 추적으로 개인 코칭 및 표창 프로그램 연계", "매우 만족 고객에게 리뷰/추천인 프로그램 제안 자동화"])

case(doc, "048", "다국어 CS 자동 번역 대응",
"외국어로 수신된 지원 요청을 자동으로 번역하고 한국어로 처리한 후 원래 언어로 번역하여 응답합니다.",
"Gmail Trigger, OpenAI, Gmail",
"""Gmail Trigger (새 CS 이메일)
  │
  ▼
OpenAI (언어 감지)
  │ "이 텍스트의 언어는?"
  ▼
[IF] 언어 != 한국어?
  │ YES
  ▼
OpenAI (한국어로 번역)
  │
  ▼
Slack (CS팀: 번역된 문의 공유)
  │ - 원문 + 번역문 함께 제공
  │ - 담당자 답변 작성 요청
  ▼
[Wait] 담당자 Slack 스레드 답변
  │
  ▼
OpenAI (원래 언어로 번역)
  │
  ▼
Gmail (번역된 답변 발송)
  └─ "Dear {{name}}, [translated response]"
  ▼
[완료]""",
["GPT-4o 번역 품질이 DeepL 수준이므로 별도 번역 API 불필요", "자주 묻는 언어(영어, 일본어, 중국어)는 사전 번역 템플릿 준비", "공식 문서/약관 번역은 전문 번역가 검토 병행 권장"])

case(doc, "049", "반품/환불 처리 자동화",
"반품/환불 요청 수신 시 정책 확인, 승인 처리, 창고 알림, 환불 진행까지 자동화합니다.",
"Gmail Trigger / Webhook, Code, Gmail, Slack",
"""Gmail Trigger (반품@ 이메일)
  │
  ▼
OpenAI (반품 정보 추출)
  │ - 주문번호, 반품 사유, 요청 유형
  ▼
HTTP Request (주문 시스템 API)
  │ - 주문일, 금액, 배송 현황 확인
  ▼
Code 노드 (반품 정책 자동 검토)
  │ - 구매 후 30일 이내? ✓
  │ - 미개봉/정상 상품? ✓
  │ - 세일 상품 제외? ✓
  ▼
[IF] 자동 승인 조건 충족?
  │ YES                    │ NO
  ▼                        ▼
Gmail (반품 승인 + 레이블 발송) Slack (CS 매니저 검토 요청)
  ↓
창고 관리 시스템 (반품 등록)
  ↓
결제 시스템 API (환불 처리)
  ↓
Gmail (환불 완료 안내)""",
["자동 승인 기준을 명확히 설정하여 오승인 방지", "반품 사유 데이터 누적으로 제품 품질 개선 피드백 제공", "환불 처리 완료까지 평균 시간 KPI로 추적"])

case(doc, "050", "고객 등급 자동 업데이트",
"고객의 구매 금액, 구매 빈도, 활성도를 기반으로 자동으로 등급을 산정하고 혜택 안내를 발송합니다.",
"Schedule Trigger, DB / CRM 쿼리, Code, Gmail",
"""Schedule Trigger (매월 1일 오전 3시)
  │
  ▼
DB 쿼리 (고객별 지난 12개월 구매 데이터)
  │
  ▼
Code 노드 (등급 계산)
  │ Bronze: < 50만원
  │ Silver: 50~200만원
  │ Gold: 200~500만원
  │ VIP: > 500만원
  ▼
[Loop] 등급 변경 고객 처리
  │
  ├─ 등급 상승: Gmail (축하 이메일 + 혜택 안내)
  │             "🎉 Gold 등급으로 승급되었습니다!"
  │
  └─ 등급 하락: Gmail (유지 혜택 안내)
                "등급 유지를 위한 특별 혜택"
  ▼
CRM (고객 등급 업데이트)
  ▼
Slack (마케팅팀: 등급별 고객 현황 리포트)""",
["등급 하락 이메일은 긍정적 프레이밍으로 작성 (유지 기회 강조)", "VIP 등급 달성 시 담당 AM(Account Manager) 자동 배정", "등급별 혜택 정보는 별도 설정 파일로 관리하여 쉽게 변경"])

case(doc, "051", "실시간 채팅 → 티켓 자동 변환",
"라이브 채팅에서 해결되지 않은 대화를 자동으로 지원 티켓으로 변환하고 요약 정보를 추가합니다.",
"Intercom/Zendesk Chat Webhook, OpenAI, Zendesk",
"""Intercom Webhook (채팅 종료 이벤트)
  │
  ▼
[IF] 미해결 대화?
  │ YES
  ▼
OpenAI (대화 요약)
  │ - 주요 문제점 3줄 요약
  │ - 고객 감정 상태
  │ - 필요한 후속 조치
  ▼
Zendesk (티켓 생성)
  │ - 제목: AI 생성 요약
  │ - 본문: 전체 채팅 기록
  │ - 우선순위: 감정 기반 자동 설정
  ▼
Slack (담당 팀 알림)
  └─ "채팅 → 티켓 변환: #{{ticketId}}"
  ▼
Gmail (고객 확인 이메일)
  └─ "문의가 접수되었습니다. 티켓 #{{id}}"
  ▼
[완료]""",
["채팅 기록 전체 보존으로 담당자가 컨텍스트 파악 용이", "감정 분석 결과를 티켓 우선순위에 반영하여 불만 고객 우선 처리", "채팅 미해결률 추적으로 챗봇 성능 개선 피드백 활용"])

case(doc, "052", "고객 민원 에스컬레이션 자동화",
"반복 문의, 장기 미해결 티켓, VIP 고객 불만을 감지하여 자동으로 상위 관리자에게 에스컬레이션합니다.",
"Schedule Trigger, Zendesk API, Slack, Gmail",
"""Schedule Trigger (매 2시간)
  │
  ▼
Zendesk API (에스컬레이션 조건 확인)
  │
  ├─ [조건 1] SLA 위반 임박 (2시간 이내)
  ├─ [조건 2] 동일 고객 3건 이상 티켓
  ├─ [조건 3] VIP 고객 미해결 24시간+
  └─ [조건 4] 24시간 미응답 티켓
  ▼
[IF] 에스컬레이션 조건 충족?
  │ YES
  ▼
Switch (에스컬레이션 레벨)
  ├─ L1 → Slack (팀장 DM)
  ├─ L2 → Slack (CS 매니저) + Gmail
  └─ L3 → Slack (#경영진채널) + 전화 알림
  ▼
Zendesk (에스컬레이션 태그 + 우선순위 상향)
  ▼
[완료]""",
["에스컬레이션 기준을 명확히 문서화하여 담당자 혼선 방지", "에스컬레이션된 티켓은 별도 추적하여 해결까지 시간 측정", "반복 에스컬레이션 패턴 분석으로 근본 원인 해결 유도"])

case(doc, "053", "챗봇 → 상담원 원활한 핸드오프",
"AI 챗봇이 해결 못 하는 복잡한 문의를 상담원에게 컨텍스트와 함께 전달하여 원활한 이관을 지원합니다.",
"Webhook (챗봇 이관 요청), OpenAI, Zendesk, Slack",
"""챗봇 Webhook (이관 요청 이벤트)
  │ - 챗봇 대화 기록
  │ - 이관 사유
  │ - 고객 정보
  ▼
OpenAI (대화 컨텍스트 요약)
  │ "담당 상담원을 위한 브리핑:
  │  고객: {{name}}, 문제: {{issue}},
  │  지금까지 시도: {{attempts}},
  │  고객 감정: {{emotion}}"
  ▼
Zendesk (티켓 생성 + AI 요약 첨부)
  │
  ▼
Slack (가용 상담원 찾기)
  │ - 현재 활성 티켓 수 기준
  │ - 해당 카테고리 전문 상담원 우선
  ▼
상담원 Zendesk 배정
  ▼
채팅창 (고객에게 이관 안내)
  └─ "전문 상담원 {{agent}}님이 연결됩니다"
  ▼
[완료]""",
["이관 시 고객이 처음부터 반복 설명하지 않도록 컨텍스트 완전 전달", "이관 대기 시간이 3분 초과 시 고객에게 예상 대기 시간 안내", "이관 후 첫 응답 시간(FART) 측정으로 상담원 성과 관리"])

# ─── CH6: 재무/회계 자동화 (054~063) ─────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 6: 재무/회계 자동화", 1)
par(doc, "반복적인 재무 업무를 자동화하여 처리 속도를 높이고 인적 오류를 줄이며 재무 데이터의 실시간 가시성을 확보합니다.")

case(doc, "054", "청구서 자동 처리 (OCR + AI)",
"이메일로 수신된 청구서 PDF를 OCR로 읽고 AI로 핵심 정보를 추출하여 회계 시스템에 자동 입력합니다.",
"Gmail Trigger, Google Drive, HTTP Request (OCR API), OpenAI, QuickBooks",
"""Gmail Trigger (재무@ + PDF 첨부)
  │
  ▼
Google Drive (청구서 PDF 저장)
  │
  ▼
HTTP Request (Google Vision OCR API)
  │ PDF → 텍스트 변환
  ▼
OpenAI (청구서 정보 추출)
  │ JSON 출력:
  │ { vendor, invoice_no, date,
  │   due_date, items[], total, tax }
  ▼
[IF] 신뢰도 > 90%?
  │ YES                    │ NO
  ▼                        ▼
QuickBooks (전표 자동 생성)  Slack (담당자 수동 확인 요청)
  ↓
Gmail (발신자에게 수신 확인)
  ↓
Google Sheets (청구서 대장 기록)""",
["OCR 정확도 < 90% 케이스는 반드시 수동 검토 단계 유지", "벤더별 청구서 포맷을 사전 학습하면 추출 정확도 향상", "중복 청구서 방지: 인보이스 번호 중복 체크 로직 필수"])

case(doc, "055", "경비 정산 자동 처리",
"직원이 영수증 사진을 업로드하면 OCR로 금액을 추출하고 정책 검토 후 승인/거절 처리합니다.",
"Webhook (앱 업로드), OCR API, OpenAI, Gmail, Google Sheets",
"""Webhook (경비 신청 앱)
  │ - 사진: 영수증 이미지
  │ - 입력: 항목, 프로젝트 코드
  ▼
HTTP Request (OCR API)
  │ 영수증 → 금액, 날짜, 상호 추출
  ▼
Code 노드 (경비 정책 검토)
  │ - 금액 한도 확인 (항목별)
  │ - 날짜 유효성 확인
  │ - 중복 신청 확인
  ▼
[IF] 자동 승인 조건?
  │ YES (< 5만원, 정책 내)  │ NO (> 5만원 or 예외)
  ▼                          ▼
Google Sheets (자동 승인)  Gmail (팀장 승인 요청)
Gmail (신청자 승인 안내)     └─ 승인/거절 링크 포함
  ↓
급여 시스템 (다음 월급에 반영)""",
["경비 항목별 한도 정책을 설정 파일로 관리하여 쉽게 변경", "AI 승인 결정에 감사 로그를 반드시 유지", "분기별 경비 패턴 분석으로 비용 절감 기회 식별"])

case(doc, "056", "미수금 자동 추적 및 독촉",
"지급 기한이 지난 미수금을 자동으로 감지하여 단계별 독촉 메시지를 발송합니다.",
"Schedule Trigger, QuickBooks/Xero API, Gmail",
"""Schedule Trigger (매일 오전 8시)
  │
  ▼
QuickBooks API (기간 초과 청구서 조회)
  │ - D+1~15: 1차 안내
  │ - D+16~30: 2차 요청
  │ - D+31~60: 3차 강조
  │ - D+61+: 법무팀 이관
  ▼
[Loop] 각 미수금 건 처리
  │
  ▼
Switch (연체 기간)
  ├─ D+1~15: Gmail (친근한 리마인더)
  │   "청구서 마감일이 지났습니다. 확인 부탁드립니다"
  ├─ D+16~30: Gmail (정중한 요청)
  │   "결제 처리를 도와드릴 방법이 있을까요?"
  ├─ D+31~60: Gmail (공식 독촉) + Slack (담당자)
  └─ D+61+: Slack (법무팀) + Gmail (내용증명 예고)
  ▼
Google Sheets (독촉 이력 기록)""",
["독촉 어조는 단계별로 점진적으로 강화 (우호 → 공식 → 법적)", "결제 링크를 이메일에 직접 포함하여 결제 마찰 최소화", "장기 거래처는 결제 어려움 사전 파악 후 분할 납부 협의"])

case(doc, "057", "월말 재무 마감 자동화",
"매월 말 재무 데이터를 자동으로 수집하고 마감 체크리스트를 완료하며 경영진에게 요약 리포트를 발송합니다.",
"Schedule Trigger, QuickBooks API, Google Sheets, Code, Gmail",
"""Schedule Trigger (매월 마지막 영업일 17시)
  │
  ▼
QuickBooks API (당월 재무 데이터 추출)
  │ - 매출, 비용, 손익
  │ - 현금 흐름
  │ - 미수금/미지급금 현황
  ▼
Code 노드 (전월 대비 분석 + KPI 계산)
  │ - 매출 성장률
  │ - 비용 효율성
  │ - 현금 보유 일수
  ▼
Google Sheets (월별 재무 데이터 기록)
  │
  ▼
Code 노드 (경영진 요약 리포트 생성)
  │
  ▼
Gmail (CFO + CEO 리포트 발송)
  └─ "{{month}}월 재무 현황 요약"
  ▼
Slack (#경영진채널 알림)""",
["마감 전 외상매출금/미지급금 대사(reconciliation) 자동화 병행", "전년 동월 대비 성장률 포함으로 맥락 있는 분석 제공", "이상값 감지: 예산 대비 20% 이상 차이 시 별도 하이라이트"])

case(doc, "058", "예산 대비 지출 모니터링",
"부서별 예산 소진 현황을 실시간으로 추적하여 초과 위험 시 사전 경고를 발송합니다.",
"Schedule Trigger, Google Sheets / ERP, Code, Slack, Gmail",
"""Schedule Trigger (매주 월요일)
  │
  ▼
Google Sheets (부서별 예산 + 실제 지출 데이터 조회)
  │
  ▼
Code 노드 (예산 소진율 계산)
  │ 소진율 = 실제지출 / 예산 × 100
  │ 기간조정 소진율 = 경과시간 / 총기간 × 100
  ▼
[Loop] 각 부서 처리
  │
  ├─ 소진율 > 기간조정 + 20%
  │   → Slack (부서장) "⚠️ 예산 초과 위험"
  │
  ├─ 소진율 > 90%
  │   → Slack + Gmail (CFO 보고)
  │
  └─ 정상: 기록만
  ▼
매월 1일: 전월 예산 실적 리포트""",
["부서장에게 주간 소진율 현황 자동 발송으로 셀프 관리 유도", "초과 위험 부서에 잔여 예산 사용 가이드라인 자동 첨부", "연간 예산 재배분 검토를 위한 분기별 예산 효율성 분석"])

case(doc, "059", "세금 신고 준비 자동화",
"세금 신고 시즌에 필요한 데이터를 자동으로 수집하고 정리하여 세무 담당자에게 전달합니다.",
"Schedule Trigger, QuickBooks API, Google Sheets, Gmail",
"""Schedule Trigger (분기 마지막 월 1일)
  │
  ▼
QuickBooks API (분기 세무 데이터 추출)
  │ - 매출 (세금계산서 발행분)
  │ - 매입 (공제 가능 항목)
  │ - 직원 급여 데이터
  │ - 기타 공제 항목
  ▼
Code 노드 (부가세 신고 데이터 정리)
  │ - 매출세액, 매입세액 집계
  │ - 납부세액 계산
  ▼
Google Sheets (세무 신고 워크시트 자동 작성)
  │
  ▼
Gmail (세무사에게 자료 전달)
  │ - Google Sheets 링크
  │ - 신고 마감일 명시
  ▼
Schedule (마감 D-7, D-3, D-1 리마인더)""",
["세무 데이터는 반드시 암호화된 채널로 전송", "자동 계산 결과는 최종 신고 전 세무사 검토 필수", "전기 대비 세액 변동이 큰 경우 원인 분석 메모 자동 추가"])

case(doc, "060", "결제 이상 거래 감지",
"결제 데이터를 실시간 모니터링하여 이상 패턴(금액 이상, 반복 소액 결제 등)을 감지하고 담당자에게 알립니다.",
"Webhook (결제 이벤트), Code, Slack, Gmail",
"""결제 시스템 Webhook (결제 이벤트)
  │
  ▼
Code 노드 (이상 거래 패턴 분석)
  │ 감지 규칙:
  │ - 동일 카드 1시간 내 5건 이상
  │ - 일반 패턴 대비 금액 300% 이상
  │ - 블랙리스트 IP/카드
  │ - 야간/휴일 대규모 거래
  ▼
[IF] 이상 패턴 감지?
  │ YES
  ▼
[이상 심각도에 따라]
  ├─ 경고: Slack (재무팀 알림)
  ├─ 의심: 결제 일시 보류 + Slack
  └─ 확실: 결제 즉시 차단 + Slack + Gmail (고객)
  ▼
Google Sheets (이상 거래 로그 기록)
  ▼
[조사 완료 후] 결제 재처리 or 최종 거부""",
["오탐률(False Positive) 관리가 매우 중요: 정상 거래 차단 최소화", "이상 거래 패턴 데이터 축적으로 ML 모델 고도화 가능", "금융감독원 가이드라인에 따른 이상 거래 탐지 기준 준수"])

case(doc, "061", "전자세금계산서 자동 발행",
"매출 발생 시 자동으로 국세청 전자세금계산서 시스템에 발행을 요청하고 거래처에 이메일로 전송합니다.",
"Webhook (매출 이벤트) / ERP, HTTP Request (국세청 API), Gmail",
"""ERP Webhook (매출 확정 이벤트)
  │
  ▼
Code 노드 (세금계산서 데이터 구성)
  │ - 공급자/공급받는자 정보
  │ - 품목, 금액, 세액 계산
  │ - 거래 날짜
  ▼
HTTP Request (국세청 e세로 API)
  │ - 전자세금계산서 발행 요청
  │ - 국세청 승인번호 수신
  ▼
Google Sheets (발행 이력 기록)
  │
  ▼
Gmail (거래처에 세금계산서 발송)
  │ - PDF 첨부 또는 국세청 링크
  ▼
회계 시스템 (매출 전표 자동 생성)
  ▼
[완료]""",
["국세청 API 연동은 공인 전자세금계산서 서비스 업체 통해 구현 권장", "취소/수정 세금계산서도 동일 플로우로 자동화 가능", "발행 실패 시 즉시 재무팀 알림 + 수동 처리 안내"])

case(doc, "062", "급여 이체 자동화",
"매월 급여 지급일에 급여 데이터를 기반으로 은행 이체 파일을 생성하고 담당자 승인 후 자동 처리합니다.",
"Schedule Trigger, Google Sheets, Code, Gmail, Slack",
"""Schedule Trigger (매월 25일 오전 9시)
  │
  ▼
Google Sheets (당월 급여 확정 데이터 읽기)
  │
  ▼
Code 노드 (급여 이체 파일 생성)
  │ - 은행별 이체 파일 포맷 (KEB, KB 등)
  │ - 총 이체 금액 계산
  │ - 이체 명세서 생성
  ▼
Gmail (CFO + 인사팀장에게 승인 요청)
  │ - 이체 파일 + 명세서 첨부
  │ - 승인 링크 포함
  ▼
[Wait] 승인
  │
  ▼
[IF] CFO + 인사팀장 모두 승인?
  │ YES
  ▼
은행 API / 인터넷뱅킹 (이체 파일 업로드)
  ▼
Gmail (직원 전체: 급여 이체 완료 안내)""",
["급여 이체는 반드시 2인 이상 승인 (4-eyes principle) 구현", "이체 후 개별 급여명세서 자동 발송 워크플로우와 연계", "이체 오류 발생 시 즉시 감지 및 수동 처리 안내"])

case(doc, "063", "재무 대시보드 자동 업데이트",
"다양한 재무 데이터 소스를 취합하여 경영진 재무 대시보드를 매일 자동으로 갱신합니다.",
"Schedule Trigger, QuickBooks/ERP API, Google Sheets, Data Studio",
"""Schedule Trigger (매일 오전 7시)
  │
  ▼
[병렬 데이터 수집]
  ├─ QuickBooks (매출, 비용, 손익)
  ├─ 은행 API (현금 잔고, 입출금)
  ├─ CRM (미수금, 계약 파이프라인)
  └─ 급여 시스템 (인건비 현황)
  ▼
Code 노드 (KPI 계산)
  │ - 일일 매출, 월 누적
  │ - Runway (현금 소진 예상일)
  │ - 비용 효율성 지표
  ▼
Google Sheets (대시보드 데이터 시트 업데이트)
  │ (Looker Studio가 이 시트 참조)
  ▼
Slack (#경영진채널)
  └─ "📊 오늘 재무 현황: 매출 {{amount}}"
  ▼
[완료]""",
["데이터 갱신 실패 시 전날 데이터 유지 + 갱신 실패 알림 발송", "경영진 모바일에서 쉽게 볼 수 있도록 Looker Studio 모바일 최적화", "주간/월간 트렌드 차트 자동 생성으로 시각적 인사이트 제공"])

# ─── CH7: DevOps/IT 자동화 (064~074) ─────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 7: DevOps/IT 자동화", 1)
par(doc, "개발 파이프라인, 인시던트 대응, 인프라 관리를 자동화하여 DevOps 팀의 효율을 극대화합니다.")

case(doc, "064", "GitHub PR 자동 리뷰 요청",
"GitHub에서 PR이 생성되면 자동으로 리뷰어를 배정하고 Slack에 알림을 발송하며 CI 상태를 추적합니다.",
"GitHub Trigger, Code, Slack, GitHub API",
"""GitHub Trigger (PR 생성됨)
  │
  ▼
Code 노드 (리뷰어 자동 선택)
  │ - 변경 파일의 CODEOWNERS 확인
  │ - 최근 비활성 팀원 제외
  │ - 라운드로빈 방식으로 균등 배분
  ▼
GitHub API (리뷰어 배정)
  │
  ▼
Slack (리뷰어 DM + 팀 채널)
  │ "🔍 PR 리뷰 요청: {{title}}
  │  파일 {{count}}개, +{{additions}}/-{{deletions}}"
  ▼
Schedule (D+1: 미리뷰 PR 리마인더)
  │
  ▼
GitHub Trigger (PR 병합됨)
  │
  ▼
Slack (팀 채널 배포 알림)
  └─ "✅ PR 병합: {{title}} → {{branch}}"
  ▼
[완료]""",
["CODEOWNERS 파일 관리로 도메인별 전문가 자동 리뷰 보장", "WIP(Work In Progress) PR은 리뷰 요청에서 제외", "리뷰 사이클 타임 추적으로 병목 파악 및 개선"])

case(doc, "065", "서버 모니터링 → 자동 대응",
"서버 CPU, 메모리, 디스크 사용률을 모니터링하고 임계값 초과 시 자동 스케일링이나 프로세스 재시작을 처리합니다.",
"HTTP Request (모니터링 Webhook), Code, Slack, SSH",
"""Datadog/CloudWatch Webhook (알림)
  │
  ▼
Code 노드 (알림 파싱 + 심각도 분류)
  │ P1: CPU > 95% or 서비스 다운
  │ P2: CPU > 80% or 메모리 > 85%
  │ P3: 디스크 > 70%
  ▼
Switch (심각도)
  ├─ P1 → [즉시 대응]
  │         Slack + PagerDuty 호출
  │         AWS Auto Scaling 트리거
  │         실패 시 인스턴스 교체
  │
  ├─ P2 → [자동 완화]
  │         Slack 알림
  │         프로세스 재시작 시도
  │         10분 후 재확인
  │
  └─ P3 → [예방적 알림]
            Slack (운영팀)
            로그 정리 스크립트 실행
  ▼
인시던트 로그 기록""",
["자동 대응 전 항상 Slack 알림 먼저 발송하여 팀 인지 보장", "자동 스케일링 비용 급증 방지를 위해 최대 인스턴스 수 제한", "야간/주말 자동 대응 범위 확대로 엔지니어 수동 개입 최소화"])

case(doc, "066", "배포 파이프라인 자동화",
"Git 태그 생성 시 자동으로 빌드, 테스트, 스테이징 배포, 승인 후 프로덕션 배포까지 진행합니다.",
"GitHub Trigger, Jenkins/GitHub Actions Webhook, Slack",
"""GitHub Trigger (태그 생성: v*)
  │
  ▼
GitHub Actions (자동 빌드 + 테스트)
  │
  ▼
Webhook (빌드 완료 이벤트)
  │ 성공              │ 실패
  ▼                   ▼
스테이징 자동 배포   Slack (빌드 실패 알림)
  ↓                   + GitHub에 코멘트
스모크 테스트 실행
  ↓
Slack (QA팀 + 릴리즈 매니저)
"✅ 스테이징 배포 완료. 프로덕션 배포 승인 요청"
  ↓
[Wait] 승인 클릭
  ↓
프로덕션 배포 실행
  ↓
Slack (전체 팀 배포 완료 알림)
"🚀 v{{version}} 프로덕션 배포 완료"
  ↓
모니터링 강화 (30분)""",
["배포 전 DB 마이그레이션 성공 여부 체크 단계 추가 권장", "롤백 워크플로우도 동일하게 자동화하여 빠른 복구 가능", "배포 시간대 제한 (주말/공휴일 배포 금지) 설정"])

case(doc, "067", "보안 취약점 자동 스캔",
"코드 저장소에서 주기적으로 보안 취약점을 스캔하고 발견된 취약점을 Jira에 자동 등록합니다.",
"Schedule Trigger, GitHub API, HTTP Request (Snyk API), Jira, Slack",
"""Schedule Trigger (매일 오전 2시)
  │
  ▼
GitHub API (활성 저장소 목록 조회)
  │
  ▼
[Loop] 각 저장소 처리
  │
  ▼
HTTP Request (Snyk API)
  │ - 의존성 취약점 스캔
  │ - 코드 취약점 분석
  ▼
Code 노드 (심각도별 분류)
  │ Critical, High, Medium, Low
  ▼
Jira (Critical/High: 즉시 이슈 생성)
  │ - 취약점 설명, CVE ID
  │ - 영향 범위, 수정 방법
  ▼
Slack (#보안채널)
  └─ "🔒 보안 취약점 발견: {{count}}건
     Critical: {{crit}}, High: {{high}}"
  ▼
Email (보안팀 주간 요약 리포트)""",
["Critical 취약점은 발견 즉시 개발팀 DM으로 긴급 알림", "취약점 수정 기한 SLA 설정: Critical 24시간, High 7일", "False Positive 제외 목록을 관리하여 노이즈 감소"])

case(doc, "068", "개발 환경 자동 프로비저닝",
"신규 개발자 합류 시 개발 환경 설정, 저장소 접근 권한, 개발 도구 설치를 자동화합니다.",
"Webhook (HR 신규 개발자 등록), GitHub API, Slack, Jira",
"""HR 시스템 Webhook (신규 개발자 등록)
  │
  ▼
[병렬 처리]
  ├─ GitHub API (팀/저장소 권한 부여)
  ├─ Jira API (프로젝트 접근 권한)
  ├─ Slack (개발팀 초대 + 환영 메시지)
  ├─ AWS IAM (개발 환경 계정 생성)
  └─ 이메일 (개발 환경 설정 가이드 발송)
       - Docker 설치 가이드
       - 로컬 환경 설정 스크립트
       - 주요 내부 도구 링크
  ▼
Jira (온보딩 태스크 생성)
  └─ "개발 환경 설정 완료" 체크리스트
  ▼
Slack (팀장에게 완료 알림)
  └─ "{{이름}} 개발 환경 설정 완료"
  ▼
[완료]""",
["권한은 최소 권한 원칙으로 필요한 저장소/프로젝트만 부여", "퇴사 시 모든 접근 권한 자동 회수 워크플로우와 연계", "개발 환경 설정 스크립트를 저장소에 관리하여 최신 유지"])

case(doc, "069", "인시던트 대응 자동화",
"서비스 장애 발생 시 자동으로 인시던트 채널을 생성하고 관련자를 소집하며 포스트모템 문서를 준비합니다.",
"PagerDuty Webhook, Slack API, Confluence",
"""PagerDuty Webhook (인시던트 발생)
  │
  ▼
Code 노드 (인시던트 심각도 파싱)
  │
  ▼
Slack API (인시던트 전용 채널 자동 생성)
  │ 채널명: #inc-{{YYYYMMDD}}-{{service}}
  │
  ▼
[관련자 자동 소집]
  ├─ 서비스 오너
  ├─ 온콜 엔지니어
  ├─ 인시던트 매니저
  └─ (P1) 부서장 + CTO
  ▼
Confluence (인시던트 문서 자동 생성)
  │ - 템플릿 기반
  │ - 발생 시각, 영향 범위 자동 기입
  ▼
Schedule (15분마다: Slack 상태 업데이트)
  ▼
[인시던트 해결 Webhook]
  │
  ▼
Slack (해결 선언 + 요약)
  ↓
Confluence (포스트모템 체크리스트 추가)
  ↓
Schedule (D+3: 포스트모템 회의 리마인더)""",
["인시던트 채널을 별도 생성하여 커뮤니케이션 집중화", "타임라인 자동 기록으로 포스트모템 작성 시간 80% 절감", "인시던트 유형별 런북(Runbook) 링크 자동 첨부"])

case(doc, "070", "SSL 인증서 만료 모니터링",
"보유한 도메인의 SSL 인증서 만료일을 주기적으로 확인하여 만료 30일 전부터 단계적 알림을 발송합니다.",
"Schedule Trigger, HTTP Request (SSL 체크), Slack, Email",
"""Schedule Trigger (매일 오전 6시)
  │
  ▼
Google Sheets (모니터링 도메인 목록 읽기)
  │
  ▼
[Loop] 각 도메인 처리
  │
  ▼
HTTP Request (SSL 인증서 만료일 조회)
  │ openssl s_client 또는 SSL Labs API
  ▼
Code 노드 (만료까지 남은 일수 계산)
  │
  ▼
Switch (남은 일수)
  ├─ D-30~21: Slack (인프라팀 안내)
  ├─ D-20~8:  Slack + Email (경고)
  ├─ D-7~2:   Slack + Email (긴급)
  └─ D-1~만료: PagerDuty (긴급 호출)
  ▼
Google Sheets (만료일 이력 기록)""",
["Let's Encrypt 자동 갱신 설정 + 갱신 확인 워크플로우 연계", "만료된 SSL은 서비스 장애보다 고객 신뢰 손상이 더 큰 문제", "도메인 목록을 Google Sheets로 관리하여 신규 추가 용이"])

case(doc, "071", "코드 배포 후 성능 자동 검증",
"프로덕션 배포 후 자동으로 성능 지표를 측정하고 이상 감지 시 즉시 롤백을 트리거합니다.",
"GitHub Actions Webhook, HTTP Request (APM), Code, Slack, GitHub API",
"""배포 완료 Webhook
  │
  ▼
Wait (5분: 워밍업 대기)
  │
  ▼
[병렬 성능 측정]
  ├─ Datadog API (응답 시간, 에러율)
  ├─ HTTP Request (주요 엔드포인트 헬스체크)
  └─ Google Analytics (실시간 사용자 지표)
  ▼
Code 노드 (배포 전 기준값과 비교)
  │ - 응답 시간 > 기준 150%?
  │ - 에러율 > 기준 200%?
  ▼
[IF] 성능 이상 감지?
  │ YES                    │ NO
  ▼                        ▼
Slack (긴급 알림)          Slack (성능 정상 확인)
GitHub API (롤백 트리거)   모니터링 강화 해제
  ↓
배포 담당자 즉시 호출""",
["자동 롤백은 매우 신중하게 적용 — 첫 번째는 알림만 발송 권장", "성능 기준값은 최근 30일 데이터로 자동 업데이트", "배포 후 24시간 집중 모니터링 기간 설정"])

case(doc, "072", "데이터베이스 백업 자동화",
"주요 데이터베이스의 정기 백업을 자동으로 실행하고 백업 성공 여부를 확인하여 알림을 발송합니다.",
"Schedule Trigger, SSH, AWS S3, Slack",
"""Schedule Trigger
  ├─ 매일 오전 3시 (일간 백업)
  └─ 매주 일요일 (주간 전체 백업)
  │
  ▼
SSH (DB 서버 접속)
  │ - pg_dump / mysqldump 실행
  │ - 백업 파일 압축 + 암호화
  ▼
AWS S3 (백업 파일 업로드)
  │ - 폴더: /backups/{{db_name}}/{{date}}/
  │ - 일간: 7일 보관
  │ - 주간: 4주 보관
  │ - 월간: 12개월 보관
  ▼
Code 노드 (백업 파일 무결성 검증)
  │ - MD5 체크섬 확인
  │ - 파일 크기 이상 감지
  ▼
[IF] 백업 성공?
  │ YES                    │ NO
  ▼                        ▼
Slack (백업 완료 기록)   Slack + PagerDuty (긴급)
  ▼
오래된 백업 파일 자동 삭제 (보관 정책)""",
["백업 복구 테스트를 분기마다 자동으로 실행하여 복구 가능성 검증", "백업 파일 암호화 키는 별도 안전한 키 관리 서비스에 보관", "RTO/RPO 목표 기반으로 백업 주기와 보관 기간 설정"])

case(doc, "073", "클라우드 비용 최적화 자동화",
"AWS/GCP 클라우드 비용을 매일 분석하여 낭비 요소(미사용 리소스, 오버프로비저닝)를 감지하고 최적화합니다.",
"Schedule Trigger, AWS Cost Explorer API, Slack, Email",
"""Schedule Trigger (매일 오전 7시)
  │
  ▼
AWS Cost Explorer API (전일 비용 분석)
  │ - 서비스별, 팀별 비용
  │ - 예산 대비 현황
  ▼
AWS API (낭비 리소스 감지)
  │ - 미사용 EC2 인스턴스
  │ - 연결 없는 EBS 볼륨
  │ - 오래된 EIP (탄력적 IP)
  │ - 유휴 로드밸런서
  ▼
Code 노드 (절감 기회 계산)
  │ - 각 리소스 월간 비용 추정
  ▼
Slack (인프라팀 + 재무팀)
  └─ "💰 낭비 리소스 발견: {{count}}건
     예상 절감 가능: ${{amount}}/월"
  ▼
Jira (최적화 태스크 자동 생성)
  └─ 우선순위: 금액 기준 내림차순""",
["자동 삭제는 반드시 사람 승인 후 실행 (실수로 운영 리소스 삭제 방지)", "Reserved Instance 구매 추천 리포트 월간 자동 생성", "팀별 비용 할당 태그 적용으로 부서 비용 책임 명확화"])

case(doc, "074", "API 사용량 및 쿼터 모니터링",
"외부 API 사용량을 실시간 추적하여 쿼터 초과 위험 시 사전 알림을 발송하고 자동으로 대체 방안을 활성화합니다.",
"Schedule Trigger, HTTP Request (API 상태), Code, Slack",
"""Schedule Trigger (매시간)
  │
  ▼
[Loop] 모니터링 API 목록
  │ (OpenAI, Google Maps, Stripe 등)
  ▼
HTTP Request (각 API 사용량 조회)
  │
  ▼
Code 노드 (쿼터 소진율 계산)
  │ 일일/월간 한도 대비 현재 사용량
  ▼
Switch (소진율)
  ├─ > 90%: Slack (긴급) + 속도 제한 활성화
  ├─ > 75%: Slack (경고) + 캐싱 강화
  └─ 정상:  기록만
  ▼
[IF] 100% 도달 → 예비 API 키로 자동 전환
  │ (멀티 API 키 로테이션)
  ▼
Google Sheets (사용량 이력 기록)
  ▼
월간 API 비용 리포트 자동 생성""",
["중요 API는 최소 2개 키 준비하여 쿼터 초과 시 자동 전환", "API 비용 급증은 코드 버그(무한 호출)일 수 있으므로 로그 분석", "API 사용량 트렌드 분석으로 플랜 업그레이드 최적 시점 결정"])

# ─── CH8: 데이터/AI 자동화 (075~084) ─────────────────────────────────────────
doc.add_page_break()
sc(doc, "Chapter 8: 데이터 및 AI 자동화", 1)
par(doc, "데이터 파이프라인 자동화와 AI 기반 분석으로 데이터에서 빠르게 인사이트를 추출합니다.")

case(doc, "075", "ETL 파이프라인 자동화",
"여러 소스(DB, API, 파일)에서 데이터를 추출하여 변환하고 데이터 웨어하우스에 적재합니다.",
"Schedule Trigger, MySQL/PostgreSQL, HTTP Request, Code, BigQuery",
"""Schedule Trigger (매일 오전 1시)
  │
  ▼
[병렬 데이터 추출]
  ├─ MySQL (거래 데이터)
  ├─ PostgreSQL (사용자 데이터)
  ├─ HTTP Request (외부 API 데이터)
  └─ Google Sheets (수동 입력 데이터)
  ▼
Code 노드 (데이터 변환)
  │ - 스키마 통일
  │ - 데이터 정제 (null 처리, 타입 변환)
  │ - 중복 제거
  │ - 파생 컬럼 계산
  ▼
[데이터 품질 검증]
  │ - 행 수 검증 (전일 대비 ±30% 이상 변동?)
  │ - NULL 비율 확인
  │ - 참조 무결성 확인
  ▼
[IF] 품질 기준 통과?
  │ YES
  ▼
BigQuery (데이터 적재)
  ▼
Slack (파이프라인 완료 알림)""",
["데이터 품질 게이트를 통과하지 못하면 적재 중단 + 알림", "증분 로드(Incremental Load) 구현으로 처리 시간과 비용 절감", "파이프라인 실행 메타데이터(시작/종료/행수) 별도 테이블에 기록"])

case(doc, "076", "AI 기반 데이터 이상값 감지",
"비즈니스 핵심 지표를 매일 모니터링하여 통계적 이상값이 감지되면 즉시 담당자에게 알립니다.",
"Schedule Trigger, BigQuery / DB, OpenAI, Slack",
"""Schedule Trigger (매일 오전 8시)
  │
  ▼
BigQuery (어제 핵심 지표 조회)
  │ - 일간 매출, DAU, 전환율
  │ - 오류율, 응답 시간
  ▼
Code 노드 (통계적 이상값 감지)
  │ - Z-score 분석 (|Z| > 2.5)
  │ - 전주 동요일 대비 ±20% 이상
  │ - 연속 3일 하락 트렌드
  ▼
[IF] 이상값 감지?
  │ YES
  ▼
OpenAI (이상값 원인 가설 생성)
  │ "매출이 23% 급락한 원인은?
  │  (컨텍스트: 최근 이벤트, 시즌, 배포)"
  ▼
Slack (#데이터채널)
  └─ "📉 이상 감지: {{metric}} {{변화율}}
     가능한 원인: {{hypothesis}}"
  ▼
Jira (조사 태스크 자동 생성)""",
["이상값 감지 후 false alarm이 많으면 임계값 조정 필요", "원인 가설은 AI 제안이므로 반드시 사람이 검증", "이상값 감지 이력 누적으로 계절성 패턴 학습 가능"])

case(doc, "077", "자동 리포트 생성 및 배포",
"비즈니스 데이터를 분석하여 주간/월간 경영 리포트를 자동 생성하고 관계자에게 배포합니다.",
"Schedule Trigger, BigQuery, Code, Google Slides, Gmail",
"""Schedule Trigger (매주 월요일 오전 8시)
  │
  ▼
BigQuery (지난 주 핵심 지표 조회)
  │
  ▼
Code 노드 (지표 계산 + 인사이트 생성)
  │ - KPI 달성율
  │ - 전주 대비 변화
  │ - 트렌드 분석
  ▼
OpenAI (주요 발견사항 요약)
  │ "3개 핵심 인사이트와 1개 액션 아이템"
  ▼
Google Slides API (리포트 슬라이드 업데이트)
  │ - 수치 자동 업데이트
  │ - 차트 데이터 갱신
  ▼
Google Drive (PDF 변환)
  │
  ▼
Gmail (경영진 + 팀장 배포)
  └─ 요약 텍스트 + PDF 첨부
  ▼
Slack (#경영진채널: 슬라이드 링크)""",
["Google Slides 템플릿을 잘 설계하면 자동화 품질이 크게 향상", "리포트 수신자 목록을 Google Sheets로 관리하여 변경 용이", "리포트 열람률 추적으로 실제 활용도 측정"])

case(doc, "078", "자연어 질의 → SQL 자동 변환",
"비기술직 사용자가 Slack에 자연어로 데이터를 요청하면 AI가 SQL로 변환하여 쿼리 결과를 반환합니다.",
"Slack Trigger, OpenAI, BigQuery / PostgreSQL, Slack",
"""Slack Trigger (@데이터봇 멘션)
  │ "지난 주 상품별 매출 순위 알려줘"
  ▼
OpenAI (자연어 → SQL 변환)
  │ System: "다음 스키마를 참고하여 SQL을 작성하세요:
  │          {{db_schema}}"
  │ SQL 생성 + 설명 포함
  ▼
Code 노드 (SQL 안전성 검증)
  │ - SELECT만 허용
  │ - 민감 테이블 접근 차단
  │ - 실행 시간 제한 설정
  ▼
BigQuery (쿼리 실행)
  │
  ▼
Code 노드 (결과 포매팅)
  │ - 표 형식 또는 차트
  │ - 최대 20행 표시
  ▼
Slack (결과 + 사용된 SQL 코드 공유)
  └─ "결과: {{table}}
     사용된 쿼리: {{sql}}"
  ▼
[완료]""",
["DB 스키마 정보는 민감도에 따라 공개 범위 제한", "실행 전 SQL을 사용자에게 먼저 보여주고 확인 요청 가능", "쿼리 실행 비용이 큰 경우 사전 비용 알림 기능 추가"])

case(doc, "079", "데이터 품질 자동 검증",
"데이터 파이프라인 각 단계에서 품질 규칙을 자동으로 검증하고 위반 사항을 데이터 엔지니어에게 보고합니다.",
"Schedule Trigger, BigQuery, Code, Slack, Jira",
"""Schedule Trigger (데이터 적재 후 자동 트리거)
  │
  ▼
BigQuery (데이터 품질 규칙 검증 쿼리 실행)
  │ 검증 항목:
  │ - NULL 비율 > 임계값?
  │ - 중복 키 존재?
  │ - 참조 무결성 위반?
  │ - 값 범위 이상 (음수 금액 등)?
  │ - 행 수 급변 감지?
  ▼
Code 노드 (실패 항목 집계)
  │
  ▼
[IF] 품질 규칙 위반?
  │ 심각                   │ 경미
  ▼                        ▼
Slack (긴급 + 데이터 파이프라인 중단) Slack (경고 알림)
Jira (P1 이슈 생성)                   Jira (P3 이슈 생성)
  ▼
데이터 품질 대시보드 업데이트""",
["데이터 품질 규칙은 비즈니스 팀과 협의하여 정의", "품질 위반 이력 분석으로 근본 원인 제거에 집중", "다운스트림 영향 분석: 어떤 리포트/대시보드가 영향받는지 자동 파악"])

case(doc, "080", "실시간 사용자 행동 분석",
"웹사이트 사용자 행동 이벤트를 실시간으로 처리하여 이상 패턴(이탈 급증, 오류 급증)을 즉시 감지합니다.",
"Webhook (이벤트 스트림), Code, BigQuery, Slack",
"""Webhook (사용자 행동 이벤트)
  │ 이벤트: page_view, click, error, purchase
  ▼
Code 노드 (이벤트 집계 — 5분 단위)
  │ - 방문자 수, 페이지뷰
  │ - 에러 발생률
  │ - 이탈률
  │ - 전환율
  ▼
[IF] 이상 패턴 감지?
  │ - 에러율 > 5%
  │ - 이탈율 급증 (> 전일 150%)
  │ - 전환율 급락 (< 전일 50%)
  ▼
Slack (#데이터채널)
  └─ "⚠️ 이상 패턴: 에러율 {{rate}}%
     영향 추정 사용자: {{count}}명"
  ▼
BigQuery (이벤트 적재 — 분석용)
  ▼
실시간 대시보드 업데이트""",
["이벤트 볼륨이 클 경우 Kafka/PubSub와 n8n 연동 고려", "개인정보 포함 이벤트는 수집 전 마스킹 처리 필수", "이상 감지 → 배포 이력 자동 연계로 코드 원인 빠른 파악"])

case(doc, "081", "ML 모델 성능 모니터링",
"운영 중인 머신러닝 모델의 예측 성능을 자동으로 추적하여 드리프트(성능 저하)를 감지합니다.",
"Schedule Trigger, HTTP Request (ML Platform API), Code, Slack",
"""Schedule Trigger (매일 오전 6시)
  │
  ▼
ML Platform API (어제 예측 성능 지표 조회)
  │ - 정확도, 정밀도, 재현율
  │ - 예측값 분포 (Output Drift)
  │ - 입력 데이터 분포 (Data Drift)
  ▼
Code 노드 (기준선 대비 드리프트 감지)
  │ - PSI (Population Stability Index)
  │ - KS 검정 통계량
  ▼
[IF] 드리프트 감지?
  │ YES
  ▼
Slack (ML 엔지니어 알림)
  └─ "📉 모델 드리프트 감지: {{model_name}}
     PSI: {{psi_score}} (임계값: 0.2)"
  ▼
Jira (모델 재학습 태스크 생성)
  ▼
[심각한 경우] 모델 자동 롤백 (이전 버전으로)""",
["드리프트 감지 임계값은 모델별 특성에 맞게 설정", "재학습 파이프라인 자동 트리거와 연계하여 완전 자동화 가능", "A/B 테스트 프레임워크와 연계하여 신규 모델 안전하게 배포"])

case(doc, "082", "데이터 카탈로그 자동 업데이트",
"데이터베이스 스키마 변경을 감지하여 자동으로 데이터 카탈로그와 문서를 업데이트합니다.",
"Schedule Trigger, DB 메타데이터 API, OpenAI, Confluence",
"""Schedule Trigger (매일 오전 4시)
  │
  ▼
DB API (현재 스키마 메타데이터 조회)
  │ - 테이블 목록, 컬럼, 타입, 설명
  ▼
Code 노드 (전일 대비 스키마 변경 감지)
  │ - 신규 테이블/컬럼
  │ - 삭제된 테이블/컬럼
  │ - 타입 변경
  ▼
[IF] 스키마 변경 감지?
  │ YES
  ▼
OpenAI (변경된 컬럼/테이블 설명 자동 생성)
  │ "이 테이블의 목적을 한국어로 설명하세요"
  ▼
Confluence (데이터 카탈로그 자동 업데이트)
  │
  ▼
Slack (#데이터채널)
  └─ "📋 DB 스키마 변경 감지: {{count}}건
     카탈로그 업데이트 완료"
  ▼
관련 파이프라인 담당자에게 변경 알림""",
["스키마 변경이 하위 파이프라인에 미치는 영향 자동 분석 추가", "자동 생성 설명은 데이터 오너 검토 후 최종 확정 권장", "컬럼 삭제는 즉시 알림 + 1개월 후 실제 삭제 정책 적용"])

case(doc, "083", "고객 세그먼트 자동 분류",
"고객 행동 데이터를 분석하여 자동으로 마케팅 세그먼트를 분류하고 CRM을 업데이트합니다.",
"Schedule Trigger, BigQuery, Code, HubSpot",
"""Schedule Trigger (매주 월요일 오전 3시)
  │
  ▼
BigQuery (지난 30일 고객 행동 데이터 조회)
  │ - 구매 금액, 빈도, 최근성 (RFM)
  │ - 기능 사용 패턴
  │ - 지원 문의 빈도
  ▼
Code 노드 (RFM 세그먼트 계산)
  │ Champions: R높음, F높음, M높음
  │ Loyal: F높음, M보통
  │ At Risk: R낮음, F높았던
  │ Hibernating: R낮음, F낮음
  │ Lost: 최근 6개월 미활성
  ▼
[Loop] 각 고객 처리
  │
  ▼
HubSpot (세그먼트 태그 업데이트)
  │
  ▼
[IF] 세그먼트 변경?
  │ YES (예: Champions → At Risk)
  ▼
Slack (CS팀 알림: 이탈 위험 VIP)
  ▼
마케팅 캠페인 자동 트리거 (세그먼트별)""",
["세그먼트 정의는 비즈니스 특성에 맞게 커스터마이징 필요", "세그먼트 이동 추적으로 마케팅 효과 측정 가능", "하이퍼 개인화 캠페인은 세그먼트를 더 세분화하여 적용"])

case(doc, "084", "AI 기반 수요 예측 자동화",
"과거 판매 데이터와 외부 요인을 분석하여 수요를 예측하고 재고 담당자에게 발주 추천을 제공합니다.",
"Schedule Trigger, BigQuery, OpenAI, Google Sheets, Slack",
"""Schedule Trigger (매주 월요일 오전 7시)
  │
  ▼
BigQuery (과거 52주 판매 데이터 조회)
  │
  ▼
HTTP Request (외부 데이터 수집)
  │ - 날씨 예보 (기상청 API)
  │ - 공휴일 캘린더
  │ - 트렌드 데이터
  ▼
OpenAI (수요 예측 분석)
  │ 입력: 과거 데이터 + 외부 요인
  │ 출력: 다음 4주 상품별 예상 판매량
  │       신뢰 구간 포함
  ▼
Code 노드 (발주 추천 계산)
  │ 발주량 = 예측수요 + 안전재고 - 현재재고
  ▼
Google Sheets (예측 결과 + 발주 추천 업데이트)
  │
  ▼
Slack (재고팀 알림)
  └─ "📦 다음 주 발주 추천: {{count}}개 품목
     긴급 발주 필요: {{urgent}}개"
  ▼
[IF] 긴급 발주 품목?
  │ YES
  ▼
발주 시스템 (자동 발주 초안 생성)
  └─ 담당자 최종 승인 후 실행""",
["예측 모델은 실제 결과와 주기적으로 비교하여 성능 보정", "계절성/프로모션 효과를 모델에 반영하여 예측 정확도 향상", "발주 자동화는 발주 금액 한도 설정 후 승인 프로세스 유지"])

doc.save(PATH)
print("[OK] gen_vol1_revised_p3.py 완료 - Ch5~Ch8 케이스043-084 추가")
print("[완료] n8n_cases_full.docx - 총 84케이스 ASCII 순서도 포함 완성")
