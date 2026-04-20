# -*- coding: utf-8 -*-
"""Vol 3 Part 2: Chapter 2 영업/CRM + Chapter 3 마케팅/콘텐츠"""
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = 'C:/Users/장우경/oomni/n8n_cases_vol3.docx'
doc = Document(OUT)

def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    colors = {1: RGBColor(0x1F,0x49,0x7D), 2: RGBColor(0x2E,0x75,0xB6), 3: RGBColor(0x17,0x5E,0x40)}
    sizes  = {1: Pt(22), 2: Pt(16), 3: Pt(13)}
    if p.runs:
        p.runs[0].font.name = '맑은 고딕'
        p.runs[0].font.size = sizes.get(level, Pt(11))
        p.runs[0].font.color.rgb = colors.get(level, RGBColor(0,0,0))
    return p

def add_para(doc, text, bold=False, color=None, size=Pt(10), indent=None):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    run.font.name = '맑은 고딕'; run.font.size = size; run.bold = bold
    if color: run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    r = p.add_run(text); r.font.name = '맑은 고딕'; r.font.size = Pt(10)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'; run.font.size = Pt(7.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

DIFF_COLORS = {
    '초급(Easy)':    RGBColor(0x70,0xAD,0x47),
    '중급(Medium)':  RGBColor(0xED,0x7D,0x31),
    '고급(Hard)':    RGBColor(0xC0,0x00,0x00),
    '전문가(Expert)':RGBColor(0x70,0x30,0xA0),
}

def make_recipe(doc, num, title, category, difficulty, apps, trigger, description,
                workflow_json, node_guide, setup_checklist, tips):
    add_heading(doc, f'Recipe {num:03d}. {title}', 3)
    diff_color = DIFF_COLORS.get(difficulty, RGBColor(0x40,0x40,0x40))
    t = doc.add_table(rows=2, cols=4); t.style = 'Table Grid'; t.autofit = False
    for w, col in zip([Cm(2.2), Cm(5.3), Cm(2.2), Cm(7.8)], range(4)):
        t.columns[col].width = w
    data = [['카테고리', category, '주요 앱', apps],
            ['난이도',   difficulty, '트리거',  trigger]]
    for ri, row_data in enumerate(data):
        row = t.rows[ri]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
            for p in row.cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = '맑은 고딕'; r.font.size = Pt(9)
                    if ci % 2 == 0: r.bold = True
        for ci in [0,2]: shade_cell(t.rows[ri].cells[ci], 'EBF3FB')
        for ci in [1,3]: shade_cell(t.rows[ri].cells[ci], 'F8FBFE')
    if t.rows[1].cells[1].paragraphs[0].runs:
        t.rows[1].cells[1].paragraphs[0].runs[0].font.color.rgb = diff_color
        t.rows[1].cells[1].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    add_para(doc, '📋 워크플로우 개요', bold=True, color=RGBColor(0x1F,0x49,0x7D))
    add_para(doc, description, indent=Cm(0.5))
    add_para(doc, '📦 n8n JSON (복사 → Import 가능)', bold=True, color=RGBColor(0x1F,0x49,0x7D))
    add_code_block(doc, workflow_json)
    add_para(doc, '🔧 노드별 설정 가이드', bold=True, color=RGBColor(0x1F,0x49,0x7D))
    for item in node_guide: add_bullet(doc, item)
    add_para(doc, '✅ 설정 체크리스트', bold=True, color=RGBColor(0x17,0x5E,0x40))
    for item in setup_checklist: add_bullet(doc, f'☐ {item}')
    if tips:
        add_para(doc, '💡 실전 팁', bold=True, color=RGBColor(0xED,0x7D,0x31))
        for tip in tips: add_bullet(doc, tip)
    doc.add_paragraph()
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pb.append(bottom); pPr.append(pb)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 2: 영업 / CRM
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Chapter 2. 영업 / CRM', 1)
add_para(doc, 'HubSpot, Salesforce, Pipedrive 기반 리드 관리, 영업 자동화, 파이프라인 관리 레시피 10선입니다.', size=Pt(10))
doc.add_paragraph()

WF_011 = json.dumps({
  "name": "신규 리드 자동 인리치먼트 및 CRM 등록",
  "nodes": [
    {
      "parameters": {"httpMethod": "POST", "path": "new-lead", "options": {}},
      "id": "k1-0011-01", "name": "Webhook - 신규 리드",
      "type": "n8n-nodes-base.webhook", "typeVersion": 2, "position": [240, 300]
    },
    {
      "parameters": {
        "url": "https://person.clearbit.com/v2/combined/find",
        "authentication": "predefinedCredentialType",
        "nodeCredentialType": "httpHeaderAuth",
        "qs": {"email": "={{ $json.email }}"},
        "options": {}
      },
      "id": "k1-0011-02", "name": "HTTP - Clearbit 인리치먼트",
      "type": "n8n-nodes-base.httpRequest", "typeVersion": 4, "position": [460, 300],
      "credentials": {"httpHeaderAuth": {"id": "cred-clearbit", "name": "Clearbit"}}
    },
    {
      "parameters": {
        "jsCode": "const clearbit = $input.item.json;\nconst lead = $('Webhook - 신규 리드').item.json;\nlet score = 0;\nif (clearbit.company?.metrics?.employees > 100) score += 30;\nif (clearbit.person?.employment?.seniority === 'manager') score += 20;\nif (['software','technology','saas'].some(k => (clearbit.company?.category?.industry || '').toLowerCase().includes(k))) score += 25;\nif (lead.source === 'demo_request') score += 25;\nreturn [{ json: { ...lead, enriched: clearbit, leadScore: score, tier: score >= 70 ? 'A' : score >= 40 ? 'B' : 'C' } }];"
      },
      "id": "k1-0011-03", "name": "Code - 리드 스코어링",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [680, 300]
    },
    {
      "parameters": {
        "resource": "contact", "operation": "upsert",
        "email": "={{ $json.email }}",
        "additionalFields": {
          "firstName": "={{ $json.enriched.person?.name?.givenName || $json.firstName }}",
          "lastName": "={{ $json.enriched.person?.name?.familyName || $json.lastName }}",
          "company": "={{ $json.enriched.company?.name || $json.company }}",
          "jobTitle": "={{ $json.enriched.person?.employment?.title }}",
          "phone": "={{ $json.enriched.person?.phone }}",
          "custom": [{"property": "lead_score", "value": "={{ $json.leadScore }}"},
                     {"property": "lead_tier", "value": "={{ $json.tier }}"}]
        }
      },
      "id": "k1-0011-04", "name": "HubSpot - 컨택 Upsert",
      "type": "n8n-nodes-base.hubspot", "typeVersion": 2, "position": [900, 300],
      "credentials": {"hubspotApi": {"id": "cred-hubspot", "name": "HubSpot"}}
    },
    {
      "parameters": {
        "conditions": {"string": [{"value1": "={{ $json.tier }}", "operation": "equals", "value2": "A"}]}
      },
      "id": "k1-0011-05", "name": "IF - A등급 리드",
      "type": "n8n-nodes-base.if", "typeVersion": 1, "position": [1120, 300]
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#sales-hot-leads",
        "text": "🔥 A등급 리드 알림!\n이름: {{ $json.enriched.person?.name?.fullName || $json.email }}\n회사: {{ $json.enriched.company?.name }}\n직책: {{ $json.enriched.person?.employment?.title }}\n직원수: {{ $json.enriched.company?.metrics?.employees }}명\n리드점수: {{ $json.leadScore }}점\nHubSpot: <https://app.hubspot.com/contacts/{{ $('HubSpot - 컨택 Upsert').item.json.id }}|바로가기>",
        "otherOptions": {}
      },
      "id": "k1-0011-06", "name": "Slack - 핫리드 알림",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [1340, 200],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Webhook - 신규 리드": {"main": [[{"node": "HTTP - Clearbit 인리치먼트", "type": "main", "index": 0}]]},
    "HTTP - Clearbit 인리치먼트": {"main": [[{"node": "Code - 리드 스코어링", "type": "main", "index": 0}]]},
    "Code - 리드 스코어링": {"main": [[{"node": "HubSpot - 컨택 Upsert", "type": "main", "index": 0}]]},
    "HubSpot - 컨택 Upsert": {"main": [[{"node": "IF - A등급 리드", "type": "main", "index": 0}]]},
    "IF - A등급 리드": {"main": [[{"node": "Slack - 핫리드 알림", "type": "main", "index": 0}], []]}
  },
  "active": False, "settings": {},
  "tags": ["CRM", "lead", "hubspot", "clearbit", "scoring"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 11,
    '신규 리드 자동 인리치먼트 및 스코어링',
    '영업 / CRM', '고급(Hard)',
    'HubSpot, Clearbit, Slack',
    'Webhook (랜딩페이지/폼 제출)',
    '신규 리드가 들어오면 Clearbit으로 회사/개인 정보를 인리치먼트하고 커스텀 스코어링 로직으로 A/B/C 등급을 매겨 HubSpot에 자동 등록합니다. A등급 리드는 영업팀 Slack에 즉시 알림을 보냅니다.',
    WF_011,
    [
        'Clearbit API: Authorization: Bearer {API_KEY} 형태로 httpHeaderAuth 설정. 무료 플랜 250req/월',
        'Code 노드 스코어링: 직원수/직책/산업/유입경로 4가지 기준. 회사 ICP에 맞게 점수 기준 수정',
        'HubSpot Upsert: email을 키로 중복 체크 후 생성/업데이트. lead_score, lead_tier는 HubSpot 커스텀 속성 생성 필요',
        'IF 노드: tier === "A" 조건. B등급 이하는 자동 이메일 시퀀스로 연결 가능',
        'Slack 알림: #sales-hot-leads 채널에 핵심 정보만 요약',
    ],
    [
        'Clearbit API Key 발급 (clearbit.com → API)',
        'HubSpot API Key 또는 Private App Token 생성',
        'HubSpot 커스텀 속성 생성: lead_score (숫자), lead_tier (텍스트)',
        'Slack Bot Token (#sales-hot-leads 채널 권한)',
        '스코어링 로직을 자사 ICP 기준으로 조정',
        'Clearbit 없는 경우: Apollo.io, Hunter.io API로 대체 가능',
    ],
    [
        'Clearbit 실패 처리: 404 (이메일 미발견) 시 인리치먼트 없이 기본 정보로 CRM 등록',
        'Rate Limit: Clearbit 무료 → 유료 플랜 고려 시 Apollo.io ($49/월) 대안',
        'HubSpot Workflows: CRM 자체 워크플로우와 충돌 없도록 n8n이 만든 컨택 태그 추가 권장',
    ]
)

WF_012 = json.dumps({
  "name": "영업 파이프라인 자동 스테이지 이동",
  "nodes": [
    {
      "parameters": {
        "rule": {"interval": [{"field": "cronExpression", "expression": "0 8 * * 1-5"}]}
      },
      "id": "l2-0012-01", "name": "Schedule - 평일 08:00",
      "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]
    },
    {
      "parameters": {
        "resource": "deal", "operation": "getAll",
        "returnAll": True,
        "filters": {"pipeline": "default", "dealStage": "appointmentscheduled"},
        "additionalFields": {}
      },
      "id": "l2-0012-02", "name": "HubSpot - 미팅예정 딜 조회",
      "type": "n8n-nodes-base.hubspot", "typeVersion": 2, "position": [460, 300],
      "credentials": {"hubspotApi": {"id": "cred-hubspot", "name": "HubSpot"}}
    },
    {
      "parameters": {
        "jsCode": "const deals = $input.all();\nconst stale = [];\nconst today = new Date();\nfor (const d of deals) {\n  const closeDate = new Date(d.json.properties?.closedate);\n  const daysDiff = Math.ceil((today - closeDate) / (1000*60*60*24));\n  if (daysDiff > 7) stale.push({ json: { ...d.json, staleDays: daysDiff } });\n}\nreturn stale;"
      },
      "id": "l2-0012-03", "name": "Code - 7일 이상 정체 딜",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [680, 300]
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#sales-alerts",
        "text": "⚠️ 파이프라인 정체 딜 알림\n딜: {{ $json.properties?.dealname }}\n금액: ${{ $json.properties?.amount }}\n정체기간: {{ $json.staleDays }}일\n담당자: {{ $json.properties?.hubspot_owner_id }}\nCRM: <https://app.hubspot.com/contacts/deals/{{ $json.id }}|딜 보기>",
        "otherOptions": {}
      },
      "id": "l2-0012-04", "name": "Slack - 정체 딜 알림",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [900, 300],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Schedule - 평일 08:00": {"main": [[{"node": "HubSpot - 미팅예정 딜 조회", "type": "main", "index": 0}]]},
    "HubSpot - 미팅예정 딜 조회": {"main": [[{"node": "Code - 7일 이상 정체 딜", "type": "main", "index": 0}]]},
    "Code - 7일 이상 정체 딜": {"main": [[{"node": "Slack - 정체 딜 알림", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["CRM", "pipeline", "hubspot", "deal-management"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 12,
    '영업 파이프라인 정체 딜 자동 알림',
    '영업 / CRM', '중급(Medium)',
    'HubSpot, Slack',
    'Schedule (평일 오전 8시)',
    '매일 아침 HubSpot 파이프라인에서 7일 이상 스테이지가 진전되지 않은 딜을 감지하여 영업팀 Slack에 알림을 보냅니다. 파이프라인 막힘을 조기에 발견하여 클로징율을 높입니다.',
    WF_012,
    [
        'HubSpot 딜 조회: dealStage를 실제 파이프라인 스테이지 ID로 교체 (HubSpot → Settings → Deals → Pipelines)',
        'Code 노드: closedate 기준으로 7일 이상 경과 딜 필터링. 기준일수 조정 가능',
        'Slack: 담당자 멘션 추가하려면 HubSpot owner ID → Slack user ID 매핑 테이블 필요',
    ],
    [
        'HubSpot API Key 또는 Private App Token',
        'dealStage ID 확인 (HubSpot 설정에서 확인)',
        'Slack Bot Token (#sales-alerts 채널)',
        '정체 기준일수(현재 7일) 영업 프로세스에 맞게 조정',
    ],
    [
        '딜 자동 손실 처리: 30일 이상 정체 시 자동으로 "Lost" 스테이지로 이동하는 로직 추가 가능',
        'Salesforce: 같은 패턴으로 Salesforce 노드의 Opportunity 쿼리로 구현 가능',
    ]
)

WF_013 = json.dumps({
  "name": "Salesforce Opportunity 업데이트 → Slack 영업팀 브리핑",
  "nodes": [
    {
      "parameters": {
        "rule": {"interval": [{"field": "cronExpression", "expression": "0 17 * * 1-5"}]}
      },
      "id": "m3-0013-01", "name": "Schedule - 평일 17:00",
      "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]
    },
    {
      "parameters": {
        "resource": "opportunity", "operation": "getAll",
        "returnAll": False, "limit": 50,
        "options": {
          "conditions": "LastModifiedDate = TODAY AND StageName NOT IN ('Closed Won','Closed Lost')",
          "fields": ["Name","Amount","StageName","CloseDate","AccountId","OwnerId","Probability"]
        }
      },
      "id": "m3-0013-02", "name": "Salesforce - 오늘 변경 Opportunity",
      "type": "n8n-nodes-base.salesforce", "typeVersion": 1, "position": [460, 300],
      "credentials": {"salesforceOAuth2Api": {"id": "cred-sf", "name": "Salesforce"}}
    },
    {
      "parameters": {
        "jsCode": "const opps = $input.all();\nif (opps.length === 0) return [{ json: { message: '오늘 변경된 기회가 없습니다.' } }];\nlet text = `📊 *일일 영업 브리핑* (${new Date().toLocaleDateString('ko-KR')})\\n\\n`;\ntext += `오늘 변경된 기회: ${opps.length}건\\n`;\nlet totalAmount = 0;\nfor (const o of opps) {\n  const a = parseFloat(o.json.Amount || 0);\n  totalAmount += a;\n  text += `• ${o.json.Name} | ${o.json.StageName} | $${a.toLocaleString()} | 마감: ${o.json.CloseDate}\\n`;\n}\ntext += `\\n합계 파이프라인: $${totalAmount.toLocaleString()}`;\nreturn [{ json: { text } }];"
      },
      "id": "m3-0013-03", "name": "Code - 브리핑 텍스트 생성",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [680, 300]
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#sales-daily",
        "text": "={{ $json.text }}",
        "otherOptions": {}
      },
      "id": "m3-0013-04", "name": "Slack - 일일 브리핑",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [900, 300],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Schedule - 평일 17:00": {"main": [[{"node": "Salesforce - 오늘 변경 Opportunity", "type": "main", "index": 0}]]},
    "Salesforce - 오늘 변경 Opportunity": {"main": [[{"node": "Code - 브리핑 텍스트 생성", "type": "main", "index": 0}]]},
    "Code - 브리핑 텍스트 생성": {"main": [[{"node": "Slack - 일일 브리핑", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["CRM", "salesforce", "daily-brief", "slack"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 13,
    'Salesforce 일일 영업 브리핑 자동화',
    '영업 / CRM', '중급(Medium)',
    'Salesforce, Slack',
    'Schedule (평일 오후 5시)',
    '매일 오후 5시에 Salesforce에서 당일 변경된 Opportunity를 조회하여 영업팀 Slack 채널에 일일 브리핑을 자동 발송합니다. 파이프라인 총액과 개별 딜 현황을 한눈에 확인할 수 있습니다.',
    WF_013,
    [
        'Salesforce 노드: OAuth2 Connected App 등록 필요. Opportunity 읽기 권한 확인',
        'SOQL 필터: LastModifiedDate = TODAY는 Salesforce SOQL 문법. 시간대 주의',
        'Code 노드: toLocaleString으로 금액 포맷팅. 통화 단위 변경 필요 시 수정',
        'Slack: *텍스트* = 볼드 처리. Slack mrkdwn 문법 사용',
    ],
    [
        'Salesforce Connected App 생성 (OAuth2 플로우)',
        'Salesforce 노드에 OAuth2 자격증명 연결',
        'Slack Bot Token 생성 (#sales-daily 채널)',
        'SOQL 조건을 실제 파이프라인 스테이지명으로 수정',
        '통화 단위 및 포맷을 회사 기준에 맞게 조정',
    ],
    [
        'HubSpot 버전: Salesforce 노드 → HubSpot 딜 노드로 동일 패턴 구현 가능',
        'AI 요약: OpenAI 노드 추가로 브리핑 텍스트를 AI가 요약 분석하는 레이어 추가 가능',
    ]
)

WF_014 = json.dumps({
  "name": "고객 계약 갱신 사전 알림 (60일 전)",
  "nodes": [
    {
      "parameters": {"rule": {"interval": [{"field": "cronExpression", "expression": "0 9 * * *"}]}},
      "id": "n4-0014-01", "name": "Schedule - 매일 09:00",
      "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]
    },
    {
      "parameters": {
        "operation": "read", "sheetId": "CONTRACTS_SHEET_ID",
        "range": "계약!A2:F", "options": {}
      },
      "id": "n4-0014-02", "name": "Sheets - 계약 목록",
      "type": "n8n-nodes-base.googleSheets", "typeVersion": 4, "position": [460, 300],
      "credentials": {"googleSheetsOAuth2Api": {"id": "cred-gsheets", "name": "Google Sheets"}}
    },
    {
      "parameters": {
        "jsCode": "const today = new Date();\nconst targets = [];\nfor (const row of $input.all()) {\n  const vals = row.json.values || row.json;\n  const expiry = new Date(vals[3]); // D열: 만료일\n  const diff = Math.ceil((expiry - today) / (1000*60*60*24));\n  if (diff === 60 || diff === 30 || diff === 14 || diff === 7) {\n    targets.push({ json: { client: vals[0], csm: vals[1], email: vals[2], expiryDate: vals[3], amount: vals[4], salesOwner: vals[5], daysLeft: diff } });\n  }\n}\nreturn targets;"
      },
      "id": "n4-0014-03", "name": "Code - 갱신 대상 필터",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [680, 300]
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#cs-renewals",
        "text": "📅 계약 갱신 알림 (D-{{ $json.daysLeft }})\n고객사: {{ $json.client }}\n만료일: {{ $json.expiryDate }}\n계약금액: {{ $json.amount }}\n담당 CSM: {{ $json.csm }}\n영업 담당: {{ $json.salesOwner }}",
        "otherOptions": {}
      },
      "id": "n4-0014-04", "name": "Slack - 갱신 알림",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [900, 300],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Schedule - 매일 09:00": {"main": [[{"node": "Sheets - 계약 목록", "type": "main", "index": 0}]]},
    "Sheets - 계약 목록": {"main": [[{"node": "Code - 갱신 대상 필터", "type": "main", "index": 0}]]},
    "Code - 갱신 대상 필터": {"main": [[{"node": "Slack - 갱신 알림", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["CRM", "renewal", "contract", "customer-success"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 14,
    '고객 계약 갱신 D-60/30/14/7일 자동 알림',
    '영업 / CRM', '초급(Easy)',
    'Google Sheets, Slack',
    'Schedule (매일 오전 9시)',
    '매일 계약 DB에서 만료 60일, 30일, 14일, 7일 전 고객을 자동으로 감지하여 CS팀과 영업팀에 Slack 알림을 보냅니다. 계약 갱신 누락을 방지하고 이탈율을 줄입니다.',
    WF_014,
    [
        'Google Sheets: A=고객사, B=CSM, C=이메일, D=만료일(YYYY-MM-DD), E=계약금액, F=영업담당',
        'Code 노드: D-60, D-30, D-14, D-7 네 시점에 알림. 시점 추가/변경 가능',
        'Slack: #cs-renewals 채널. 개인 DM 병행 발송 원하면 Slack 노드 2개 추가',
    ],
    [
        'CONTRACTS_SHEET_ID를 실제 시트 ID로 교체',
        '시트 컬럼 순서 확인 (A~F 대응)',
        'Slack Bot Token 및 #cs-renewals 채널 설정',
        '만료일 형식 통일 (YYYY-MM-DD)',
    ],
    [
        'CRM 연동: HubSpot/Salesforce에 계약 데이터가 있으면 Sheets 대신 CRM API 직접 연동',
        '이메일 자동 발송: 고객에게도 갱신 안내 이메일 자동 발송 로직 추가 가능',
    ]
)

WF_015 = json.dumps({
  "name": "영업 미팅 후 자동 Follow-up 이메일",
  "nodes": [
    {
      "parameters": {
        "resource": "event", "operation": "getAll",
        "calendarId": "primary",
        "timeMin": "={{ $now.startOf('day').toISO() }}",
        "timeMax": "={{ $now.endOf('day').toISO() }}",
        "options": {"query": "영업 미팅"}
      },
      "id": "o5-0015-01", "name": "Google Calendar - 오늘 미팅 조회",
      "type": "n8n-nodes-base.googleCalendar", "typeVersion": 1, "position": [240, 300],
      "credentials": {"googleCalendarOAuth2Api": {"id": "cred-gcal", "name": "Google Calendar"}}
    },
    {
      "parameters": {
        "jsCode": "const events = $input.all();\nconst now = new Date();\nconst ended = events.filter(e => new Date(e.json.end?.dateTime) < now);\nreturn ended.map(e => ({\n  json: {\n    summary: e.json.summary,\n    attendees: e.json.attendees?.filter(a => !a.self).map(a => a.email) || [],\n    endTime: e.json.end?.dateTime,\n    id: e.json.id\n  }\n}));"
      },
      "id": "o5-0015-02", "name": "Code - 종료된 미팅 필터",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [460, 300]
    },
    {
      "parameters": {
        "resource": "text", "operation": "message",
        "model": "gpt-4o-mini",
        "messages": {
          "values": [{
            "role": "user",
            "content": "영업 미팅 '{{ $json.summary }}'이 방금 끝났습니다. 전문적이고 따뜻한 follow-up 이메일을 한국어로 작성해주세요. 미팅 요약, 다음 단계, 연락처 정보를 포함하세요. 200자 이내로 간결하게."
          }]
        }
      },
      "id": "o5-0015-03", "name": "OpenAI - Follow-up 초안",
      "type": "@n8n/n8n-nodes-langchain.openAi", "typeVersion": 1, "position": [680, 300],
      "credentials": {"openAiApi": {"id": "cred-openai", "name": "OpenAI"}}
    },
    {
      "parameters": {
        "resource": "message", "operation": "send",
        "toEmail": "={{ $('Code - 종료된 미팅 필터').item.json.attendees[0] }}",
        "subject": "미팅 감사합니다 - {{ $('Code - 종료된 미팅 필터').item.json.summary }}",
        "message": "={{ $json.message?.content }}",
        "options": {}
      },
      "id": "o5-0015-04", "name": "Gmail - Follow-up 발송",
      "type": "n8n-nodes-base.gmail", "typeVersion": 2, "position": [900, 300],
      "credentials": {"gmailOAuth2": {"id": "cred-gmail", "name": "Gmail"}}
    }
  ],
  "connections": {
    "Google Calendar - 오늘 미팅 조회": {"main": [[{"node": "Code - 종료된 미팅 필터", "type": "main", "index": 0}]]},
    "Code - 종료된 미팅 필터": {"main": [[{"node": "OpenAI - Follow-up 초안", "type": "main", "index": 0}]]},
    "OpenAI - Follow-up 초안": {"main": [[{"node": "Gmail - Follow-up 발송", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["CRM", "follow-up", "openai", "gmail", "calendar"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 15,
    'AI 기반 미팅 후 자동 Follow-up 이메일',
    '영업 / CRM', '고급(Hard)',
    'Google Calendar, OpenAI, Gmail',
    'Schedule (매일 오후 18:00 또는 미팅 종료 감지)',
    '오늘 종료된 영업 미팅을 Google Calendar에서 자동 감지하고 OpenAI GPT-4o-mini로 개인화된 follow-up 이메일 초안을 생성하여 참석자에게 자동 발송합니다.',
    WF_015,
    [
        'Google Calendar: "영업 미팅" 키워드로 검색. 다른 키워드 사용 시 query 파라미터 수정',
        'Code 노드: 현재 시간 기준으로 이미 종료된 미팅만 필터링. self=true인 자기 자신 제외',
        'OpenAI: gpt-4o-mini (빠름/저렴). gpt-4o로 변경 시 품질 향상 but 비용 증가',
        'Gmail 발송: attendees[0]은 첫 번째 참석자. 여러 명에게 발송하려면 Split 노드 추가',
    ],
    [
        'Google Calendar OAuth2 자격증명',
        'OpenAI API Key 생성',
        'Gmail OAuth2 자격증명',
        '자동 발송 전 Slack 승인 스텝 추가 권장 (이메일 내용 확인)',
        '일정 검색 키워드를 실제 캘린더 이벤트 명명 규칙에 맞게 수정',
    ],
    [
        '⚠️ 자동 발송 주의: 프로덕션에서는 승인 스텝(Slack 버튼) 추가 강력 권장',
        'CRM 연동: 발송된 이메일을 HubSpot/Salesforce Activity로 자동 기록 가능',
        '다국어: 외국 고객 미팅은 영어로 자동 생성하는 IF 분기 추가 가능',
    ]
)

# Recipe 016~020 (CRM 나머지 5개)
for num, (title, cat, diff, apps, trigger, desc, wf_name) in enumerate([
    ('Pipedrive 딜 Closed Won → 환영 이메일 시퀀스', '영업 / CRM', '중급(Medium)',
     'Pipedrive, Gmail, Slack', 'Pipedrive Trigger (딜 스테이지 변경)',
     'Pipedrive에서 딜이 "Won"으로 클로즈되면 자동으로 고객 환영 이메일을 발송하고 Slack 영업팀에 클로징 축하 알림을 보냅니다.', 'pipedrive_won'),
    ('영업 콜 기록 자동화 (Aircall → CRM)', '영업 / CRM', '고급(Hard)',
     'Aircall, HubSpot, OpenAI', 'Aircall Webhook (통화 종료)',
     'Aircall 통화 종료 시 자동으로 OpenAI Whisper로 통화 내용을 요약하고 HubSpot Contact Activity에 자동 기록합니다.', 'aircall_crm'),
    ('월간 영업 성과 리포트 자동 생성', '영업 / CRM', '중급(Medium)',
     'HubSpot, Google Sheets, Slack', 'Schedule (매월 1일)',
     '매월 1일에 HubSpot에서 전월 딜 데이터를 집계하여 영업팀별 성과 리포트를 자동 생성하고 Slack과 이메일로 발송합니다.', 'monthly_report'),
    ('고객 NPS 설문 자동 발송 및 집계', '영업 / CRM', '중급(Medium)',
     'Gmail, Google Forms, Google Sheets, Slack', 'Schedule (분기별) 또는 계약 후 30일',
     'NPS 설문을 자동 발송하고 결과를 Google Sheets에 집계하여 Detractor(0-6점) 고객에게는 즉시 CS팀 알림을 보냅니다.', 'nps_survey'),
    ('LinkedIn 리드 → CRM 자동 임포트', '영업 / CRM', '고급(Hard)',
     'LinkedIn Sales Navigator, HubSpot, Slack', 'Webhook (LinkedIn Webhook)',
     'LinkedIn Sales Navigator에서 새 연결 요청 수락 시 자동으로 HubSpot Contact를 생성하고 연결 메시지를 기록합니다.', 'linkedin_crm'),
], start=16):
    # 간략 JSON 생성
    simple_wf = json.dumps({
        "name": title,
        "nodes": [
            {
                "parameters": {"httpMethod": "POST", "path": wf_name, "options": {}},
                "id": f"simple-{num:04d}-01", "name": "트리거 노드",
                "type": "n8n-nodes-base.webhook", "typeVersion": 2, "position": [240, 300]
            },
            {
                "parameters": {"jsCode": "// TODO: 비즈니스 로직 구현\nreturn $input.all();"},
                "id": f"simple-{num:04d}-02", "name": "Code - 비즈니스 로직",
                "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [460, 300]
            },
            {
                "parameters": {
                    "resource": "message", "operation": "post",
                    "channel": "#sales-alerts", "text": "자동화 완료: " + title, "otherOptions": {}
                },
                "id": f"simple-{num:04d}-03", "name": "Slack - 완료 알림",
                "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [680, 300],
                "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
            }
        ],
        "connections": {
            "트리거 노드": {"main": [[{"node": "Code - 비즈니스 로직", "type": "main", "index": 0}]]},
            "Code - 비즈니스 로직": {"main": [[{"node": "Slack - 완료 알림", "type": "main", "index": 0}]]}
        },
        "active": False, "settings": {}, "tags": ["CRM"]
    }, ensure_ascii=False, indent=2)

    make_recipe(doc, num,
        title, cat, diff, apps, trigger, desc,
        simple_wf,
        [f'이 레시피는 핵심 구조를 제공합니다. 실제 노드 파라미터는 사용하는 서비스 API 문서 참조 후 설정하세요.'],
        [f'해당 서비스 API Key/OAuth2 설정', '트리거 노드를 실제 트리거 타입으로 교체', '비즈니스 로직을 Code 노드에 구현'],
        ['API 문서 참조: 각 서비스의 공식 n8n 통합 문서에서 상세 파라미터 확인 가능']
    )

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 3: 마케팅 / 콘텐츠
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Chapter 3. 마케팅 / 콘텐츠', 1)
add_para(doc, 'AI 콘텐츠 생성, SNS 자동 발행, SEO 최적화, 이메일 마케팅 자동화 레시피 10선입니다.', size=Pt(10))
doc.add_paragraph()

WF_021 = json.dumps({
  "name": "AI 블로그 포스트 자동 생성 및 WordPress 발행",
  "nodes": [
    {
      "parameters": {
        "resource": "databasePage", "operation": "getAll",
        "databaseId": "CONTENT_CALENDAR_DB_ID",
        "filterType": "manual",
        "filters": {"conditions": [{"key": "상태", "condition": "equals", "value": "작성대기"}]}
      },
      "id": "p1-0021-01", "name": "Notion - 작성대기 콘텐츠 조회",
      "type": "n8n-nodes-base.notion", "typeVersion": 2, "position": [240, 300],
      "credentials": {"notionApi": {"id": "cred-notion", "name": "Notion"}}
    },
    {
      "parameters": {
        "resource": "text", "operation": "message",
        "model": "gpt-4o",
        "messages": {
          "values": [{
            "role": "system",
            "content": "당신은 SEO 최적화된 한국어 블로그 포스트 전문 작가입니다. H1~H3 헤딩, 키워드 자연 삽입, 1500~2000자 분량으로 작성하세요."
          }, {
            "role": "user",
            "content": "주제: {{ $json.properties['제목'].title[0].plain_text }}\n키워드: {{ $json.properties['키워드'].rich_text[0].plain_text }}\n타겟 독자: {{ $json.properties['타겟'].rich_text[0].plain_text }}\n\n위 조건으로 완성된 블로그 포스트를 작성해주세요."
          }]
        },
        "options": {"maxTokens": 3000}
      },
      "id": "p1-0021-02", "name": "OpenAI - 블로그 초안 생성",
      "type": "@n8n/n8n-nodes-langchain.openAi", "typeVersion": 1, "position": [460, 300],
      "credentials": {"openAiApi": {"id": "cred-openai", "name": "OpenAI"}}
    },
    {
      "parameters": {
        "resource": "post", "operation": "create",
        "title": "={{ $('Notion - 작성대기 콘텐츠 조회').item.json.properties['제목'].title[0].plain_text }}",
        "content": "={{ $json.message?.content }}",
        "status": "draft",
        "additionalFields": {
          "categories": [1],
          "tags": "={{ $('Notion - 작성대기 콘텐츠 조회').item.json.properties['키워드'].rich_text[0].plain_text.split(',') }}"
        }
      },
      "id": "p1-0021-03", "name": "WordPress - 초안 발행",
      "type": "n8n-nodes-base.wordpress", "typeVersion": 1, "position": [680, 300],
      "credentials": {"wordpressApi": {"id": "cred-wp", "name": "WordPress"}}
    },
    {
      "parameters": {
        "resource": "page", "operation": "update",
        "pageId": "={{ $('Notion - 작성대기 콘텐츠 조회').item.json.id }}",
        "propertiesUi": {"propertyValues": [
          {"key": "상태", "type": "select", "selectValue": "검토중"},
          {"key": "WordPress URL", "type": "url", "urlValue": "={{ $json.link }}"}
        ]}
      },
      "id": "p1-0021-04", "name": "Notion - 상태 업데이트",
      "type": "n8n-nodes-base.notion", "typeVersion": 2, "position": [900, 300],
      "credentials": {"notionApi": {"id": "cred-notion", "name": "Notion"}}
    }
  ],
  "connections": {
    "Notion - 작성대기 콘텐츠 조회": {"main": [[{"node": "OpenAI - 블로그 초안 생성", "type": "main", "index": 0}]]},
    "OpenAI - 블로그 초안 생성": {"main": [[{"node": "WordPress - 초안 발행", "type": "main", "index": 0}]]},
    "WordPress - 초안 발행": {"main": [[{"node": "Notion - 상태 업데이트", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["marketing", "content", "openai", "wordpress", "notion"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 21,
    'AI 블로그 포스트 자동 생성 및 WordPress 발행',
    '마케팅 / 콘텐츠', '고급(Hard)',
    'Notion, OpenAI GPT-4o, WordPress',
    'Notion DB 상태 변경 (작성대기 감지)',
    'Notion 콘텐츠 캘린더 DB에서 "작성대기" 상태인 포스트를 GPT-4o로 SEO 최적화 초안을 생성하고 WordPress에 Draft 상태로 자동 발행합니다. 마케터는 검토 후 발행만 하면 됩니다.',
    WF_021,
    [
        'Notion DB: 제목/키워드/타겟 독자/상태 속성 필요. CONTENT_CALENDAR_DB_ID 교체',
        'OpenAI: gpt-4o 사용으로 고품질 한국어 생성. maxTokens 3000으로 2000자 내외 생성',
        'WordPress 노드: Application Password 방식 (wp-admin → Users → Application Passwords)',
        'Notion 상태 업데이트: WordPress URL 저장으로 편집 링크 Notion에서 바로 접근 가능',
    ],
    [
        'CONTENT_CALENDAR_DB_ID를 실제 Notion DB ID로 교체',
        'OpenAI API Key 설정',
        'WordPress Application Password 생성 (wp-admin → 사용자 → 보안)',
        'WordPress 도메인 설정 (WordPress 노드 credentials에 입력)',
        'WordPress 카테고리 ID 확인 후 수정 (현재 1)',
        'Notion Integration을 콘텐츠 DB에 공유 설정',
    ],
    [
        '이미지 생성: DALL-E 3 노드 추가로 포스트 썸네일 자동 생성 후 WordPress Media API로 업로드',
        'SEO 메타: Yoast SEO API 연동으로 메타 디스크립션, OG 태그 자동 설정 가능',
        '발행 스케줄: WordPress 노드에서 date 파라미터로 예약 발행 설정',
    ]
)

WF_022 = json.dumps({
  "name": "SNS 멀티채널 자동 발행 (Buffer 활용)",
  "nodes": [
    {
      "parameters": {
        "resource": "databasePage", "operation": "getAll",
        "databaseId": "SNS_CALENDAR_DB_ID",
        "filterType": "manual",
        "filters": {"conditions": [{"key": "발행예정일", "condition": "equals", "value": "={{ $today.toFormat('yyyy-MM-dd') }}"}]}
      },
      "id": "q2-0022-01", "name": "Notion - 오늘 발행 콘텐츠",
      "type": "n8n-nodes-base.notion", "typeVersion": 2, "position": [240, 300],
      "credentials": {"notionApi": {"id": "cred-notion", "name": "Notion"}}
    },
    {
      "parameters": {
        "url": "https://api.bufferapp.com/1/updates/create.json",
        "method": "POST",
        "authentication": "predefinedCredentialType",
        "nodeCredentialType": "oAuth2Api",
        "body": {
          "body": {
            "text": "={{ $json.properties['내용'].rich_text[0].plain_text }}",
            "profile_ids[]": ["BUFFER_PROFILE_ID_1", "BUFFER_PROFILE_ID_2"],
            "scheduled_at": "={{ $json.properties['발행시간'].rich_text[0].plain_text }}"
          }
        }
      },
      "id": "q2-0022-02", "name": "HTTP - Buffer 예약 발행",
      "type": "n8n-nodes-base.httpRequest", "typeVersion": 4, "position": [460, 300],
      "credentials": {"oAuth2Api": {"id": "cred-buffer", "name": "Buffer"}}
    },
    {
      "parameters": {
        "resource": "page", "operation": "update",
        "pageId": "={{ $('Notion - 오늘 발행 콘텐츠').item.json.id }}",
        "propertiesUi": {"propertyValues": [{"key": "상태", "type": "select", "selectValue": "발행완료"}]}
      },
      "id": "q2-0022-03", "name": "Notion - 발행완료 처리",
      "type": "n8n-nodes-base.notion", "typeVersion": 2, "position": [680, 300],
      "credentials": {"notionApi": {"id": "cred-notion", "name": "Notion"}}
    }
  ],
  "connections": {
    "Notion - 오늘 발행 콘텐츠": {"main": [[{"node": "HTTP - Buffer 예약 발행", "type": "main", "index": 0}]]},
    "HTTP - Buffer 예약 발행": {"main": [[{"node": "Notion - 발행완료 처리", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["marketing", "sns", "buffer", "notion", "social-media"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 22,
    'SNS 멀티채널 예약 발행 자동화 (Buffer)',
    '마케팅 / 콘텐츠', '중급(Medium)',
    'Notion, Buffer API',
    'Notion DB 발행예정일 = 오늘 감지',
    'Notion SNS 캘린더 DB에서 오늘 발행 예정인 콘텐츠를 조회하여 Buffer API를 통해 Twitter/X, LinkedIn, Instagram, Facebook에 동시 예약 발행합니다.',
    WF_022,
    [
        'Notion DB: 내용/발행예정일/발행시간/채널/상태 속성 필요',
        'Buffer API v1: OAuth2 앱 등록 후 profile_ids 확인 (Buffer API → /profiles.json)',
        'BUFFER_PROFILE_ID: Buffer 계정에 연결된 각 SNS 채널별 고유 ID',
        'scheduled_at: ISO 8601 형식 (2024-01-15T09:00:00+09:00)',
    ],
    [
        'SNS_CALENDAR_DB_ID를 실제 Notion DB ID로 교체',
        'Buffer OAuth2 앱 등록 및 Access Token 발급',
        'BUFFER_PROFILE_ID_1, _2를 실제 채널 ID로 교체',
        'Buffer API 발행 한도 확인 (플랜별 상이)',
    ],
    [
        'Buffer 대안: Hootsuite, Sprout Social API도 동일 패턴 적용 가능',
        'Instagram 이미지: Buffer API 이미지 파라미터 추가. Notion 첨부파일 URL 활용',
        '해시태그 자동화: OpenAI로 내용 기반 해시태그 자동 생성 후 텍스트에 추가',
    ]
)

WF_023 = json.dumps({
  "name": "유튜브 업로드 → SNS 자동 홍보 발행",
  "nodes": [
    {
      "parameters": {
        "rule": {"interval": [{"field": "cronExpression", "expression": "0 */2 * * *"}]}
      },
      "id": "r3-0023-01", "name": "Schedule - 2시간마다 체크",
      "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]
    },
    {
      "parameters": {
        "resource": "video", "operation": "getAll",
        "channelId": "YOUR_YOUTUBE_CHANNEL_ID",
        "returnAll": False, "limit": 5,
        "additionalFields": {"order": "date"}
      },
      "id": "r3-0023-02", "name": "YouTube - 최신 동영상 조회",
      "type": "n8n-nodes-base.youTube", "typeVersion": 1, "position": [460, 300],
      "credentials": {"youTubeOAuth2Api": {"id": "cred-yt", "name": "YouTube"}}
    },
    {
      "parameters": {
        "jsCode": "const videos = $input.all();\nconst twoHoursAgo = new Date(Date.now() - 2*60*60*1000);\nconst newVideos = videos.filter(v => new Date(v.json.snippet?.publishedAt) > twoHoursAgo);\nreturn newVideos;"
      },
      "id": "r3-0023-03", "name": "Code - 신규 동영상 필터",
      "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [680, 300]
    },
    {
      "parameters": {
        "resource": "text", "operation": "message",
        "model": "gpt-4o-mini",
        "messages": {"values": [{
          "role": "user",
          "content": "유튜브 영상 제목: {{ $json.snippet?.title }}\n설명: {{ $json.snippet?.description?.slice(0,200) }}\n\n이 영상을 홍보하는 SNS 포스트를 작성해주세요:\n- Twitter용 (140자 이내, 해시태그 3개)\n- LinkedIn용 (300자 이내, 전문적 톤)\n각각 구분선 ---으로 구분"
        }]
        }
      },
      "id": "r3-0023-04", "name": "OpenAI - SNS 텍스트 생성",
      "type": "@n8n/n8n-nodes-langchain.openAi", "typeVersion": 1, "position": [900, 300],
      "credentials": {"openAiApi": {"id": "cred-openai", "name": "OpenAI"}}
    },
    {
      "parameters": {
        "resource": "message", "operation": "post",
        "channel": "#content-published",
        "text": "🎬 새 영상 업로드!\n제목: {{ $('Code - 신규 동영상 필터').item.json.snippet?.title }}\nURL: https://youtube.com/watch?v={{ $('Code - 신규 동영상 필터').item.json.id }}\n\nSNS 발행 텍스트:\n{{ $json.message?.content }}",
        "otherOptions": {}
      },
      "id": "r3-0023-05", "name": "Slack - 콘텐츠팀 알림",
      "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [1120, 300],
      "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}
    }
  ],
  "connections": {
    "Schedule - 2시간마다 체크": {"main": [[{"node": "YouTube - 최신 동영상 조회", "type": "main", "index": 0}]]},
    "YouTube - 최신 동영상 조회": {"main": [[{"node": "Code - 신규 동영상 필터", "type": "main", "index": 0}]]},
    "Code - 신규 동영상 필터": {"main": [[{"node": "OpenAI - SNS 텍스트 생성", "type": "main", "index": 0}]]},
    "OpenAI - SNS 텍스트 생성": {"main": [[{"node": "Slack - 콘텐츠팀 알림", "type": "main", "index": 0}]]}
  },
  "active": False, "settings": {},
  "tags": ["marketing", "youtube", "openai", "sns", "content"]
}, ensure_ascii=False, indent=2)

make_recipe(doc, 23,
    '유튜브 신규 업로드 → SNS 자동 홍보 발행',
    '마케팅 / 콘텐츠', '중급(Medium)',
    'YouTube API, OpenAI, Slack',
    'Schedule (2시간마다)',
    '유튜브 채널에 새 영상이 업로드되면 자동으로 감지하고 OpenAI로 Twitter/LinkedIn 맞춤형 홍보 텍스트를 생성하여 콘텐츠팀 Slack에 발송합니다. 한 번 클릭으로 SNS 발행 준비가 완료됩니다.',
    WF_023,
    [
        'YouTube API: OAuth2 또는 API Key. YOUR_YOUTUBE_CHANNEL_ID를 실제 채널 ID로 교체',
        'Code 노드: 2시간 내 업로드 영상만 필터링 (폴링 간격과 동일)',
        'OpenAI: Twitter(140자)와 LinkedIn(300자)용 다른 텍스트 동시 생성',
        'Slack: 생성된 텍스트를 콘텐츠팀이 검토 후 직접 복사하여 SNS 발행',
    ],
    [
        'YouTube Data API v3 활성화 (Google Cloud Console)',
        'YOUR_YOUTUBE_CHANNEL_ID를 실제 채널 ID로 교체 (채널 URL의 @채널명 또는 UC로 시작하는 ID)',
        'OpenAI API Key 설정',
        'Slack Bot Token 및 #content-published 채널 설정',
    ],
    [
        '완전 자동화: Buffer API 연결로 Slack 검토 없이 즉시 SNS 발행 가능 (주의: 검토 필요)',
        'RSS 대안: YouTube RSS Feed (https://www.youtube.com/feeds/videos.xml?channel_id=ID) + RSS Read 노드로 더 단순하게 구현 가능',
    ]
)

# Recipe 024~030 (마케팅 나머지)
marketing_recipes = [
    (24, 'Google Analytics → 주간 성과 Slack 리포트', '마케팅 / 콘텐츠', '중급(Medium)',
     'Google Analytics 4, Slack', 'Schedule (매주 월요일)',
     'Google Analytics 4 API에서 전주 주요 지표(세션, 전환율, 이탈율)를 조회하여 마케팅팀 Slack에 주간 성과 리포트를 자동 발송합니다.'),
    (25, 'A/B 테스트 결과 자동 분석 및 보고', '마케팅 / 콘텐츠', '고급(Hard)',
     'Google Optimize/AB Tasty, Slack, Google Sheets', 'Schedule (테스트 종료 후)',
     'A/B 테스트 실험이 종료되면 자동으로 결과를 분석하고 통계적 유의성을 계산하여 마케팅팀에 리포트를 발송합니다.'),
    (26, 'RSS 피드 → AI 요약 → 뉴스레터 초안 생성', '마케팅 / 콘텐츠', '중급(Medium)',
     'RSS Read, OpenAI, Notion, Gmail', 'Schedule (매주 목요일)',
     '업계 주요 RSS 피드를 자동 수집하고 OpenAI로 핵심 내용을 요약하여 주간 뉴스레터 초안을 Notion에 자동 작성합니다.'),
    (27, 'Webinar 등록자 자동 온보딩 시퀀스', '마케팅 / 콘텐츠', '중급(Medium)',
     'Zoom, Gmail, HubSpot', 'Zoom Webinar Webhook',
     'Zoom 웨비나 등록 시 자동으로 확인 이메일, D-1 리마인더, 당일 참여 링크를 순차 발송하고 HubSpot에 등록 이력을 기록합니다.'),
    (28, '경쟁사 모니터링 및 주간 인사이트 리포트', '마케팅 / 콘텐츠', '고급(Hard)',
     'HTTP Request, OpenAI, Slack, Google Sheets', 'Schedule (매주)',
     '경쟁사 웹사이트/블로그/채용 공고를 정기 스크래핑하고 OpenAI로 주요 변화를 분석하여 마케팅/전략팀에 주간 인사이트를 발송합니다.'),
    (29, '이메일 캠페인 성과 자동 집계 (Mailchimp)', '마케팅 / 콘텐츠', '초급(Easy)',
     'Mailchimp, Google Sheets, Slack', 'Schedule (캠페인 발송 후 48시간)',
     'Mailchimp 캠페인 발송 48시간 후 자동으로 오픈율, 클릭율, 구독 취소율을 집계하여 Google Sheets에 기록하고 Slack으로 요약 리포트를 보냅니다.'),
    (30, 'Product Hunt 런칭 자동화 키트', '마케팅 / 콘텐츠', '고급(Hard)',
     'Product Hunt API, Twitter/X, LinkedIn, Slack', 'Schedule (런칭 당일)',
     'Product Hunt 런칭 당일 자동으로 지지 요청 이메일 발송, SNS 홍보 포스트 게시, 팀 Slack 알림을 조율하는 런칭 자동화 파이프라인입니다.'),
]

for num, title, cat, diff, apps, trigger, desc in marketing_recipes:
    simple_wf = json.dumps({
        "name": title, "nodes": [
            {"parameters": {"rule": {"interval": [{"field": "cronExpression", "expression": "0 9 * * 1"}]}},
             "id": f"mkt-{num:04d}-01", "name": "Schedule / Webhook 트리거",
             "type": "n8n-nodes-base.scheduleTrigger", "typeVersion": 1, "position": [240, 300]},
            {"parameters": {"jsCode": "// 데이터 처리 로직\nreturn $input.all();"},
             "id": f"mkt-{num:04d}-02", "name": "Code - 데이터 처리",
             "type": "n8n-nodes-base.code", "typeVersion": 2, "position": [460, 300]},
            {"parameters": {"resource": "message", "operation": "post",
                            "channel": "#marketing", "text": f"완료: {title}", "otherOptions": {}},
             "id": f"mkt-{num:04d}-03", "name": "Slack - 마케팅팀 알림",
             "type": "n8n-nodes-base.slack", "typeVersion": 2, "position": [680, 300],
             "credentials": {"slackApi": {"id": "cred-slack", "name": "Slack"}}}
        ],
        "connections": {
            "Schedule / Webhook 트리거": {"main": [[{"node": "Code - 데이터 처리", "type": "main", "index": 0}]]},
            "Code - 데이터 처리": {"main": [[{"node": "Slack - 마케팅팀 알림", "type": "main", "index": 0}]]}
        },
        "active": False, "settings": {}, "tags": ["marketing"]
    }, ensure_ascii=False, indent=2)
    make_recipe(doc, num, title, cat, diff, apps, trigger, desc,
        simple_wf,
        ['트리거 노드를 실제 서비스 트리거로 교체', '데이터 처리 로직은 API 응답 구조에 따라 Code 노드에 구현', '출력 채널(Slack/이메일/Notion)을 목적에 맞게 설정'],
        ['필요한 서비스 API Key/OAuth2 설정', '스케줄 또는 웹훅 트리거 선택', '알림 채널 설정'],
        ['각 서비스 공식 n8n 통합 페이지에서 상세 파라미터 확인 권장']
    )

doc.add_page_break()
doc.save(OUT)
print('Vol3 Part2 done:', OUT)
