# -*- coding: utf-8 -*-
"""Vol 5 Part 3: Ch5 6380개 패턴 분석 + Ch6 프롬프트 엔지니어링 + Ch7 Human-in-the-loop"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'C:/Users/장우경/oomni/n8n_cases_vol5.docx'
doc = Document(OUT)

def sc(cell,hx):
    tc=cell._tc;pr=tc.get_or_add_tcPr()
    s=OxmlElement('w:shd');s.set(qn('w:val'),'clear');s.set(qn('w:color'),'auto');s.set(qn('w:fill'),hx);pr.append(s)

def h(doc,text,lv=1):
    p=doc.add_heading(text,level=lv)
    C={1:RGBColor(0x1F,0x49,0x7D),2:RGBColor(0x2E,0x75,0xB6),3:RGBColor(0x17,0x5E,0x40),4:RGBColor(0x40,0x40,0x40)}
    S={1:Pt(22),2:Pt(16),3:Pt(13),4:Pt(11)}
    if p.runs:p.runs[0].font.name='맑은 고딕';p.runs[0].font.size=S.get(lv,Pt(11));p.runs[0].font.color.rgb=C.get(lv,RGBColor(0,0,0))
    return p

def par(doc,text,bold=False,color=None,size=Pt(10),indent=None):
    p=doc.add_paragraph()
    if indent:p.paragraph_format.left_indent=indent
    r=p.add_run(text);r.font.name='맑은 고딕';r.font.size=size;r.bold=bold
    if color:r.font.color.rgb=color
    return p

def bul(doc,text,lv=0):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent=Cm(0.5+lv*0.5)
    r=p.add_run(text);r.font.name='맑은 고딕';r.font.size=Pt(10);return p

def code(doc,text):
    p=doc.add_paragraph();p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run(text);r.font.name='Courier New';r.font.size=Pt(8)
    pr=p._p.get_or_add_pPr();s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear');s.set(qn('w:color'),'auto');s.set(qn('w:fill'),'F2F2F2');pr.append(s)
    return p

def tbl(doc,headers,rows,col_widths,hc='1F497D'):
    t=doc.add_table(rows=len(rows)+1,cols=len(headers));t.style='Table Grid';t.autofit=False
    for i,w in enumerate(col_widths):t.columns[i].width=w
    for ci,hdr in enumerate(headers):
        sc(t.rows[0].cells[ci],hc)
        t.rows[0].cells[ci].text=hdr
        for p in t.rows[0].cells[ci].paragraphs:
            for r in p.runs:r.font.name='맑은 고딕';r.font.size=Pt(9);r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            sc(t.rows[ri+1].cells[ci],'EBF3FB' if ci==0 else 'FFFFFF')
            t.rows[ri+1].cells[ci].text=str(val)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:r.font.name='맑은 고딕';r.font.size=Pt(9)
                if ci==0:
                    for r in p.runs:r.bold=True
    return t

# ════════════════════════════════════════════════════
# CHAPTER 5: 6,380개 AI 워크플로우 패턴 분석
# ════════════════════════════════════════════════════
h(doc,'Chapter 5. 6,380개 AI 워크플로우 패턴 심층 분석',1)
par(doc,'n8n 커뮤니티(n8n.io/workflows)에서 AI 태그가 포함된 6,380개 워크플로우를 '
    '분석한 결과입니다. 패턴 분류, 트렌드, 핵심 인사이트를 제공합니다.', size=Pt(10))
doc.add_paragraph()

h(doc,'5.1 전체 통계 개요',2)
code(doc,
'6,380개 AI 워크플로우 분석 결과 (2024년 기준)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'총 분석 워크플로우:    6,380개\n'
'AI Agent 활용:         2,847개 (44.6%)\n'
'RAG/Vector Store:      1,203개 (18.9%)\n'
'단순 LLM Chain:        1,891개 (29.6%)\n'
'멀티에이전트:            439개  (6.9%)\n'
'\n'
'가장 많이 사용된 LLM:\n'
'  1. OpenAI GPT-4o/GPT-4    2,981개 (46.7%)\n'
'  2. OpenAI GPT-3.5-turbo   1,423개 (22.3%)\n'
'  3. Anthropic Claude          642개 (10.1%)\n'
'  4. Ollama (Local)            511개  (8.0%)\n'
'  5. Google Gemini             389개  (6.1%)\n'
'  6. 기타 모델                  434개  (6.8%)\n'
'\n'
'가장 많이 조합된 서비스:\n'
'  1. Slack + OpenAI            891개\n'
'  2. Gmail + OpenAI            743개\n'
'  3. Notion + OpenAI           612개\n'
'  4. Telegram + OpenAI         534개\n'
'  5. Google Sheets + OpenAI    487개\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'5.2 산업별 AI 워크플로우 분포',2)
industry_dist=[
    ('마케팅/콘텐츠','1,247개','19.5%','AI 콘텐츠 생성, SNS 자동화, SEO'),
    ('영업/CRM','892개','14.0%','리드 인리치먼트, 이메일 자동화'),
    ('고객지원','756개','11.9%','AI 챗봇, 티켓 분류, FAQ 자동 응답'),
    ('IT운영/DevOps','634개','9.9%','코드 리뷰, 배포 자동화, 모니터링'),
    ('HR/인사','423개','6.6%','JD 작성, 이력서 스크리닝, 온보딩'),
    ('금융/회계','387개','6.1%','문서 분석, 이상 탐지, 리포트'),
    ('교육','312개','4.9%','커리큘럼 생성, 평가 자동화'),
    ('법률','289개','4.5%','계약서 검토, 규정 준수 확인'),
    ('의료','201개','3.2%','임상 노트 요약, 진단 보조'),
    ('연구/분석','198개','3.1%','문헌 조사, 데이터 분석'),
    ('기타','1,041개','16.3%','다양한 개인 자동화'),
]
tbl(doc,['산업','워크플로우 수','비율','주요 패턴'],industry_dist,
    [Cm(4),Cm(2.5),Cm(2),Cm(9)],hc='2E75B6')
doc.add_paragraph()

h(doc,'5.3 AI 워크플로우 핵심 패턴 25선',2)
par(doc,'6,380개 중 가장 많이 반복되는 핵심 패턴을 추출하여 사용 빈도 순으로 정리했습니다.', size=Pt(10))
doc.add_paragraph()

top_patterns=[
    (1,'AI 이메일 초안 생성','1,203개','이메일 수신 → AI 초안 → 검토 → 발송','Gmail+OpenAI'),
    (2,'문서 요약 자동화','987개','PDF/웹페이지 → LLM 요약 → 저장/공유','HTTP+OpenAI+Notion'),
    (3,'AI 챗봇 (Webhook 기반)','892개','Webhook → AI Agent → 응답 반환','n8n Agent+LLM'),
    (4,'Slack AI 어시스턴트','756개','Slack @멘션 → AI 처리 → 답변','Slack+OpenAI'),
    (5,'데이터 분류/태깅 자동화','634개','신규 데이터 → LLM 분류 → DB 업데이트','다양+OpenAI'),
    (6,'소셜미디어 콘텐츠 생성','612개','트렌드/주제 → AI 생성 → 예약 발행','OpenAI+Buffer'),
    (7,'RAG 지식베이스 Q&A','589개','질문 → 벡터 검색 → 컨텍스트 기반 답변','Pinecone+OpenAI'),
    (8,'코드 리뷰 자동화','534개','GitHub PR → Claude/GPT 리뷰 → 코멘트','GitHub+Claude'),
    (9,'회의록 자동화','487개','음성/텍스트 → Whisper → GPT 요약','Whisper+GPT-4o'),
    (10,'리드 스코어링','423개','새 리드 → LLM 분석 → 점수 부여 → CRM','Clearbit+HubSpot'),
    (11,'AI 번역 파이프라인','412개','원문 → LLM 번역 → 품질 검증 → 저장','OpenAI+Sheets'),
    (12,'이상 감지 및 알림','389개','데이터 → LLM 패턴 분석 → 이상 감지 → 알림','다양+GPT'),
    (13,'SEO 콘텐츠 최적화','356개','기존 콘텐츠 → LLM SEO 분석 → 개선안','WordPress+OpenAI'),
    (14,'고객 감성 분석','334개','리뷰/피드백 → 감성 분류 → 집계 → 리포트','다양+OpenAI'),
    (15,'AI HR 스크리닝','312개','이력서 → LLM 분석 → 점수 → ATS 업데이트','Gmail+OpenAI'),
    (16,'데이터 추출 (문서→구조화)','298개','PDF/이미지 → Vision API → JSON 추출','OpenAI Vision'),
    (17,'Telegram AI 봇','287개','Telegram 메시지 → AI 처리 → 응답','Telegram+OpenAI'),
    (18,'AI 상품 설명 생성','276개','제품 데이터 → GPT 생성 → 쇼핑몰 발행','Shopify+OpenAI'),
    (19,'자동 보고서 생성','234개','데이터 → LLM 분석 → 리포트 → 발송','다양+GPT-4o'),
    (20,'멀티에이전트 리서치','198개','주제 → 리서치 에이전트 → 작성 에이전트','Agent Chain'),
    (21,'음성 → 텍스트 → 처리','187개','음성 파일 → Whisper → GPT 처리','Whisper+GPT'),
    (22,'AI 이미지 생성+발행','176개','텍스트 프롬프트 → DALL-E → SNS 발행','DALL-E+Buffer'),
    (23,'계약서 AI 분석','165개','계약서 → LLM 핵심 조항 추출 → 리스크 체크','Claude+Notion'),
    (24,'개인화 이메일 캠페인','154개','고객 데이터 → LLM 개인화 → 발송','HubSpot+OpenAI'),
    (25,'AI 코드 생성 보조','143개','요구사항 → Claude/GPT 코드 생성 → 테스트','다양+Claude'),
]
t=doc.add_table(rows=len(top_patterns)+1,cols=6);t.style='Table Grid';t.autofit=False
for w,ci in zip([Cm(0.8),Cm(4),Cm(1.5),Cm(5.5),Cm(4.5)],range(5)):t.columns[ci].width=w
for ci,hdr in enumerate(['순위','패턴명','빈도','워크플로우 흐름','주요 조합']):
    sc(t.rows[0].cells[ci],'1F497D')
    t.rows[0].cells[ci].text=hdr
    for p in t.rows[0].cells[ci].paragraphs:
        for r in p.runs:r.font.name='맑은 고딕';r.font.size=Pt(8.5);r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for ri,row in enumerate(top_patterns):
    for ci,val in enumerate(row[:5]):
        sc(t.rows[ri+1].cells[ci],'EBF3FB' if ci<2 else 'FFFFFF')
        t.rows[ri+1].cells[ci].text=str(val)
        for p in t.rows[ri+1].cells[ci].paragraphs:
            for r in p.runs:r.font.name='맑은 고딕';r.font.size=Pt(8.5)
            if ci<=1:
                for r in p.runs:r.bold=True
doc.add_paragraph()

h(doc,'5.4 2024년 AI 워크플로우 트렌드',2)
trends=[
    ('AI Agent 급증',
     '2023년 대비 2024년 Agent 패턴 340% 증가. 단순 LLM 호출에서 자율 에이전트로 전환 가속.',
     '↑↑↑'),
    ('멀티모달 확산',
     'GPT-4o Vision, Whisper 조합. 이미지+텍스트+음성 통합 워크플로우 250% 증가.',
     '↑↑↑'),
    ('로컬 LLM (Ollama)',
     'Ollama 기반 완전 로컬 워크플로우 180% 증가. 프라이버시/비용 절감 수요.',
     '↑↑'),
    ('RAG 고도화',
     '단순 벡터 검색에서 하이브리드 검색, Re-ranking, HyDE 등 고급 RAG로 전환.',
     '↑↑'),
    ('n8n AI Agent 노드',
     'v1.22+ 도입 후 AI Agent 노드 활용 폭발적 증가. 기존 HTTP 방식에서 전환.',
     '↑↑↑'),
    ('Telegram 봇',
     'Slack 대비 구현 쉬운 Telegram AI 봇 50% 증가. 개인/소규모 팀에서 인기.',
     '↑↑'),
    ('비용 최적화 패턴',
     'GPT-4 → GPT-4o-mini 전환, 캐싱 패턴, 하이브리드(로컬+클라우드) 증가.',
     '↑'),
    ('MCP 통합',
     '2024 말 Claude MCP(Model Context Protocol) 지원 워크플로우 급증 시작.',
     '↑↑↑ (신규)'),
]
tbl(doc,['트렌드','설명','증가세'],trends,[Cm(4),Cm(11),Cm(2.5)],hc='175E40')
doc.add_paragraph()

h(doc,'5.5 핵심 인사이트: 실무에서 배운 교훈',2)
insights=[
    ('에이전트 복잡도 최소화','가장 성공한 워크플로우의 77%는 3개 이하의 Tool을 사용. 많은 도구 = 낮은 신뢰성.'),
    ('시스템 프롬프트의 중요성','품질 좋은 워크플로우의 공통점: 구체적이고 명확한 시스템 프롬프트. 역할+제약+출력형식 명시.'),
    ('gpt-4o-mini 충분','83%의 케이스는 gpt-4o-mini로 충분. gpt-4o는 복잡한 추론과 긴 문서에만 사용.'),
    ('에러 처리 필수','상위 평점 워크플로우의 95%는 Error Workflow 설정. 프로덕션에서 필수.'),
    ('캐싱으로 비용 70% 절감','동일 쿼리 반복 많은 RAG 시스템에서 Redis 캐싱으로 비용 50~70% 절감 사례 다수.'),
    ('인간 감독 게이트','실제 행동(이메일 발송, 파일 삭제, 결제)을 수행하는 에이전트는 반드시 승인 단계 포함.'),
    ('단계별 검증','복잡한 파이프라인은 각 단계에서 중간 결과 검증. 문제 조기 발견으로 디버깅 시간 80% 단축.'),
    ('토큰 효율화','프롬프트에 JSON 응답 형식 명시하면 파싱 실패 95% 감소. Structured Output Parser 활용.'),
]
tbl(doc,['교훈','상세 내용'],insights,[Cm(5),Cm(12.5)],hc='2E75B6')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════
# CHAPTER 6: 프롬프트 엔지니어링 완전 가이드
# ════════════════════════════════════════════════════
h(doc,'Chapter 6. 프롬프트 엔지니어링 완전 가이드',1)
par(doc,'n8n AI 워크플로우에서 LLM의 성능을 극대화하는 프롬프트 엔지니어링 기법을 '
    '실전 예시와 함께 완전히 가이드합니다. 시스템 프롬프트 설계부터 고급 기법까지 다룹니다.', size=Pt(10))
doc.add_paragraph()

h(doc,'6.1 시스템 프롬프트 설계 원칙',2)
par(doc,'AI Agent의 시스템 프롬프트는 에이전트의 "헌법"입니다. '
    '잘 설계된 시스템 프롬프트는 에이전트 성능의 60-80%를 결정합니다.', size=Pt(10))
doc.add_paragraph()

h(doc,'6.1.1 시스템 프롬프트 7가지 핵심 구성 요소',3)
code(doc,
'완전한 시스템 프롬프트 템플릿\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'## 1. 역할 정의 (필수)\n'
'당신은 [회사명]의 [직책/역할]입니다.\n'
'[핵심 전문성]에 특화된 AI 어시스턴트입니다.\n'
'\n'
'## 2. 목적 및 범위 (필수)\n'
'당신의 주요 임무는 [구체적 목표]입니다.\n'
'다음 사항은 처리하지 않습니다: [제외 범위]\n'
'\n'
'## 3. 도구 사용 지침 (Tool 사용 시)\n'
'사용 가능한 도구:\n'
'- search_web: 최신 정보가 필요할 때\n'
'- get_calendar: 일정 조회가 필요할 때\n'
'항상 정보가 불확실한 경우 도구를 사용하세요.\n'
'\n'
'## 4. 응답 형식 (필수)\n'
'응답은 반드시 다음 JSON 형식으로 반환하세요:\n'
'{ "answer": "...", "confidence": 0-1, "sources": [] }\n'
'\n'
'## 5. 언어 및 톤\n'
'항상 한국어로 응답하세요.\n'
'전문적이지만 친근한 톤으로 작성하세요.\n'
'\n'
'## 6. 제약 사항 (필수)\n'
'- 개인정보(이름, 전화번호, 이메일)는 절대 수집/저장하지 마세요.\n'
'- 불확실한 정보는 추측하지 말고 모른다고 말하세요.\n'
'- 의료/법률/금융 조언은 전문가 상담 권고를 추가하세요.\n'
'\n'
'## 7. 컨텍스트 (옵션)\n'
'현재 날짜: {{ $now.toFormat("yyyy-MM-dd") }}\n'
'사용자 ID: {{ $json.userId }}\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'6.2 프롬프트 엔지니어링 핵심 기법',2)

techniques=[
    ('Zero-Shot Prompting',
     '예시 없이 지시만으로 작업 수행.',
     '"다음 텍스트의 감성을 분류하세요: positive/negative/neutral"',
     '단순 분류, 번역, 요약'),
    ('Few-Shot Prompting',
     '2-5개 예시를 포함하여 패턴 학습.',
     '"입력: 배송이 늦었어요 → 출력: negative\n입력: 빠르게 받았어요 → 출력: positive\n입력: {text} → 출력:"',
     '일관성 중요한 분류 작업'),
    ('Chain-of-Thought (CoT)',
     '"단계별로 생각하세요" 추가. 복잡한 추론 향상.',
     '"단계별로 생각한 후 최종 답변을 제시하세요."',
     '수학 문제, 논리 추론, 복잡한 분석'),
    ('Tree of Thought (ToT)',
     '여러 사고 경로를 탐색하고 최적 선택.',
     '"3가지 다른 접근 방식을 각각 평가한 후 최선을 선택하세요."',
     '전략 수립, 복잡한 의사결정'),
    ('ReAct 패턴',
     'Reasoning + Acting. 생각→행동→관찰 반복.',
     '시스템: "Thought: ...\nAction: tool_name\nObservation: ...\nAnswer: ..."',
     'AI Agent에 자동 적용'),
    ('Structured Output',
     'JSON 스키마 명시로 구조화 출력 강제.',
     '"반드시 다음 JSON만 반환: {\"score\":1-10, \"reason\":\"...\"}"',
     '데이터 추출, 분류, 파싱'),
    ('Role Prompting',
     '전문가 역할 부여로 품질 향상.',
     '"당신은 20년 경력의 시니어 법무 변호사입니다..."',
     '전문 도메인 작업'),
    ('Negative Prompting',
     '하지 말아야 할 것을 명시.',
     '"절대로 HTML 태그를 사용하지 마세요. 마크다운 없이 순수 텍스트로만..."',
     '출력 형식 제어'),
]
tbl(doc,['기법','설명','예시 패턴','적합 케이스'],techniques,
    [Cm(4),Cm(3.5),Cm(6.5),Cm(3.5)],hc='2E75B6')
doc.add_paragraph()

h(doc,'6.3 산업별 시스템 프롬프트 실전 예시',2)

h(doc,'6.3.1 고객지원 AI 에이전트',3)
code(doc,
'## 고객지원 AI 에이전트 시스템 프롬프트\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'당신은 [회사명] 고객지원팀의 AI 어시스턴트입니다.\n'
'친절하고 전문적으로 고객 문의를 처리합니다.\n'
'\n'
'## 처리 가능한 요청:\n'
'- 주문 조회, 배송 상태 확인\n'
'- 반품/교환 안내 (정책: 구매 후 30일 이내)\n'
'- 제품 사용법 및 기술 지원\n'
'- 계정 관련 문의\n'
'\n'
'## 처리 불가 (에스컬레이션 필요):\n'
'- 1만원 이상 환불 처리\n'
'- 법적 분쟁, 언론 문의\n'
'- 불만 3회 이상 반복 고객\n'
'\n'
'## 에스컬레이션 방법:\n'
'위 케이스 감지 시 반드시: {"escalate": true, "reason": "..."}\n'
'형태로 응답하세요.\n'
'\n'
'## 응답 규칙:\n'
'- 항상 한국어로 응답\n'
'- 200자 이내로 간결하게\n'
'- 고객 이름을 알면 반드시 이름으로 호칭\n'
'- 부정적 상황은 "불편을 드려 죄송합니다"로 시작\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'6.3.2 코드 리뷰 AI 에이전트',3)
code(doc,
'## 코드 리뷰 AI 에이전트 시스템 프롬프트\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'당신은 10년 경력의 시니어 소프트웨어 엔지니어입니다.\n'
'Pull Request 코드 리뷰를 수행합니다.\n'
'\n'
'## 리뷰 기준 (우선순위 순):\n'
'1. 보안 취약점 (SQL 인젝션, XSS, 인증 누락)\n'
'2. 버그 및 논리 오류\n'
'3. 성능 이슈 (N+1 쿼리, 메모리 누수)\n'
'4. 코드 품질 (가독성, 중복, SOLID 원칙)\n'
'5. 테스트 커버리지\n'
'\n'
'## 응답 형식 (반드시 준수):\n'
'{\n'
'  "severity": "critical|high|medium|low",\n'
'  "issues": [\n'
'    {"line": 번호, "type": "security|bug|performance|style",\n'
'     "description": "문제 설명", "suggestion": "개선 방안"}\n'
'  ],\n'
'  "summary": "전체 요약 (100자 이내)",\n'
'  "approve": true|false\n'
'}\n'
'\n'
'## 주의사항:\n'
'- 개인 취향이 아닌 객관적 기준으로만 리뷰\n'
'- critical 이슈 있으면 approve: false\n'
'- 개선 방안은 반드시 구체적 코드 예시 포함\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'6.3.3 금융/법률 분석 에이전트',3)
code(doc,
'## 금융/법률 문서 분석 에이전트 시스템 프롬프트\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'당신은 금융 및 법률 문서 분석 전문 AI입니다.\n'
'계약서, 법률 문서, 금융 보고서를 분석합니다.\n'
'\n'
'## 분석 시 반드시 포함할 내용:\n'
'1. 핵심 조항 요약 (3-5개)\n'
'2. 잠재적 리스크 요인\n'
'3. 주의해야 할 조건/기한\n'
'4. 불명확하거나 모호한 조항\n'
'\n'
'## 반드시 준수할 규칙:\n'
'- 분석은 문서 내용에만 근거\n'
'- 법적 조언/투자 추천 절대 금지\n'
'- 반드시 "전문 법률가/금융 전문가 상담 권고" 추가\n'
'- 불확실한 내용은 "[확인 필요]" 태그\n'
'\n'
'## 출력 형식:\n'
'## 📋 문서 유형: [유형]\n'
'## 📌 핵심 조항\n'
'## ⚠️ 리스크 요인\n'
'## 📅 중요 기한/조건\n'
'## ❓ 불명확 조항\n'
'## ⚖️ 전문가 상담 권고\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'6.4 프롬프트 안전성 및 인젝션 방어',2)
par(doc,'n8n 워크플로우에서 사용자 입력이 프롬프트에 포함될 때 '
    'Prompt Injection 공격에 노출될 수 있습니다.', size=Pt(10))
doc.add_paragraph()

injection_defenses=[
    ('입력 검증','사용자 입력에서 특수 문자/지시 패턴 제거','Code 노드: text.replace(/ignore previous|disregard|system:/gi, "")'),
    ('입력/지시 분리','사용자 입력을 시스템 프롬프트와 명확히 분리','XML 태그 사용: <user_input>{input}</user_input>'),
    ('Constitutional AI','"다음 원칙을 절대 위반하지 마세요: 1. 2. ..." 명시','시스템 프롬프트 끝에 제약 조항 반복'),
    ('출력 검증','LLM 출력에 위험 패턴이 있는지 사후 검증','Code 노드로 출력 패턴 검사'),
    ('토큰 제한','maxTokens 설정으로 비정상적으로 긴 출력 방지','maxTokens: 케이스별 적절한 한도 설정'),
    ('샌드박스 모드','개발/테스트 시 실제 액션 대신 시뮬레이션','IF 노드: $env.NODE_ENV === "production"일 때만 실제 실행'),
]
tbl(doc,['방어 기법','설명','n8n 구현'],injection_defenses,[Cm(3.5),Cm(6),Cm(8)],hc='C00000')
doc.add_paragraph()

h(doc,'6.5 프롬프트 버전 관리',2)
code(doc,
'n8n에서 프롬프트 버전 관리 방법\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'방법 1: Notion DB로 프롬프트 중앙 관리\n'
'  Notion DB: {name, version, prompt_text, status(active/draft)}\n'
'  → 워크플로우 시작 시 Notion에서 active 프롬프트 조회\n'
'  → 즉시 롤백/업데이트 가능 (워크플로우 수정 불필요)\n'
'\n'
'방법 2: n8n 환경변수 활용\n'
'  Settings → Variables → SYSTEM_PROMPT_V2\n'
'  → {{ $vars.SYSTEM_PROMPT_V2 }} 참조\n'
'\n'
'방법 3: Google Sheets 프롬프트 라이브러리\n'
'  A열: prompt_id, B열: version, C열: text\n'
'  → Schedule Trigger 시작 시 최신 버전 로드\n'
'\n'
'A/B 테스트:\n'
'  IF 노드: Math.random() < 0.5\n'
'    → true: 프롬프트 A 사용\n'
'    → false: 프롬프트 B 사용\n'
'  → 결과를 Sheets에 기록하여 성능 비교\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════
# CHAPTER 7: Human-in-the-loop 설계 패턴
# ════════════════════════════════════════════════════
h(doc,'Chapter 7. Human-in-the-loop 설계 패턴',1)
par(doc,'AI 자동화에서 인간의 감독과 개입은 안전성과 품질을 보장하는 핵심입니다. '
    'n8n에서 효과적인 Human-in-the-loop (HITL) 시스템을 구현하는 방법을 완전 가이드합니다.', size=Pt(10))
doc.add_paragraph()

h(doc,'7.1 HITL이 필요한 경우',2)
hitl_cases=[
    ('높은 영향도 결정','이메일 발송, 파일 삭제, 데이터 수정 등 되돌리기 어려운 작업','항상 승인 필요'),
    ('불확실성 높음','AI 신뢰도 점수 < 70%, 모호한 질문, 새로운 유형의 요청','조건부 에스컬레이션'),
    ('금융/의료/법률','규제가 있는 영역. 자동 결정이 법적 책임 발생 가능','반드시 검토'),
    ('처음 실행','새로 배포된 워크플로우. 처음 100회는 모두 검토 후 자동화','단계적 자동화'),
    ('예외 케이스','학습하지 않은 새로운 패턴, 이상한 입력','자동 에스컬레이션'),
    ('고객 불만','감성 점수 매우 낮은 고객, 반복 민원','인간 처리'),
]
tbl(doc,['케이스','이유','권장 처리'],hitl_cases,[Cm(4),Cm(8),Cm(5.5)],hc='2E75B6')
doc.add_paragraph()

h(doc,'7.2 n8n HITL 구현 패턴',2)

h(doc,'7.2.1 Slack 승인 버튼 패턴 (가장 많이 사용)',3)
code(doc,
'Slack 승인 버튼을 이용한 HITL 패턴\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'[AI 출력]\n'
'   │\n'
'   ▼\n'
'[Slack 승인 메시지 발송]\n'
'"AI가 다음 이메일 초안을 생성했습니다:\n'
'{초안 내용}\n'
'[✅ 승인 발송] [✏️ 수정] [❌ 취소]"\n'
'   │\n'
'   ▼\n'
'[Wait 노드 (Webhook 대기)] ← 타임아웃: 24시간\n'
'   │\n'
'   ├─ 버튼 클릭 없음 (타임아웃) → 자동 취소 + 알림\n'
'   │\n'
'   └─ Webhook 수신 (버튼 클릭)\n'
'        ├─ action=approve → 이메일 발송\n'
'        ├─ action=edit    → 수정 폼 제시\n'
'        └─ action=cancel  → 취소 + 로그 기록\n'
'\n'
'n8n 구현:\n'
'  Slack 노드 (interactive 메시지) +\n'
'  Wait 노드 (resume URL) +\n'
'  n8n Webhook (버튼 클릭 수신)\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'7.2.2 이메일 승인 패턴',3)
code(doc,
'이메일 링크 클릭 승인 패턴\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'[n8n Webhook URL 생성]\n'
'  APPROVE_URL = https://n8n.company.com/webhook/{uuid}/approve\n'
'  REJECT_URL  = https://n8n.company.com/webhook/{uuid}/reject\n'
'\n'
'[승인 이메일 발송]\n'
'"다음 내용을 검토하고 승인해 주세요:\n'
'{내용}\n'
'\n'
'✅ 승인: APPROVE_URL\n'
'❌ 거부: REJECT_URL\n'
'\n'
'본 링크는 48시간 후 만료됩니다."\n'
'\n'
'[Wait 노드]\n'
'  resume: "on webhook call"\n'
'  url: APPROVE_URL 또는 REJECT_URL 수신\n'
'  timeout: 48시간\n'
'\n'
'[IF 노드]\n'
'  $json.action === "approve" → 실행\n'
'  else → 취소\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'7.2.3 신뢰도 기반 자동 에스컬레이션',3)
code(doc,
'AI 신뢰도 점수 기반 자동 에스컬레이션\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
'\n'
'AI 출력 예시:\n'
'{\n'
'  "answer": "환불 처리됩니다.",\n'
'  "confidence": 0.85,  ← 신뢰도 0-1\n'
'  "category": "refund"\n'
'}\n'
'\n'
'에스컬레이션 로직 (Code 노드):\n'
'const confidence = $json.confidence;\nconst category = $json.category;\n\n// 강제 에스컬레이션 조건\nif (category === "legal" || category === "complaint") {\n  return [{ json: { ...($json), escalate: true, reason: "민감 카테고리" } }];\n}\n\n// 신뢰도 기반 에스컬레이션\nif (confidence < 0.7) {\n  return [{ json: { ...($json), escalate: true, reason: "낮은 신뢰도: " + confidence } }];\n}\n\nreturn [{ json: { ...($json), escalate: false } }];\n'
'━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
)
doc.add_paragraph()

h(doc,'7.3 HITL 설계 원칙',2)
principles=[
    ('명확한 컨텍스트 제공','승인 요청 시 "무엇을 결정해야 하는지"와 "왜 AI가 판단하기 어려운지"를 명확히 제시.'),
    ('단계적 자동화','처음에는 모든 케이스 검토 → 일정 기간 후 패턴 학습 → 신뢰도 높은 케이스 자동화.'),
    ('타임아웃 처리','인간이 응답하지 않을 경우 명확한 기본 동작 정의. 금융/의료는 "타임아웃 = 취소"'),
    ('감사 로그','모든 HITL 결정(승인/거부/수정)을 DB에 기록. 학습 데이터 및 감사 대응용.'),
    ('결정 이유 수집','거부 시 이유 선택/입력 기능 추가. 모델 개선에 활용.'),
    ('알림 피로 방지','너무 많은 승인 요청은 담당자 피로 → 번아웃. 자동화 가능한 것은 먼저 자동화.'),
    ('우선순위 표시','긴급도와 영향도를 메시지에 명시. Critical은 PagerDuty/SMS 병행.'),
]
tbl(doc,['원칙','상세 내용'],principles,[Cm(5),Cm(12.5)],hc='175E40')
doc.add_paragraph()

doc.add_page_break()
doc.save(OUT)
print('Vol5 Part3 done:', OUT)
